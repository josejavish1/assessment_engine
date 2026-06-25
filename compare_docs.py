
from docx import Document
import sys

def get_doc_info(path):
    try:
        doc = Document(path)
        print(f"--- Document: {path} ---")
        print(f"Paragraphs: {len(doc.paragraphs)}")
        print(f"Sections: {len(doc.sections)}")
        print(f"Tables: {len(doc.tables)}")
        print("\nText Sample (First 2000 chars):")
        full_text = "\n".join([p.text for p in doc.paragraphs])
        print(full_text[:2000])
        print("\n" + "="*50 + "\n")
    except Exception as e:
        print(f"Error reading {path}: {e}")

if __name__ == "__main__":
    for path in sys.argv[1:]:
        get_doc_info(path)
