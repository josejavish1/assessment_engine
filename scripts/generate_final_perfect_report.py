import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


def set_cell_shading(cell, color):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def generate_final_perfect_report():
    json_path = Path("working/eurovision/client_intelligence.json")
    with open(json_path, "r", encoding="utf-8-sig") as f:
        data = json.load(f)

    doc = Document()

    # --- CABECERA CORPORATIVA ---
    title = doc.add_heading("INFORME TÉCNICO-ESTRATÉGICO: EUROVISION SERVICES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(
        "AUDITORÍA DE HALLAZGOS PARA RESPUESTA A RFP", 1
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")
    meta = doc.add_paragraph()
    meta.add_run("Cliente: ").bold = True
    meta.add_run("Eurovision Services (DUBAG Group)\n")
    meta.add_run("Asunto: ").bold = True
    meta.add_run("Inteligencia de Modernización de Infraestructura\n")
    meta.add_run("Fecha: ").bold = True
    meta.add_run("08 de junio de 2026\n")
    meta.add_run("Nivel: ").bold = True
    meta.add_run("CONFIDENCIAL / PARTNER NTT DATA ONLY").font.color.rgb = RGBColor(
        200, 0, 0
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

    # --- 1. ESTRATEGIA Y LIDERAZGO ---
    doc.add_heading("1. Estrategia y Agenda de Liderazgo", level=1)
    doc.add_paragraph(
        "Eurovision Services opera actualmente como una entidad privada bajo el fondo DUBAG Group desde 2023. "
        "Su agenda está marcada por la transición hacia la independencia total de la EBU y la expansión de su cartera global."
    )

    agenda = data.get("business_context", {}).get("ceo_agenda", {}).get("summary", "")
    # Clean text
    clean_agenda = agenda.replace("###", "").replace("**", "").replace("*", "•")
    doc.add_paragraph(clean_agenda)

    # --- 2. HALLAZGOS TÉCNICOS (THE CORE STACK) ---
    doc.add_heading("2. Mapa Tecnológico Forense (Technical Inventory)", level=1)
    doc.add_paragraph(
        "Este inventario representa el stack tecnológico real detectado en las infraestructuras de Eurovision Services:"
    )

    stack = data.get("technology_context", {}).get("group_level_stack", [])
    if stack:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Tecnología / Dominio"
        hdr[1].text = "Fabricante / Socio"
        hdr[2].text = "Fuente de Verificación"
        for cell in hdr:
            set_cell_shading(cell, "4F81BD")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for s in stack:
            row = table.add_row().cells
            row[0].text = str(s.get("technology", "N/A"))
            row[1].text = str(s.get("vendor", "N/A"))
            row[2].text = s.get("source", {}).get("uri", "N/A").split("/")[-1]

    # --- 3. MÉTRICAS Y BENCHMARKS ---
    doc.add_heading("3. Métricas de Ingeniería y Negocio", level=1)
    doc.add_paragraph("Cifras clave para el dimensionamiento de la propuesta de RFP:")

    metrics = data.get("technology_context", {}).get("field_metrics", [])
    if metrics:
        m_table = doc.add_table(rows=1, cols=3)
        m_table.style = "Table Grid"
        hdr = m_table.rows[0].cells
        hdr[0].text = "KPI / Parámetro"
        hdr[1].text = "Valor"
        hdr[2].text = "Contexto de Origen"
        for cell in hdr:
            set_cell_shading(cell, "C0504D")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for m in metrics:
            row = m_table.add_row().cells
            row[0].text = str(m.get("name", "N/A"))
            row[1].text = f"{m.get('value')} {m.get('unit')}"
            row[2].text = m.get("source", {}).get("uri", "N/A").split("/")[-1]

    # --- 4. REFERENCIAS Y EVIDENCIAS ---
    doc.add_heading("4. Bóveda de Referencias (Evidence Vault)", level=1)
    doc.add_paragraph(
        "Trazabilidad completa de las fuentes utilizadas para este informe:"
    )

    for i, claim in enumerate(data.get("claims", []), 1):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[REF-{i:02d}] ").bold = True
        p.add_run(f"URI: {claim.get('url')}\n")
        p.add_run(
            f"    Hash: {claim.get('content_hash')[:15]}... | Captura: {claim.get('captured_at')[:10]}"
        )

    doc.add_paragraph("\n--- FIN DEL INFORME ---")

    save_path = "working/eurovision/Informe_Maestro_Eurovision_Services_RFP_V2.docx"
    doc.save(save_path)
    print(f"PERFECT REPORT GENERATED: {save_path}")


if __name__ == "__main__":
    generate_final_perfect_report()
