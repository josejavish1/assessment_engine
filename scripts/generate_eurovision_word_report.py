import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """Set shading color for a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def generate_report():
    json_path = Path("working/eurovision/client_intelligence.json")
    with open(json_path, "r", encoding="utf-8-sig") as f:
        data = json.load(f)

    doc = Document()
    
    # 1. Header & Title
    title = doc.add_heading('Informe de Inteligencia Estratégica: Eurovision Services', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_paragraph()
    meta.add_run('Asunto: ').bold = True
    meta.add_run('Soporte para RFP de Modernización de Infraestructura\n')
    meta.add_run('Fecha: ').bold = True
    meta.add_run('08 de junio de 2026\n')
    meta.add_run('Clasificación: ').bold = True
    meta.add_run('CONFIDENCIAL / PARTNER NTT DATA ONLY').font.color.rgb = RGBColor(200, 0, 0)
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("_" * 80)

    # 2. Executive Summary
    doc.add_heading('1. Resumen Estratégico y Agenda del CEO', level=1)
    ceo_summary = data.get("business_context", {}).get("ceo_agenda", {}).get("summary", "N/A")
    # Clean markdown artifacts from summary
    clean_prose = ceo_summary.replace("###", "").replace("**", "").replace("*", "•")
    doc.add_paragraph(clean_prose)

    # 3. Technical Stack - THE PERFECT TABLE
    doc.add_heading('2. Mapa Tecnológico Verificado (Technical Footprint)', level=1)
    doc.add_paragraph('Inventario detallado de la infraestructura OT e IT detectada:')

    stack = data.get("technology_context", {}).get("group_level_stack", [])
    if stack:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        headers = ['Dominio / Tecnología', 'Fabricante', 'Fuente / Evidencia', 'Contexto']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(hdr_cells[i], "D9D9D9")

        for item in stack:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get("technology", "N/A"))
            row_cells[1].text = str(item.get("vendor", "N/A"))
            
            source = item.get("source", {})
            uri = source.get("uri", "N/A")
            row_cells[2].text = uri.split("/")[-1] if "/" in uri else uri
            
            row_cells[3].text = source.get("paragraph_snippet", "N/A")[:150] + "..."

    # 4. Metrics Table
    doc.add_heading('3. Métricas Operacionales y de Rendimiento', level=1)
    metrics = data.get("technology_context", {}).get("field_metrics", [])
    if metrics:
        m_table = doc.add_table(rows=1, cols=3)
        m_table.style = 'Table Grid'
        
        hdr_cells = m_table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        hdr_cells[2].text = 'Unidad'
        for i in range(3):
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(hdr_cells[i], "D9D9D9")

        for m in metrics:
            row_cells = m_table.add_row().cells
            row_cells[0].text = str(m.get("name", "N/A"))
            row_cells[1].text = str(m.get("value", "N/A"))
            row_cells[2].text = str(m.get("unit", "N/A"))

    # 5. Regulatory Context
    doc.add_heading('4. Marco Regulatorio y Cumplimiento', level=1)
    regs = data.get("regulatory_context", [])
    if regs:
        r_table = doc.add_table(rows=1, cols=3)
        r_table.style = 'Table Grid'
        
        hdr_cells = r_table.rows[0].cells
        hdr_cells[0].text = 'Regulación / Directiva'
        hdr_cells[1].text = 'Impacto'
        hdr_cells[2].text = 'Referencia'
        for i in range(3):
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(hdr_cells[i], "D9D9D9")

        for r in regs:
            row_cells = r_table.add_row().cells
            row_cells[0].text = str(r.get("name", "N/A"))
            row_cells[1].text = str(r.get("impacto", "Crítico"))
            row_cells[2].text = r.get("source", {}).get("uri", "N/A").split("/")[-1]

    # 6. Conclusion
    doc.add_paragraph("\n")
    doc.add_paragraph("Este informe ha sido generado automáticamente por el Motor de Inteligencia Soberana V34.0 (NTT DATA).")
    
    save_path = "working/eurovision/Informe_Estrategico_Eurovision_Services_RFP.docx"
    doc.save(save_path)
    print(f"WORD REPORT GENERATED: {save_path}")

if __name__ == "__main__":
    generate_report()
