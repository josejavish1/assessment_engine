"""Renders a DOCX tower blueprint from a structured JSON payload."""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt, RGBColor

from domain.maturity_band import (
    ANNEX_MATURITY_BANDS,
    resolve_maturity_band,
)
from domain.schemas.blueprint import BlueprintPayload
from infrastructure.client_intelligence import (
    load_client_intelligence_legacy_view,
)
from infrastructure.docx_render_utils import (
    add_body_paragraph as _orig_add_body_paragraph,
)
from infrastructure.docx_render_utils import (
    add_heading_paragraph,
    autofit_table_to_contents,
    clear_paragraph,
    finalize_table,
    set_cell_text,
    shade_cell,
)
from infrastructure.runtime_paths import (
    resolve_tower_annex_template_path,
)
from infrastructure.text_utils import clean_text_for_word

BASE_TEXT_COLOR = RGBColor(46, 64, 77)
ROOT = Path(__file__).resolve().parents[3]
DEFAULT_TEMPLATE_PATH = resolve_tower_annex_template_path()


from typing import Any, cast


def load_json(path: Path) -> dict[str, Any]:
    """Deserialize a UTF-8 encoded JSON file from a given path."""
    return cast(dict[str, Any], json.loads(path.read_text(encoding="utf-8-sig")))


def add_spacer(doc, points=12) -> Any:
    """Add a vertical spacer of a specified height to a document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)


def clean_text_for_render(text) -> Any:
    """Formats input data into a string suitable for rendering.

    If the input `text` is a dictionary, it is converted to a string based on
    its keys. If both 'name' and 'description' keys are present, the output
    is formatted as 'name: description'. Otherwise, all dictionary values are
    string-coerced and joined by ' - '.

    For all non-dictionary input types, processing is delegated to the
    `clean_text_for_word` function.

    Args:
        text (Any): The input data to format. Can be a dictionary or any type
            compatible with `clean_text_for_word`.

    Returns:
        Any: A string representation of the input if it is a dictionary, or
            the processed output from `clean_text_for_word` otherwise.
    """
    if isinstance(text, dict):
        if "name" in text and "description" in text:
            return f"{text['name']}: {text['description']}"
        return " - ".join(str(v) for v in text.values())
    return clean_text_for_word(text)


def _safe_float(value: object) -> float | None:
    if value in (None, ""):
        return None
    try:
        return float(str(value).split()[0].replace(",", "."))
    except (TypeError, ValueError):
        return None


def _resolve_annex_band(value: object) -> str:
    score = _safe_float(value)
    if score is None:
        return ""
    return resolve_maturity_band(score, ANNEX_MATURITY_BANDS)["label"]


def add_body_paragraph(doc, text, color_rgb=BASE_TEXT_COLOR) -> Any:
    """Add a cleaned, justified, and styled body paragraph to a document."""
    text = clean_text_for_render(text)
    return _orig_add_body_paragraph(
        doc,
        text,
        space_after=12,
        color_rgb=color_rgb,
        justify=True,
    )


def add_bullet_p(container, text, color_rgb=BASE_TEXT_COLOR) -> Any:
    r"""{'docstring': "Adds a formatted bullet point paragraph to a python-docx container.\n\n    The function sanitizes the input text and adds it as a justified paragraph.\n    It first attempts to apply the document's 'List Bullet' style. If the\n    style is not available, it manually synthesizes a bullet point using a '•'\n    character run and a hanging indent.\n\n    If the input text contains a colon (':'), the substring preceding the colon\n    is rendered in a bold font to serve as a label.\n\n    Args:\n        container (typing.Any): A `python-docx` object that can host a\n            paragraph, such as a `Document` or `_Cell` instance.\n        text (str): The string content for the bullet point.\n        color_rgb (docx.shared.RGBColor): The RGB color to apply to the text\n            and bullet symbol. Defaults to `BASE_TEXT_COLOR`.\n\n    Returns:\n        docx.text.paragraph.Paragraph: The newly created and formatted paragraph\n            object."}."""
    text = clean_text_for_render(text)
    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)

    try:
        p.style = "List Bullet"
    except KeyError:
        p.paragraph_format.left_indent = Pt(20)
        p.paragraph_format.first_line_indent = Pt(-15)
        bullet = p.add_run("• ")
        bullet.font.size = Pt(10.5)
        if color_rgb:
            bullet.font.color.rgb = color_rgb

    if ":" in text:
        prefix, rest = text.split(":", 1)
        prefix_run = p.add_run(prefix + ":")
        prefix_run.bold = True
        prefix_run.font.size = Pt(10.5)
        if color_rgb:
            prefix_run.font.color.rgb = color_rgb

        rest_run = p.add_run(rest)
        rest_run.font.size = Pt(10.5)
        if color_rgb:
            rest_run.font.color.rgb = color_rgb
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        if color_rgb:
            run.font.color.rgb = color_rgb
    return p


def create_page_number_footer(section) -> Any:
    """Adds a "Page X of Y" footer to a `docx.section.Section` object.

    Configures the section to use a different footer for the first page. The
    primary footer (used for all subsequent pages) is then modified. Its first
    paragraph is cleared, centered, and repopulated with a combination of
    static text and dynamic fields to produce the string
    "Página {PAGE} de {NUMPAGES}".

    The dynamic `PAGE` and `NUMPAGES` fields are generated by directly
    constructing the underlying OpenXML elements. This ensures that Microsoft
    Word correctly interprets them as page counters. The input `section` object
    is modified in-place.

    Args:
        section (docx.section.Section): The document section to modify.

    Returns:
        None.

    Raises:
        IndexError: If the section's primary footer does not contain at least one
            paragraph to modify.
    """
    section.different_first_page_header_footer = True
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph(paragraph)

    def add_field(p, field_code):
        r"""Injects a complex field into a `python-docx` paragraph.

        This function constructs the low-level OpenXML elements required for a complex
        field, such as a page number, by directly manipulating the underlying XML of a
        `run` element within the specified paragraph. It assembles the required
        `w:fldChar` (begin, separate, end) and `w:instrText` tags that constitute a
        valid field instruction for Microsoft Word. A default, non-configurable font
        style (9pt Arial, gray) is applied to the field.

        The paragraph object is modified in-place.

        Args:
            p (docx.paragraph.Paragraph): The paragraph object to which the field will
                be added.
            field_code (str): The instruction string that defines the field's behavior,
                e.g., "PAGE \\* MERGEFORMAT" or "NUMPAGES".
        """
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(127, 127, 127)
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(ns.qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)
        instr_text = OxmlElement("w:instrText")
        instr_text.set(ns.qn("xml:space"), "preserve")
        instr_text.text = field_code
        run._r.append(instr_text)
        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(ns.qn("w:fldCharType"), "separate")
        run._r.append(fld_char_sep)
        text = OxmlElement("w:t")
        text.text = "0"
        run._r.append(text)
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(ns.qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)

    prefix = paragraph.add_run("Página ")
    prefix.font.size = Pt(9)
    prefix.font.name = "Arial"
    prefix.font.color.rgb = RGBColor(127, 127, 127)
    add_field(paragraph, "PAGE")
    middle = paragraph.add_run(" de ")
    middle.font.size = Pt(9)
    middle.font.name = "Arial"
    middle.font.color.rgb = RGBColor(127, 127, 127)
    add_field(paragraph, "NUMPAGES")


def clear_document_body(doc) -> None:
    """Clears all content from a document's body while preserving section properties.

    This function directly manipulates the underlying lxml element representation
    of the document's body (`w:body`). It iterates through and removes all child
    elements, such as paragraphs (`w:p`) and tables (`w:tbl`), while explicitly
    sparing the final section properties element (`w:sectPr`).

    This process is useful for resetting a document to a blank state while
    retaining its page layout, margins, orientation, and other section-level
    formatting, making it suitable for use as a template. The modification is
    performed in-place.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to modify.

    Returns:
        None: The document is modified in-place.
    """
    body = doc._body._element
    for child in list(body):
        if (
            child.tag
            != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
        ):
            body.remove(child)


def resolve_client_dir(payload_path: Path, payload_data: dict) -> Path:
    """Resolve the client directory as the grandparent directory of the payload path."""
    # The physical directory structure is the source of truth for determining the tower identifier, as the payload_path contains the canonical location (e.g., working/redeia_v3/T2/payload.json).
    return payload_path.parents[1]


def load_client_intelligence(client_dir: Path) -> dict:
    """Load client intelligence data from 'client_intelligence.json' within a directory."""
    path = client_dir / "client_intelligence.json"
    return load_client_intelligence_legacy_view(path)


def load_annex_data(client_dir: Path, tower_code: str) -> dict:
    """Loads the approved annex data payload for a specific tower from a JSON file.

    Constructs a path to a tower-specific JSON file using the convention:
    `{client_dir}/{TOWER_CODE}/approved_annex_{tower_code}.template_payload.json`.
    If the file exists, its contents are deserialized from JSON. If the file
    is not found, an empty dictionary is returned as a default.

    Args:
        client_dir (pathlib.Path): The root directory for the client's data.
        tower_code (str): The unique identifier for the tower. Case variations of
            this code are used to construct the final file path.

    Returns:
        dict: A dictionary containing the parsed annex data, or an empty
        dictionary if the corresponding file does not exist.

    Raises:
        json.JSONDecodeError: If the target annex file is present but contains
            malformed JSON that cannot be deserialized.
    """
    tower_dir = client_dir / tower_code.upper()
    path = tower_dir / f"approved_annex_{tower_code.lower()}.template_payload.json"
    if path.exists():
        return load_json(path)
    return {}


def _list_or_default(value, default=None) -> Any:
    if isinstance(value, list):
        return value
    return default or []


def _string_or_default(value, default="") -> Any:
    if isinstance(value, str):
        return value
    return default


def _derive_executive_snapshot(data: dict, annex_data: dict) -> dict:
    snapshot = data.get("executive_snapshot") or {}
    annex_summary = annex_data.get("executive_summary", {})
    annex_sections = annex_data.get("sections", {})
    risks = annex_sections.get("risks", {}).get("risks", [])
    todo_items = annex_sections.get("todo", {}).get("priority_initiatives", [])
    target_capabilities = annex_sections.get("tobe", {}).get("target_capabilities", [])

    if isinstance(annex_summary.get("summary_body"), list):
        summary_body = " ".join(str(x) for x in annex_summary.get("summary_body", []))
    else:
        summary_body = _string_or_default(annex_summary.get("summary_body"))

    return {
        "bottom_line": _string_or_default(snapshot.get("bottom_line"))
        or summary_body
        or _string_or_default(annex_summary.get("headline"))
        or "La torre requiere una transformación priorizada para reducir riesgo y habilitar el negocio.",
        "decisions": _list_or_default(snapshot.get("decisions"))
        or [
            clean_text_for_render(item.get("initiative", ""))
            for item in todo_items[:4]
            if item.get("initiative")
        ]
        or ["Validar el backlog priorizado y su secuencia de ejecución."],
        "cost_of_inaction": _string_or_default(snapshot.get("cost_of_inaction"))
        or "Mantener el estado actual prolonga el riesgo operativo, la deuda técnica y la incapacidad de escalar con control.",
        "structural_risks": _list_or_default(snapshot.get("structural_risks"))
        or [
            clean_text_for_render(item.get("risk", ""))
            for item in risks[:4]
            if item.get("risk")
        ],
        "business_impact": _string_or_default(snapshot.get("business_impact"))
        or "La modernización de la torre reduce exposición al riesgo y mejora la capacidad de ejecución del negocio.",
        "operational_benefits": _list_or_default(snapshot.get("operational_benefits"))
        or [clean_text_for_render(item) for item in target_capabilities[:4]],
        "transformation_complexity": _string_or_default(
            snapshot.get("transformation_complexity")
        )
        or "La transformación exige coordinación de arquitectura, operación y gobierno, pero es abordable por fases.",
    }


def _derive_cross_capabilities_analysis(data: dict) -> dict:
    cca = data.get("cross_capabilities_analysis") or {}
    pillars = _list_or_default(data.get("pillars_analysis"))
    low_score_pillars = [
        p.get("pilar_name", "") for p in pillars if float(p.get("score", 0) or 0) < 3.0
    ]

    if cca:
        return {
            "common_deficiency_patterns": _list_or_default(
                cca.get("common_deficiency_patterns")
            ),
            "transformation_paradigm": _string_or_default(
                cca.get("transformation_paradigm")
            )
            or "La torre requiere una evolución por dominios, priorizando estabilización, gobierno y posterior industrialización.",
            "critical_technical_debt": _string_or_default(
                cca.get("critical_technical_debt")
            )
            or "La deuda técnica limita la resiliencia y la capacidad de operar con consistencia.",
        }

    deficiency_patterns = []
    if low_score_pillars:
        deficiency_patterns.append(
            "Las mayores brechas se concentran en: "
            + ", ".join(low_score_pillars)
            + "."
        )

    return {
        "common_deficiency_patterns": deficiency_patterns
        or [
            "Persisten carencias repetidas en estandarización, automatización y gobierno técnico."
        ],
        "transformation_paradigm": "La transformación debe ejecutarse de forma incremental, combinando quick wins con capacidades fundacionales de largo recorrido.",
        "critical_technical_debt": "La deuda técnica acumulada incrementa el riesgo operativo y reduce la velocidad de adopción de nuevos servicios.",
    }


def _derive_roadmap(data: dict) -> list[dict]:
    roadmap = _list_or_default(data.get("roadmap"))
    if roadmap:
        return roadmap

    initiatives = []
    for pillar in _list_or_default(data.get("pillars_analysis")):
        for project in _list_or_default(pillar.get("projects_todo")):
            name = clean_text_for_render(project.get("name", ""))
            sizing = str(project.get("sizing", "")).lower()
            initiatives.append((name, sizing))

    quick_wins = [name for name, sizing in initiatives if name and "s" in sizing][:4]
    medium_term = [name for name, sizing in initiatives if name and "m" in sizing][:4]
    strategic = [name for name, sizing in initiatives if name and "l" in sizing][:4]

    fallback = [name for name, _ in initiatives[:4] if name]

    waves = [
        {"wave": "Wave 1", "projects": quick_wins or fallback[:2]},
        {"wave": "Wave 2", "projects": medium_term or fallback[2:4]},
        {"wave": "Wave 3", "projects": strategic or fallback[:2]},
    ]
    return [wave for wave in waves if wave["projects"]]


def normalize_blueprint_payload_dict(data: dict, annex_data: dict) -> dict:
    """Normalizes and enriches a raw tower blueprint dictionary.

    This function produces a new, normalized blueprint by creating a shallow copy
    of the input `data` dictionary. It then derives values for the
    'executive_snapshot', 'cross_capabilities_analysis', and 'roadmap' keys.
    It also ensures that optional fields like 'external_dependencies' and
    'pillars_analysis' are normalized to lists.

    Args:
        data (dict): The primary dictionary containing the raw blueprint data.
        annex_data (dict): A supplementary dictionary with data used for enrichment.

    Returns:
        dict: A new dictionary representing the normalized blueprint. The original
            input dictionaries are not mutated.

    Raises:
        KeyError: If a required key is missing from `data` or `annex_data` and is
            needed by an internal derivation function.
    """
    normalized = dict(data)
    normalized["executive_snapshot"] = _derive_executive_snapshot(data, annex_data)
    normalized["cross_capabilities_analysis"] = _derive_cross_capabilities_analysis(
        data
    )
    normalized["roadmap"] = _derive_roadmap(data)
    normalized["external_dependencies"] = _list_or_default(
        data.get("external_dependencies")
    )
    normalized["pillars_analysis"] = _list_or_default(data.get("pillars_analysis"))
    return normalized


def load_payload(
    payload_path: Path, annex_data: dict | None = None
) -> BlueprintPayload:
    """Load, normalize, and validate a blueprint payload from a JSON file.

    Reads a JSON file from the specified path, merges its content with optional
    supplementary data, and validates the resulting structure against the
    `BlueprintPayload` schema.

    Args:
        payload_path: The file system path to the JSON file containing the
            blueprint payload.
        annex_data: An optional dictionary of supplementary data. If provided,
            its key-value pairs are merged into the data loaded from the JSON
            file. Existing keys in the file data will be overwritten by keys
            from `annex_data`. Defaults to None.

    Returns:
        A validated `BlueprintPayload` Pydantic model instance representing
        the complete and normalized blueprint data.

    Raises:
        FileNotFoundError: The file at `payload_path` does not exist or is not
            accessible.
        json.JSONDecodeError: The content of the file at `payload_path` is not a
            valid JSON document.
        pydantic.ValidationError: The merged data structure does not conform to
            the `BlueprintPayload` model schema.
    """
    raw_data = load_json(payload_path)
    normalized_data = normalize_blueprint_payload_dict(raw_data, annex_data or {})
    return BlueprintPayload.model_validate(normalized_data)


def render_cover(doc, payload: BlueprintPayload) -> Any:
    """Constructs the cover page for a technology maturity report.

    This function populates a `docx.Document` object with a formatted cover page, including a title, client name, version details, and a confidentiality disclaimer. Each element is styled with specific fonts, sizes, colors, and paragraph spacing. A page break is appended to ensure subsequent content begins on a new page.

    Args:
        doc (docx.document.Document): An instance of a python-docx Document to be modified in place.
        payload (BlueprintPayload): A data object containing metadata for the report, such as `tower_name`, `client_name`, `tower_code`, and `transformation_horizon`.

    Returns:
        None: The `doc` object is modified by side effect.
    """
    meta = payload.document_meta
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    title_p = doc.add_paragraph()
    run = title_p.add_run(f"{meta.tower_name}\nInforme de Madurez Tecnológica")
    run.font.size = Pt(34)
    run.font.name = "Georgia"
    run.font.color.rgb = RGBColor(0, 114, 188)
    add_spacer(doc, 30)

    client_p = doc.add_paragraph()
    client_run = client_p.add_run(meta.client_name.upper())
    client_run.font.size = Pt(24)
    client_run.font.name = "Arial"
    client_run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(100)
    version_p = doc.add_paragraph()
    version_text = (
        f"Torre Técnica: {meta.tower_code}\nHorizonte: {meta.transformation_horizon}"
    )
    version_run = version_p.add_run(version_text)
    version_run.font.size = Pt(14)
    version_run.font.name = "Arial"
    version_run.font.color.rgb = RGBColor(127, 127, 127)

    add_spacer(doc, 150)
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    title_run = disclaimer.add_run("Confidencialidad: ")
    title_run.bold = True
    title_run.font.size = Pt(8)
    title_run.font.color.rgb = RGBColor(127, 127, 127)
    text_run = disclaimer.add_run(
        "Este documento técnico es propiedad de NTT DATA y el cliente. Contiene "
        "información estratégica y de arquitectura sujeta a acuerdos de confidencialidad."
    )
    text_run.font.size = Pt(8)
    text_run.font.color.rgb = RGBColor(127, 127, 127)
    doc.add_page_break()


def render_snapshot_page(
    doc,
    payload: BlueprintPayload,
    client_intelligence: dict,
    annex_data: dict,
):
    """Constructs and appends the Executive Snapshot section to a document object.

    This function generates the 'Executive Snapshot' section of a report. It
    begins by creating a summary table with the current score, maturity level,
    and target maturity. It then adds a series of subsections detailing the
    business context, material risks, cost of inaction, expected business
    impact, operational benefits, transformation complexity, and key decisions.
    Data for these sections is sourced from the `payload`,
    `client_intelligence`, and `annex_data` arguments.

    Args:
        doc (docx.document.Document): The document object to which the snapshot
            section will be appended.
        payload (BlueprintPayload): An object containing the primary data for the
            executive snapshot, including narratives for risks, impacts, and
            decisions.
        client_intelligence (dict): A dictionary containing client-specific
            contextual data. Expected keys include 'ceo_agenda',
            'regulatory_frameworks', and 'technological_drivers'.
        annex_data (dict): A dictionary containing supplementary data. The function
            accesses `annex_data['executive_summary']` for 'global_score',
            'global_band', and 'target_maturity' values.

    Raises:
        AttributeError: If `payload` or its `executive_snapshot` attribute is
            missing required fields (e.g., `bottom_line`, `decisions`).
    """
    snap = payload.executive_snapshot
    exec_sum = annex_data.get("executive_summary", {})
    global_score_val = exec_sum.get("global_score", "N/A")
    global_band_val = exec_sum.get("global_band") or _resolve_annex_band(
        global_score_val
    )
    target_score_val = exec_sum.get("target_maturity", "N/A")

    add_heading_paragraph(doc, "1. Executive Snapshot (Resumen ejecutivo)", level=1)

    table = doc.add_table(rows=2, cols=3)
    finalize_table(table)
    headers = ["SCORE ACTUAL", "NIVEL DE MADUREZ", "MADUREZ OBJETIVO"]
    values = [global_score_val, global_band_val or "N/A", target_score_val]

    for i, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            header,
            bold=True,
            font_size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table.rows[0].cells[i], "0072BC")
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_text(
            table.rows[1].cells[i],
            values[i],
            font_size=14,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    if tbl_pr is not None:
        widths = tbl_pr.xpath("w:tblW")
        tbl_w = widths[0] if widths else OxmlElement("w:tblW")
        if not widths:
            tbl_pr.append(tbl_w)
        tbl_w.set(ns.qn("w:type"), "pct")
        tbl_w.set(ns.qn("w:w"), "5000")

    add_spacer(doc, 15)
    add_body_paragraph(doc, snap.bottom_line, color_rgb=BASE_TEXT_COLOR)

    business_angles = []
    ceo_agenda = clean_text_for_render(client_intelligence.get("ceo_agenda", ""))
    if ceo_agenda:
        business_angles.append("Agenda de negocio prioritaria: " + ceo_agenda)
    regulatory = client_intelligence.get("regulatory_frameworks", []) or []
    if regulatory:
        business_angles.append(
            "Presión regulatoria material: " + ", ".join(str(x) for x in regulatory)
        )
    drivers = client_intelligence.get("technological_drivers", []) or []
    for text in drivers:
        lowered = str(text).lower()
        if any(
            token in lowered
            for token in ["adquis", "m&a", "expansi", "ia", "pacient", "eficien"]
        ):
            business_angles.append(clean_text_for_render(text))

    if not business_angles:
        business_angles.append(
            "La transformación tecnológica propuesta asegura la competitividad, fortalece la resiliencia operativa y garantiza el cumplimiento normativo."
        )

    add_heading_paragraph(doc, "Por qué importa al negocio", level=2)
    for item in business_angles[:4]:
        add_bullet_p(doc, item)

    if snap.structural_risks:
        add_heading_paragraph(doc, "Riesgos de negocio más materiales", level=2)
        for risk in snap.structural_risks:
            add_bullet_p(doc, risk)

    add_heading_paragraph(doc, "Coste de Inacción (Do Nothing)", level=2)
    add_body_paragraph(doc, snap.cost_of_inaction, color_rgb=BASE_TEXT_COLOR)

    if snap.business_impact:
        add_heading_paragraph(doc, "Impacto Esperado en Negocio", level=2)
        add_body_paragraph(doc, snap.business_impact, color_rgb=BASE_TEXT_COLOR)

    if snap.operational_benefits:
        add_heading_paragraph(doc, "Beneficios Operativos Target", level=2)
        for benefit in snap.operational_benefits:
            add_bullet_p(doc, benefit)

    if snap.transformation_complexity:
        add_heading_paragraph(doc, "Complejidad de la Transformación", level=2)
        add_body_paragraph(
            doc,
            snap.transformation_complexity,
            color_rgb=BASE_TEXT_COLOR,
        )

    add_heading_paragraph(doc, "Decisiones prioritarias", level=2)
    for decision in snap.decisions:
        add_bullet_p(doc, decision)


def render_cross_capabilities_analysis(doc, payload: BlueprintPayload) -> Any:
    """Renders the cross-capabilities analysis section into a document object.

    This function populates a document object with content from the
    `cross_capabilities_analysis` attribute of the provided payload. If this
    attribute is not present or is falsy, the function returns immediately without
    modifying the document. Otherwise, it adds a structured section including
    headings for the transformation paradigm, critical technical debt, and
    common deficiency patterns, followed by their corresponding content and a
    bulleted list.

    Args:
        doc (Any): The document object to be populated. It is modified in-place.
        payload (BlueprintPayload): The data transfer object containing the
            cross-capabilities analysis data.

    Returns:
        None.
    """
    cca = payload.cross_capabilities_analysis
    if not cca:
        return

    add_heading_paragraph(doc, "Análisis Transversal de Capacidades", level=1)
    add_heading_paragraph(doc, "El Paradigma de Transformación", level=2)
    add_body_paragraph(doc, cca.transformation_paradigm, color_rgb=BASE_TEXT_COLOR)

    add_heading_paragraph(doc, "Deuda Técnica Crítica", level=2)
    add_body_paragraph(doc, cca.critical_technical_debt, color_rgb=BASE_TEXT_COLOR)

    add_heading_paragraph(doc, "Patrones Comunes de Deficiencia", level=2)
    for item in cca.common_deficiency_patterns:
        add_bullet_p(doc, item)


def render_consolidated_asis(doc, payload: BlueprintPayload) -> Any:
    """Renders a consolidated 'AS-IS' technology diagnosis section into a document.

    This function adds a main heading for the consolidated diagnosis. It then
    iterates through each technology pillar in the provided payload, creating a
    subsection and a summary table for each. The table details the technical
    capability, specific findings or evidence, and the resulting business risk.

    A key feature is the conditional background shading of the 'Business Risk'
    cell to visually indicate severity. The cell is shaded light red for high-
    impact risks (e.g., containing 'crítico', 'alto') and light amber for
    medium-impact risks (e.g., containing 'medio').

    Args:
        doc: The `python-docx` Document object to which the content will be added.
        payload (BlueprintPayload): An object containing the aggregated technology
            pillar analysis and as-is health check data.

    Returns:
        None. The function modifies the input `doc` object in-place.
    """
    add_heading_paragraph(doc, "Diagnóstico Tecnológico Consolidado (AS-IS)", level=1)
    
    for pilar in payload.pillars_analysis:
        add_heading_paragraph(doc, f"Capacidad: {pilar.pilar_name}", level=2)
        
        table = doc.add_table(rows=1, cols=3)
        finalize_table(table)
        headers = ["Capacidad Técnica", "Hallazgo / Evidencia", "Riesgo de Negocio"]
        for i, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=10)
            shade_cell(table.rows[0].cells[i], "D9EAF7")

        for row_data in pilar.health_check_asis:
            row = table.add_row()
            set_cell_text(row.cells[0], row_data.target_state, bold=True, font_size=10)
            set_cell_text(row.cells[1], row_data.risk_observed, font_size=10)
            
            #
            impact_text = row_data.impact
            impact_lower = impact_text.lower()
            if any(k in impact_lower for k in ["crítico", "alto", "crítica", "alta", "severo", "material"]):
                shade_cell(row.cells[2], "FADBD8") # Defines the RGB value for the Light Red color state.
            elif any(k in impact_lower for k in ["medio", "moderado"]):
                shade_cell(row.cells[2], "FCF3CF") # Defines the RGB value for the Light Amber color state.
                
            set_cell_text(row.cells[2], impact_text, font_size=10)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BASE_TEXT_COLOR
        autofit_table_to_contents(table)
        add_spacer(doc, 10)


def render_fair_risk_heatmap(doc, payload: BlueprintPayload) -> Any:
    r"""{'docstring': 'Renders a FAIR quantitative risk heatmap and details table into a document.\n\n    This function processes Factor Analysis of Information Risk (FAIR) data from\n    the provided payload to generate two primary components within the given\n    Word document: a risk matrix and a detailed risk register.\n\n    First, a 5x5 heatmap is constructed where rows represent Threat Event\n    Frequency (TEF) and columns represent Loss Magnitude (LM). Each cell is\n    color-coded based on the calculated risk score (TEF * LM), mapping to a\n    severity scale from low (green) to critical (dark red). Cells are\n    populated with unique identifiers (e.g., "R01", "R02") for each risk that\n    falls within that specific TEF/LM coordinate.\n\n    Second, a detailed table is generated that lists each identified risk. This\n    table includes the risk\'s identifier, its domain, a description of the\n    finding, its coordinates on the heatmap, and its calculated Annualized Loss\n    Expectancy (ALE). The risks in this table are sorted in descending order of\n    their ALE to prioritize the most significant financial exposures.\n\n    If the payload contains no findings with valid FAIR data (where both TEF > 0\n    and LM > 0), the function appends a notification message to the document\n    instead of the heatmap and table.\n\n    Args:\n        doc (docx.document.Document): The python-docx Document object to which\n            the generated content will be appended.\n        payload (BlueprintPayload): A data object containing analysis results. It is\n            expected to have a `pillars_analysis` attribute, which is an\n            iterable of objects each containing a `health_check_asis` list of\n            findings. Each finding may have `threat_event_frequency`,\n            `loss_magnitude`, and `fair_ale_score` attributes.\n\n    Returns:\n        None: The function modifies the input `doc` object in place.'}."""
    add_heading_paragraph(doc, "Matriz de Riesgo Cuantitativa (FAIR)", level=1)
    add_body_paragraph(
        doc,
        "La siguiente matriz consolida la expectativa de pérdida anualizada (ALE) de las vulnerabilidades detectadas en el estado actual, basada en factores de frecuencia de la amenaza (TEF) y magnitud de pérdida (LM). "
        "Nota: La métrica ALE representa la exposición aislada de cada riesgo. Varios riesgos pueden compartir la misma causa raíz, por lo que su sumatorio refleja el volumen total de amenaza bajo gestión, no un impacto simultáneo garantizado.",
        color_rgb=BASE_TEXT_COLOR
    )
    
    # Initializes a 5x5 matrix for risk assessment. Rows represent TEF (Frequency) on a scale of 5 to 1, and columns represent LM (Magnitude) on a scale of 1 to 5.
    matrix = {r: {c: [] for c in range(1, 6)} for r in range(5, 0, -1)}
    
    has_fair_data = False
    risk_details = []
    risk_counter = 1
    
    for pilar in payload.pillars_analysis:
        for finding in pilar.health_check_asis:
            tef = getattr(finding, "threat_event_frequency", 0)
            lm = getattr(finding, "loss_magnitude", 0)
            ale = getattr(finding, "fair_ale_score", 0.0)
            if tef > 0 and lm > 0:
                has_fair_data = True
                r_id = f"R{risk_counter:02d}"
                risk_counter += 1
                matrix[tef][lm].append(r_id)
                risk_details.append({
                    "id": r_id,
                    "pilar": pilar.pilar_name,
                    "finding": finding.risk_observed,
                    "target_state": finding.target_state,
                    "tef": tef,
                    "lm": lm,
                    "ale": ale,
                    "risk_score": tef * lm
                })
                
    if not has_fair_data:
        add_body_paragraph(doc, "No hay datos cuantitativos FAIR suficientes para renderizar el mapa de calor.", color_rgb=BASE_TEXT_COLOR)
        return
        
    table = doc.add_table(rows=6, cols=6)
    finalize_table(table)
    
    #
    set_cell_text(table.rows[0].cells[0], "Frecuencia \\ Impacto", bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in range(1, 6):
        set_cell_text(table.rows[0].cells[c], f"Nivel {c}", bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[c], "F2F2F2")
        
    for r_idx, r in enumerate(range(5, 0, -1), 1):
        set_cell_text(table.rows[r_idx].cells[0], f"Freq {r}", bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[r_idx].cells[0], "F2F2F2")
        for c in range(1, 6):
            cell = table.rows[r_idx].cells[c]
            risk_score = r * c
            color = "D9F2D9" # Defines the RGB value for the Green color state.
            if risk_score >= 20: color = "C00000" # Defines the RGB value for the Dark Red color state.
            elif risk_score >= 15: color = "FF0000" # Defines the RGB value for the Red color state.
            elif risk_score >= 10: color = "FFC000" # Defines the RGB value for the Amber color state.
            elif risk_score >= 5: color = "FFFF00" # Defines the RGB value for the Yellow color state.
            
            shade_cell(cell, color)
            
            items = matrix[r][c]
            if items:
                set_cell_text(cell, ", ".join(items), font_size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                if risk_score >= 15: # A specific text color (white) is required for high contrast against Red or Dark Red backgrounds to maintain readability.
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                set_cell_text(cell, "", font_size=8)

    autofit_table_to_contents(table)
    add_spacer(doc, 15)

    #
    # Sorts risks in descending order by Annualized Loss Expectancy (ALE) to prioritize the most significant threats.
    risk_details.sort(key=lambda x: x["ale"], reverse=True)
    
    add_heading_paragraph(doc, "Detalle de Exposición al Riesgo (FAIR)", level=2)
    risk_table = doc.add_table(rows=1, cols=5)
    finalize_table(risk_table)
    headers = ["ID", "Dominio", "Hallazgo / Brecha", "Coordenadas (Freq x Imp)", "Exposición de Riesgo (ALE)"]
    for i, header in enumerate(headers):
        set_cell_text(risk_table.rows[0].cells[i], header, bold=True, font_size=9)
        shade_cell(risk_table.rows[0].cells[i], "D9EAF7")
        
    for detail in risk_details:
        row = risk_table.add_row()
        set_cell_text(row.cells[0], detail["id"], font_size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], detail["pilar"], font_size=8)
        set_cell_text(row.cells[2], f"{detail['target_state']}: {detail['finding']}", font_size=8)
        set_cell_text(row.cells[3], f"TEF: {detail['tef']} | LM: {detail['lm']}", font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Determines the appropriate cell background color based on the calculated Annualized Loss Expectancy (ALE) risk level.
        risk_score = detail["risk_score"]
        color = "D9F2D9" # Defines the RGB value for the Green color state.
        if risk_score >= 20: color = "C00000" # Defines the RGB value for the Dark Red color state.
        elif risk_score >= 15: color = "FF0000" # Defines the RGB value for the Red color state.
        elif risk_score >= 10: color = "FFC000" # Defines the RGB value for the Amber color state.
        elif risk_score >= 5: color = "FFFF00" # Defines the RGB value for the Yellow color state.
        
        ale_val = detail["ale"]
        set_cell_text(row.cells[4], f"{ale_val:,.2f} €".replace(",", "X").replace(".", ",").replace("X", "."), font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        shade_cell(row.cells[4], color)
        if risk_score >= 15: # A specific text color (white) is required for high contrast against Red or Dark Red backgrounds to maintain readability.
            for paragraph in row.cells[4].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                            
    autofit_table_to_contents(risk_table)
    add_spacer(doc, 15)

def render_consolidated_tobe(doc, payload: BlueprintPayload) -> Any:
    """Populates a document object with the 'TO-BE' architecture section.

    This function constructs the target architecture overview by adding a main
    heading, a paragraph on the transformation paradigm, a static statement
    on the sovereign agnostic model, and a list of unified design principles.
    Content is primarily sourced from the `payload` object, and some static
    text is rendered in Spanish.

    Args:
        doc: A document-like object that exposes methods for content
            manipulation, such as `add_heading_paragraph`, `add_body_paragraph`,
            and `add_bullet_p`.
        payload (BlueprintPayload): A data object containing the blueprint
            definition. The function accesses
            `payload.cross_capabilities_analysis.transformation_paradigm` and
            `payload.design_principles`.

    Returns:
        None. This function modifies the `doc` object in-place.

    Raises:
        AttributeError: If `payload` is missing the `cross_capabilities_analysis`
            attribute.
    """
    add_heading_paragraph(doc, "Arquitectura Objetivo y Visión Soberana (TO-BE)", level=1)
    
    if payload.cross_capabilities_analysis:
        add_body_paragraph(doc, payload.cross_capabilities_analysis.transformation_paradigm, color_rgb=BASE_TEXT_COLOR)

    add_body_paragraph(
        doc,
        "La arquitectura objetivo propuesta se rige por un modelo soberano agnóstico, evitando el vendor lock-in e integrando las infraestructuras críticas con los ecosistemas Cloud de mayor innovación tecnológica (e.g. AWS).",
        color_rgb=BASE_TEXT_COLOR
    )
    
    if hasattr(payload, "design_principles") and payload.design_principles:
        add_heading_paragraph(doc, "Principios de Diseño Unificados", level=2)
        for principle in payload.design_principles:
            add_bullet_p(doc, principle)


def render_deep_project_charters(doc, payload: BlueprintPayload) -> Any:
    r"""{'docstring': 'Renders a technical appendix of detailed project charters into a document.\n\n    Generates and appends a series of formatted tables, representing detailed\n    project charters, to a document object. The function processes projects\n    from the input payload, first grouping them by their `transformation_typology`.\n    For each project, it aggregates comprehensive details including an executive\n    summary, S.M.A.R.T. objectives, scope definitions, technical dependencies,\n    financial estimates, Work Breakdown Structure (WBS), and risk analysis.\n\n    Financial data (OPEX) is sanitized to remove internal margin calculations\n    before rendering. The function is robust to variations in the payload\n    structure, using default values for missing attributes and handling potential\n    type inconsistencies in WBS data.\n\n    Args:\n        doc: A document object, typically from a library like `python-docx`, to\n            which the appendix is added. This object is modified in-place.\n        payload (BlueprintPayload): A data object containing the project blueprint,\n            including architectural pillars, project definitions, dependencies,\n            and financial data.\n\n    Returns:\n        None. The function modifies the `doc` object directly.\n\n    Raises:\n        AttributeError: If the `payload` object or its nested structures do not\n            conform to the expected schema (e.g., missing `pillars_analysis`).'}."""
    doc.add_page_break()
    add_heading_paragraph(doc, "Anexo Técnico: Fichas Profundas de Proyectos (Charters)", level=1)
    
    # Projects are first grouped by their transformation typology to structure the subsequent rendering stages.
    from collections import defaultdict
    grouped_projects = defaultdict(list)
    
    for pilar in payload.pillars_analysis:
        for project in pilar.projects_todo:
            typology = getattr(project, "transformation_typology", "Iniciativas Core")
            grouped_projects[typology].append((pilar.pilar_name, project))
            
    #
    project_idx = 1
    for typology, projects_in_group in grouped_projects.items():
        add_heading_paragraph(doc, f"Vector de Transformación: {typology}", level=2)
        add_spacer(doc, 10)
        
        for pilar_name, project in projects_in_group:
            add_heading_paragraph(doc, f"Iniciativa {project_idx}: {project.initiative}", level=3)
            
            deps = [d.depends_on for d in payload.external_dependencies if d.project == project.initiative]
            deps_text = " • ".join(deps) if deps else "Independiente (Habilitador Fase 0)"

            #
            mitigated_risk_text = "N/A"
            if getattr(project, "mitigates_risk_id", None):
                for p_ana in payload.pillars_analysis:
                    for hc in p_ana.health_check_asis:
                        if hc.node_id == project.mitigates_risk_id:
                            mitigated_risk_text = f"[{hc.target_state}] {hc.risk_observed}"
                            break

            # Performs FinOps resolution by sanitizing financial data to exclude internal cost margins from client-facing documents.
            capex = getattr(project, "capex_estimate", "N/A")
            opex = getattr(project, "opex_estimate", "N/A")
            
            # The OPEX value is sanitized to strip any internal margin calculations, ensuring only the client-facing cost is presented.
            # The financial data is filtered to retain only the final selling price for this context.
            clean_opex = opex
            if "(Margen" in opex:
                clean_opex = opex.split("(Margen")[0].strip()
            
            finops_text = f"Inversión de Implantación: {clean_opex}\nInversión de Licencias/Hardware: {capex}"

            #
            wbs_items = getattr(project, "wbs_breakdown", [])
            wbs_text = "Desglose no disponible."
            if wbs_items:
                try:
                    wbs_text = "\n".join([f"• {w.task_name} (Perfil: {w.required_profile})" for w in wbs_items])
                except Exception:
                    # Handles potential inconsistencies in data types (dict vs. object) during payload parsing to ensure robust attribute access.
                    wbs_text = "\n".join([f"• {w.get('task_name')} (Perfil: {w.get('required_profile')})" for w in wbs_items if isinstance(w, dict)])

            #
            proj_desc = getattr(project, "project_description", "N/A") or "N/A"
            smart_obj = getattr(project, "smart_objectives", "N/A") or "N/A"
            in_scope_list = getattr(project, "in_scope", []) or []
            out_scope_list = getattr(project, "out_of_scope", []) or []
            gov_roles_list = getattr(project, "governance_roles", []) or []
            crit_risks_list = getattr(project, "critical_risks", []) or []
            
            in_scope = "\n".join([f"• {item}" for item in in_scope_list]) or "N/A"
            out_scope = "\n".join([f"• {item}" for item in out_scope_list]) or "N/A"
            gov_roles = "\n".join([f"• {item}" for item in gov_roles_list]) or "N/A"
            crit_risks = "\n".join([f"• {item}" for item in crit_risks_list]) or "N/A"
            roi_just = getattr(project, "roi_justification", "N/A") or "N/A"

            rows = [
                ("Dominio Arquitectónico", pilar_name),
                ("Descripción Ejecutiva", proj_desc),
                ("Objetivo SMART", smart_obj),
                ("Riesgo AS-IS Mitigado (Trazabilidad)", mitigated_risk_text),
                ("Alcance (In-Scope)", in_scope),
                ("Fuera de Alcance (Out-of-Scope)", out_scope),
                ("Dependencias Técnicas (Predecesores)", deps_text),
                (
                    "Entregables Técnicos Duros (DoD)",
                    "\n".join([f"• {item}" for item in project.deliverables]),
                ),
                (
                    "Work Breakdown Structure (LLD Fases)",
                    wbs_text
                ),
                (
                    "Sizing & Cronograma",
                    f"Complejidad Técnica: {project.sizing}\nHorizonte de Ejecución: {project.duration}",
                ),
                ("Estimación FinOps & TCO (3 Años)", finops_text),
                ("Justificación de ROI & Valor", roi_just),
                ("Perfiles y Gobernanza (RACI)", gov_roles),
                ("Riesgos de Ejecución y Mitigación", crit_risks),
            ]
            
            project_table = doc.add_table(rows=len(rows)+1, cols=2)
            finalize_table(project_table)
            
            merged = project_table.rows[0].cells[0].merge(project_table.rows[0].cells[1])
            set_cell_text(merged, "FICHA TÉCNICA DE PROYECTO DE INGENIERÍA", bold=True, font_size=11)
            shade_cell(merged, "0072BC")
            for run in merged.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

            for idx, (label, value) in enumerate(rows, 1):
                set_cell_text(
                    project_table.rows[idx].cells[0], label, bold=True, font_size=9
                )
                shade_cell(project_table.rows[idx].cells[0], "F2F2F2")
                set_cell_text(project_table.rows[idx].cells[1], value, font_size=9.5)
                for run in project_table.rows[idx].cells[1].paragraphs[0].runs:
                    run.font.color.rgb = BASE_TEXT_COLOR
            
            autofit_table_to_contents(project_table)
            add_spacer(doc, 15)
            project_idx += 1


def render_roadmap_page(doc, payload: BlueprintPayload, output_path: Path) -> Any:
    """Constructs and appends a strategic roadmap page to a document object.

    This function builds a comprehensive roadmap section within a provided
    `python-docx` document. It first attempts to generate and embed a visual
    Gantt chart from the payload data, using a temporary image file as an
    intermediary. Following the chart, it constructs a detailed table that
    maps each engineering initiative to its execution wave and lists its
    technical dependencies. Dependencies are programmatically classified as
    either internal to the engineering tower or cross-tower, and are
    color-coded within the table to visually distinguish their origin.

    Args:
        doc (docx.document.Document): The document object to which the roadmap
            content will be appended.
        payload (BlueprintPayload): A data structure containing all necessary data,
            including roadmap waves, project definitions, and dependency mappings.
        output_path (pathlib.Path): The filesystem path for the final output file.
            This path is used to derive a location for transient assets, such
            as the generated Gantt chart image.

    Returns:
        None. The function modifies the `doc` object in place.

    Raises:
        ImportError: If the `application.generate_mermaid_gantt` dependency
            cannot be imported.
        FileNotFoundError: If Gantt chart generation is reported as successful but
            the resulting image file cannot be found on disk.
        AttributeError: If the `payload` object lacks expected attributes, such
            as `roadmap` or `external_dependencies`, indicating malformed input.
    """
    add_heading_paragraph(doc, "Roadmap Estratégico y Matriz de Dependencias (SOTA)", level=1)
    
    add_body_paragraph(
        doc,
        "La siguiente matriz establece la ruta crítica de ejecución (Gantt). Las iniciativas no son compartimentos estancos; su secuenciación matemática garantiza que los habilitadores técnicos (dependencias) se desplieguen antes de la construcción de las capas superiores.",
        color_rgb=BASE_TEXT_COLOR
    )
    
    # Attempts to generate a Mermaid syntax Gantt chart; this operation may fail and should be handled gracefully.
    import os
    import tempfile

    from application.generate_mermaid_gantt import generate_mermaid_gantt
    
    payload_path = output_path.parent / f"blueprint_{payload.document_meta.tower_code.lower()}_payload.json"
    png_path = Path(tempfile.gettempdir()) / f"gantt_{payload_path.name}.png"
    if payload_path.exists() and generate_mermaid_gantt(payload_path, png_path):
        doc.add_picture(str(png_path), width=Inches(6.5))
        add_spacer(doc, 10)
        try:
            os.remove(png_path)
        except OSError:
            pass

    table = doc.add_table(rows=1, cols=3)
    finalize_table(table)
    headers = ["Horizonte / Fase", "Iniciativa de Ingeniería", "Dependencias Técnicas (Predecesores)"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=10)
        shade_cell(table.rows[0].cells[i], "D9EAF7")

    # Pre-processes and collects all internal projects into a lookup structure to optimize dependency resolution.
    internal_projects = set()
    for pilar in payload.pillars_analysis:
        for proj in pilar.projects_todo:
            internal_projects.add(proj.initiative)

    for wave in payload.roadmap:
        for project_name in wave.projects:
            #
            deps = []
            for d in payload.external_dependencies:
                if d.project == project_name:
                    dep_name = d.depends_on
                    tag = "[Interna]" if dep_name in internal_projects else "[Cross-Tower]"
                    deps.append(f"{tag} {dep_name}")
            
            deps_text = " • ".join(deps) if deps else "Independiente (Habilitador Fase 0)"
            
            row = table.add_row()
            set_cell_text(row.cells[0], wave.wave, bold=True, font_size=9.5)
            set_cell_text(row.cells[1], project_name, font_size=9.5)
            set_cell_text(row.cells[2], deps_text, font_size=9)
            
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BASE_TEXT_COLOR
                        if "[Cross-Tower]" in run.text:
                            run.font.color.rgb = RGBColor(192, 0, 0) # Defines the RGB value for Dark Red, used to signify cross-tower project dependencies.
                        elif "[Interna]" in run.text:
                            run.font.color.rgb = RGBColor(0, 112, 192) # Defines the RGB value for Blue, used to signify internal project dependencies.
                        
    autofit_table_to_contents(table)
    add_spacer(doc, 15)


def render_maturity_profile(doc, annex_data: dict) -> Any:
    r"""{'docstring': "Renders a detailed maturity profile section into a `docx.document.Document`.\n\nConstructs and appends a formatted maturity profile to the provided document. The profile consists of an introductory text, an optional radar chart image, a table detailing pillar scores and executive summaries, and an AS-IS analysis including strengths, gaps, and operational impacts. If `annex_data` is empty, a placeholder message is added instead.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` Document object to be modified in-place.\n    annex_data (dict): Data for the maturity profile. An empty dictionary results in a placeholder message. Expected structure includes top-level keys 'pillar_score_profile' and 'sections'. The 'pillar_score_profile' dictionary should contain textual introductions and a list of 'pillars' dictionaries. The 'sections' dictionary should contain an 'asis' dictionary with lists of 'strengths', 'gaps', and 'operational_impacts'.\n\nReturns:\n    None. The function modifies the `doc` object directly.\n\nRaises:\n    ValueError: If the image file path provided under the 'radar_chart' key within `annex_data` is invalid or refers to an unsupported image format."}."""
    add_heading_paragraph(doc, "Perfil de madurez por pilar", level=1)

    if not annex_data:
        add_body_paragraph(
            doc,
            "No se ha encontrado el perfil de madurez detallado para esta torre.",
            color_rgb=BASE_TEXT_COLOR,
        )
        return

    profile = annex_data.get("pillar_score_profile", {})
    asis = annex_data.get("sections", {}).get("asis", {})

    add_body_paragraph(doc, profile.get("profile_intro", ""), color_rgb=BASE_TEXT_COLOR)
    note = add_body_paragraph(
        doc,
        profile.get("scoring_method_note", ""),
        color_rgb=BASE_TEXT_COLOR,
    )
    if note:
        for run in note.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(127, 127, 127)

    radar_path = profile.get("radar_chart", "")
    if radar_path and Path(radar_path).exists():
        add_heading_paragraph(doc, "Gráfico radial", level=2)
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(radar_path, width=Inches(4.35))
        add_spacer(doc, 10)

    pillars = profile.get("pillars", [])
    add_heading_paragraph(doc, "Detalle compacto por pilar", level=2)
    if pillars:
        table = doc.add_table(rows=1, cols=4)
        finalize_table(table)
        headers = ["Pilar", "Score", "Nivel", "Lectura ejecutiva"]
        for i, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=10)
            shade_cell(table.rows[0].cells[i], "D9EAF7")

        for pillar in pillars:
            row = table.add_row()
            maturity_band = pillar.get("maturity_band") or _resolve_annex_band(
                pillar.get("score_display")
            )
            set_cell_text(
                row.cells[0], pillar.get("pillar_label", ""), bold=True, font_size=9.5
            )
            set_cell_text(
                row.cells[1],
                pillar.get("score_display", ""),
                font_size=9.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text(row.cells[2], maturity_band, font_size=9.5)
            set_cell_text(
                row.cells[3], pillar.get("executive_reading", ""), font_size=9.5
            )
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BASE_TEXT_COLOR
        autofit_table_to_contents(table)
        add_spacer(doc, 10)

    add_heading_paragraph(doc, "AS-IS resumido", level=2)
    add_body_paragraph(doc, asis.get("narrative", ""), color_rgb=BASE_TEXT_COLOR)

    strengths = asis.get("strengths", [])
    gaps = asis.get("gaps", [])
    if strengths or gaps:
        sg_table = doc.add_table(rows=2, cols=2)
        finalize_table(sg_table)
        set_cell_text(
            sg_table.rows[0].cells[0], "Fortalezas clave", bold=True, font_size=10
        )
        shade_cell(sg_table.rows[0].cells[0], "D9EAF7")
        set_cell_text(
            sg_table.rows[0].cells[1], "Brechas clave", bold=True, font_size=10
        )
        shade_cell(sg_table.rows[0].cells[1], "D9EAF7")

        strengths_cell = sg_table.rows[1].cells[0]
        strengths_cell.text = ""
        for item in strengths:
            p = strengths_cell.add_paragraph(f"• {item}")
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.first_line_indent = Pt(-10)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = BASE_TEXT_COLOR

        gaps_cell = sg_table.rows[1].cells[1]
        gaps_cell.text = ""
        for item in gaps:
            p = gaps_cell.add_paragraph(f"• {item}")
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.first_line_indent = Pt(-10)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = BASE_TEXT_COLOR

        autofit_table_to_contents(sg_table)
        add_spacer(doc, 10)

    impacts = asis.get("operational_impacts", [])
    if impacts:
        add_heading_paragraph(doc, "Implicaciones operativas clave", level=2)
        for impact in impacts:
            p = doc.add_paragraph(f"• {impact}")
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.first_line_indent = Pt(-10)
            for run in p.runs:
                run.font.color.rgb = BASE_TEXT_COLOR


def render_conclusion(doc, annex_data: dict) -> Any:
    """Appends a formatted conclusion section to a python-docx Document.

    This function adds a primary level-1 heading ("5. Conclusión") and populates
    the section with content extracted from the `annex_data` dictionary.
    Subsections are added with level-2 headings for the final assessment,
    executive message, priority focus areas, and a closing statement. Priority
    focus areas are rendered as a bulleted list with specific indentation.

    If the required conclusion data is not found within `annex_data`, a
    placeholder message is inserted into the document instead.

    Args:
        doc (docx.document.Document): The document object to which the conclusion
            section will be appended. This object is modified in-place.
        annex_data (dict): A dictionary containing the report data. The function
            expects a specific nested structure for the conclusion content, as
            shown below:
            {
                'sections': {
                    'conclusion': {
                        'final_assessment': str,
                        'executive_message': str,
                        'priority_focus_areas': list[str],
                        'closing_statement': str
                    }
                }
            }

    Returns:
        None. The function modifies the `doc` object directly.

    Raises:
        AttributeError: If `doc` is not a valid `docx.document.Document` object
            lacking the required `add_paragraph` method.
        TypeError: If the 'priority_focus_areas' key is present in the
            conclusion data but its value is not an iterable.
    """
    add_heading_paragraph(doc, "5. Conclusión", level=1)

    if not annex_data:
        add_body_paragraph(
            doc,
            "No se ha encontrado información de conclusión para esta torre.",
            color_rgb=BASE_TEXT_COLOR,
        )
        return

    conclusion = annex_data.get("sections", {}).get("conclusion", {})

    if conclusion.get("final_assessment"):
        add_heading_paragraph(doc, "Evaluación final", level=2)
        add_body_paragraph(
            doc, conclusion.get("final_assessment"), color_rgb=BASE_TEXT_COLOR
        )

    if conclusion.get("executive_message"):
        add_heading_paragraph(doc, "Mensaje para el responsable técnico", level=2)
        add_body_paragraph(
            doc, conclusion.get("executive_message"), color_rgb=BASE_TEXT_COLOR
        )

    priority_areas = conclusion.get("priority_focus_areas", [])
    if priority_areas:
        add_heading_paragraph(doc, "Áreas de foco prioritarias", level=2)
        for area in priority_areas:
            p = doc.add_paragraph(f"• {area}")
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.first_line_indent = Pt(-10)
            for run in p.runs:
                run.font.color.rgb = BASE_TEXT_COLOR

    if conclusion.get("closing_statement"):
        add_heading_paragraph(doc, "Próximos pasos", level=2)
        add_body_paragraph(
            doc, conclusion.get("closing_statement"), color_rgb=BASE_TEXT_COLOR
        )


def render_blueprint(
    payload: BlueprintPayload,
    output_path: Path,
    client_dir: Path,
    template_path: Path = DEFAULT_TEMPLATE_PATH,
) -> Path:
    r"""{'docstring': 'Renders a tower blueprint Word document from a structured data payload.\n\nOrchestrates the generation of a multi-section Word document from a\ntemplate. The function loads a specified .docx template, clears its\nexisting body content, and then sequentially populates it by invoking\ndedicated rendering functions for each document section. It aggregates data\nfrom the primary `payload` with supplementary client-specific information\n(e.g., annex data) before saving the final document to the output path.\n\nArgs:\n    payload: An object containing the structured data for all blueprint\n        sections.\n    output_path: The file system path where the generated .docx file will be\n        saved. Parent directories are created automatically if they do not\n        exist.\n    client_dir: The directory path containing client-specific supplementary\n        data files, such as annex and intelligence reports.\n    template_path: Path to the base .docx template file. If not provided,\n        a default system path is used.\n\nReturns:\n    The path to the newly created document, identical to `output_path`.\n\nRaises:\n    FileNotFoundError: If the template file or required client data files\n        do not exist at their specified paths.\n    IOError: If writing the output file to `output_path` fails due to\n        file system errors, such as insufficient permissions.'}."""
    annex_data = load_annex_data(client_dir, payload.document_meta.tower_code)
    client_intelligence = load_client_intelligence(client_dir)

    doc = Document(str(template_path))
    clear_document_body(doc)
    if doc.sections:
        create_page_number_footer(doc.sections[0])

    render_cover(doc, payload)
    render_snapshot_page(doc, payload, client_intelligence, annex_data)
    render_maturity_profile(doc, annex_data)
    render_cross_capabilities_analysis(doc, payload)

    render_consolidated_asis(doc, payload)
    render_fair_risk_heatmap(doc, payload)
    render_consolidated_tobe(doc, payload)
    render_roadmap_page(doc, payload, output_path)
    render_deep_project_charters(doc, payload)

    render_conclusion(doc, annex_data)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main(argv: list[str] | None = None) -> None:
    r"""{'docstring': 'Renders a tower blueprint from a command-line specified payload file.\n\nOrchestrates the blueprint rendering pipeline. This function acts as the main\nscript entry point, parsing a source payload path and a destination output\npath from command-line arguments. It then loads and resolves all required\ndata (payload, client configuration, annex data) before calling the final\nrendering function.\n\nThe process terminates with exit code 1 if the command-line arguments are\ninvalid.\n\nArgs:\n    argv: An optional list of command-line arguments, defaulting to `sys.argv`.\n        The list must contain the script name, a source payload path, and a\n        destination output path.\n\nReturns:\n    None. A blueprint file is written to the destination path and a success\n    message is printed to standard output.'}."""
    if len(argv if argv is not None else sys.argv) != 3:
        sys.exit(1)

    payload_path = Path((argv if argv is not None else sys.argv)[1])
    output_path = Path((argv if argv is not None else sys.argv)[2])
    client_dir = resolve_client_dir(payload_path, load_json(payload_path))
    annex_data = load_annex_data(client_dir, payload_path.parent.name)
    payload = load_payload(payload_path, annex_data=annex_data)

    render_blueprint(
        payload=payload,
        output_path=output_path,
        client_dir=client_dir,
    )
    print(f"Blueprint de Torre renderizado: {output_path}")


if __name__ == "__main__":
    main()
