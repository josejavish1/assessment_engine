import glob
import json
import os
import re
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from assessment_engine.scripts.lib.text_utils import format_currency_custom


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    p = cell.paragraphs[0]
    p.alignment = align
    for run in p.runs:
        run.font.bold = bold
        run.font.size = Pt(font_size)

def autofit_table_to_contents(table):
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1)

def extract_number(text):
    if not text: return 0.0
    # Try to find something like 163.745,45
    # Remove dots (thousands), replace comma with dot
    match = re.search(r'([\d\.]+,\d+)', text)
    if match:
        clean = match.group(1).replace(".", "").replace(",", ".")
        try:
            return float(clean)
        except Exception:
            return 0.0
    return 0.0

def render_consolidated_todo(working_dir: str, output_path: str):
    doc = Document()
    
    payloads = []
    search_pattern = os.path.join(working_dir, "T*", "blueprint_*_payload.json")
    for file_path in glob.glob(search_pattern):
        with open(file_path, 'r', encoding='utf-8-sig') as f:
            payloads.append(json.load(f))
            
    if not payloads:
        print("No payloads found for consolidated TO-DO.")
        return
        
    client_name = payloads[0].get("document_meta", {}).get("client_name", "Cliente")
    meta_lang = payloads[0].get("document_meta", {}).get("language", "es").lower()
    
    # Cargar marca y localización declarativa SOTA
    brand_path = Path("engine_config/brand_profile.json")
    brand = {}
    if brand_path.exists():
        with open(brand_path, "r", encoding="utf-8") as bf:
            brand = json.load(bf)
    styling = brand.get("styling", {})
    color_blue = styling.get("primary_color_hex", "0072BC")
    
    locales_path = Path("engine_config/locales.json")
    locales_data = {}
    if locales_path.exists():
        with open(locales_path, "r", encoding="utf-8-sig") as lf:
            locales_data = json.load(lf)
    vocab = locales_data.get(meta_lang, locales_data.get("es", {}))
    
    # 1. Merge Roadmap Waves & Dependencies
    merged_waves = {} # e.g. "Wave 1: 0-6m" -> ["proj1", "proj2"]
    merged_deps = []
    all_projects = []
    
    for data in payloads:
        merged_deps.extend(data.get("external_dependencies", []))
        
        for w in data.get("roadmap", []):
            w_name = w.get("wave", "Unknown Wave")
            if w_name not in merged_waves:
                merged_waves[w_name] = []
            merged_waves[w_name].extend(w.get("projects", []))
            
        for pilar in data.get("pillars_analysis", []):
            for proj in pilar.get("projects_todo", []):
                proj["tower"] = data.get("document_meta", {}).get("tower_name")
                proj["pilar"] = pilar.get("pilar_name")
                all_projects.append(proj)
                
    # Sort merged waves to reconstruct a single roadmap array
    sorted_wave_names = sorted(merged_waves.keys())
    unified_roadmap = [{"wave": wn, "projects": merged_waves[wn]} for wn in sorted_wave_names]
    
    # Create Unified Payload for Mermaid
    unified_payload = {
        "roadmap": unified_roadmap,
        "external_dependencies": merged_deps,
        "pillars_analysis": [
            {"projects_todo": all_projects}
        ]
    }
    
    # Calculate global financials
    total_opex = 0.0
    total_capex = 0.0
    for proj in all_projects:
        opex_text = proj.get("opex_estimate", "")
        capex_text = proj.get("capex_estimate", "")
        total_opex += extract_number(opex_text)
        total_capex += extract_number(capex_text)
    
    # Cover
    company = brand.get("company_name", "NTT DATA")
    classification = brand.get("default_classification", "Confidencial")
    doc.add_heading(f'Roadmap Estratégico y Financiero TO-DO: {client_name}', 0)
    doc.add_paragraph(f"Clasificación: {classification} ({company} Tier-1 Architecture Assessment)")
    
    # 1. Financial Summary
    doc.add_heading('1. Resumen Financiero Consolidado (TCO)', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    headers = ["Concepto", "Descripción", "Estimación (€)"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=10)
        shade_cell(table.rows[0].cells[i], color_blue)
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    row = table.add_row()
    set_cell_text(row.cells[0], "OPEX de Implantación")
    set_cell_text(row.cells[1], "Servicios profesionales de ingeniería para el diseño, despliegue y estabilización.")
    set_cell_text(row.cells[2], format_currency_custom(total_opex, vocab, meta_lang), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    
    row = table.add_row()
    set_cell_text(row.cells[0], "CAPEX Estimado")
    set_cell_text(row.cells[1], "Hardware, licencias perpetuas o suscripciones asociadas (Estimación Mínima).")
    set_cell_text(row.cells[2], format_currency_custom(total_capex, vocab, meta_lang), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    
    doc.add_paragraph("\nNota: Las estimaciones financieras son paramétricas (Bottom-Up) basadas en Unit Economics y el Work Breakdown Structure.")
    doc.add_page_break()
    
    # 2. Global Roadmap
    doc.add_heading('2. Master Plan de Transformación (Gantt Global)', level=1)
    doc.add_paragraph("Ruta crítica consolidada con interdependencias Cross-Tower.")
    
    # Generate Mermaid
    from application.generate_mermaid_gantt import generate_mermaid_gantt
    
    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json', encoding='utf-8-sig') as f:
        json.dump(unified_payload, f, ensure_ascii=False)
        temp_payload_path = f.name
        
    png_path = Path(tempfile.gettempdir()) / "global_gantt.png"
    if generate_mermaid_gantt(Path(temp_payload_path), png_path):
        doc.add_picture(str(png_path), width=Inches(7.0))
        try:
            os.remove(png_path)
            os.remove(temp_payload_path)
        except OSError:
            pass
            
    doc.add_page_break()
    
    # 3. Project Charters
    doc.add_heading('3. Fichas de Proyecto (Board-Ready Charters)', level=1)
    
    for proj in all_projects:
        doc.add_heading(proj.get("name", "Proyecto Desconocido"), level=2)
        
        # We reuse the logic from render_deep_project_charters
        # Resolve mitigated risk text
        mitigated_risk_text = "N/A"
        if proj.get("mitigates_risk_id"):
            mitigated_risk_text = f"ID: {proj.get('mitigates_risk_id')}"
                
        # Resolve FinOps
        capex = proj.get("capex_estimate", "N/A")
        opex = proj.get("opex_estimate", "N/A")
        clean_opex = opex.split("(Margen")[0].strip() if "(Margen" in opex else opex
        finops_text = f"Implantación: {clean_opex}\nLicencias/HW: {capex}"

        # Resolve WBS
        wbs_items = proj.get("wbs_breakdown", [])
        wbs_text = "Desglose no disponible."
        if wbs_items:
            try:
                wbs_text = "\n".join([f"• {w.get('task_name')} (Perfil: {w.get('required_profile')})" for w in wbs_items if isinstance(w, dict)])
            except Exception:
                pass

        proj_desc = proj.get("project_description", "N/A") or "N/A"
        smart_obj = proj.get("smart_objectives", "N/A") or "N/A"
        in_scope = "\n".join([f"• {item}" for item in proj.get("in_scope", [])]) or "N/A"
        out_scope = "\n".join([f"• {item}" for item in proj.get("out_of_scope", [])]) or "N/A"
        gov_roles = "\n".join([f"• {item}" for item in proj.get("governance_roles", [])]) or "N/A"
        crit_risks = "\n".join([f"• {item}" for item in proj.get("critical_risks", [])]) or "N/A"
        roi_just = proj.get("roi_justification", "N/A") or "N/A"
        
        # Get dependencies
        deps = [d.get("depends_on") for d in merged_deps if d.get("project") == proj.get("name")]
        deps_text = " • ".join(deps) if deps else "Independiente (Habilitador Fase 0)"

        rows = [
            ("Torre Tecnológica", proj.get("tower", "")),
            ("Descripción Ejecutiva", proj_desc),
            ("Objetivo SMART", smart_obj),
            ("Riesgo AS-IS Mitigado", mitigated_risk_text),
            ("Alcance (In-Scope)", in_scope),
            ("Fuera de Alcance (Out-of-Scope)", out_scope),
            ("Dependencias Técnicas", deps_text),
            ("Work Breakdown Structure (LLD Fases)", wbs_text),
            ("Sizing & Cronograma", f"Talla: {proj.get('sizing')}\nDuración: {proj.get('duration')}"),
            ("Estimación FinOps & TCO", finops_text),
            ("Justificación de ROI & Valor", roi_just),
            ("Perfiles y Gobernanza (RACI)", gov_roles),
            ("Riesgos de Ejecución y Mitigación", crit_risks),
        ]
        
        project_table = doc.add_table(rows=len(rows)+1, cols=2)
        project_table.style = 'Table Grid'
        
        merged = project_table.rows[0].cells[0].merge(project_table.rows[0].cells[1])
        set_cell_text(merged, "FICHA TÉCNICA DE PROYECTO DE INGENIERÍA", bold=True, font_size=10)
        shade_cell(merged, color_blue)
        for r in merged.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

        for idx, (label, value) in enumerate(rows, 1):
            set_cell_text(project_table.rows[idx].cells[0], label, bold=True, font_size=9)
            shade_cell(project_table.rows[idx].cells[0], "F2F2F2")
            set_cell_text(project_table.rows[idx].cells[1], value, font_size=9)
            
        doc.add_page_break()
            
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ Documento Global TO-DO generado en: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python render_togaf_consolidated_todo.py <working_dir> <output_doc.docx>")
        sys.exit(1)
    render_consolidated_todo(sys.argv[1], sys.argv[2])
