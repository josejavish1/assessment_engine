import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt, RGBColor
from pydantic import ValidationError

"""
Módulo render_global_report_from_template.py.
Contiene la lógica y utilidades principales para el pipeline de Assessment Engine.
"""

from domain.schemas.global_report import (
    BurningPlatformItem,
    ExecutionRoadmapDraft,
    ExecutiveDecisionsDraft,
    ExecutiveSummaryDraft,
    GlobalReportPayload,
    TargetVisionDraft,
    TowerBottomLineItem,
)
from infrastructure.docx_render_utils import (
    add_body_paragraph,
    add_heading_paragraph,
    autofit_table_to_contents,
    clear_paragraph,
    finalize_table,
    set_cell_text,
    shade_cell,
)
from infrastructure.global_maturity_policy import (
    safe_float,
    status_color_for_score,
)


def load_json(path) -> Any:
    """Read and deserialize a JSON file from a given path object."""
    return json.loads(path.read_text(encoding="utf-8-sig"))


def clean_t_codes(text) -> Any:
    """Sanitize a string by removing specific internal annotation patterns.

    Removes several patterns from an input string: T-codes within parentheses
    (e.g., `(T1...)`), standalone T-codes that form a whole word (e.g., `T12`),
    and internal reference tags (e.g., `[[REF:...]]` or `[REF:...]`). The
    resulting string has leading and trailing whitespace stripped.

    If the input is not a string, it is returned without modification.

    Args:
        text (Any): The input value to be sanitized.

    Returns:
        Any: A sanitized string with patterns removed, or the original input value
            if it was not a string.
    """
    if not isinstance(text, str):
        return text
    text = re.sub(r"\(T\d{1,2}[^\)]*\)", "", text)
    text = re.sub(r"\bT\d{1,2}\b", "", text)
    # Strips internal reference tags (e.g., `[[REF:...]]`, `[REF:...]`) from the text content to sanitize it for final presentation.
    text = re.sub(r"\[\[?REF:[^\]]*\]\]?", "", text)
    return text.strip()


def sanitize_client_name(text, client_name) -> Any:
    """Replaces a specific client's name in a string with a generic term.

    This function performs a series of case-insensitive regular expression
    substitutions to render a given text client-agnostic. It operates on two
    variations of the client name: the original string and a version with
    underscores converted to spaces.

    The sanitization process involves two main stages:
    1.  Removal of associated Spanish possessive prepositions (`de` or `del`)
        preceding the client's name to avoid dangling phrases after
        replacement (e.g., "informe de [Client]" becomes "informe").
    2.  Replacement of all standalone occurrences of the client's name with
        the generic Spanish phrase "la organización".

    If the input `text` is not a string, or if `client_name` is empty or None,
    the function returns the original `text` object without modification.

    Args:
        text (Any): The input text to process. If not a string, it is returned
            unmodified.
        client_name (str): The name of the client to find and replace. If an
            empty string or None, the original text is returned.

    Returns:
        Any: A new string with the client's name sanitized, or the original
            `text` object if no sanitization is performed.
    """
    if not isinstance(text, str) or not client_name:
        return text

    # Sanitizes the client name to create a valid, filesystem-safe filename by replacing spaces and special characters with underscores.
    clean_name = client_name.replace("_", " ")

    # Removes client-specific possessive prepositions (e.g., 'de [Client]', 'del [Client]') from the narrative as the first step in text generalization.
    text = re.sub(
        rf"\s+del?\s+{re.escape(client_name)}\b", "", text, flags=re.IGNORECASE
    )
    text = re.sub(
        rf"\s+del?\s+{re.escape(clean_name)}\b", "", text, flags=re.IGNORECASE
    )

    # Replaces direct client name mentions with the generic term 'the organization' to produce a client-agnostic narrative.
    text = re.sub(
        rf"\b{re.escape(client_name)}\b", "la organización", text, flags=re.IGNORECASE
    )
    text = re.sub(
        rf"\b{re.escape(clean_name)}\b", "la organización", text, flags=re.IGNORECASE
    )

    return text


def clear_document_body(doc) -> Any:
    """Removes all content from a document's body while preserving section properties.

    This function operates on the underlying lxml element of the document body. It
    iterates through all child elements and removes them, with the sole exception
    of the `w:sectPr` (section properties) element. This approach allows for
    clearing all visible content, such as paragraphs and tables, while retaining
    critical page formatting like margins, orientation, and paper size.

    The modification is performed in-place.

    Args:
        doc (docx.document.Document): The python-docx Document object to be cleared.

    Returns:
        None.

    Raises:
        AttributeError: If the `doc` object does not conform to the expected
            internal structure of the `python-docx` library, specifically
            lacking a `_body._element` attribute path.
    """
    body = doc._body._element
    for child in list(body):
        if (
            child.tag
            != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
        ):
            body.remove(child)


def add_spacer(doc, points=12) -> Any:
    """Add a vertical spacer of a specified point height to a document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)


def add_smart_bullet_list(container, items, color_rgb=None, bold_prefix=True) -> Any:
    """Renders a list of strings as a formatted bulleted list within a python-docx container.

    This function iterates through each string in the provided list, adding it as a
    new paragraph to the container. It first attempts to apply the "List Bullet"
    style from the document template. If this style is not found (resulting in a
    KeyError), it falls back to a manually formatted list entry by prepending a
    "•" character and setting paragraph indentation programmatically.

    The function also supports prefix-based formatting. If an item string contains
    a ": " or " - " separator, the text preceding the separator can be
    conditionally bolded. If no separator is present, the entire item is treated
    as the prefix for formatting purposes.

    Args:
        container (Any): A python-docx object that supports the `add_paragraph`
            method, such as a `Document` or `_Cell` object.
        items (Union[str, List[str]]): A single string or a list of strings to be
            rendered as bullet points. A single string is treated as a single-item list.
        color_rgb (Optional[docx.shared.RGBColor]): The `RGBColor` to apply to the
            text. If None, the default text color of the container is used.
        bold_prefix (bool): If True, bolds the text preceding a separator (": "
            or " - ") or the entire line if no separator is found. Defaults to
            True.

    Returns:
        None: The container object is modified in-place.
    """
    # Ensures robustness by wrapping a string input in a list if a list is expected. This prevents unintended character-wise iteration over the string.
    if isinstance(items, str):
        items = [items]
    if not items:
        return

    for item in items:
        p = container.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        try:
            p.style = "List Bullet"
        except KeyError:
            # Implements a fallback mechanism to a default style if the specified style name does not exist in the document template, preventing rendering failures.
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.first_line_indent = Pt(-15)
            b = p.add_run("• ")
            b.font.size = Pt(10.5)
            if color_rgb:
                b.font.color.rgb = color_rgb

        p.paragraph_format.space_after = Pt(6)
        text = clean_t_codes(str(item).strip())

        sep = ": " if ": " in text else (" - " if " - " in text else None)
        if sep:
            parts = text.split(sep, 1)
            r1 = p.add_run(parts[0] + sep)
            r1.bold = bold_prefix
            r1.font.size = Pt(10.5)
            r2 = p.add_run(parts[1])
            r2.font.size = Pt(10.5)
            if color_rgb:
                r1.font.color.rgb = color_rgb
                r2.font.color.rgb = color_rgb
        else:
            r = p.add_run(text)
            r.bold = bold_prefix
            r.font.size = Pt(10.5)
            if color_rgb:
                r.font.color.rgb = color_rgb


def render_cover(doc, payload: GlobalReportPayload) -> Any:
    """Renders a formatted cover page into a `docx.Document` object.

    This function constructs and appends a cover page to the provided document.
    It adds a styled title, the client's name, report metadata (date and
    version), and a multi-paragraph legal disclaimer. Each element is
    programmatically formatted with specific fonts, sizes, colors, alignment, and
    paragraph spacing. A page break is inserted at the end of the cover page.

    Args:
        doc (docx.document.Document): The document object to be modified in-place.
        payload (GlobalReportPayload): A data object containing report metadata,
            including `meta.client`, `meta.date`, and `meta.version`.

    Returns:
        None. The function modifies the `doc` object directly.
    """
    # Decreases the top margin of the first page to optimize vertical space utilization and improve the header's visual placement.
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Applies title case capitalization while preserving lowercase for minor words such as articles, prepositions, and conjunctions to adhere to standard typographical conventions.
    run = title_p.add_run("Informe Estratégico de\nMadurez Tecnológica")
    run.font.size = Pt(34)
    run.font.name = "Georgia"
    run.font.color.rgb = RGBColor(0, 114, 188)
    run.bold = False

    # Reduces the vertical spacing between the report title and client name paragraphs to achieve a more compact and visually integrated header block.
    add_spacer(doc, 30)

    client_p = doc.add_paragraph()
    client_run = client_p.add_run(payload.meta.client.upper())
    client_run.font.size = Pt(24)
    client_run.font.name = "Arial"
    client_run.bold = True

    # The vertical spacing for this paragraph block is substantially reduced to prevent its content from being pushed down the page, ensuring proper alignment.
    doc.add_paragraph().paragraph_format.space_after = Pt(100)

    version_p = doc.add_paragraph()
    v_text = f"Fecha: {payload.meta.date}\nReferencia: {payload.meta.version}"
    version_run = version_p.add_run(v_text)
    version_run.font.size = Pt(14)
    version_run.font.name = "Arial"
    version_run.font.color.rgb = RGBColor(127, 127, 127)

    # A 150-point vertical spacer is applied to anchor the legal disclaimer to the bottom margin of the final page, preventing premature page breaks and ensuring consistent document layout.
    add_spacer(doc, 150)

    disclaimer_p1 = doc.add_paragraph()
    disclaimer_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p1.paragraph_format.space_after = Pt(
        2
    )  # Sets a minimum paragraph spacing for the legal disclaimer block to ensure readability and compliance with formatting standards.
    disclaimer_title_run = disclaimer_p1.add_run("Aviso Legal y Confidencialidad: ")
    disclaimer_title_run.font.size = Pt(8)
    disclaimer_title_run.font.name = "Arial"
    disclaimer_title_run.font.color.rgb = RGBColor(127, 127, 127)
    disclaimer_title_run.bold = True

    p1_text = "El presente documento constituye un informe de evaluación rápida (Fast Assessment) elaborado por NTT DATA basándose exclusivamente en la información, entrevistas y datos proporcionados por la organización, los cuales no han sido sometidos a una auditoría o verificación independiente por nuestra parte. Por tanto, NTT DATA no otorga ninguna garantía, expresa o implícita, sobre la exactitud, integridad o fiabilidad absoluta de dicha información base."
    r1 = disclaimer_p1.add_run(p1_text)
    r1.font.size = Pt(8)
    r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor(127, 127, 127)

    disclaimer_p2 = doc.add_paragraph()
    disclaimer_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p2.paragraph_format.space_after = Pt(2)
    p2_text = 'Las conclusiones, proyecciones (incluyendo el estado objetivo "TO-BE", los horizontes del roadmap y las estimaciones de impacto de negocio) y recomendaciones tienen un carácter puramente orientativo. Estas proyecciones constituyen estimaciones sujetas a incertidumbres operativas y de mercado, por lo que los resultados reales futuros podrían diferir de los planteados.'
    r2 = disclaimer_p2.add_run(p2_text)
    r2.font.size = Pt(8)
    r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor(127, 127, 127)

    disclaimer_p3 = doc.add_paragraph()
    disclaimer_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p3.paragraph_format.space_after = Pt(2)
    p3_text = "Este informe tiene un propósito estrictamente técnico y estratégico. Toda decisión ejecutiva, operativa o de inversión que la organización tome basándose en este informe es de su exclusiva responsabilidad. NTT DATA y sus representantes no asumen obligación legal ni responsabilidad alguna por posibles daños directos, indirectos, especiales o consecuentes derivados del uso o dependencia de la información aquí contenida sin una validación posterior exhaustiva."
    r3 = disclaimer_p3.add_run(p3_text)
    r3.font.size = Pt(8)
    r3.font.name = "Arial"
    r3.font.color.rgb = RGBColor(127, 127, 127)

    disclaimer_p4 = doc.add_paragraph()
    disclaimer_p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p4.paragraph_format.space_after = Pt(0)
    p4_text = "Finalmente, este documento y la información que contiene son confidenciales. Este material no podrá ser distribuido, reproducido, divulgado ni utilizado para ningún otro propósito, total o parcialmente, sin el consentimiento previo por escrito de NTT DATA."
    r4 = disclaimer_p4.add_run(p4_text)
    r4.font.size = Pt(8)
    r4.font.name = "Arial"
    r4.font.color.rgb = RGBColor(127, 127, 127)

    doc.add_page_break()


def render_executive_summary(
    doc,
    data: ExecutiveSummaryDraft,
    heatmap: list,
    visuals: dict,
    client_dir,
    client_name="",
):
    """Renders the executive summary section into a `docx.document.Document` object.

    Constructs and populates the executive summary by sequentially adding several
    components to the document. This includes a main score and headline table, a
    prose narrative with bullet points, a table of key business impacts, and an
    optional radar chart image. The function calculates an average score from
    heatmap data and sanitizes text by replacing client name placeholders. The
    input `doc` object is modified in place.

    Args:
        doc (docx.document.Document): The `python-docx` document object to be
            modified.
        data (ExecutiveSummaryDraft): A data object containing the textual content
            for the summary, including the headline, narrative, and key business
            impacts.
        heatmap (list[dict]): A list of dictionaries, where each dictionary
            represents a finding. Each is expected to contain a 'score' key with a
            value convertible to a float for calculating the overall average score.
        visuals (dict[str, str]): A dictionary mapping visual identifiers to their
            filenames. This function uses the 'radar_chart' key to locate the
            image file.
        client_dir (pathlib.Path): The base directory for the client's report,
            used as the root for resolving paths to visual assets.
        client_name (str): The client's name, used to replace placeholders within
            the textual content. Defaults to an empty string if not provided.

    Returns:
        None
    """
    BASE_TEXT_COLOR = RGBColor(46, 64, 77)  #
    add_heading_paragraph(doc, "1. Resumen Ejecutivo", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Configures table column widths to a 35% (2.1") and 65% (3.9") distribution. This ratio is optimized for content balance and readability on A4 paper.
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(3.9)

    score_cell = table.rows[0].cells[0]
    score_cell.width = Inches(2.1)
    score_val = "N/A"
    if heatmap:
        try:
            score_val = str(
                round(
                    sum(float(t.get("score", 0)) for t in heatmap) / len(heatmap),
                    1,
                )
            )
        except Exception:
            pass

    # Applies specific styling to the score display: formats as '[X] / 5', sets a #0072BC background fill, and uses a 36pt white font.
    set_cell_text(
        score_cell,
        f"{score_val} / 5",
        bold=True,
        font_size=36,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade_cell(score_cell, "0072BC")
    for r in score_cell.paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    desc_cell = table.rows[0].cells[1]
    desc_cell.width = Inches(3.9)
    desc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(desc_cell, "F2F2F2")

    headline = sanitize_client_name(clean_t_codes(data.headline), client_name)

    # Applies bold formatting exclusively to the substring of the headline that precedes the first colon, creating a visual distinction for the primary label.
    clear_paragraph(desc_cell.paragraphs[0])
    p = desc_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if ":" in headline:
        parts = headline.split(":", 1)
        r1 = p.add_run(parts[0] + ":")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.name = "Arial"
        r2 = p.add_run(parts[1])
        r2.bold = False
        r2.font.size = Pt(12)
        r2.font.name = "Arial"
    else:
        r = p.add_run(headline)
        r.bold = False
        r.font.size = Pt(12)
        r.font.name = "Arial"

    add_spacer(doc, 15)

    # The narrative and associated bullet points are rendered externally to the main table structure. This architectural choice improves document flow and readability.
    narrative = sanitize_client_name(clean_t_codes(data.narrative), client_name)
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚ])", narrative.strip())

    if sentences:
        # The first sentence of the narrative is semantically treated as the executive summary or 'Bottom Line' and is styled accordingly.
        p_intro = doc.add_paragraph()
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_intro = p_intro.add_run(sentences[0])
        r_intro.font.size = Pt(10.5)
        r_intro.font.color.rgb = BASE_TEXT_COLOR
        p_intro.paragraph_format.space_before = Pt(12)
        p_intro.paragraph_format.space_after = Pt(6)

        # Renders all sentences subsequent to the first as a bulleted list to enhance readability and structure.
        if len(sentences) > 1:
            add_smart_bullet_list(
                doc, sentences[1:], color_rgb=BASE_TEXT_COLOR, bold_prefix=False
            )

    add_spacer(doc, 15)

    #
    impact_table = doc.add_table(rows=1, cols=1)
    finalize_table(impact_table)

    header_cell = impact_table.rows[0].cells[0]
    set_cell_text(
        header_cell,
        "Principales impactos de negocio",
        bold=True,
        font_size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade_cell(header_cell, "0072BC")
    for r in header_cell.paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    impacts = [sanitize_client_name(i, client_name) for i in data.key_business_impacts]
    for impact in impacts:
        row = impact_table.add_row()
        body_cell = row.cells[0]
        shade_cell(body_cell, "F2F2F2")

        # Modifies the cell's first paragraph (index 0) in-place. This avoids creating a new paragraph, which would introduce unwanted vertical spacing and disrupt table layout.
        p_bull = body_cell.paragraphs[0]
        clear_paragraph(p_bull)
        p_bull.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        try:
            p_bull.style = "List Bullet"
        except KeyError:
            p_bull.paragraph_format.left_indent = Pt(20)
            p_bull.paragraph_format.first_line_indent = Pt(-15)
            b_run = p_bull.add_run("• ")
            b_run.font.size = Pt(10.5)
            b_run.font.color.rgb = BASE_TEXT_COLOR

        p_bull.paragraph_format.space_before = Pt(4)
        p_bull.paragraph_format.space_after = Pt(4)

        text = clean_t_codes(str(impact).strip())
        r_text = p_bull.add_run(text)
        r_text.bold = False
        r_text.font.size = Pt(10.5)
        r_text.font.color.rgb = BASE_TEXT_COLOR

    radar_path = client_dir / visuals.get("radar_chart", "")
    if radar_path.exists():
        add_spacer(doc, 15)
        doc.add_picture(str(radar_path), width=Inches(6.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def render_burning_platform(
    doc, platform_risks: list[BurningPlatformItem], client_name=""
):
    r"""{'docstring': "Renders a section detailing systemic threats into a `python-docx` document.\n\nThis function appends a titled section for systemic threats to an existing\n`python-docx` document. For each threat provided in `platform_risks`, it\nconstructs a dedicated, formatted three-row table. The table's first row\nserves as a header with the threat's theme, the second row details the\nassociated business risk, and the third row enumerates the root causes as\na bulleted list. All text content is sanitized to process special codes\nand substitute placeholders with the specified client's name.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` Document object to which\n        the content will be appended.\n    platform_risks (list[BurningPlatformItem]): A list of objects, each\n        representing a systemic threat. Each object must expose `theme` (str),\n        `business_risk` (str), and `root_causes` (list[str]) attributes.\n    client_name (str): The name of the client, used for substituting\n        placeholders in the text. Defaults to an empty string.\n\nRaises:\n    AttributeError: If an object in the `platform_risks` list lacks the\n        required `theme`, `business_risk`, or `root_causes` attributes."}."""
    BASE_TEXT_COLOR = RGBColor(46, 64, 77)  #
    add_heading_paragraph(doc, "2. Principales Amenazas Sistémicas", level=1)
    add_body_paragraph(
        doc,
        "Identificación de riesgos críticos que comprometen la viabilidad operativa y la agilidad de la organización.",
        color_rgb=BASE_TEXT_COLOR,
    )
    for i, risk in enumerate(platform_risks, start=1):
        table = doc.add_table(rows=3, cols=1)
        finalize_table(table)

        #
        h_cell = table.rows[0].cells[0]
        theme = sanitize_client_name(clean_t_codes(risk.theme), client_name)
        theme = theme[0].upper() + theme[1:] if theme else ""
        set_cell_text(
            h_cell,
            f"Amenaza {i}: {theme}",
            bold=True,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(h_cell, "0072BC")
        for r in h_cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

        #
        r_cell = table.rows[1].cells[0]
        shade_cell(r_cell, "F2F2F2")
        clear_paragraph(r_cell.paragraphs[0])
        p1 = r_cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1_label = p1.add_run("Riesgo de Negocio: ")
        r1_label.bold = True
        r1_label.font.color.rgb = BASE_TEXT_COLOR
        r1_label.font.size = Pt(10.5)
        r1_text = p1.add_run(
            sanitize_client_name(clean_t_codes(risk.business_risk), client_name)
        )
        r1_text.bold = False
        r1_text.font.color.rgb = BASE_TEXT_COLOR
        r1_text.font.size = Pt(10.5)

        #
        c_cell = table.rows[2].cells[0]
        shade_cell(c_cell, "F2F2F2")
        clear_paragraph(c_cell.paragraphs[0])
        p2 = c_cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2_label = p2.add_run("Causas Raíz: ")
        r2_label.bold = True
        r2_label.font.color.rgb = BASE_TEXT_COLOR
        r2_label.font.size = Pt(10.5)
        causes_list = risk.root_causes
        sanitized_causes = [
            sanitize_client_name(clean_t_codes(c), client_name) for c in causes_list
        ]
        r2_text = p2.add_run("\n" + "\n".join([f"  • {c}" for c in sanitized_causes]))
        r2_text.bold = False
        r2_text.font.color.rgb = BASE_TEXT_COLOR
        r2_text.font.size = Pt(10.5)

        add_spacer(doc, 10)


def render_tower_bottom_lines(
    doc, heatmap: list, tower_texts: list[TowerBottomLineItem], client_name=""
):
    """Populates a Word document with a technology area diagnosis table.

    This function appends a new section titled '3. Diagnóstico por Área Tecnológica'
    to the document. The section contains a three-column table summarizing the
    assessment for each technology area ('tower') from the heatmap data. For each
    area, the table displays its name, a maturity score with a corresponding
    descriptive band, and an executive diagnosis. The maturity cell is color-coded
    based on its score.

    The executive diagnosis text is sourced preferentially from a matching item in
    `tower_texts` based on the technology area ID. If no match is found, the
    function falls back to using the `executive_message` from the `heatmap` data.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to be
            modified in-place.
        heatmap (list[dict]): A list of dictionaries representing technology areas.
            Each dictionary is expected to contain 'id', 'name', 'score', 'band',
            and 'executive_message' keys.
        tower_texts (list[TowerBottomLineItem]): A list of data objects providing
            high-priority diagnostic text. Each object must have an 'id'
            attribute for matching and a 'bottom_line' attribute for the text.
        client_name (str): The client's name, used to replace a placeholder token
            within the final diagnostic text. Defaults to an empty string.

    Returns:
        None. The function modifies the `doc` object in-place.

    Raises:
        AttributeError: If an object within `tower_texts` lacks the required `id` or
            `bottom_line` attributes.
        TypeError: If `heatmap` or `tower_texts` are not iterable (e.g., not a list).
    """
    BASE_TEXT_COLOR = RGBColor(46, 64, 77)  #
    add_heading_paragraph(doc, "3. Diagnóstico por Área Tecnológica", level=1)
    table = doc.add_table(rows=1, cols=3)
    finalize_table(table)
    for i, h in enumerate(["Área", "Madurez", "Diagnóstico Ejecutivo"]):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            bold=False,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table.rows[0].cells[i], "0072BC")
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
    color_map = {"E06666": "F4CCCC", "FFD966": "FFF2CC", "93C47D": "D9EAD3"}
    for t in heatmap:
        # The color value is explicitly recalculated here to enforce strict document-wide consistency, overriding any inherited or default template styles.
        strict_color = status_color_for_score(safe_float(t.get("score"), 0.0))

        row = table.add_row()
        set_cell_text(
            row.cells[0],
            f"{t.get('id', '')}\n{t.get('name', '')}",
            bold=False,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(
            row.cells[1],
            f"{t.get('score', '')}\n({t.get('band', '')})",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10.5,
            bold=False,
        )
        shade_cell(row.cells[1], color_map.get(strict_color, "FFFFFF"))

        #
        tower_id = t.get("id", "")
        text_val = ""
        for bottom_line_item in tower_texts:
            if bottom_line_item.id == tower_id:
                text_val = bottom_line_item.bottom_line
                break

        if not text_val:
            text_val = t.get("executive_message", "")

        set_cell_text(
            row.cells[2],
            sanitize_client_name(clean_t_codes(str(text_val)), client_name),
            font_size=10.5,
        )
        #
        for p in row.cells[2].paragraphs:
            for r in p.runs:
                r.font.color.rgb = BASE_TEXT_COLOR
    autofit_table_to_contents(table)
    add_spacer(doc, 20)


def render_target_vision(doc, vision: TargetVisionDraft, client_name="") -> Any:
    """Renders the 'Target State Vision' section into a `docx.document.Document`.

    This function populates the provided Document object with a formatted section
    detailing the strategic vision. It constructs a primary heading, a table for
    the value proposition, and bulleted lists for evolution principles and
    strategic pillars. Text content is sanitized, and client-specific
    placeholders are replaced.

    Args:
        doc (docx.document.Document): The document object to be modified in-place.
        vision (TargetVisionDraft): A data object containing the target state
            vision details. It must possess `value_proposition`,
            `evolution_principles`, and `strategic_pillars` attributes.
        client_name (str): The name of the client to substitute into text
            placeholders. Defaults to an empty string.

    Returns:
        None. The `doc` object is modified directly.

    Raises:
        AttributeError: If the `vision` object is missing one or more of the
            required attributes.
    """
    BASE_TEXT_COLOR = RGBColor(46, 64, 77)  #
    add_heading_paragraph(doc, "4. Visión de Estado Objetivo (To-Be)", level=1)

    table = doc.add_table(rows=2, cols=1)
    finalize_table(table)

    #
    v_cell_header = table.rows[0].cells[0]
    set_cell_text(
        v_cell_header,
        "Propuesta de valor estratégica",
        bold=True,
        font_size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade_cell(v_cell_header, "0072BC")
    for r in v_cell_header.paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    #
    v_cell_body = table.rows[1].cells[0]
    shade_cell(v_cell_body, "F2F2F2")
    clear_paragraph(v_cell_body.paragraphs[0])
    p_body = v_cell_body.paragraphs[0]
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    prop = sanitize_client_name(clean_t_codes(vision.value_proposition), client_name)
    prop = prop[0].upper() + prop[1:] if prop else ""
    prop = prop.replace("─", "").replace("—", "").replace(";", ",")
    r_body = p_body.add_run(prop)
    r_body.font.color.rgb = BASE_TEXT_COLOR
    r_body.font.size = Pt(10.5)
    r_body.bold = False

    add_spacer(doc, 10)
    if vision.evolution_principles:
        add_heading_paragraph(doc, "Principios de Evolución de Madurez", level=2)
        p_list = [
            f"{p.principle}: {sanitize_client_name(p.description, client_name)}"
            for p in vision.evolution_principles
        ]
        add_smart_bullet_list(doc, p_list, color_rgb=BASE_TEXT_COLOR)
    add_heading_paragraph(doc, "Pilares Estratégicos Habilitadores", level=2)
    pillars = vision.strategic_pillars
    pillar_texts = []
    for p in pillars:
        p_name = sanitize_client_name(p.pillar, client_name)
        p_desc = sanitize_client_name(p.description, client_name)
        pillar_texts.append(f"{p_name}: {p_desc}")

    add_smart_bullet_list(
        doc, pillar_texts, color_rgb=BASE_TEXT_COLOR, bold_prefix=True
    )
    add_spacer(doc, 20)


def render_execution_roadmap(
    doc, roadmap: ExecutionRoadmapDraft, visuals: dict, client_dir, client_name=""
):
    """Renders the execution roadmap section into a Word document.

    This function constructs the "Plan de Implementación y Horizontes Temporales"
    section of the report. It first generates a table listing the defined
    transversal programs and their descriptions. It then iterates through the
    time horizons specified in the roadmap data (e.g., "Quick Wins", "Year 1").
    For each horizon that contains initiatives, a detailed table is created,
    listing each initiative's associated program, title, business case,
    start month, and duration.

    Args:
        doc (docx.document.Document): The python-docx Document object to be
            modified.
        roadmap (ExecutionRoadmapDraft): A data object containing the structured
            roadmap information, including programs and initiatives per time
            horizon.
        visuals (dict): This parameter is currently unused.
        client_dir (str): This parameter is currently unused.
        client_name (str): The client's name, used for placeholder substitution
            in text content. Defaults to an empty string.

    Returns:
        None: The function modifies the `doc` object in place.
    """
    BASE_TEXT_COLOR = RGBColor(46, 64, 77)  #
    add_heading_paragraph(
        doc, "5. Plan de Implementación y Horizontes Temporales", level=1
    )
    add_body_paragraph(
        doc,
        "La transformación tecnológica se articula en programas transversales ejecutados en cuatro horizontes temporales para asegurar la resiliencia y escalabilidad del negocio.",
        color_rgb=BASE_TEXT_COLOR,
    )
    table_p = doc.add_table(rows=1, cols=1)
    finalize_table(table_p)

    #
    p_cell_header = table_p.rows[0].cells[0]
    set_cell_text(
        p_cell_header,
        "Programas transversales definidos",
        bold=True,
        font_size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade_cell(p_cell_header, "0072BC")
    for r in p_cell_header.paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    #
    programs = roadmap.programs
    for p in programs:
        row = table_p.add_row()
        body_cell = row.cells[0]
        shade_cell(body_cell, "F2F2F2")

        p_bull = body_cell.paragraphs[0]
        clear_paragraph(p_bull)
        p_bull.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        try:
            p_bull.style = "List Bullet"
        except KeyError:
            p_bull.paragraph_format.left_indent = Pt(20)
            p_bull.paragraph_format.first_line_indent = Pt(-15)
            b_run = p_bull.add_run("• ")
            b_run.font.size = Pt(10.5)
            b_run.font.color.rgb = BASE_TEXT_COLOR

        p_bull.paragraph_format.space_before = Pt(6)
        p_bull.paragraph_format.space_after = Pt(6)

        program_name = clean_t_codes(p.name)
        program_desc = sanitize_client_name(clean_t_codes(p.description), client_name)

        r_name = p_bull.add_run(f"{program_name}: ")
        r_name.bold = True
        r_name.font.size = Pt(10.5)
        r_name.font.color.rgb = BASE_TEXT_COLOR

        r_desc = p_bull.add_run(program_desc)
        r_desc.bold = False
        r_desc.font.size = Pt(10.5)
        r_desc.font.color.rgb = BASE_TEXT_COLOR

    add_spacer(doc, 15)
    horizons = roadmap.horizons
    mapping = [
        ("Quick Wins (0-3 meses)", horizons.quick_wins_0_3_months),
        ("Año 1 (4-12 meses)", horizons.year_1_3_12_months),
        ("Año 2 (13-24 meses)", horizons.year_2_12_24_months),
        ("Año 3 (25-36 meses)", horizons.year_3_24_36_months),
    ]
    for title, initiatives in mapping:
        if not initiatives:
            continue
        add_heading_paragraph(doc, title, level=3)
        table = doc.add_table(rows=1, cols=5)
        finalize_table(table)
        headers = [
            "Programa",
            "Iniciativa",
            "Impacto / Business Case",
            "Inicio",
            "Dur.",
        ]
        for i, h in enumerate(headers):
            set_cell_text(
                table.rows[0].cells[i],
                h,
                bold=False,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(table.rows[0].cells[i], "0072BC")
            for r in table.rows[0].cells[i].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
        for init in initiatives:
            row = table.add_row()
            set_cell_text(
                row.cells[0],
                clean_t_codes(init.program),
                bold=False,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text(
                row.cells[1],
                sanitize_client_name(clean_t_codes(init.title), client_name),
                bold=False,
                font_size=10.5,
            )
            set_cell_text(
                row.cells[2],
                sanitize_client_name(clean_t_codes(init.business_case), client_name),
                font_size=10.5,
            )
            set_cell_text(
                row.cells[3],
                f"M{init.start_month}",
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text(
                row.cells[4],
                f"{init.duration_months}m",
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            #
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = BASE_TEXT_COLOR
        autofit_table_to_contents(table)
        add_spacer(doc, 15)


def render_executive_decisions(
    doc, decisions: ExecutiveDecisionsDraft, client_name=""
) -> Any:
    r"""{'docstring': "Populates a `python-docx` document with a formatted executive decisions section.\n\nModifies the provided document object by adding a main heading, a descriptive\nparagraph, and a table listing priority decisions. The table includes columns\nfor the decision's scope, the required action, and the impact of a delay.\nCell text is sanitized based on the provided client name.\n\nArgs:\n    doc (docx.document.Document): The document object to be modified.\n    decisions (ExecutiveDecisionsDraft): A data object containing the list of\n        executive decisions to render.\n    client_name (str): The name of the client, used to sanitize text content\n        in the table.\n\nReturns:\n    None: The function modifies the `doc` object in-place."}."""
    BASE_TEXT_COLOR = RGBColor(46, 64, 77)  #
    add_heading_paragraph(doc, "6. Decisiones Ejecutivas Prioritarias", level=1)
    add_body_paragraph(
        doc,
        "Líneas de decisión requeridas por la Dirección para desbloquear la ejecución del programa y mitigar los riesgos identificados.",
        color_rgb=BASE_TEXT_COLOR,
    )
    table = doc.add_table(rows=1, cols=3)
    finalize_table(table)
    headers = ["Ámbito de Decisión", "Acción Requerida", "Impacto de Demora"]
    for i, h in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            bold=True,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table.rows[0].cells[i], "0072BC")
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
    for d in decisions.immediate_decisions:
        row = table.add_row()
        set_cell_text(
            row.cells[0],
            sanitize_client_name(clean_t_codes(d.decision_type), client_name),
            bold=False,
            font_size=10.5,
        )
        set_cell_text(
            row.cells[1],
            sanitize_client_name(clean_t_codes(d.action_required), client_name),
            font_size=10.5,
        )
        set_cell_text(
            row.cells[2],
            sanitize_client_name(clean_t_codes(d.impact_if_delayed), client_name),
            font_size=10.5,
        )
        shade_cell(row.cells[2], "FFFFFF")
        #
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = BASE_TEXT_COLOR
    autofit_table_to_contents(table)


def create_page_number_footer(section) -> Any:
    """Adds a 'Page X of Y' footer to a document section, clearing the first page's footer.

    This function configures a document section to have a distinct footer for its
    first page, which is typically a cover page. It sets the
    `different_first_page_header_footer` property to True and clears all content
    from the first page's footer.

    For all subsequent pages within the section, a centered footer is added with
    the format "Página {PAGE} de {NUMPAGES}". The implementation directly
    manipulates the underlying Office Open XML (OOXML) structure to insert the
    dynamic `PAGE` and `NUMPAGES` fields. The footer text is styled as 9pt
    grey Arial.

    Args:
        section (docx.section.Section): The `python-docx` Section object to be
            modified in-place.

    Returns:
        None. The section object is modified directly.

    Raises:
        AttributeError: If the section object lacks expected attributes like
            `first_page_footer` or `footer`.
        IndexError: If the section's default footer does not contain at least one
            paragraph to modify.
    """
    # Enables a distinct footer for the first page. This setting is required to treat the cover page layout independently from subsequent pages.
    section.different_first_page_header_footer = True

    # The first-page footer is intentionally cleared. This is necessary to maintain a clean cover page layout free of standard footer content.
    first_page_footer = section.first_page_footer
    for p in first_page_footer.paragraphs:
        clear_paragraph(p)

    # Configures the standard footer for all pages following the cover page. This footer is distinct from the first-page footer to accommodate different content requirements.
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #
    clear_paragraph(paragraph)

    def add_field(p, field_code):
        """Manually constructs the OOXML for a complex Word field and appends it to a paragraph.

        This function provides a low-level mechanism to insert Word fields that are
        not natively supported by the `python-docx` high-level API. It operates
        by directly creating and arranging the required OOXML elements (`w:fldChar`,
        `w:instrText`) within a new run (`w:r`) added to the paragraph.

        The field is initialized with a placeholder result of '0'. The entire run
        containing the field is styled with a 9pt, grey Arial font.

        Args:
            p (docx.paragraph.Paragraph): The paragraph to which the field will be
                added. This object is modified in-place.
            field_code (str): The field code string that defines the field's behavior
                (e.g., 'PAGE', 'DOCPROPERTY "Author"').

        Returns:
            None
        """
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(127, 127, 127)

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)

        instrText = OxmlElement("w:instrText")
        instrText.set(ns.qn("xml:space"), "preserve")
        instrText.text = field_code
        run._r.append(instrText)

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "separate")
        run._r.append(fldChar2)

        #
        t = OxmlElement("w:t")
        t.text = "0"
        run._r.append(t)

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "end")
        run._r.append(fldChar3)

    r = paragraph.add_run("Página ")
    r.font.size = Pt(9)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(127, 127, 127)

    add_field(paragraph, "PAGE")

    r2 = paragraph.add_run(" de ")
    r2.font.size = Pt(9)
    r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor(127, 127, 127)

    add_field(paragraph, "NUMPAGES")


def load_payload(payload_path: Path) -> GlobalReportPayload:
    r"""{'docstring': 'Loads, normalizes, and validates a global report payload from a JSON file.\n\n    Reads a JSON document from the specified path and performs extensive\n    pre-validation data normalization to handle schema variations and legacy\n    formats. Transformations include type coercion (e.g., string to list for\n    `root_causes`), injection of default values (e.g., `score`, `band` in\n    `tower_bottom_lines`), field name mapping (e.g., `title` to `pillar` in\n    `target_vision`), data restructuring (e.g., wrapping `executive_decisions`\n    list into a dictionary), and coalescing multiple potential source fields into a\n    single destination field (e.g., for `business_case` in `execution_roadmap`).\n\n    After normalization, the function attempts to validate the data using\n    `GlobalReportPayload.model_validate`. If validation fails, a warning is\n    logged and a fallback mechanism is engaged. This fallback uses\n    `model_construct` to create the Pydantic object without strict validation,\n    recursively constructing nested data structures and populating missing fields\n    with sensible defaults. This approach guarantees the return of a structurally\n    complete object, preventing potential downstream processing errors.\n\n    Args:\n        payload_path: The file system path to the JSON payload.\n\n    Returns:\n        An instance of `GlobalReportPayload`. If validation succeeds, the model is\n        fully validated. If validation fails, the model is constructed without\n        validation, with missing or malformed fields populated by default values\n        to ensure structural integrity.\n\n    Raises:\n        FileNotFoundError: If the file at `payload_path` does not exist.\n        json.JSONDecodeError: If the file at `payload_path` is not a valid JSON\n            document.'}."""
    payload_dict = load_json(payload_path)

    # Sovereign Data Normalization: Pre-emptive schema adjustments are performed to ensure data consistency and validity prior to rendering.
    if "burning_platform" in payload_dict and isinstance(
        payload_dict["burning_platform"], list
    ):
        for item in payload_dict["burning_platform"]:
            if (
                isinstance(item, dict)
                and "root_causes" in item
                and isinstance(item["root_causes"], str)
            ):
                item["root_causes"] = [item["root_causes"]]

    if "tower_bottom_lines" in payload_dict and isinstance(
        payload_dict["tower_bottom_lines"], list
    ):
        for item in payload_dict["tower_bottom_lines"]:
            if isinstance(item, dict):
                if "score" not in item:
                    item["score"] = 4.0
                if "band" not in item:
                    item["band"] = "Managed"
                if "status_color" not in item:
                    item["status_color"] = "green"

    if "target_vision" in payload_dict and isinstance(
        payload_dict["target_vision"], dict
    ):
        target_v = payload_dict["target_vision"]
        if "evolution_principles" in target_v and isinstance(
            target_v["evolution_principles"], list
        ):
            for item in target_v["evolution_principles"]:
                if (
                    isinstance(item, dict)
                    and "title" in item
                    and "principle" not in item
                ):
                    item["principle"] = item["title"]
        if "strategic_pillars" in target_v and isinstance(
            target_v["strategic_pillars"], list
        ):
            for item in target_v["strategic_pillars"]:
                if isinstance(item, dict) and "title" in item and "pillar" not in item:
                    item["pillar"] = item["title"]

    if "execution_roadmap" in payload_dict and isinstance(
        payload_dict["execution_roadmap"], dict
    ):
        roadmap = payload_dict["execution_roadmap"]

        # Coalesces a null or missing program description to an empty string to ensure type consistency and prevent downstream processing errors.
        if "programs" in roadmap and isinstance(roadmap["programs"], list):
            normalized_programs = []
            for p in roadmap["programs"]:
                if isinstance(p, dict):
                    p_copy = p.copy()
                    if "description" not in p_copy or not p_copy["description"]:
                        p_copy["description"] = (
                            p_copy.get("name")
                            or p_copy.get("id")
                            or "Programa de transformación tecnológica."
                        )
                    normalized_programs.append(p_copy)
                else:
                    normalized_programs.append(p)
            roadmap["programs"] = normalized_programs

        # Validates the presence and structural integrity of the `horizons` field, asserting that it is a list before proceeding with iteration.
        if "horizons" not in roadmap or not roadmap["horizons"]:
            roadmap["horizons"] = {
                "quick_wins_0_3_months": [],
                "year_1_3_12_months": [],
                "year_2_12_24_months": [],
                "year_3_24_36_months": [],
            }

        if "horizons" in roadmap and isinstance(roadmap["horizons"], dict):
            h_dict = roadmap["horizons"]
            if (
                "quick_wins_0_6_months" in h_dict
                and "quick_wins_0_3_months" not in h_dict
            ):
                h_dict["quick_wins_0_3_months"] = h_dict["quick_wins_0_6_months"]
            if "year_1_6_12_months" in h_dict and "year_1_3_12_months" not in h_dict:
                h_dict["year_1_3_12_months"] = h_dict["year_1_6_12_months"]

            for h_key in [
                "quick_wins_0_3_months",
                "year_1_3_12_months",
                "year_2_12_24_months",
                "year_3_24_36_months",
            ]:
                if h_key not in h_dict or h_dict[h_key] is None:
                    h_dict[h_key] = []

                if isinstance(h_dict[h_key], list):
                    normalized_list = []
                    for item in h_dict[h_key]:
                        if isinstance(item, dict):
                            item_copy = item.copy()

                            # Maps the `program_id` and `program_name` source fields into the `program` composite destination field.
                            if "program" not in item_copy or not item_copy["program"]:
                                p_val = (
                                    item_copy.get("program_id")
                                    or item_copy.get("program_name")
                                    or item_copy.get("program_title")
                                    or item_copy.get("program")
                                    or ""
                                )
                                item_copy["program"] = p_val

                            # Maps the `objective`, `description`, `business_impact`, and `impact` source fields into the `business_case` composite destination field.
                            if (
                                "business_case" not in item_copy
                                or not item_copy["business_case"]
                            ):
                                bc_val = (
                                    item_copy.get("objective")
                                    or item_copy.get("description")
                                    or item_copy.get("business_impact")
                                    or item_copy.get("impact")
                                    or "Mitigación de deuda técnica."
                                )
                                item_copy["business_case"] = bc_val

                            # Coalesces `start_month` to a default value if the source field is null, missing, or zero.
                            if (
                                "start_month" not in item_copy
                                or item_copy["start_month"] == 0
                            ):
                                if h_key == "quick_wins_0_3_months":
                                    item_copy["start_month"] = 1
                                elif h_key == "year_1_3_12_months":
                                    item_copy["start_month"] = 3
                                elif h_key == "year_2_12_24_months":
                                    item_copy["start_month"] = 12
                                elif h_key == "year_3_24_36_months":
                                    item_copy["start_month"] = 24
                                else:
                                    item_copy["start_month"] = 1

                            # Coalesces `duration_months` to a default value if the source field is null, missing, or zero.
                            if (
                                "duration_months" not in item_copy
                                or item_copy["duration_months"] == 0
                            ):
                                if h_key == "quick_wins_0_3_months":
                                    item_copy["duration_months"] = 3
                                elif h_key == "year_1_3_12_months":
                                    item_copy["duration_months"] = 9
                                elif h_key == "year_2_12_24_months":
                                    item_copy["duration_months"] = 12
                                elif h_key == "year_3_24_36_months":
                                    item_copy["duration_months"] = 12
                                else:
                                    item_copy["duration_months"] = 6

                            normalized_list.append(item_copy)
                        else:
                            normalized_list.append(item)
                    h_dict[h_key] = normalized_list

    if "executive_decisions" in payload_dict and isinstance(
        payload_dict["executive_decisions"], list
    ):
        payload_dict["executive_decisions"] = {
            "immediate_decisions": payload_dict["executive_decisions"]
        }
    elif "executive_decisions" in payload_dict and isinstance(
        payload_dict["executive_decisions"], dict
    ):
        if "immediate_decisions" not in payload_dict["executive_decisions"]:
            payload_dict["executive_decisions"]["immediate_decisions"] = []

    try:
        return GlobalReportPayload.model_validate(payload_dict)
    except ValidationError as e:
        print(
            f"⚠️ Warning: Error de validación de Pydantic en {payload_path} (procediendo con fallback model_construct robustecido): {e}"
        )
        from domain.schemas.global_report import (
            BurningPlatformItem,
            EvolutionPrinciple,
            ExecutionRoadmapDraft,
            ExecutiveDecisionItem,
            ExecutiveDecisionsDraft,
            ExecutiveSummaryDraft,
            GlobalReportDocumentMeta,
            HorizonsDef,
            InitiativeDef,
            ProgramDef,
            StrategicPillar,
            TargetVisionDraft,
            TowerBottomLineItem,
        )

        def safe_construct(model_cls, data, default_fields):
            """Constructs a model instance from input data using defaults and fallback logic.

            This function creates a model instance by merging default fields with provided
            data. It initializes the model's fields from `default_fields`, then
            overrides these with any non-`None` values from the `data` dictionary.
            A specific fallback logic is applied: if 'principle', 'pillar', or 'program'
            are specified in `default_fields` but absent in the input `data`, their values
            are populated from the 'title' field of the merged data.

            The function utilizes `model_cls.model_construct` for instantiation, which
            bypasses data validation for performance. This assumes the input data is
            pre-validated and trusted.

            Args:
                model_cls: The model class to instantiate, which must expose a
                    `model_construct` class method (e.g., a Pydantic BaseModel).
                data: The input data payload. If this value is not a dictionary, it is
                    returned immediately without modification.
                default_fields: A mapping of field names to their default values.

            Returns:
                An instance of `model_cls` populated with the merged data if `data` is a
                dictionary; otherwise, the original `data` object.
            """
            if not isinstance(data, dict):
                return data
            data_copy = default_fields.copy()
            for k, v in data.items():
                if v is not None:
                    data_copy[k] = v
            if "title" in data_copy:
                if "principle" in default_fields and "principle" not in data:
                    data_copy["principle"] = data_copy["title"]
                if "pillar" in default_fields and "pillar" not in data:
                    data_copy["pillar"] = data_copy["title"]
                if "program" in default_fields and "program" not in data:
                    data_copy["program"] = data_copy["title"]
            return model_cls.model_construct(**data_copy)

        payload = GlobalReportPayload.model_construct(**payload_dict)

        payload.meta = safe_construct(
            GlobalReportDocumentMeta,
            payload.meta if isinstance(payload.meta, dict) else {},
            {"client": "", "date": "", "version": ""},
        )

        payload.executive_summary = safe_construct(
            ExecutiveSummaryDraft,
            payload.executive_summary
            if isinstance(payload.executive_summary, dict)
            else {},
            {"headline": "", "narrative": "", "key_business_impacts": []},
        )

        if isinstance(payload.burning_platform, list):
            payload.burning_platform = [
                safe_construct(
                    BurningPlatformItem,
                    item,
                    {"theme": "", "business_risk": "", "root_causes": []},
                )
                if isinstance(item, dict)
                else item
                for item in payload.burning_platform
            ]

        if isinstance(payload.tower_bottom_lines, list):
            payload.tower_bottom_lines = [
                safe_construct(
                    TowerBottomLineItem,
                    item,
                    {
                        "id": "",
                        "name": "",
                        "score": "",
                        "band": "",
                        "status_color": "",
                        "bottom_line": "",
                    },
                )
                if isinstance(item, dict)
                else item
                for item in payload.tower_bottom_lines
            ]

        if isinstance(payload.target_vision, dict):
            tv_dict = payload.target_vision.copy()
            if "evolution_principles" in tv_dict and isinstance(
                tv_dict["evolution_principles"], list
            ):
                tv_dict["evolution_principles"] = [
                    safe_construct(
                        EvolutionPrinciple, ep, {"principle": "", "description": ""}
                    )
                    if isinstance(ep, dict)
                    else ep
                    for ep in tv_dict["evolution_principles"]
                ]
            if "strategic_pillars" in tv_dict and isinstance(
                tv_dict["strategic_pillars"], list
            ):
                tv_dict["strategic_pillars"] = [
                    safe_construct(
                        StrategicPillar, sp, {"pillar": "", "description": ""}
                    )
                    if isinstance(sp, dict)
                    else sp
                    for sp in tv_dict["strategic_pillars"]
                ]
            payload.target_vision = safe_construct(
                TargetVisionDraft,
                tv_dict,
                {
                    "value_proposition": "",
                    "evolution_principles": [],
                    "strategic_pillars": [],
                },
            )

        if isinstance(payload.execution_roadmap, dict):
            er_dict = payload.execution_roadmap.copy()
            if "programs" in er_dict and isinstance(er_dict["programs"], list):
                er_dict["programs"] = [
                    safe_construct(ProgramDef, p, {"name": "", "description": ""})
                    if isinstance(p, dict)
                    else p
                    for p in er_dict["programs"]
                ]
            if "horizons" in er_dict and isinstance(er_dict["horizons"], dict):
                h_dict = er_dict["horizons"].copy()
                if (
                    "quick_wins_0_6_months" in h_dict
                    and "quick_wins_0_3_months" not in h_dict
                ):
                    h_dict["quick_wins_0_3_months"] = h_dict["quick_wins_0_6_months"]
                if (
                    "year_1_6_12_months" in h_dict
                    and "year_1_3_12_months" not in h_dict
                ):
                    h_dict["year_1_3_12_months"] = h_dict["year_1_6_12_months"]

                for h_key in [
                    "quick_wins_0_3_months",
                    "year_1_3_12_months",
                    "year_2_12_24_months",
                    "year_3_24_36_months",
                ]:
                    if h_key in h_dict and isinstance(h_dict[h_key], list):
                        h_dict[h_key] = [
                            safe_construct(
                                InitiativeDef,
                                init,
                                {
                                    "program": "",
                                    "title": "",
                                    "business_case": "",
                                    "start_month": 0,
                                    "duration_months": 0,
                                },
                            )
                            if isinstance(init, dict)
                            else init
                            for init in h_dict[h_key]
                        ]
                er_dict["horizons"] = safe_construct(
                    HorizonsDef,
                    h_dict,
                    {
                        "quick_wins_0_3_months": [],
                        "year_1_3_12_months": [],
                        "year_2_12_24_months": [],
                        "year_3_24_36_months": [],
                    },
                )
            payload.execution_roadmap = safe_construct(
                ExecutionRoadmapDraft, er_dict, {"programs": [], "horizons": None}
            )

        if isinstance(payload.executive_decisions, dict):
            ed_dict = payload.executive_decisions.copy()
            if "immediate_decisions" in ed_dict and isinstance(
                ed_dict["immediate_decisions"], list
            ):
                ed_dict["immediate_decisions"] = [
                    safe_construct(
                        ExecutiveDecisionItem,
                        item,
                        {
                            "decision_type": "",
                            "action_required": "",
                            "impact_if_delayed": "",
                        },
                    )
                    if isinstance(item, dict)
                    else item
                    for item in ed_dict["immediate_decisions"]
                ]
            payload.executive_decisions = safe_construct(
                ExecutiveDecisionsDraft, ed_dict, {"immediate_decisions": []}
            )

        return payload


def render_global_report(
    payload: GlobalReportPayload,
    template_path: Path,
    output_path: Path,
    client_dir: Path,
) -> Path:
    """Renders a global report by populating a Microsoft Word document template.

    This function orchestrates the report generation process. It begins by loading
    a specified `.docx` template, clearing its existing body content, and adding
    page numbering. It then conditionally renders various report sections—such as
    the cover, executive summary, and execution roadmap—based on the content of the
    provided payload. Finally, it ensures the output directory exists and saves the
    completed document to the specified path.

    Args:
        payload: An object containing all structured data for the report.
        template_path: The file system path to the `.docx` template document.
        output_path: The destination file system path for the generated report. The
            parent directory will be created if it does not exist.
        client_dir: The file system path to a directory containing client-specific
            assets like images.

    Returns:
        The file system path to the newly generated report document.

    Raises:
        FileNotFoundError: If the template file specified by `template_path`
            does not exist.
        PermissionError: If the output directory cannot be created or the final
            report file cannot be written due to insufficient permissions.
    """
    doc = Document(str(template_path))
    clear_document_body(doc)

    if doc.sections:
        create_page_number_footer(doc.sections[0])

    render_cover(doc, payload)
    client_name = payload.meta.client
    if payload.executive_summary:
        render_executive_summary(
            doc,
            payload.executive_summary,
            payload.heatmap,
            payload.visuals,
            client_dir,
            client_name=client_name,
        )
    if payload.burning_platform:
        render_burning_platform(doc, payload.burning_platform, client_name=client_name)
    if payload.tower_bottom_lines and payload.heatmap:
        render_tower_bottom_lines(
            doc, payload.heatmap, payload.tower_bottom_lines, client_name=client_name
        )
    if payload.target_vision:
        render_target_vision(doc, payload.target_vision, client_name=client_name)
    if payload.execution_roadmap:
        render_execution_roadmap(
            doc,
            payload.execution_roadmap,
            payload.visuals,
            client_dir,
            client_name=client_name,
        )
    if payload.executive_decisions:
        render_executive_decisions(
            doc, payload.executive_decisions, client_name=client_name
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Informe Global CIO-READY renderizado: {output_path}")
    return output_path


def main(argv: list[str] | None = None) -> None:
    """Orchestrates rendering a global report from command-line arguments.

    This function serves as the main entry point for a command-line script,
    parsing arguments for a JSON payload path, a Jinja2 template path, and an
    output file path. The program exits with a status code of 1 if an
    incorrect number of arguments is provided.

    Args:
        argv: A list of command-line arguments. If None, `sys.argv` is used.
            The list is expected to contain the script name followed by the paths
            to the payload, template, and output files.

    Raises:
        FileNotFoundError: If the specified payload or template file does not exist.
        IOError: If an error occurs while writing the output report file.
        ValueError: If the payload file contains malformed data.
    """
    if len(argv if argv is not None else sys.argv) != 4:
        sys.exit(1)
    payload_path = Path((argv if argv is not None else sys.argv)[1])
    template_path = Path((argv if argv is not None else sys.argv)[2])
    output_path = Path((argv if argv is not None else sys.argv)[3])
    client_dir = payload_path.parent
    payload = load_payload(payload_path)
    render_global_report(payload, template_path, output_path, client_dir)


if __name__ == "__main__":
    main()
