import csv
import html
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from infrastructure.docx_render_utils import (
    autofit_table_to_contents,
    finalize_table,
    set_cell_text,
)
from infrastructure.text_utils import clean_text_for_word

# Define constants for branding and document styling.
COLOR_BLUE = "0072BC"
COLOR_HEADER_BG = "D9EAF7"
COLOR_ROW_ALT = "F2F2F2"


def format_currency_custom(value: float, thousands_sep: str, decimal_sep: str) -> str:
    """Formats a float into a string with custom thousands and decimal separators.

    The input value is first formatted to a string with exactly two decimal
    places. The integer part is then grouped by thousands using the provided
    separator, and the standard decimal point is replaced with the custom
    decimal separator.

    Args:
        value: The numeric value to format.
        thousands_sep: The character to use as the thousands separator.
        decimal_sep: The character to use as the decimal separator.

    Returns:
        The formatted numeric string.
    """
    parts = f"{value:.2f}".split(".")
    integer_part = parts[0]
    decimal_part = parts[1]
    reversed_integer = integer_part[::-1]
    grouped = [reversed_integer[i : i + 3] for i in range(0, len(reversed_integer), 3)]
    formatted_integer = thousands_sep.join(grouped)[::-1]
    return f"{formatted_integer}{decimal_sep}{decimal_part}"


# Use an internationalization (i18n) dictionary to support multiple locales.
LANG_VOCAB = {
    "es": {
        "title": "Informe de Situación Actual (AS-IS)",
        "project": "Consultoría para el rediseño de infraestructura tecnológica",
        "client_lbl": "Cliente:\n",
        "proj_lbl": "\nProyecto:\n",
        "toc_title": "Contenido",
        "dashboard_title": "Cuadro de Mando de Madurez Global de la Torre",
        "dashboard_intro": "La evaluación integral del estado actual consolida una valoración cuantitativa que sitúa a la torre en su banda de madurez correspondiente, justificando el grado de adopción de prácticas de resiliencia:",
        "score_lbl": "PUNTUACIÓN AS-IS",
        "maturity_lbl": "NIVEL DE MADUREZ CUALITATIVO",
        "resilience_reading": "Lectura de Resiliencia de la Plataforma:",
        "radar_title": "Perfil Visual de Madurez Ponderada",
        "justification_table_headers": [
            "Pilar / Capacidad Evaluada",
            "Score AS-IS",
            "Análisis de Brecha y Justificación de Nota",
        ],
        "platform_overview_title": "Descripción de la Plataforma de Infraestructura Actual",
        "swot_title": "Fortalezas y Brechas Clave",
        "swot_headers": [
            "Fortalezas Clave de la Plataforma",
            "Brechas y Deudas Operativas Críticas",
        ],
        "transversal_title": "Análisis Transversal de Capacidades",
        "risk_matrix_title": "Matriz de Riesgo Cuantitativa (FAIR)",
        "risk_intro_1": "El análisis cuantitativo del estado actual establece que las deudas técnicas operativas y las vulnerabilidades detectadas exponen ",
        "risk_intro_2": " a una Expectativa de Pérdida Anualizada (ALE) real de ",
        "risk_intro_3": " anuales bajo el estándar internacional FAIR. Esta matriz de mosaico consolida la distribución de esta exposición a partir de la frecuencia de amenazas (TEF) y magnitud de pérdidas (LM) de cada hallazgo forense:",
        "exposure_summary_title": "Resumen de Exposición y Mapa de Calor",
        "tef_title": "Criterios de Frecuencia de la Amenaza (TEF)",
        "tef_headers": ["Nivel TEF", "Denominación", "Frecuencia Anualizada Estimada"],
        "lm_title": "Criterios de Magnitud de Pérdida Directa/Indirecta (LM)",
        "lm_headers": [
            "Nivel LM",
            "Clasificación de Impacto",
            "Límite de Pérdida Financiera",
        ],
        "top_risks_title": "Top de Riesgos Prioritarios de la Torre",
        "top_risks_intro": "Top de riesgos priorizados por su impacto directo en la resiliencia del servicio y el coste de inacción (ALE):",
        "motivo_priorizacion": "Motivo de priorización: ",
        "top_risks_fallbacks": {
            "compute": "El modelo reactivo en el cómputo on-premise arriesga directamente la continuidad del servicio y bloquea el despliegue del Gemelo Digital y la IA estratégica.",
            "container": "La falta de plataforma corporativa de contenedores estanca la innovación, genera silos aislados en desarrollo y multiplica la deuda de modernización.",
            "cloud": "La remediación manual de las Landing Zones impide verificar el cumplimiento normativo en tiempo real, exponiendo al negocio a severas sanciones regulatorias.",
            "automation": "La dependencia de tickets y validaciones manuales retarda semanas el time-to-market y genera cuellos de botella severos en la provisión.",
            "default": "La fragmentación e inconsistencia en las monitorizaciones imposibilita correlar alertas y predecir fallas complejas que detienen las operaciones.",
        },
        "detailed_risks_title": "Registro Detallado de Hallazgos Forenses por Pilar",
        "detailed_risks_pilar_title": "Análisis de Vulnerabilidades: ",
        "detailed_risks_headers": [
            "ID",
            "Vulnerabilidad y Evidencias de Auditoría",
            "Exposición y Riesgo de Negocio",
        ],
        "next_steps_title": "Siguientes Pasos y Coste de Inacción",
        "appendix_a_title": "Apéndice A: Lista de Abreviaturas",
        "appendix_a_intro": "A continuación, se define el glosario de términos técnicos y acrónimos utilizados en este diagnóstico técnico actual:",
        "appendix_a_headers": ["Abreviatura", "Significado / Descripción Técnica"],
        "appendix_b_title": "Apéndice B: Cláusula de Limitación de Responsabilidad",
        "appendix_c_title": "Apéndice C: Registro de Custodia de Fuentes de Información",
        "appendix_c_intro": "Para garantizar la transparencia, veracidad e inmutabilidad de la información recopilada en este diagnóstico técnico, a continuación se listan y codifican las fuentes documentales utilizadas bajo custodia de auditoría:",
        "appendix_c_headers": [
            "Código de Referencia",
            "Documento Fuente / Origen de Datos",
            "Descripción y Ámbito de Custodia",
        ],
        "vulnerability_lbl": "Vulnerabilidad Identificada:\n",
        "evidence_lbl": "Evidencia Forense Literal:\n",
        "impact_lbl": "Impacto Operativo:\n",
        "exposure_lbl": "Exposición Cuantitativa (FAIR):\n",
        "ale_proyectado": "ALE Proyectado: ",
        "v_lbl": "Versión ",
        "disclaimer_text_1": "Este informe de diagnóstico técnico (el Documento) representa una evaluación retrospectiva del estado actual (AS-IS) de la torre evaluada basándose exclusivamente en la información, telemetrías y respuestas de autoevaluación proporcionadas por el cliente de buena fe. Las recomendaciones, valoraciones cualitativas e inventarios de infraestructura presentados reflejan un diagnóstico de la situación actual y no constituyen garantías de rendimiento, seguridad continua, ni compromisos de disponibilidad de servicios por parte de la consultora.",
        "disclaimer_text_2": "Las proyecciones cuantitativas de riesgo financiero y la Expectativa de Pérdida Anualizada (ALE) calculadas a través de la metodología O-FAIR y simulaciones de Monte Carlo son estimaciones estadísticas de probabilidad matemática basadas en modelos de incertidumbre calibrados de forma estándar para el sector de infraestructuras críticas. Dichas cifras se proporcionan únicamente con carácter prioritario e ilustrativo de coste de inacción, y no deben ser interpretadas en ningún caso como auditorías contables formales de pérdidas, garantías de impacto directo en el balance financiero, ni compromisos de pasivos exigibles legalmente.",
        "disclaimer_text_3": "Las referencias a normativas legales y de cumplimiento (tales como la directiva europea NIS2 o el Esquema Nacional de Seguridad ENS - Categoría Alta) se enmarcan como guías y recomendaciones de buena práctica de ingeniería tecnológica para mitigar la superficie de exposición, y no constituyen de ninguna manera asesoramiento jurídico formal. La interpretación e implantación final de los requisitos legales de cumplimiento recaen de forma estricta e intransferible en los responsables de cumplimiento, seguridad y gobierno de la organización cliente.",
        "bib_cues": "Cuestionarios consolidados de autoevaluación técnica de la torre de ",
        "bib_cues_desc": ", recopilando las respuestas oficiales del equipo de ingeniería de ",
        "bib_contexto": "Dossier estratégico de contexto de arquitectura de ",
        "bib_contexto_desc": " que define la volumetría del patrimonio, centros de datos y criticidad operativa.",
        "bib_minutas": "Minutas y registros tomados durante las sesiones presenciales de contexto de arquitectura mantenidas entre los líderes técnicos de ",
        "bib_minutas_desc": " y el equipo consultor.",
        "default_overview_text": "El inventario consolidado y la evaluación de la plataforma de ",
    },
    "en": {
        "title": "Current State Assessment Report (AS-IS)",
        "project": "Technology infrastructure redesign consultancy",
        "client_lbl": "Client:\n",
        "proj_lbl": "\nProject:\n",
        "toc_title": "Contents",
        "dashboard_title": "Global Tower Maturity Dashboard",
        "dashboard_intro": "The comprehensive current state assessment consolidates a quantitative valuation that positions the tower in its corresponding maturity band, justifying the adoption of resilience practices:",
        "score_lbl": "AS-IS SCORE",
        "maturity_lbl": "QUALITATIVE MATURITY LEVEL",
        "resilience_reading": "Platform Resilience Reading:",
        "radar_title": "Weighted Maturity Visual Profile",
        "justification_table_headers": [
            "Pillar / Evaluated Capability",
            "AS-IS Score",
            "Gap Analysis and Score Justification",
        ],
        "platform_overview_title": "Current Infrastructure Platform Description",
        "swot_title": "Key Strengths and Gaps",
        "swot_headers": [
            "Key Platform Strengths",
            "Critical Gaps and Operational Debts",
        ],
        "transversal_title": "Cross-Cutting Capabilities Analysis",
        "risk_matrix_title": "Quantitative Risk Matrix (FAIR)",
        "risk_intro_1": "The quantitative current state analysis establishes that the technical operational debts and identified vulnerabilities expose ",
        "risk_intro_2": " to a real Annual Loss Exposure (ALE) of ",
        "risk_intro_3": " annually under the international FAIR standard. This mosaic matrix consolidates the distribution of this exposure based on the Threat Event Frequency (TEF) and Loss Magnitude (LM) of each forensic finding:",
        "exposure_summary_title": "Exposure Summary and Heatmap",
        "tef_title": "Threat Event Frequency Criteria (TEF)",
        "tef_headers": ["TEF Level", "Denomination", "Estimated Annualized Frequency"],
        "lm_title": "Direct/Indirect Loss Magnitude Criteria (LM)",
        "lm_headers": ["LM Level", "Impact Classification", "Financial Loss Limit"],
        "top_risks_title": "Top Prioritized Risks for the Tower",
        "top_risks_intro": "Top risks prioritized by their direct impact on service resilience and the cost of inaction (ALE):",
        "motivo_priorizacion": "Prioritization rationale: ",
        "top_risks_fallbacks": {
            "compute": "The reactive model in on-premise compute directly risks service continuity and blocks the deployment of the Digital Twin and strategic AI.",
            "container": "The lack of an enterprise container platform stalls innovation, creates isolated development silos, and multiplies modernization debt.",
            "cloud": "The manual remediation of Landing Zones prevents verifying compliance in real-time, exposing the business to severe regulatory penalties.",
            "automation": "Dependence on manual tickets and approvals delays time-to-market by weeks and generates severe provisioning bottlenecks.",
            "default": "Fragmentation and inconsistency in monitoring systems make it impossible to correlate alerts and predict complex outages.",
        },
        "detailed_risks_title": "Detailed Registry of Forensic Findings by Pillar",
        "detailed_risks_pilar_title": "Vulnerability Analysis: ",
        "detailed_risks_headers": [
            "ID",
            "Vulnerability and Audit Evidence",
            "FAIR Exposure and Business Risk",
        ],
        "next_steps_title": "Next Steps and Cost of Inaction",
        "appendix_a_title": "Appendix A: List of Abbreviations",
        "appendix_a_intro": "The following defines the glossary of technical terms and acronyms used in this current technical diagnosis:",
        "appendix_a_headers": ["Abbreviation", "Significado / Technical Description"],
        "appendix_b_title": "Appendix B: Limitation of Liability Disclaimer",
        "appendix_c_title": "Appendix C: Registry of Custody of Information Sources",
        "appendix_c_intro": "To ensure the transparency, veracity, and immutability of the information collected in this technical diagnosis, the document sources used under audit custody are listed and coded below:",
        "appendix_c_headers": [
            "Reference Code",
            "Source Document / Data Origin",
            "Custody Scope and Description",
        ],
        "vulnerability_lbl": "Identified Vulnerability:\n",
        "evidence_lbl": "Literal Forensic Evidence:\n",
        "impact_lbl": "Operational Impact:\n",
        "exposure_lbl": "Quantitative Exposure (FAIR):\n",
        "ale_proyectado": "Projected ALE: ",
        "v_lbl": "Version ",
        "disclaimer_text_1": "This technical diagnosis report (the Document) represents a retrospective evaluation of the current state (AS-IS) of the evaluated tower based exclusively on the information, telemetries, and self-assessment responses provided by the client in good faith. The recommendations, qualitative assessments, and infrastructure inventories presented reflect a current diagnosis and do not constitute guarantees of performance, continuous security, or service availability commitments by the consultant.",
        "disclaimer_text_2": "The quantitative financial risk projections and the Annual Loss Exposure (ALE) calculated through the O-FAIR methodology and Monte Carlo simulations are statistical probability estimates based on standard uncertainty models for the critical infrastructure sector. These figures are provided solely for prioritization and illustrative purposes of cost of inaction, and should not be interpreted in any case as formal balance sheet accounting audits, asset guarantees, or legally enforceable liabilities.",
        "disclaimer_text_3": "References to legal and compliance regulations (such as the European directive NIS2 or the national security scheme ENS - High Category) are framed as guidelines and best engineering practices to mitigate the attack surface, and do not constitute formal legal advice. The final interpretation and implementation of compliance requirements remain strictly and non-transferably with the compliance, security, and governance officers of the client organization.",
        "bib_cues": "Consolidated technical self-assessment questionnaires of the ",
        "bib_cues_desc": " tower, gathering official responses from the engineering team of ",
        "bib_contexto": "Strategic architecture context dossier of ",
        "bib_contexto_desc": " defining inventory volumes, data centers, and operational criticality.",
        "bib_minutas": "Minutes and records taken during the on-site architecture sessions between the technical leaders of ",
        "bib_minutas_desc": " and the consulting team.",
        "default_overview_text": "The consolidated inventory and evaluation of the ",
    },
}


def shade_cell(cell, color_hex):
    """Applies a background color shading to a table cell.

    This function directly manipulates the underlying Open Office XML (OOXML)
    of the cell's properties to set its background fill color.

    Args:
        cell (docx.table._Cell): The table cell object to modify.
        color_hex (str): A six-character RRGGBB hexadecimal string representing
            the background color (e.g., 'FFC000').
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_border_white(cell):
    r"""{'docstring': 'Set a 2.0 pt white border on a table cell.\n\nDirectly manipulates the underlying Office Open XML (OOXML) representation\nof a table cell to apply a single, white, 2.0 pt border to all four sides.\nThis effect is typically used to create visual separation between cells in\na heatmap-style table, giving them a tiled appearance.\n\nArgs:\n    cell (docx.table.Cell): The target table cell object whose borders will\n        be modified in-place.\n\nRaises:\n    AttributeError: If the provided `cell` object does not conform to the\n        expected internal OOXML structure.'}."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")

    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(
            qn("w:sz"), "16"
        )  # The OpenXML value of 16 corresponds to a 2.0 pt border thickness (16 / 8 = 2.0 pt).
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "FFFFFF")  #
        tcBorders.append(border)

    tcPr.append(tcBorders)


def set_cell_text_custom(
    cell, text, bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=None
):
    r"""Populates a `docx` table cell with richly formatted, multi-line text.

    This function processes an input string by clearing the target cell and then
    inserting the string as a series of formatted paragraphs. It applies several
    formatting rules:

    - Sanitization: Removes markdown emphasis characters (`*`, `_`, `**`, `__`).
    - Paragraphs: Each line in the input string (separated by `\n` or `\n`)
      becomes a separate paragraph in the cell.
    - Bullet Points: Lines are prefixed with bullet points (`•` for top-level,
      `-` for indented sub-points).
    - Hierarchy: Indentation is applied to lines that start with leading
      whitespace or a hyphen.
    - Label Highlighting: Lines containing a colon (e.g., "Gaps: ...") are split,
      with the text before the colon rendered in bold.
    - Citation Formatting: Source citations matching the pattern `[Ref: ...]` are
      extracted and appended with a distinct style (smaller, italic, gray font).

    For single-line input without newlines, processing is delegated to the
    `set_cell_text` helper function.

    Args:
        cell (docx.table._Cell): The `python-docx` table cell object to modify.
            Any existing content in the cell will be cleared.
        text (str): The content to insert into the cell. Newline characters are
            interpreted as paragraph breaks.
        bold (bool): If True, applies bold formatting to entire lines that do not
            contain a formatting label (i.e., a colon). Defaults to False.
        font_size (int): The font size in points (Pt). Defaults to 9.
        align (WD_ALIGN_PARAGRAPH): An alignment constant from the
            `docx.enum.text.WD_ALIGN_PARAGRAPH` enumeration. Defaults to
            `WD_ALIGN_PARAGRAPH.LEFT`.
        color_rgb (docx.shared.RGBColor or None): An `RGBColor` object for the
            primary text color. If None, the style's default color is used.
            Defaults to None.

    Returns:
        None. The `cell` object is modified in-place.
    """
    # Sanitization: Remove markdown emphasis characters from body text.
    clean_txt = (
        str(text).replace("**", "").replace("__", "").replace("*", "").replace("_", "")
    )

    # Handle multi-line input by splitting the text into separate paragraphs at each newline character.
    if "\n" in clean_txt or "\\n" in clean_txt:
        # Remove the default empty paragraph that `python-docx` adds on cell creation to avoid unwanted vertical space.
        for p in list(cell.paragraphs):
            cell._tc.remove(p._element)

        raw_lines = clean_txt.replace("\\n", "\n").split("\n")
        for line in raw_lines:
            # Preserve leading whitespace for subsequent hierarchy detection; it will be stripped later.
            is_subpoint = line.startswith("   ") or line.strip().startswith("-")

            clean_line = line.strip()
            if not clean_line:
                continue

            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.alignment = align

            # Clear any existing list formatting from the paragraph to ensure a consistent state before applying a new style.
            if clean_line.startswith("•") or clean_line.startswith("-"):
                clean_line = clean_line[1:].strip()

            if is_subpoint:
                # Define the style for an indented, secondary bullet point, per requirement 2.
                p.paragraph_format.left_indent = Inches(0.25)
                bullet_prefix = "- "
            else:
                #
                p.paragraph_format.left_indent = Inches(0.0)
                bullet_prefix = "• "

            # Implement detection and extraction of user-provided citations (e.g., '[Ref: ...]').
            ref_text = ""
            if "[Ref:" in clean_line:
                clean_line, ref_part = clean_line.split("[Ref:", 1)
                ref_text = "[Ref:" + ref_part

            # Detect and apply bold formatting to leading labels (e.g., 'Gaps:') to enhance scannability.
            if ":" in clean_line:
                label, desc = clean_line.split(":", 1)
                run_lbl = p.add_run(f"{bullet_prefix}{label.strip()}: ")
                run_lbl.bold = True
                run_lbl.font.name = "Arial"
                run_lbl.font.size = Pt(font_size)
                if color_rgb:
                    run_lbl.font.color.rgb = color_rgb

                run_desc = p.add_run(desc.strip())
                run_desc.font.name = "Arial"
                run_desc.font.size = Pt(font_size)
                if color_rgb:
                    run_desc.font.color.rgb = color_rgb
            else:
                run = p.add_run(f"{bullet_prefix}{clean_line.strip()}")
                run.bold = bold
                run.font.name = "Arial"
                run.font.size = Pt(font_size)
                if color_rgb:
                    run.font.color.rgb = color_rgb

            # Append the source reference to the paragraph with distinct formatting (italic, gray, smaller font size).
            if ref_text:
                run_ref = p.add_run(f" {ref_text.strip()}")
                run_ref.font.name = "Arial"
                run_ref.font.size = Pt(font_size - 1)  #
                run_ref.font.italic = True
                run_ref.font.color.rgb = RGBColor(120, 130, 140)  #
    else:
        #
        set_cell_text(cell, clean_txt, bold=bold, align=align, font_size=font_size)
        p = cell.paragraphs[0]
        for r in p.runs:
            if color_rgb:
                r.font.color.rgb = color_rgb


def add_body_paragraph(
    doc,
    text,
    bold=False,
    italic=False,
    space_after=6,
    text_color_rgb=(46, 64, 77),
    style="Normal",
) -> Any:
    r"""{'docstring': "Adds a sanitized and formatted paragraph to a `docx.Document` object.\n\n    The input text is first sanitized by removing markdown emphasis markers\n    (`*`, `_`, `**`, `__`). The function then adds the cleaned text as a new,\n    justified paragraph formatted with a consistent 'Arial' 10pt font.\n\n    For bulleted styles (e.g., 'List Bullet', 'Bullet'), any text preceding\n    the first colon is automatically rendered in bold to enhance readability;\n    this behavior overrides the `bold` parameter.\n\n    If a specified style name does not exist in the document's template, the\n    function falls back to the 'Normal' style. For a non-existent bulleted\n    style, it also manually prepends a '•' character to simulate a list item.\n\n    Args:\n        doc (docx.document.Document): The document instance to which the\n            paragraph will be added.\n        text (str): The raw string content for the paragraph. Markdown emphasis\n            characters will be removed.\n        bold (bool, optional): If True, applies bold formatting to the entire\n            paragraph. This parameter is ignored for bulleted styles containing\n            a colon. Defaults to False.\n        italic (bool, optional): If True, applies italic formatting to the\n            entire paragraph. Defaults to False.\n        space_after (int, optional): The spacing in points (Pt) to apply after\n            the paragraph. Defaults to 6.\n        text_color_rgb (tuple[int, int, int], optional): An RGB tuple specifying\n            the font color. Each value must be in the range 0-255. Defaults\n            to (46, 64, 77).\n        style (str, optional): The name of the paragraph style to apply from\n            the document template. Defaults to 'Normal'.\n\n    Returns:\n        docx.text.paragraph.Paragraph: The newly created paragraph object that was\n            added to the document.\n\n    Raises:\n        ValueError: If any value in the `text_color_rgb` tuple is outside the\n            valid 0-255 range."}."""
    # Sanitization: Remove markdown emphasis characters from body text.
    clean_txt = (
        text.replace("**", "").replace("__", "").replace("*", "").replace("_", "")
    )

    is_bullet = style == "List Bullet" or style == "Bullet"
    target_style = "Bullet" if is_bullet else style

    try:
        p = doc.add_paragraph(style=target_style)
    except KeyError:
        p = doc.add_paragraph(style="Normal")
        if is_bullet:
            run_b = p.add_run("• ")
            run_b.font.name = "Arial"
            run_b.font.size = Pt(10)
            run_b.font.color.rgb = RGBColor(*text_color_rgb)

    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Programmatically bold the leading words of each bullet point to improve scannability.
    if is_bullet and ":" in clean_txt:
        parts = clean_txt.split(":", 1)
        run_bold = p.add_run(parts[0] + ":")
        run_bold.bold = True
        run_bold.font.name = "Arial"
        run_bold.font.size = Pt(10)
        run_bold.font.color.rgb = RGBColor(*text_color_rgb)

        run_rest = p.add_run(parts[1])
        run_rest.font.name = "Arial"
        run_rest.font.size = Pt(10)
        run_rest.font.color.rgb = RGBColor(*text_color_rgb)
    else:
        run = p.add_run(clean_txt)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*text_color_rgb)
    return p


def add_heading(doc, text, level, primary_color_rgb=(0, 114, 188)):
    r"""{'docstring': "Adds a sanitized and styled heading to a document.\n\n    Sanitizes the input text by removing markdown emphasis characters ('*', '_')\n    before adding it as a heading. The function applies specific paragraph\n    and font styles based on the heading level. Paragraph formatting includes\n    spacing adjustments and ensuring headings are not separated from subsequent\n    paragraphs. Level 1 headings are configured to start on a new page via the\n    'Page break before' style property, preventing the creation of extraneous\n    blank pages.\n\n    Args:\n        doc (docx.document.Document): The document object to which the heading\n            will be added.\n        text (str): The raw heading text from which markdown emphasis will be\n            removed.\n        level (int): The heading level, where 0 represents a title and 1-9\n            represent heading levels.\n        primary_color_rgb (Tuple[int, int, int]): An RGB tuple for the text\n            color. Defaults to (0, 114, 188).\n\n    Returns:\n        docx.text.paragraph.Paragraph: The newly created and styled heading\n            paragraph object.\n\n    Raises:\n        ValueError: If `level` is not an integer between 0 and 9, or if any\n            value in `primary_color_rgb` is outside the 0-255 range."}."""
    # Sanitization: Remove markdown emphasis characters from heading text.
    clean_txt = (
        text.replace("**", "").replace("__", "").replace("*", "").replace("_", "")
    )

    h = doc.add_heading(clean_txt, level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True

    # Each chapter styled with 'Heading 1' is configured to begin on a new page.
    # Using the 'Page break before' style property prevents the creation of empty pages that can result from explicit `add_page_break()` calls.
    if level == 1:
        h.paragraph_format.page_break_before = True

    # Define color palettes and dimensional parameters for document elements.
    for run in h.runs:
        run.font.color.rgb = RGBColor(*primary_color_rgb)
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
    return h


def add_spacer(doc, points=12):
    """Add a vertical spacer to a document via an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    return p


def add_toc(doc):
    r"""Inserts an Office Open XML field code to generate a Table of Contents (TOC).

    This function directly manipulates the underlying OOXML structure of the
    document to add a complex field for a TOC. It constructs the necessary
    `w:fldChar` (field character) and `w:instrText` (instruction text)
    elements. The generated field code is `TOC \\o "1-3" \\h \\z \\u`, which
    instructs a word processing application to build a hyperlinked TOC using
    heading levels 1 through 3.

    The inserted TOC is a placeholder that must be updated by the end-user in a
    host application (e.g., by right-clicking the field and selecting 'Update
    Field' in Microsoft Word) to populate its content.

    Args:
        doc (docx.document.Document): The `python-docx` document object to which
            the TOC field will be added.

    Returns:
        None. The `doc` object is modified in-place.

    Raises:
        AttributeError: If `doc` is not a valid `docx.document.Document` object
            or lacks the required methods for manipulation.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)


def set_update_fields(doc):
    """Configures the document to automatically update fields upon opening.

    Injects a `w:updateFields` element with its `w:val` attribute set to "true"
    into the document's core settings XML part (`/word/settings.xml`). This
    instructs consuming applications, such as Microsoft Word, to regenerate all
    dynamic fields (e.g., Table of Contents, page number references) when the
    document is next opened.

    This function modifies the `doc` object in-place.

    Args:
        doc (docx.document.Document): The `python-docx` document object to be
            modified.

    Returns:
        None
    """
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def parse_and_append_markdown_section(
    doc,
    md_path: Path,
    p_color_rgb,
    text_color_rgb,
    start_header=None,
    end_header=None,
    skip_level_1=False,
):
    r"""{'docstring': "Parses a specified section of a Markdown file and appends it to a DOCX document.\n\nReads a Markdown file and converts its syntax into formatted DOCX elements.\nThe function handles headings (levels 1-4), bullet points (`*` or `-`), and\nblockquotes (`>`). HTML entities within the source file are unescaped.\n\nParsing can be constrained to a specific section by providing `start_header`\nand `end_header` substrings. These are matched case-insensitively against\nlevel 2 headings to start and stop the inclusion process. If the specified\nMarkdown file does not exist, the function returns silently. Certain\nhardcoded placeholder lines for tables are also ignored.\n\nArgs:\n    doc (docx.document.Document): The document object to which content will be\n        appended. This object is modified in-place.\n    md_path (pathlib.Path): The path to the source Markdown file.\n    p_color_rgb (tuple[int, int, int]): An RGB tuple for primary heading styles.\n    text_color_rgb (tuple[int, int, int]): An RGB tuple for body text styles.\n    start_header (Optional[str]): A case-insensitive substring that, when found\n        in a level 2 heading, starts content parsing. If None, parsing\n        begins from the file's start. Defaults to None.\n    end_header (Optional[str]): A case-insensitive substring that, when found\n        in a level 2 heading, terminates content parsing. Defaults to None.\n    skip_level_1 (bool): If True, level 1 headings are skipped and not added\n        to the document. Defaults to False.\n\nReturns:\n    None. The `doc` object is modified in-place.\n\nRaises:\n    IOError: If the file at `md_path` exists but cannot be opened or read."}."""
    if not md_path.exists():
        return

    with open(md_path, "r", encoding="utf-8") as f:
        lines = [html.unescape(line) for line in f.readlines()]

    recording = start_header is None
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            continue

        if (
            start_header
            and cleaned.startswith("## ")
            and start_header.lower() in cleaned.lower()
        ):
            recording = True
            continue

        if (
            end_header
            and cleaned.startswith("## ")
            and end_header.lower() in cleaned.lower()
        ):
            recording = False
            break

        if not recording:
            continue

        # Skip processing for lines that are markers for table placeholders.
        if (
            "--- TABLA COMPARATIVA" in cleaned
            or "FORTALEZAS_CLAVE:" in cleaned
            or "BRECHAS_CLAVE:" in cleaned
            or "--- FIN TABLA" in cleaned
        ):
            continue

        # Assemble and render Section 1: Chapter Headings.
        if cleaned.startswith("# "):
            if skip_level_1:
                continue
            add_heading(doc, cleaned[2:], level=1, primary_color_rgb=p_color_rgb)
        elif cleaned.startswith("## "):
            add_heading(doc, cleaned[3:], level=2, primary_color_rgb=p_color_rgb)
        elif cleaned.startswith("### "):
            add_heading(doc, cleaned[4:], level=3, primary_color_rgb=p_color_rgb)
        elif cleaned.startswith("#### "):
            add_heading(doc, cleaned[5:], level=4, primary_color_rgb=p_color_rgb)
        elif cleaned.startswith("* ") or cleaned.startswith("- "):
            add_body_paragraph(
                doc, cleaned[2:], style="Bullet", text_color_rgb=text_color_rgb
            )
        elif cleaned.startswith("> "):
            add_body_paragraph(
                doc,
                cleaned[2:],
                italic=True,
                style="Intense Quote",
                text_color_rgb=text_color_rgb,
            )
        else:
            add_body_paragraph(doc, cleaned, text_color_rgb=text_color_rgb)


def add_appendix_heading(doc, text: str, primary_color_rgb: Any) -> Any:
    """Adds a styled paragraph to a document, formatted as an appendix heading.

    This function creates and appends a new paragraph with specific formatting
    attributes suitable for an appendix title. The applied styling includes:
      - Font: Century Gothic, 16pt, Bold.
      - Color: Set according to `primary_color_rgb`.
      - Paragraph Spacing: 18pt before and 6pt after.
      - Pagination: Configured to keep the heading with the subsequent paragraph.

    Args:
        doc (docx.document.Document): The `python-docx` document object to modify.
        text (str): The title text for the appendix heading.
        primary_color_rgb (Union[docx.shared.RGBColor, tuple[int, int, int]]):
            The RGB color for the heading text, provided as a `python-docx`
            RGBColor object or a 3-tuple of integers (0-255).

    Returns:
        docx.text.paragraph.Paragraph: The newly created and appended paragraph
            object containing the formatted appendix heading.

    Raises:
        TypeError: If `primary_color_rgb` is a tuple that does not contain
            exactly three integer elements.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Century Gothic"
    run.font.size = Pt(16)
    run.font.bold = True
    if isinstance(primary_color_rgb, tuple):
        run.font.color.rgb = RGBColor(*primary_color_rgb)
    else:
        run.font.color.rgb = primary_color_rgb
    return p


def extract_key_bullets_from_md(md_path: Path, section_marker: str) -> list[str]:
    """Extracts bullet point text from a designated section of a Markdown file.

    Parses a Markdown file to locate a section beginning after a specific marker.
    The function initiates content extraction upon finding a line containing
    `section_marker`. It then collects all subsequent lines formatted as bullet
    points (i.e., starting with `* ` or `- `), stripping the marker itself
    from the captured text. The process also applies HTML entity unescaping to
    each line read from the source file.

    Extraction is terminated if a line begins with "BRECHAS_CLAVE:",
    "FORTALEZAS_CLAVE:", contains "--- FIN TABLA", or if the end of the file is
    reached.

    Args:
        md_path (pathlib.Path): The file system path to the input Markdown file.
        section_marker (str): A substring used to identify the line that immediately
            precedes the target bullet points.

    Returns:
        list[str]: A list containing the text of each extracted bullet point.
            Returns an empty list if the file does not exist, the section marker
            is not found, or the section contains no qualifying bullet points.
    """
    if not md_path.exists():
        return []
    with open(md_path, "r", encoding="utf-8") as f:
        lines = [html.unescape(line) for line in f.readlines()]
    bullets = []
    recording = False
    for line in lines:
        cleaned = line.strip()
        if section_marker in cleaned:
            recording = True
            continue
        if recording:
            if (
                cleaned.startswith("BRECHAS_CLAVE:")
                or cleaned.startswith("FORTALEZAS_CLAVE:")
                or "--- FIN TABLA" in cleaned
            ):
                break
            if cleaned.startswith("* ") or cleaned.startswith("- "):
                bullets.append(cleaned[2:])
    return bullets


def compile_docx(tower_dir: str, output_path: str):
    """Assembles a DOCX report from modular source files and assessment data.

    Generates a styled DOCX report by orchestrating content from a structured
    directory of source files. The function integrates corporate branding profiles,
    internationalization (i18n) locales, assessment data (JSON, CSV), and
    narrative sections (Markdown) into a Word template. The final document
    includes a cover page, table of contents, executive summary, maturity
    analysis, quantitative risk matrix, and appendices.

    Args:
        tower_dir: The root directory path for a specific assessment. This
            directory must contain an `asis_modules` subdirectory with Markdown
            and CSV content, along with root-level JSON files for the assessment
            payload (e.g., `blueprint_*.json`) and scoring results
            (`scoring_output.json`).
        output_path: The full destination file path for the compiled DOCX document.
            The parent directory will be created if it does not exist.

    Returns:
        None. The document is generated and saved to the specified `output_path`.

    Raises:
        FileNotFoundError: If required input files or directories are not found.
        json.JSONDecodeError: If a JSON configuration or data file is malformed.
        ValueError: If data within a source file cannot be parsed into the
            expected numeric type.
    """
    tower_dir_obj = Path(tower_dir)
    modules_dir = tower_dir_obj / "asis_modules"

    print(f"📦 Compilando Anexo Técnico AS-IS Word desde módulos en: {modules_dir}")

    # Internationalization (i18n) settings are loaded from a central configuration file to support multiple regions.
    locales_path = Path("engine_config/locales.json")
    locales = {}
    if locales_path.exists():
        with open(locales_path, "r", encoding="utf-8-sig") as lf:
            locales = json.load(lf)

    # Dynamically load corporate brand profiles based on the operational context.
    brand_path = Path("engine_config/brand_profile.json")
    with open(brand_path, "r", encoding="utf-8-sig") as bf:
        brand = json.load(bf)

    company_name = brand.get("company_name", "NTT DATA")
    classification = brand.get("default_classification", "Confidencial")

    styling = brand.get("styling", {})
    p_color_hex = styling.get("primary_color_hex", "0072BC")
    alt_row_hex = styling.get("alternate_row_color_hex", "F2F2F2")
    text_color_rgb = styling.get("text_dark_color_rgb", [46, 64, 77])

    #
    r_color = int(p_color_hex[0:2], 16)
    g_color = int(p_color_hex[2:4], 16)
    b_color = int(p_color_hex[4:6], 16)
    p_color_rgb = (r_color, g_color, b_color)

    #
    from infrastructure.runtime_paths import resolve_tower_annex_template_path

    template_path = resolve_tower_annex_template_path()
    try:
        doc = Document(str(template_path))
        print("   ├─ Plantilla Word pre-estilizada cargada correctamente.")
    except Exception:
        doc = Document()
        print("   ⚠️ No se encontró plantilla. Inicializando documento en blanco.")

    # Initialize the document by removing all placeholder paragraphs and tables from the template.
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    # The `blueprint_payload` file path is resolved dynamically to support a multi-tenant architecture.
    payload_path = next(tower_dir_obj.glob("blueprint_*_payload.json"), None)
    if not payload_path:
        payload_path = tower_dir_obj / "blueprint_t2_payload.json"  #

    scoring_path = tower_dir_obj / "scoring_output.json"

    tower_name = "Desconocida"
    client_name = "Cliente"
    global_fair_ale = 0.0
    global_score = 3.0
    global_band = "Definido"
    global_reading = "Prácticas operativas consolidadas."
    client_group = None
    custom_overview_intro = None

    #
    doc_lang = "es"  #
    doc_date = datetime.now().strftime(
        "%d %B %Y"
    )  # Use the current system date if a date is not provided in the payload.
    doc_version = "1.0"
    currency = "€"

    if payload_path.exists():
        with open(payload_path, "r", encoding="utf-8-sig") as f:
            b_data = json.load(f)
            tower_meta = b_data.get("document_meta", {})
            tower_name = tower_meta.get("tower_name", "Desconocida")
            client_name = tower_meta.get("client_name", "Cliente")
            client_group = tower_meta.get("client_parent_group")  #
            global_fair_ale = b_data.get("total_fair_ale", 0.0)
            custom_overview_intro = b_data.get("platform_overview_intro")  #

            # Load internationalization (i18n) parameters from the input payload. The implementation avoids using default or hardcoded locales.
            doc_lang = tower_meta.get("language", "es").lower()
            doc_date = tower_meta.get("date", doc_date)
            doc_version = tower_meta.get("version", doc_version)
            currency = tower_meta.get("currency", currency)

    if scoring_path.exists():
        with open(scoring_path, "r", encoding="utf-8-sig") as f:
            s_data = json.load(f)
            global_score = s_data.get("tower_score_exact", 3.0)
            global_band = s_data.get("maturity_band_from_exact", {}).get(
                "label", "Definido"
            )
            global_reading = s_data.get("maturity_band_from_exact", {}).get(
                "reading", "Gobernanza estándar."
            )

    # Resolve the translation dictionary for the current session. This implementation is designed to be tenant-agnostic.
    vocab_fallback = LANG_VOCAB.get(doc_lang, LANG_VOCAB.get("es", {}))
    vocab_loaded = locales.get(doc_lang, locales.get("es", {}))
    vocab = {**vocab_fallback, **vocab_loaded}
    org_label = (
        "holding" if client_group else ("group" if doc_lang == "en" else "organización")
    )

    #
    # Assemble and render Section 0: Corporate Cover Page.
    #
    add_spacer(doc, 40)

    p_corp = doc.add_paragraph()
    run_corp = p_corp.add_run(company_name.upper())
    run_corp.bold = True
    run_corp.font.name = "Arial"
    run_corp.font.size = Pt(11)
    run_corp.font.color.rgb = RGBColor(*p_color_rgb)

    add_spacer(doc, 80)

    p_title = doc.add_paragraph()
    run_title = p_title.add_run(vocab["title"])
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(28)
    run_title.font.color.rgb = RGBColor(*text_color_rgb)

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(f"Anexo Técnico: {tower_name}")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(*p_color_rgb)

    add_spacer(doc, 100)

    p_client = doc.add_paragraph()
    run_cli_lbl = p_client.add_run(vocab["client_lbl"])
    run_cli_lbl.bold = True
    run_cli_lbl.font.size = Pt(10)
    run_cli_lbl.font.color.rgb = RGBColor(120, 130, 140)

    # Data Sanitization: Avoid hardcoding client-specific holding company names on the cover page to prevent data leakage, per requirement 1.
    if client_group:
        run_cli = p_client.add_run(f"{client_name.upper()} ({client_group})\n")
    else:
        run_cli = p_client.add_run(f"{client_name.upper()}\n")
    run_cli.font.size = Pt(12)
    run_cli.font.color.rgb = RGBColor(*text_color_rgb)

    run_proj_lbl = p_client.add_run(vocab["proj_lbl"])
    run_proj_lbl.bold = True
    run_proj_lbl.font.size = Pt(10)
    run_proj_lbl.font.color.rgb = RGBColor(120, 130, 140)
    run_proj = p_client.add_run(vocab["project"])
    run_proj.font.size = Pt(11)
    run_proj.font.color.rgb = RGBColor(*text_color_rgb)

    add_spacer(doc, 100)

    p_footer = doc.add_paragraph()
    run_date = p_footer.add_run(
        f"{doc_date} | {vocab['v_lbl']}{doc_version} | {classification}"
    )
    run_date.font.name = "Arial"
    run_date.font.size = Pt(9.5)
    run_date.font.color.rgb = RGBColor(150, 150, 150)

    #
    doc.add_page_break()

    #
    # Assemble and render Section 0.1: Table of Contents (TOC).
    #
    add_heading(doc, vocab["toc_title"], level=4, primary_color_rgb=p_color_rgb)
    add_toc(doc)

    # Manual page breaks are not used. Page breaks between chapters are controlled by the 'Page break before' property of the 'Heading 1' style.
    # A page break is automatically inserted before this element. The underlying 'Heading 1' style in the document template is configured with the 'Page break before' property.
    # The `page_break_before` property is enabled on this style to ensure each chapter begins on a new page.

    #
    # Assemble and render Chapter 1: Executive Summary and Business Context.
    #
    parse_and_append_markdown_section(
        doc, modules_dir / "02_resumen_ejecutivo.md", p_color_rgb, text_color_rgb
    )

    #
    # Assemble and render Chapter 2: Objective, Scope, and Assessment Methodology.
    #
    add_heading(
        doc,
        vocab["justification_table_headers"][0].split(" / ")[1]
        if "/" in vocab["justification_table_headers"][0]
        else "Metodología",
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    parse_and_append_markdown_section(
        doc,
        modules_dir / "01_introduccion.md",
        p_color_rgb,
        text_color_rgb,
        skip_level_1=True,
    )

    #
    # Assemble and render Chapter 3: Maturity Profile and General Evaluation.
    #
    add_heading(doc, vocab["radar_title"], level=1, primary_color_rgb=p_color_rgb)

    # Assemble and render Section 3.1: Global Maturity Dashboard; centered, 2-column layout per requirement 2.
    add_heading(doc, vocab["dashboard_title"], level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, vocab["dashboard_intro"], text_color_rgb=text_color_rgb)

    dash_table = doc.add_table(rows=2, cols=2)
    dash_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    finalize_table(dash_table)

    # Data Sanitization: Ensure data integrity by dynamically resolving scores and qualitative bands from the canonical scoring output.
    formatted_global_score = (
        f"{global_score:.2f}".replace(".", ",")
        if doc_lang == "es"
        else f"{global_score:.2f}"
    )
    display_score_str = (
        f"{formatted_global_score} / 5,00"
        if doc_lang == "es"
        else f"{formatted_global_score} / 5.00"
    )
    display_band_str = f"{global_band}"

    #
    for i, h_txt in enumerate([vocab["score_lbl"], vocab["maturity_lbl"]]):
        set_cell_text_custom(
            dash_table.rows[0].cells[i],
            h_txt,
            bold=True,
            font_size=9,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(dash_table.rows[0].cells[i], p_color_hex)
        dash_table.rows[0].cells[i].width = Inches(3.25)  #

    set_cell_text_custom(
        dash_table.rows[1].cells[0],
        display_score_str,
        bold=True,
        font_size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text_custom(
        dash_table.rows[1].cells[1],
        display_band_str,
        bold=True,
        font_size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Apply conditional cell shading based on the final severity score. A score below 3.40 receives a yellow background.
    shade_cell(
        dash_table.rows[1].cells[0], "FFF3CD" if global_score < 3.4 else "D9F2D9"
    )
    shade_cell(
        dash_table.rows[1].cells[1], "FFF3CD" if global_score < 3.4 else "D9F2D9"
    )

    # Set data cells to expand to the full table width to maintain a fixed layout.
    dash_table.rows[1].cells[0].width = Inches(3.25)
    dash_table.rows[1].cells[1].width = Inches(3.25)  #

    # Disable the `autofit` property to enforce a fixed table width of 6.5 inches.
    dash_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_spacer(doc, 10)

    add_body_paragraph(
        doc,
        f"**{vocab['resilience_reading']}** {global_reading}",
        text_color_rgb=text_color_rgb,
    )
    add_spacer(doc, 15)

    # Assemble and render Section 3.2: Full-Page Radar Chart; 6.0-inch width per requirement 2.
    add_heading(doc, vocab["radar_title"], level=2, primary_color_rgb=p_color_rgb)
    radar_path = tower_dir_obj / "pillar_radar_chart.generated.png"
    if radar_path.exists():
        add_spacer(doc, 5)
        doc.add_picture(str(radar_path), width=Inches(6.0))  #
        p_img = doc.paragraphs[-1]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_spacer(doc, 15)

    # Assemble and render Section 3.3: Score Justification Matrix; left-aligned, omits 'TO-BE Target' column.
    csv_mat_path = modules_dir / "04_matriz_madurez.csv"
    if csv_mat_path.exists():
        mat_table = doc.add_table(rows=1, cols=3)
        mat_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        mat_table.style = "Table Grid"

        headers_mat = vocab["justification_table_headers"]
        for i, h_txt in enumerate(headers_mat):
            set_cell_text_custom(
                mat_table.rows[0].cells[i],
                h_txt,
                bold=True,
                font_size=9,
                color_rgb=RGBColor(255, 255, 255),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(mat_table.rows[0].cells[i], p_color_hex)

        with open(csv_mat_path, "r", encoding="utf-8") as cf:
            reader = csv.reader(cf, delimiter=";")
            next(reader)  #

            for r_idx, row_data in enumerate(reader):
                row = mat_table.add_row()
                set_cell_text_custom(
                    row.cells[0], row_data[0], bold=True, font_size=8.5
                )
                set_cell_text_custom(
                    row.cells[1],
                    row_data[1],
                    font_size=8.5,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                set_cell_text_custom(row.cells[2], row_data[2], font_size=8)

                if r_idx % 2 == 1:
                    for cell in row.cells:
                        shade_cell(cell, alt_row_hex)

        autofit_table_to_contents(mat_table)
        add_spacer(doc, 20)

    #
    # Assemble and render Chapter 4: Technology Diagnosis and Cross-Cutting Platform Analysis.
    #
    add_heading(
        doc, vocab["platform_overview_title"], level=1, primary_color_rgb=p_color_rgb
    )

    # Assemble and render Section 4.1: Current State Infrastructure Platform Description.
    add_heading(
        doc, vocab["platform_overview_title"], level=2, primary_color_rgb=p_color_rgb
    )

    # Data Sanitization: Decouple volumetric data from the core logic to ensure scalability, per requirement 2.
    p_vol = doc.add_paragraph()
    p_vol.paragraph_format.space_after = Pt(10)
    p_vol.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if custom_overview_intro:
        # Conditionally process the detailed overview if present in the payload. This content may be populated by an external RAG system.
        overview_text = custom_overview_intro.strip()
    else:
        # Provide a generic fallback message that is not specific to any client or technology.
        overview_text = (
            f"{vocab['default_overview_text']}{tower_name} de {client_name} se basan "
            f"en un modelo de provisión técnica y telemetrías estructuradas recogidas durante la fase de auditoría. "
            f"Este baseline de activos e infraestructura actual proporciona el punto de partida fundamental para "
            f"mapear la topología, identificar cuellos de botella y justificar las brechas operativas detectadas "
            f"con respecto a los requerimientos de resiliencia y conformidad reguladora del {org_label} [Ref: Dossier de Contexto, Pág. 3]."
        )

    run_vol = p_vol.add_run(overview_text)
    run_vol.font.name = "Arial"
    run_vol.font.size = Pt(10)
    run_vol.font.italic = True
    run_vol.font.color.rgb = RGBColor(*text_color_rgb)

    parse_and_append_markdown_section(
        doc,
        modules_dir / "03_descripcion_plataforma.md",
        p_color_rgb,
        text_color_rgb,
        skip_level_1=True,
    )

    # Assemble and render Section 4.2: Key Strengths and Gaps.
    add_heading(doc, vocab["swot_title"], level=2, primary_color_rgb=p_color_rgb)
    parse_and_append_markdown_section(
        doc,
        modules_dir / "07_conclusiones.md",
        p_color_rgb,
        text_color_rgb,
        start_header="Resumen de Situación",
        end_header="Fortalezas y Brechas",
    )

    # Assemble and render the Quick Diagnosis Table using a two-column, left-aligned layout.
    strengths = extract_key_bullets_from_md(
        modules_dir / "07_conclusiones.md", "FORTALEZAS_CLAVE:"
    )
    gaps = extract_key_bullets_from_md(
        modules_dir / "07_conclusiones.md", "BRECHAS_CLAVE:"
    )

    if strengths or gaps:
        diag_table = doc.add_table(rows=1, cols=2)
        diag_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        diag_table.style = "Table Grid"

        set_cell_text_custom(
            diag_table.rows[0].cells[0],
            vocab["swot_headers"][0],
            bold=True,
            font_size=9.5,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(diag_table.rows[0].cells[0], "28B463")  #

        set_cell_text_custom(
            diag_table.rows[0].cells[1],
            vocab["swot_headers"][1],
            bold=True,
            font_size=9.5,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(diag_table.rows[0].cells[1], "C0392B")  #

        max_rows = max(len(strengths), len(gaps))
        for r_idx in range(max_rows):
            row = diag_table.add_row()
            s_text = strengths[r_idx] if r_idx < len(strengths) else ""
            g_text = gaps[r_idx] if r_idx < len(gaps) else ""

            set_cell_text_custom(row.cells[0], s_text, font_size=8.5)
            set_cell_text_custom(row.cells[1], g_text, font_size=8.5)

            if r_idx % 2 == 1:
                shade_cell(row.cells[0], alt_row_hex)
                shade_cell(row.cells[1], alt_row_hex)

        autofit_table_to_contents(diag_table)
        add_spacer(doc, 20)

    # Assemble and render the Key Operational Implications section.
    parse_and_append_markdown_section(
        doc,
        modules_dir / "07_conclusiones.md",
        p_color_rgb,
        text_color_rgb,
        start_header="Implicaciones Operativas Clave",
        end_header="Coste de Inacción",
    )
    doc.add_paragraph()

    # Assemble and render Section 4.3: Cross-Cutting Capability Analysis.
    add_heading(doc, vocab["transversal_title"], level=2, primary_color_rgb=p_color_rgb)
    parse_and_append_markdown_section(
        doc,
        modules_dir / "05_transversal.md",
        p_color_rgb,
        text_color_rgb,
        skip_level_1=True,
    )

    #
    # Assemble and render Chapter 5: Quantitative Risk Matrix.
    #
    add_heading(doc, vocab["risk_matrix_title"], level=1, primary_color_rgb=p_color_rgb)

    # Format the Annualized Loss Expectancy (ALE) value. Localization logic is isolated in a separate function, per requirement 2.
    t_sep = vocab.get("thousands_sep", ".")
    d_sep = vocab.get("decimal_sep", ",")
    formatted_ale = format_currency_custom(global_fair_ale, t_sep, d_sep)

    add_body_paragraph(
        doc,
        f"{vocab['risk_intro_1']}{org_label}{vocab['risk_intro_2']}**{formatted_ale} {currency}**{vocab['risk_intro_3']}",
        text_color_rgb=text_color_rgb,
    )

    #
    csv_risks_path = modules_dir / "06_matriz_riesgos_fair.csv"
    risks_data_list = []
    risks_by_pilar = {}

    if csv_risks_path.exists():
        with open(csv_risks_path, "r", encoding="utf-8") as cf:
            reader = csv.reader(cf, delimiter=";")
            next(reader)
            for r_idx, row_data in enumerate(reader):
                pilar_full_name = row_data[0].split(" - ")[0]
                r_item = {
                    "id": f"RVS{r_idx + 1:02d}",
                    "pilar": pilar_full_name,
                    "capability": row_data[0].split(" - ")[1]
                    if " - " in row_data[0]
                    else row_data[0],
                    "finding": row_data[1],
                    "business_risk": row_data[2],
                    "tef": float(row_data[3]),
                    "lm": float(row_data[4]),
                    "ale": float(row_data[5]),
                    "prioritization_rationale": row_data[6]
                    if len(row_data) > 6
                    else None,  #
                }
                risks_data_list.append(r_item)

                if pilar_full_name not in risks_by_pilar:
                    risks_by_pilar[pilar_full_name] = []
                risks_by_pilar[pilar_full_name].append(r_item)

    # Assemble and render Section 5.1: 5x5 Heatmap as a centered, native table mosaic.
    if risks_data_list:
        add_heading(
            doc, vocab["exposure_summary_title"], level=2, primary_color_rgb=p_color_rgb
        )

        matrix_cells = {r: {c: [] for c in range(1, 6)} for r in range(5, 0, -1)}
        for r in risks_data_list:
            r_tef = min(5, max(1, int(round(r["tef"]))))
            r_lm = min(5, max(1, int(round(r["lm"]))))
            matrix_cells[r_tef][r_lm].append(r["id"])

        # Initialize a 6x6 table (5x5 data grid plus headers). The table is centered and set to full page width.
        heatmap_table = doc.add_table(rows=6, cols=6)
        heatmap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        finalize_table(heatmap_table)

        #
        set_cell_text_custom(
            heatmap_table.rows[0].cells[0],
            "TEF \\ LM",
            bold=True,
            font_size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(heatmap_table.rows[0].cells[0], COLOR_BLUE)
        set_cell_border_white(heatmap_table.rows[0].cells[0])
        heatmap_table.rows[0].cells[0].width = Inches(1.0)  #
        for col_idx in range(1, 6):
            set_cell_text_custom(
                heatmap_table.rows[0].cells[col_idx],
                f"LM {col_idx}",
                bold=True,
                font_size=8,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(heatmap_table.rows[0].cells[col_idx], COLOR_BLUE)
            set_cell_border_white(heatmap_table.rows[0].cells[col_idx])
            heatmap_table.rows[0].cells[col_idx].width = Inches(
                1.1
            )  # Set widths for data columns 1-5. Total width is 6.5 inches: (1.1 inches * 5) + 1.0 inch for the header column.

        # Populate the 5x5 data matrix with content, ensuring it fills the full width of the table.
        for row_idx, tef_val in enumerate(range(5, 0, -1), 1):
            set_cell_text_custom(
                heatmap_table.rows[row_idx].cells[0],
                f"TEF {tef_val}",
                bold=True,
                font_size=8,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(heatmap_table.rows[row_idx].cells[0], COLOR_BLUE)
            set_cell_border_white(heatmap_table.rows[row_idx].cells[0])
            heatmap_table.rows[row_idx].cells[0].width = Inches(1.0)

            for col_idx, lm_val in enumerate(range(1, 6), 1):
                cell = heatmap_table.rows[row_idx].cells[col_idx]

                # Configure table dimensions to create a fixed, full-width mosaic layout.
                cell.width = Inches(1.1)
                heatmap_table.rows[row_idx].height = Inches(0.85)

                cell_risks = matrix_cells[tef_val][lm_val]
                cell_text = ", ".join(cell_risks) if cell_risks else "-"

                set_cell_text_custom(
                    cell,
                    cell_text,
                    bold=True,
                    font_size=8,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                set_cell_border_white(cell)

                severity = tef_val * lm_val
                if severity >= 15:
                    shade_cell(cell, "FADBD8")  #
                elif severity >= 8:
                    shade_cell(cell, "FCF3CF")  #
                else:
                    shade_cell(cell, "D5F5E3")  #

        heatmap_table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Recalculate cell content alignment. This is required because the initial auto-fit operation may not apply alignment correctly.
        add_spacer(doc, 15)

        # Assemble and render the Methodological Scale Legend, per requirement 1.
        add_heading(
            doc, "Leyendas Metodológicas", level=3, primary_color_rgb=p_color_rgb
        )

        # Assemble and render Table 1: Threat Event Frequency (TEF) Scale; full-width per requirement 1.
        add_heading(doc, vocab["tef_title"], level=4, primary_color_rgb=p_color_rgb)
        tef_table = doc.add_table(rows=1, cols=3)
        tef_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        finalize_table(tef_table)

        for i, h_txt in enumerate(vocab["tef_headers"]):
            set_cell_text_custom(
                tef_table.rows[0].cells[i],
                h_txt,
                bold=True,
                font_size=8.5,
                color_rgb=RGBColor(255, 255, 255),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(tef_table.rows[0].cells[i], p_color_hex)

        #
        tef_table.rows[0].cells[0].width = Inches(1.0)
        tef_table.rows[0].cells[1].width = Inches(1.2)
        tef_table.rows[0].cells[2].width = Inches(4.3)  #

        tef_data = [
            (
                "TEF 1",
                "Muy Bajo" if doc_lang == "es" else "Very Low",
                "< 0,1 eventos/año (menos de una vez cada 10 años)"
                if doc_lang == "es"
                else "< 0.1 events/year (less than once every 10 years)",
            ),
            (
                "TEF 2",
                "Bajo" if doc_lang == "es" else "Low",
                "0,1 - 0,5 eventos/año (una vez cada 4 años)"
                if doc_lang == "es"
                else "0.1 - 0.5 events/year (once every 4 years)",
            ),
            (
                "TEF 3",
                "Medio" if doc_lang == "es" else "Medium",
                "0,5 - 2,0 eventos/año (una vez al año)"
                if doc_lang == "es"
                else "0.5 - 2.0 events/year (once a year)",
            ),
            (
                "TEF 4",
                "Alto" if doc_lang == "es" else "High",
                "2,0 - 10 eventos/año (una vez al trimestre)"
                if doc_lang == "es"
                else "2.0 - 10 events/year (once a quarter)",
            ),
            (
                "TEF 5",
                "Muy Alto" if doc_lang == "es" else "Very High",
                "> 10 eventos/año (una vez al mes)"
                if doc_lang == "es"
                else "> 10 events/year (once a month)",
            ),
        ]
        for t_idx, row_vals in enumerate(tef_data):
            row = tef_table.add_row()
            set_cell_text_custom(
                row.cells[0],
                row_vals[0],
                bold=True,
                font_size=8,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text_custom(
                row.cells[1], row_vals[1], font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER
            )
            set_cell_text_custom(row.cells[2], row_vals[2], font_size=8)

            #
            row.cells[0].width = Inches(1.0)
            row.cells[1].width = Inches(1.2)
            row.cells[2].width = Inches(4.3)

            if t_idx % 2 == 1:
                shade_cell(row.cells[0], alt_row_hex)
                shade_cell(row.cells[1], alt_row_hex)
                shade_cell(row.cells[2], alt_row_hex)

        # Disable the `autofit` property to enforce a fixed table width of 6.5 inches.
        add_spacer(doc, 10)

        # Assemble and render Table 2: Loss Magnitude (LM) Scale; full-width per requirement 1.
        add_heading(doc, vocab["lm_title"], level=4, primary_color_rgb=p_color_rgb)
        lm_table = doc.add_table(rows=1, cols=3)
        lm_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        finalize_table(lm_table)

        for i, h_txt in enumerate(vocab["lm_headers"]):
            set_cell_text_custom(
                lm_table.rows[0].cells[i],
                f"{h_txt} ({currency})",
                bold=True,
                font_size=8.5,
                color_rgb=RGBColor(255, 255, 255),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(lm_table.rows[0].cells[i], p_color_hex)

        #
        lm_table.rows[0].cells[0].width = Inches(1.0)
        lm_table.rows[0].cells[1].width = Inches(1.2)
        lm_table.rows[0].cells[2].width = Inches(4.3)  #

        lm_data = [
            (
                "LM 1",
                "Muy Bajo" if doc_lang == "es" else "Very Low",
                f"< 1.000 {currency}" if doc_lang == "es" else f"< 1,000 {currency}",
            ),
            (
                "LM 2",
                "Bajo" if doc_lang == "es" else "Low",
                f"1.000 {currency} - 5.000 {currency}"
                if doc_lang == "es"
                else f"1,000 {currency} - 5,000 {currency}",
            ),
            (
                "LM 3",
                "Medio" if doc_lang == "es" else "Medium",
                f"5.000 {currency} - 25.000 {currency}"
                if doc_lang == "es"
                else f"5,000 {currency} - 25,000 {currency}",
            ),
            (
                "LM 4",
                "Alto" if doc_lang == "es" else "High",
                f"25.000 {currency} - 100.000 {currency}"
                if doc_lang == "es"
                else f"25,000 {currency} - 100,000 {currency}",
            ),
            (
                "LM 5",
                "Muy Alto" if doc_lang == "es" else "Very High",
                f"100.000 {currency} - 500.000 {currency}"
                if doc_lang == "es"
                else f"100,000 {currency} - 500,000 {currency}",
            ),
        ]
        for l_idx, row_vals in enumerate(lm_data):
            row = lm_table.add_row()
            set_cell_text_custom(
                row.cells[0],
                row_vals[0],
                bold=True,
                font_size=8,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text_custom(
                row.cells[1], row_vals[1], font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER
            )
            set_cell_text_custom(row.cells[2], row_vals[2], font_size=8)

            #
            row.cells[0].width = Inches(1.0)
            row.cells[1].width = Inches(1.2)
            row.cells[2].width = Inches(4.3)

            if l_idx % 2 == 1:
                shade_cell(row.cells[0], alt_row_hex)
                shade_cell(row.cells[1], alt_row_hex)
                shade_cell(row.cells[2], alt_row_hex)

        # Disable the `autofit` property to enforce a fixed table width relative to the page margins.
        add_spacer(doc, 20)

    # Assemble and render Section 5.2: Top Priority Risks; hierarchical indentation per requirement 4.1.
    if risks_data_list:
        add_heading(
            doc, vocab["top_risks_title"], level=2, primary_color_rgb=p_color_rgb
        )
        add_body_paragraph(doc, vocab["top_risks_intro"], text_color_rgb=text_color_rgb)

        top_risks = [r for r in risks_data_list if r["ale"] >= 100000]
        top_risks.sort(key=lambda x: x["ale"], reverse=True)

        for r_item in top_risks[:4]:
            p_top = add_body_paragraph(
                doc,
                f"{r_item['id']}: {r_item['pilar']} - {r_item['capability']} (ALE: {r_item['ale']:,.0f} {currency})".replace(
                    ",", "."
                ),
                style="Bullet",
                text_color_rgb=text_color_rgb,
            )
            p_top.paragraph_format.space_after = Pt(2)

            # Apply an indented, hierarchical sub-bullet style for the prioritization rationale, per requirement 4.1.
            p_mot = doc.add_paragraph()
            p_mot.paragraph_format.left_indent = Inches(0.4)
            p_mot.paragraph_format.space_after = Pt(6)

            run_mot_lbl = p_mot.add_run(f"   • {vocab['motivo_priorizacion']}")
            run_mot_lbl.bold = True
            run_mot_lbl.font.name = "Arial"
            run_mot_lbl.font.size = Pt(9.5)
            run_mot_lbl.font.color.rgb = RGBColor(*text_color_rgb)

            # Data Sanitization: Use a universal prioritization logic resolver to avoid client-specific implementations, per requirement 3.
            mot = r_item.get("prioritization_rationale")
            if not mot:
                # If a value is not provided in the payload, generate it dynamically based on the risk impact to avoid hardcoded dependencies.
                biz_impact_txt = r_item.get(
                    "business_risk",
                    "Exposición de riesgo crítica que compromete la resiliencia operativa.",
                )
                mot = f"La persistencia de esta brecha arriesga directamente la continuidad del servicio {org_label}, exponiéndolo a un coste de inacción crítico debido a: {biz_impact_txt}"

            run_mot = p_mot.add_run(mot)
            run_mot.font.name = "Arial"
            run_mot.font.size = Pt(9.5)
            run_mot.font.color.rgb = RGBColor(*text_color_rgb)

        add_spacer(doc, 15)

    # Assemble and render Section 5.3: Detailed Risk Register by Pillar; left-aligned content per requirement 4.2.
    add_heading(
        doc, vocab["detailed_risks_title"], level=2, primary_color_rgb=p_color_rgb
    )

    for p_name, p_risks in risks_by_pilar.items():
        add_heading(
            doc,
            f"{vocab['detailed_risks_pilar_title']}{p_name}",
            level=3,
            primary_color_rgb=p_color_rgb,
        )

        table_risks = doc.add_table(rows=1, cols=3)
        table_risks.alignment = WD_TABLE_ALIGNMENT.LEFT
        finalize_table(table_risks)

        headers_table = vocab["detailed_risks_headers"]
        for i, h_txt in enumerate(headers_table):
            set_cell_text_custom(
                table_risks.rows[0].cells[i],
                h_txt,
                bold=True,
                font_size=8.5,
                color_rgb=RGBColor(255, 255, 255),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(table_risks.rows[0].cells[i], p_color_hex)

        for r_idx, r in enumerate(p_risks):
            row = table_risks.add_row()

            #
            set_cell_text_custom(
                row.cells[0],
                r["id"],
                bold=True,
                font_size=8.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(row.cells[0], alt_row_hex)

            # Assemble and render Section 2: Description and Evidence, left-aligned per requirement 4.2.
            p_desc = row.cells[1].paragraphs[0]
            p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT

            finding = r["finding"]
            biz_risk = r["business_risk"]

            parts_f = finding.split("\n\nEvidencia Forense Literal (Audit RAG):\n")

            run_obs = p_desc.add_run(vocab["vulnerability_lbl"])
            run_obs.bold = True
            p_desc.add_run(clean_text_for_word(parts_f[0]) + "\n\n")

            if len(parts_f) > 1:
                run_ev = p_desc.add_run(vocab["evidence_lbl"])
                run_ev.bold = True
                run_cite = p_desc.add_run(parts_f[1])
                run_cite.italic = True
                run_cite.font.color.rgb = RGBColor(100, 110, 120)

            for run in p_desc.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8.5)

            # Assemble and render Section 3: Business Risk and FAIR Impact, left-aligned per requirement 4.2.
            p_risk = row.cells[2].paragraphs[0]
            p_risk.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run_biz = p_risk.add_run(vocab["impact_lbl"])
            run_biz.bold = True
            p_risk.add_run(clean_text_for_word(biz_risk) + "\n\n")

            tef = r["tef"]
            lm = r["lm"]
            ale = r["ale"]

            formatted_ale_item = format_currency_custom(ale, t_sep, d_sep).split(d_sep)[
                0
            ]
            calc_txt = f"{vocab['exposure_lbl']}TEF: {tef:.1f} / 5.0\nLM: {lm:.1f} / 5.0\n\n{vocab['ale_proyectado']}{formatted_ale_item} {currency}"

            run_fair = p_risk.add_run(calc_txt)
            run_fair.bold = True
            run_fair.font.color.rgb = RGBColor(150, 0, 0)

            for run in p_risk.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8.5)

            # Apply conditional color shading to the cell based on the severity value.
            vuln = r.get("vulnerability_level", 3.0) if hasattr(r, "get") else 3.0
            bg_color = "D9F2D9"  #
            if (tef * vuln) >= 15 or ale >= 1000000:
                bg_color = "F8D7DA"  #
            elif (tef * vuln) >= 10 or ale >= 250000:
                bg_color = "FFF3CD"  #
            elif (tef * vuln) >= 5 or ale >= 50000:
                bg_color = "E2E3E5"  #
            shade_cell(row.cells[2], bg_color)

            if r_idx % 2 == 1:
                shade_cell(row.cells[1], alt_row_hex)

        autofit_table_to_contents(table_risks)
        add_spacer(doc, 15)

    #
    # Assemble and render Chapter 6: Next Steps and Cost of Inaction.
    #
    add_heading(doc, vocab["next_steps_title"], level=1, primary_color_rgb=p_color_rgb)
    parse_and_append_markdown_section(
        doc,
        modules_dir / "07_conclusiones.md",
        p_color_rgb,
        text_color_rgb,
        start_header="Coste de Inacción",
    )

    #
    # Assemble and render Appendix A: Glossary and Abbreviations.
    #
    add_appendix_heading(doc, vocab["appendix_a_title"], primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, vocab["appendix_a_intro"], text_color_rgb=text_color_rgb)

    glossary_path = Path("engine_config/abbreviations_glossary.json")
    if glossary_path.exists():
        with open(glossary_path, "r", encoding="utf-8-sig") as gf:
            glossary = json.load(gf)

        gloss_table = doc.add_table(rows=1, cols=2)
        gloss_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        finalize_table(gloss_table)

        for i, h_txt in enumerate(vocab["appendix_a_headers"]):
            set_cell_text_custom(
                gloss_table.rows[0].cells[i],
                h_txt,
                bold=True,
                font_size=9,
                color_rgb=RGBColor(255, 255, 255),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(gloss_table.rows[0].cells[i], p_color_hex)

        #
        gloss_table.rows[0].cells[0].width = Inches(1.5)
        gloss_table.rows[0].cells[1].width = Inches(5.0)  #

        for g_idx, (term, desc) in enumerate(sorted(glossary.items())):
            row = gloss_table.add_row()
            set_cell_text_custom(row.cells[0], term, bold=True, font_size=8.5)
            set_cell_text_custom(row.cells[1], desc, font_size=8)

            #
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(5.0)

            if g_idx % 2 == 1:
                shade_cell(row.cells[0], alt_row_hex)
                shade_cell(row.cells[1], alt_row_hex)

        # Disable the `autofit` property to enforce a fixed table width of 6.5 inches.

    #
    # Assemble and render Appendix B: Limitation of Liability.
    #
    add_appendix_heading(doc, vocab["appendix_b_title"], primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, vocab["disclaimer_text_1"], text_color_rgb=text_color_rgb)
    add_body_paragraph(doc, vocab["disclaimer_text_2"], text_color_rgb=text_color_rgb)
    add_body_paragraph(doc, vocab["disclaimer_text_3"], text_color_rgb=text_color_rgb)

    #
    # Assemble and render Appendix C: Information Source Custody Record.
    #
    add_appendix_heading(doc, vocab["appendix_c_title"], primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, vocab["appendix_c_intro"], text_color_rgb=text_color_rgb)

    source_table = doc.add_table(rows=1, cols=3)
    source_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    finalize_table(source_table)

    for i, h_txt in enumerate(vocab["appendix_c_headers"]):
        set_cell_text_custom(
            source_table.rows[0].cells[i],
            h_txt,
            bold=True,
            font_size=8.5,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(source_table.rows[0].cells[i], p_color_hex)

    source_table.rows[0].cells[0].width = Inches(1.5)
    source_table.rows[0].cells[1].width = Inches(2.0)
    source_table.rows[0].cells[2].width = Inches(3.0)  #

    # Data Sanitization: Prevent generation of invalid source file paths in the bibliography, per requirement 4.
    src_data = []
    source_docs = b_data.get("source_documents") if payload_path.exists() else None

    if source_docs:
        # If source file paths are provided in the payload, use them directly, overriding any dynamic generation logic.
        for doc_item in source_docs:
            src_data.append(
                (
                    doc_item.get("code", "[Doc]"),
                    doc_item.get("name", "Documento"),
                    doc_item.get("desc", "Documento bajo custodia de auditoría."),
                )
            )
    else:
        # Dynamically generate a context-aware bibliography based on the specified technology stack to avoid invalid source file references.
        src_data = [
            (
                "[Cuestionario de Autoevaluación]",
                f"preguntas_{client_name.lower()}_con_notas.txt",
                f"{vocab['bib_cues']}{tower_name}{vocab['bib_cues_desc']}{client_name}.",
            ),
            (
                "[Dossier de Contexto]",
                f"contexto_{client_name.lower()}_elite.docx",
                f"{vocab['bib_contexto']}{client_name}{vocab['bib_contexto_desc']}",
            ),
            (
                "[Minutas de Sesión]",
                f"Sesión de Contexto ({client_name})",
                f"{vocab['bib_minutas']}{client_name}{vocab['bib_minutas_desc']}",
            ),
        ]

    for s_idx, row_vals in enumerate(src_data):
        row = source_table.add_row()
        set_cell_text_custom(
            row.cells[0],
            row_vals[0],
            bold=True,
            font_size=8.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text_custom(
            row.cells[1], row_vals[1], font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER
        )
        set_cell_text_custom(row.cells[2], row_vals[2], font_size=8)

        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(3.0)

        if s_idx % 2 == 1:
            shade_cell(row.cells[0], alt_row_hex)
            shade_cell(row.cells[1], alt_row_hex)
            shade_cell(row.cells[2], alt_row_hex)

    # Insert a field code to instruct the host application to update the Table of Contents upon opening.
    set_update_fields(doc)

    #
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(
        f"🎉 ¡Anexo Técnico AS-IS COMPILADO modularmente con éxito en: {output_path}!"
    )


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(
            "Uso: python compile_asis_docx_from_modules.py <tower_dir> <output_doc.docx>"
        )
        sys.exit(1)
    compile_docx(sys.argv[1], sys.argv[2])
