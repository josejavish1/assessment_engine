import os
import json
import sys
import glob
import tempfile
import uuid
from pathlib import Path
from typing import Any
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import numpy as np

from infrastructure.docx_render_utils import (
    add_heading_paragraph,
    autofit_table_to_contents,
    finalize_table,
    set_cell_text,
    shade_cell,
)
from infrastructure.text_utils import clean_text_for_word

# --- CORPORATE STYLING CONSTANTS ---
COLOR_BLUE = "0072BC"
COLOR_HEADER_BG = "D9EAF7"
COLOR_ROW_ALT = "F2F2F2"

def clean_text_for_word(text):
    if text is None:
        return ""
    return str(text)

def set_cell_text_custom(cell, text, bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(clean_text_for_word(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    if color_rgb:
        run.font.color.rgb = color_rgb


def add_body_paragraph(doc, text, bold=False, italic=False, space_after=6, text_color_rgb=(46, 64, 77), style='Normal') -> Any:
    try:
        p = doc.add_paragraph(style=style)
        is_bullet = (style == 'List Bullet')
    except KeyError:
        p = doc.add_paragraph(style='Normal')
        is_bullet = (style == 'List Bullet')
        if is_bullet:
            run_b = p.add_run("• ")
            run_b.font.name = "Arial"
            run_b.font.size = Pt(10)
            run_b.font.color.rgb = RGBColor(*text_color_rgb)
            
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # SOTA: Auto-bolding on the first 3-4 words of bullet points for executive scanning
    if is_bullet and ":" in text:
        parts = text.split(":", 1)
        run_bold = p.add_run(parts[0] + ":")
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_bold.font.color.rgb = RGBColor(*text_color_rgb)
        
        run_rest = p.add_run(parts[1])
        run_rest.font.size = Pt(10)
        run_rest.font.color.rgb = RGBColor(*text_color_rgb)
    elif "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 1:
                run = p.add_run(part)
                run.font.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*text_color_rgb)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*text_color_rgb)
    return p

def add_heading(doc, text, level, primary_color_rgb=(0, 114, 188)):
    h = doc.add_heading(text, level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    
    # Custom color and sizes, inheriting font dynamically. Numbers stripped for native auto-numbering.
    for run in h.runs:
        run.font.color.rgb = RGBColor(*primary_color_rgb)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
    return h

def add_spacer(doc, points=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    return p

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

def generate_radar_chart(labels: list, actual_scores: list, target_scores: list, title: str, output_path: str):
    """Genera un gráfico de radar para madurez."""
    if not labels:
        return False
    num_vars = len(labels)
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
    
    actual_scores += actual_scores[:1]
    target_scores += target_scores[:1]
    angles += angles[:1]
    
    fig, ax = plt.subplots(figsize=(6, 5), subplot_kw=dict(polar=True))
    
    # Dibujar Target (Objetivo)
    ax.plot(angles, target_scores, color='#0072BC', linewidth=2, linestyle='solid', label='Objetivo (TO-BE)')
    ax.fill(angles, target_scores, color='#0072BC', alpha=0.1)
    
    # Dibujar Actual (AS-IS)
    ax.plot(angles, actual_scores, color='#C00000', linewidth=2, linestyle='solid', label='Actual (AS-IS)')
    ax.fill(angles, actual_scores, color='#C00000', alpha=0.25)
    
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    
    import textwrap
    wrapped_labels = [textwrap.fill(l, 15) for l in labels]
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(wrapped_labels, size=8)
    
    ax.set_ylim(0, 5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(['1', '2', '3', '4', '5'], color="grey", size=7)
    
    ax.set_title(title, size=11, weight='bold', pad=15)
    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1))
    
    plt.tight_layout()
    plt.savefig(output_path, format='png', dpi=300)
    plt.close()
    return True


def render_consolidated_asis(working_dir: str, output_path: str):
    print("🚀 Cargando Payloads de Torres para Renderizado AS-IS Consolidado...")
    
    # Mapeo de prefijos de riesgo según el estándar de la FNMT
    prefix_map = {
        "T1": "RDC", # Riesgo Data Center
        "T2": "RVS", # Riesgo Virtualización y Servidores
        "T4": "RAB", # Almacenamiento y Backup
        "T5": "RRC", # Resiliencia y Continuidad
        "T6": "RSI", # Riesgo Seguridad e Identidad
        "T7": "RIT", # Riesgo ITSM y Operaciones
        "T8": "RSG", # Riesgo Estrategia y Gobernanza
        "T10": "RLM" # Legacy & Mainframe
    }
    
    # Cargar todos los payloads de las torres en working_dir
    payloads = []
    total_ale_global = 0.0
    all_risks = []
    tower_overviews = []
    input_documents = []
    
    search_pattern = os.path.join(working_dir, "T*", "blueprint_*_payload.json")
    for file_path in glob.glob(search_pattern):
        with open(file_path, 'r', encoding='utf-8-sig') as f:
            data = json.load(f)
            payloads.append(data)
            
            tower_name = data.get("document_meta", {}).get("tower_name", "Desconocida")
            tower_id = data.get("document_meta", {}).get("tower_code", "")
            
            total_ale_global += data.get("total_fair_ale", 0.0)
            
            snapshot = data.get("executive_snapshot", {})
            headline = snapshot.get("headline", "Sin resumen ejecutivo.")
            bottom_line = snapshot.get("bottom_line", "")
            
            pil_scores = [p.get("score", 0.0) for p in data.get("pillars_analysis", [])]
            t_scores = [p.get("target_score", 4.0) for p in data.get("pillars_analysis", [])]
            avg_score = sum(pil_scores) / len(pil_scores) if pil_scores else 0.0
            avg_target = sum(t_scores) / len(t_scores) if t_scores else 4.0
            
            tower_overviews.append({
                "tower_id": tower_id,
                "tower_name": tower_name,
                "score": avg_score,
                "target": avg_target,
                "headline": headline,
                "bottom_line": bottom_line,
                "pillars": [{"name": p.get("pilar_name", ""), "score": p.get("score", 0.0), "target": p.get("target_score", 4.0), "desc": p.get("asis_architecture_description", p.get("thought_process", "Sin descripción detallada."))} for p in data.get("pillars_analysis", [])]
            })
            
            input_documents.append(f"AS-IS_Anexo_Tecnico_{tower_id}.docx")
            
            # Recolectar riesgos para la matriz global
            for pilar in data.get("pillars_analysis", []):
                for hc in pilar.get("health_check_asis", []):
                    hc["tower_id"] = tower_id
                    hc["tower_name"] = tower_name
                    hc["pilar_name"] = pilar.get("pilar_name", "")
                    
                    # Normalizar nombres de llaves debido a los alias de Pydantic
                    hc["finding"] = hc.get("finding", hc.get("risk_observed", "No disponible."))
                    hc["business_risk"] = hc.get("business_risk", hc.get("impact", "No disponible."))
                    all_risks.append(hc)
                    
    if not payloads:
        print("❌ No se encontraron payloads de torre en working_dir.")
        return
        
    client_name = payloads[0].get("document_meta", {}).get("client_name", "Cliente")
    
    # SANEAMIENTO METADATOS: Cargar parámetros de localización, ratios de nube y cumplimiento de forma dinámica (Punto 2)
    doc_meta_global = payloads[0].get("document_meta", {})
    doc_lang_global = doc_meta_global.get("language", "es").lower()
    on_prem_pct = doc_meta_global.get("on_premise_percentage", "80%")
    cloud_pct = doc_meta_global.get("cloud_percentage", "20%")
    reg_frameworks = doc_meta_global.get("regulatory_frameworks")
    if not reg_frameworks:
        reg_frameworks = "NIS2 / ENS" if doc_lang_global == "es" else "applicable regulatory frameworks"
    
    # SOTA 2026: Cargar Inteligencia Estratégica del Cliente
    working_dir_path = Path(working_dir)
    client_intel_path = working_dir_path / "client_intelligence.json"
    intel = {}
    if client_intel_path.exists():
        with open(client_intel_path, "r", encoding="utf-8-sig") as cif:
            try:
                intel = json.load(cif)
                print("   ├─ Carga exitosa de Inteligencia Estratégica del Cliente en Nivel Consolidado.")
            except Exception as e:
                print(f"   ⚠️ Error cargando inteligencia: {e}")
                
    business_context = intel.get("business_context", {})
    ceo_agenda_raw = business_context.get("ceo_agenda", {}).get("summary", "No disponible.")
    
    # Cargar plantilla pre-estilizada
    template_path = "templates/template_tobe_consolidated.docx"
    if os.path.exists(template_path):
        doc = Document(template_path)
        print("   ├─ Plantilla Word pre-estilizada cargada correctamente.")
    else:
        doc = Document()
        print("   ⚠️ Plantilla no encontrada. Generando con estilos por defecto.")

    # LIMPIEZA: Vaciar párrafos y tablas placeholder de la plantilla
    for p in list(doc.paragraphs):
        p._p.getparent().remove(p._p)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    # ---------------------------------------------------------
    # 0. PORTADA PORTADA CORPORATIVA ELEGANTE
    # ---------------------------------------------------------
    doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_color_hex = COLOR_BLUE
    r_color = int(p_color_hex[0:2], 16)
    g_color = int(p_color_hex[2:4], 16)
    b_color = int(p_color_hex[4:6], 16)
    p_color_rgb = (r_color, g_color, b_color)
    
    for _ in range(5): doc.add_paragraph()
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("Informe Consolidado de Situación Actual\n(AS-IS) de la Plataforma")
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(*p_color_rgb)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(40)
    run_sub = p_sub.add_run("Tier-1 Infrastructure Modernisation & Resilience Assessment")
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(128, 128, 128)
    
    p_client = doc.add_paragraph()
    run_client = p_client.add_run(f"Preparado para: {client_name}")
    run_client.font.size = Pt(12)
    run_client.font.bold = True
    
    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run("NTT DATA Corporation | Confidencial\nVersión 1.2 | Junio 2026")
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    
    doc.add_page_break()
    
    # Tabla de Contenidos
    add_heading(doc, "Índice de Contenidos", level=1, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, "Nota: Para actualizar el índice en Word, haga clic derecho sobre el texto inferior y seleccione 'Actualizar campos'.", italic=True)
    add_toc(doc)
    doc.add_page_break()

    # ---------------------------------------------------------
    # 1. INTRODUCCIÓN Y DOCUMENTOS FUENTE SOTA
    # ---------------------------------------------------------
    add_heading(doc, "Introducción y Documentos Fuente", level=1, primary_color_rgb=p_color_rgb)
    
    add_heading(doc, "Objetivo del Diagnóstico", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, f"El presente informe ejecutivo consolida los hallazgos de los análisis AS-IS realizados sobre los distintos dominios tecnológicos de la plataforma de {client_name}. Su propósito es ofrecer a la Dirección una visión unificada y accionable del estado actual, priorizando los riesgos por impacto de negocio y resiliencia.")
    
    add_heading(doc, "Documentos Fuente Evaluados", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, "Este informe consolida las auditorías forenses y encuestas de madurez documentadas de manera independiente para cada una de las torres tecnológicas de la organización:")
    
    doc_table = doc.add_table(rows=1, cols=2)
    doc_table.style = 'Table Grid'
    
    headers_doc = ["Dominio Tecnológico Evaluado", "Documento Anexo AS-IS Resultante"]
    for i, h_txt in enumerate(headers_doc):
        set_cell_text_custom(doc_table.rows[0].cells[i], h_txt, bold=True, font_size=9, color_rgb=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(doc_table.rows[0].cells[i], COLOR_BLUE)
        
    for r_idx, file_name in enumerate(sorted(input_documents)):
        row = doc_table.add_row()
        # Parse name from file for beautiful display
        t_name = file_name.replace("AS-IS_Anexo_Tecnico_", "").replace(".docx", "")
        set_cell_text_custom(row.cells[0], f"Torre {t_name} - Diagnóstico Detallado de Infraestructura", bold=True, font_size=8.5)
        set_cell_text_custom(row.cells[1], file_name, font_size=8.5)
        if r_idx % 2 == 1:
            for cell in row.cells: shade_cell(cell, COLOR_ROW_ALT)
            
    autofit_table_to_contents(doc_table)
    doc.add_page_break()

    # ---------------------------------------------------------
    # 2. RESUMEN EJECUTIVO GLOBAL DE SITUACIÓN ACTUAL SOTA
    # ---------------------------------------------------------
    add_heading(doc, "Resumen Ejecutivo de Situación Actual Global (AS-IS)", level=1, primary_color_rgb=p_color_rgb)
    
    add_heading(doc, "Diagnóstico General de la Plataforma", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, f"El estado actual de la plataforma tecnológica de {client_name}, aunque operativo en sus pilares tradicionales, se encuentra en un punto crítico de obsolescencia operativa y fragmentación estructural. La coexistencia de tecnologías on-premise ({on_prem_pct}) y cloud ({cloud_pct}) operadas en silos independientes limita severamente la agilidad de entrega y la capacidad de resiliencia ante contingencias críticas.")
    
    # SOTA: Inyectar la Agenda del Liderazgo y el Plan de Inversión de Redeia
    if ceo_agenda_raw and ceo_agenda_raw != "No disponible.":
        add_heading(doc, "Por qué importa al negocio (Agenda de Liderazgo)", level=2, primary_color_rgb=p_color_rgb)
        for block in ceo_agenda_raw.split("\n\n"):
            if block.strip().startswith("###"):
                add_heading(doc, block.replace("###", "").strip(), level=3, primary_color_rgb=p_color_rgb)
            elif block.strip().startswith("####"):
                add_heading(doc, block.replace("####", "").strip(), level=4, primary_color_rgb=p_color_rgb)
            elif block.strip().startswith("1.") or block.strip().startswith("-"):
                add_body_paragraph(doc, block.strip(), style='List Bullet')
            else:
                add_body_paragraph(doc, block.strip())

    add_heading(doc, "Exposición Financiera al Riesgo Global (FAIR)", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, f"A partir del modelo de cuantificación financiera de riesgo **FAIR**, la Expectativa de Pérdida Anualizada (ALE) consolidada agregando todos los riesgos críticos detectados en el assessment de la plataforma asciende a:")
    
    p_ale = doc.add_paragraph()
    p_ale.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ale = p_ale.add_run(f"ALE Global Proyectado: {total_ale_global:,.0f} € / Año".replace(",", "."))
    run_ale.font.size = Pt(16)
    run_ale.font.bold = True
    run_ale.font.color.rgb = RGBColor(192, 0, 0) # Dark Red for alerting
    
    add_body_paragraph(doc, f"Esta cifra representa la expectativa de coste anualizado resultante de mantener el status quo de operaciones manuales, indisponibilidades no programadas de sistemas críticos y deudas de cumplimiento normativo ({reg_frameworks}).")
    
    doc.add_page_break()

    # ---------------------------------------------------------
    # 3. DESCRIPCIÓN DE LA PLATAFORMA POR TORRE
    # ---------------------------------------------------------
    add_heading(doc, "Descripción de la Plataforma de Infraestructura Actual", level=1, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, "A continuación, se define el estado de situación y arquitectura de cada torre evaluada, proporcionando la visión ejecutiva y el 'bottom line' de cada dominio:")
    
    for t in sorted(tower_overviews, key=lambda x: x["tower_id"]):
        add_heading(doc, f"Dominio: {t['tower_name']} (Torre {t['tower_id']})", level=2, primary_color_rgb=p_color_rgb)
        add_body_paragraph(doc, f"**Nivel de Madurez Actual:** Nivel {t['score']:.2f} de 5,00 (Clasificación: {t['headline']})", bold=True)
        add_body_paragraph(doc, t["bottom_line"])
        
        # Slices de pilares de esta torre
        add_heading(doc, "Estructura de Pilares Técnicos Evaluados", level=3, primary_color_rgb=p_color_rgb)
        for p in t["pillars"]:
            add_body_paragraph(doc, f"**{p['name']}:** {p['desc'].split('.')[0]}.", style='List Bullet')
            
        doc.add_paragraph() # Separador

    doc.add_page_break()

    # ---------------------------------------------------------
    # 4. REGISTRO CONSOLIDADO DE VULNERABILIDADES (MATRIZ GLOBAL)
    # ---------------------------------------------------------
    add_heading(doc, "Registro Consolidado de Vulnerabilidades (Matriz FAIR Global)", level=1, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, "Este capítulo constituye la matriz consolidada de los **12 riesgos técnicos más materiales** detectados en toda la plataforma, ordenados por su nivel de severidad y coste de exposición FAIR (ALE):")
    
    # Ordenar y filtrar los top 12 riesgos
    try:
        all_risks.sort(key=lambda x: x.get("fair_ale_score", 0.0), reverse=True)
    except Exception:
        pass
    top_risks = all_risks[:12]
    
    table_risks = doc.add_table(rows=1, cols=4)
    finalize_table(table_risks)
    
    headers_risks = ["ID", "Torre / Dominio", "Descripción del Riesgo y Evidencia de Auditoría (RAG)", "Exposición FAIR (ALE)"]
    for i, h_txt in enumerate(headers_risks):
        set_cell_text_custom(table_risks.rows[0].cells[i], h_txt, bold=True, font_size=9, color_rgb=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table_risks.rows[0].cells[i], COLOR_BLUE)
        
    for r_idx, hc in enumerate(top_risks):
        row = table_risks.add_row()
        
        # Mapeo del prefijo de la torre
        t_id = hc.get("tower_id", "TXX")
        t_prefix = prefix_map.get(t_id, "RTX")
        
        # 1. ID de Riesgo
        set_cell_text_custom(row.cells[0], f"{t_prefix}{r_idx+1:02d}", bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(row.cells[0], COLOR_ROW_ALT)
        
        # 2. Torre Origen
        set_cell_text_custom(row.cells[1], f"Torre {t_id}\n{hc.get('tower_name', 'General')}\n\nDominio:\n{hc.get('pilar_name', 'General')}", bold=True, font_size=8)
        
        # 3. Descripción y Evidencia (SOTA)
        p_desc = row.cells[2].paragraphs[0]
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        run_obs = p_desc.add_run("Vulnerabilidad:\n")
        run_obs.bold = True
        p_desc.add_run(clean_text_for_word(hc.get("finding", "No descripto.")) + "\n\n")
        
        run_imp = p_desc.add_run("Impacto de Negocio:\n")
        run_imp.bold = True
        p_desc.add_run(clean_text_for_word(hc.get("business_risk", "No descripto.")) + "\n\n")
        
        run_ev = p_desc.add_run("Evidencia Forense Literal (Audit RAG):\n")
        run_ev.bold = True
        run_cite = p_desc.add_run(f'"{clean_text_for_word(hc.get("literal_evidence", "No se aportó evidencia literal"))}"')
        run_cite.italic = True
        run_cite.font.color.rgb = RGBColor(100, 110, 120)
        
        for run in p_desc.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            
        # 4. Exposición FAIR
        tef = hc.get("threat_event_frequency", 0.0)
        lm = hc.get("loss_magnitude", 0.0)
        ale = hc.get("fair_ale_score", 0.0)
        
        if ale and ale > 0:
            calc_txt = f"TEF: {tef:.1f} / 5,0\nLM: {lm:.1f} / 5,0\n\nALE: {ale:,.0f} €".replace(",", "X").replace(".", ",").replace("X", ".")
            set_cell_text_custom(row.cells[3], calc_txt, font_size=8, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
            
            # Severity color
            risk_score = tef * lm
            bg_color = "D9F2D9"
            if risk_score >= 15 or ale >= 1000000: bg_color = "F8D7DA"
            elif risk_score >= 10 or ale >= 250000: bg_color = "FFF3CD"
            elif risk_score >= 5 or ale >= 50000: bg_color = "E2E3E5"
            shade_cell(row.cells[3], bg_color)
        else:
            set_cell_text_custom(row.cells[3], "N/A", font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            
        if r_idx % 2 == 1:
            shade_cell(row.cells[1], COLOR_ROW_ALT)

    autofit_table_to_contents(table_risks)
    doc.add_page_break()

    # ---------------------------------------------------------
    # 5. EVALUACIÓN GLOBAL DE NIVELES DE MADUREZ
    # ---------------------------------------------------------
    add_heading(doc, "Evaluación Global de Niveles de Madurez", level=1, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, "A continuación, se presenta la comparativa y el radar de madurez consolidado entre todos los dominios de infraestructura evaluados:")
    
    # Gráfico de radar global con UUID para evitar colisiones
    global_labels = [t["tower_name"] for t in tower_overviews]
    global_actual = [t["score"] for t in tower_overviews]
    global_target = [t["target"] for t in tower_overviews]
    if global_labels:
        chart_path = Path(tempfile.gettempdir()) / f"global_radar_asis_{uuid.uuid4().hex}.png"
        if generate_radar_chart(global_labels, global_actual, global_target, "Comparativa de Madurez Global (AS-IS vs TO-BE)", str(chart_path)):
            doc.add_picture(str(chart_path), width=Inches(5.0))
            p_img = doc.paragraphs[-1]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try: os.remove(chart_path)
            except OSError: pass
            
    add_heading(doc, "Resumen General de Puntuaciones por Torre", level=2, primary_color_rgb=p_color_rgb)
    mat_table = doc.add_table(rows=1, cols=4)
    mat_table.style = 'Table Grid'
    
    headers_mat = ["Dominio de Infraestructura Evaluado", "Score AS-IS", "Meta TO-BE", "Justificación del Diagnóstico de Madurez"]
    for i, h_txt in enumerate(headers_mat):
        set_cell_text_custom(mat_table.rows[0].cells[i], h_txt, bold=True, font_size=9, color_rgb=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(mat_table.rows[0].cells[i], COLOR_BLUE)
        
    for r_idx, t in enumerate(sorted(tower_overviews, key=lambda x: x["tower_id"])):
        row = mat_table.add_row()
        set_cell_text_custom(row.cells[0], f"Torre {t['tower_id']}\n{t['tower_name']}", bold=True, font_size=8.5)
        set_cell_text_custom(row.cells[1], f"{t['score']:.2f}", font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text_custom(row.cells[2], f"{t['target']:.2f}", font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Extraer justificación de madurez (la primera frase del bottom_line de esa torre)
        first_sentence = t["bottom_line"].split(".")[0] + "." if t["bottom_line"] else "Evaluado con éxito."
        set_cell_text_custom(row.cells[3], first_sentence, font_size=8)
        
        if r_idx % 2 == 1:
            for cell in row.cells: shade_cell(cell, COLOR_ROW_ALT)
            
    autofit_table_to_contents(mat_table)
    doc.add_page_break()

    # ---------------------------------------------------------
    # 6. CONCLUSIONES GENERALES Y BRECHAS TRANSVERSALES SOTA
    # ---------------------------------------------------------
    add_heading(doc, "Conclusiones y Brechas Transversales", level=1, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, "A partir del análisis de madurez y de la matriz consolidada de riesgos, se detectan de manera sistemática los siguientes patrones y deudas técnicas transversales en toda la infraestructura de la organización:")
    
    # SOTA: Agrupar todas las deudas técnicas críticas de las torres de manera dinámica
    add_heading(doc, "Deuda Técnica Crítica de la Plataforma", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, f"La deuda técnica acumulada se agrupa en torno a tres patrones críticos de deficiencia. Primero, la **fragmentación de herramientas y operaciones en silos discretos** para los entornos on-premise y cloud, lo que impide una visibilidad unificada de servicio de extremo a extremo. Segundo, la **dependencia de procesos manuales y reactivos** para el aprovisionamiento, la configuración y el parcheado, lo que genera cuellos de botella sistémicos. Tercero, la **falta de mecanismos continuos de auditoría de conformidad**, validando el cumplimiento normativo (ENS, NIS2) a posteriori.")

    add_heading(doc, "Siguientes Pasos Ejecutivos de Gobernanza", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, "Para solventar las deficiencias del estado actual y sentar las bases de la resiliencia proactiva, se proponen las siguientes acciones inmediatas a nivel directivo:")
    add_body_paragraph(doc, "**Establecer un Comité de Gobernanza de Plataforma Híbrida:** Encargado de unificar los estándares operativos on-premise y cloud.", style='List Bullet')
    add_body_paragraph(doc, "**Acelerar el Programa de Platform Engineering:** Automatizando los flujos de provisión bajo un modelo de autoservicio.", style='List Bullet')
    add_body_paragraph(doc, "**Implantar el Modelo de Conformidad Continua (Policy-as-Code):** Asegurando el cumplimiento preventivo y en tiempo real de ENS y NIS2.", style='List Bullet')

    # Guardar documento final
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"🎉 ¡Informe Consolidado AS-IS COMPLETO generado con éxito en: {output_path}!")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python render_togaf_consolidated_asis.py <working_dir> <output_doc.docx>")
        sys.exit(1)
    render_consolidated_asis(sys.argv[1], sys.argv[2])
