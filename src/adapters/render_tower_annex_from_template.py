import os
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

"""Render the tower annex DOCX from a validated payload and template."""

from application.generate_tower_radar_chart import (
    generate_radar_chart_from_pillars,
)
from domain.schemas.annex_synthesis import AnnexPayload
from infrastructure.contract_utils import robust_load_payload
from infrastructure.docx_render_utils import (
    add_body_paragraph,
    add_bullet_list,
    add_heading_paragraph,
    add_label_value_paragraph,
    add_long_detail_table,
    apply_bullet_list_format,
    apply_paragraph_style,
    apply_table_style,
    clean_text,
    clear_cell_shading,
    clear_paragraph,
    clear_paragraph_properties,
    enable_update_fields_on_open,
    insert_field_paragraph_after_block,
    insert_paragraph_after_block,
    remove_page_break_only_paragraphs,
    render_gap_table,
    render_initiative_cards,
    render_list_at_placeholder,
    render_multi_paragraph_block,
    render_note_box,
    render_pillar_score_table,
    render_radar_chart,
    render_risks_table,
    replace_simple_placeholder,
    resolve_radar_chart_image,
)

WORD_RENDER_MODE_ENV = "AE_WORD_RENDER_MODE"


def normalize_annex_payload(payload_dict: dict) -> dict:
    """Normalize the structure of a tower annex payload dictionary.

    This function ensures that a raw data payload dictionary conforms to a
    standard structure before being used for document rendering. It creates a
    copy of the input dictionary and performs the following normalizations:
    - Ensures top-level keys ('document_meta', 'executive_summary', etc.) exist.
    - Joins list-based 'summary_body' into a single string.
    - Populates a missing 'headline' and 'key_business_impacts' with defaults.
    - Transforms a legacy 'tobe_gap' section into separate 'tobe' and 'gap'
      sections for backward compatibility.
    - Populates a missing 'domain_introduction' section using data from
      'document_meta' and 'executive_summary'.

    Args:
        payload_dict: The raw data dictionary for the annex document. It may
            have missing keys or legacy structures.

    Returns:
        A new dictionary with a standardized structure, suitable for rendering.
        The original input dictionary is not modified.
    """
    normalized = dict(payload_dict)
    meta = dict(normalized.get("document_meta") or {})
    executive_summary = dict(normalized.get("executive_summary") or {})
    sections = dict(normalized.get("sections") or {})

    summary_body = executive_summary.get("summary_body", "")
    if isinstance(summary_body, list):
        executive_summary["summary_body"] = "\n\n".join(
            part for part in summary_body if clean_text(part)
        )
    executive_summary.setdefault(
        "headline", f"Anexo ejecutivo de la torre {meta.get('tower_code', 'TX')}"
    )
    executive_summary.setdefault("key_business_impacts", [])

    if "tobe" not in sections and "tobe_gap" in sections:
        tobe_gap = dict(sections.get("tobe_gap") or {})
        sections["tobe"] = {
            "vision": tobe_gap.get("introduction", ""),
            "design_principles": tobe_gap.get("target_capabilities", []),
        }
        sections["gap"] = {
            "introduction": tobe_gap.get("introduction", ""),
            "target_capabilities": tobe_gap.get("target_capabilities", []),
            "gap_rows": tobe_gap.get("gap_rows", []),
            "closing_summary": tobe_gap.get("closing_summary", ""),
        }

    domain_introduction = dict(normalized.get("domain_introduction") or {})
    if not domain_introduction:
        domain_introduction = {
            "introduction_paragraph": executive_summary.get("summary_body", ""),
            "technological_domain": meta.get("tower_name", ""),
            "domain_objective": meta.get("tower_name", ""),
            "evaluated_capabilities": [],
            "included_components": [],
        }

    normalized["document_meta"] = meta
    normalized["executive_summary"] = executive_summary
    normalized["sections"] = sections
    normalized["domain_introduction"] = domain_introduction
    return normalized


def _replace_client_placeholders(value, client_name: str) -> Any:
    if isinstance(value, dict):
        return {
            key: _replace_client_placeholders(item, client_name)
            for key, item in value.items()
        }
    if isinstance(value, list):
        return [_replace_client_placeholders(item, client_name) for item in value]
    if isinstance(value, str):
        return value.replace("[Cliente]", client_name).replace("[cliente]", client_name)
    return value


def insert_toc_after(paragraph) -> None:
    """Insert a Table of Contents (TOC) into the document after a given paragraph.

    This function adds a new heading titled "Contents" and then inserts a Word
    field code for a standard TOC that includes heading levels 1 and 2. The new
    content is styled and formatted according to predefined styles.

    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph object after which
            the TOC section will be inserted.
    """
    heading = insert_paragraph_after_block(paragraph)
    heading.add_run("Contents")
    apply_paragraph_style(heading, "TOC Heading", "TtuloTDC")
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    toc_paragraph = insert_field_paragraph_after_block(
        heading,
        'TOC \\o "1-2" \\h \\z \\u',
        placeholder_text="Update field to build contents.",
    )
    style_existing_field_paragraph(toc_paragraph, "toc 1", "TDC1")


def rewrite_paragraph_with_style(paragraph, text: str, *styles: str) -> None:
    """Rewrite a paragraph with new text and styles.

    This function completely replaces the content and styling of a given paragraph
    object. It first clears any existing runs and paragraph properties, then
    applies the new styles, and finally adds the cleaned text. If the cleaned
    text is empty, the paragraph will remain empty but styled.

    Args:
        paragraph: The paragraph object to modify in-place.
        text: The new text content to be placed in the paragraph.
        *styles: A variable number of style names to apply to the paragraph.

    Returns:
        None.
    """
    clean_value = clean_text(text)
    clear_paragraph(paragraph)
    clear_paragraph_properties(paragraph)
    apply_paragraph_style(paragraph, *styles)
    if clean_value:
        paragraph.add_run(clean_value)


def rewrite_body_paragraph(paragraph, text: str) -> None:
    """Rewrite a paragraph with new text and apply standard body formatting.

    The function clears the paragraph's existing content and formatting, then
    populates it with the provided text. It applies a specific "NTT Body Text"
    style, falling back to the "Normal" style if the primary one is unavailable.
    It also sets specific paragraph formatting like justification, spacing, and
    font size.

    If the input `text` is empty or becomes empty after cleaning, the function
    exits without modifying the paragraph.

    Args:
        paragraph: The python-docx Paragraph object to be modified in-place.
        text: The new string content for the paragraph.
    """
    clean_value = clean_text(text)
    if not clean_value:
        return
    clear_paragraph(paragraph)
    clear_paragraph_properties(paragraph)
    if not apply_paragraph_style(
        paragraph,
        "NTT Body Text",
        "NTTBodyText",
        "Body Text",
        "Textoindependiente",
    ):
        apply_paragraph_style(paragraph, "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(clean_value)
    run.font.size = Pt(10)


def style_existing_field_paragraph(paragraph, *styles: str) -> None:
    """Clear existing formatting and apply a new set of styles to a paragraph.

    This function first removes all existing paragraph-level formatting properties.
    It then applies the given styles and sets a standard spacing of 0pt before,
    3pt after, and a line spacing of 1.0.

    Args:
        paragraph: The paragraph object to modify in-place.
        *styles: One or more style names (str) to apply to the paragraph.

    Raises:
        AttributeError: If the provided `paragraph` object does not have the
            expected `paragraph_format` attribute.
    """
    clear_paragraph_properties(paragraph)
    apply_paragraph_style(paragraph, *styles)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0


def polish_semantic_tables(doc) -> None:
    """Apply a standard style and formatting to all tables in a document.

    This function iterates through every table in the provided Word document,
    applying a consistent visual style. It sets a specific table style,
    clears cell shading, and standardizes vertical alignment, paragraph spacing,
    line spacing, text alignment, font size, and header row font color.

    Args:
        doc (docx.document.Document): The python-docx Document object to modify
            in-place.
    """
    for table in doc.tables:
        apply_table_style(
            table,
            "Tabla con cuadrícula 1 clara - Énfasis 11",
            "Tablaconcuadrcula1clara-nfasis11",
            "Smart Navy table",
            "SmartNavytable",
            "NTT Future Blue table",
            "NTTFutureBluetable",
            "Table Grid",
        )
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                clear_cell_shading(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.0
                    if row_idx == 0:
                        if paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        if row_idx == 0:
                            run.font.color.rgb = RGBColor(255, 255, 255)


def apply_semantic_annex_styles(doc, payload_dict: dict) -> None:
    """Applies semantic styles and formatting to a tower annex document in-place.

    This function iterates through all paragraphs of a python-docx Document,
    identifying titles, subtitles, and various heading levels based on their
    exact text content. It applies the corresponding Word styles to these
    elements. The function also converts text-based bullet points into
    proper list items, ensures a table of contents is present and correctly
    styled, and applies final formatting to any tables in the document.

    Args:
        doc: The python-docx Document object to be styled. This object is
            modified in-place.
        payload_dict: A dictionary containing metadata for the document under the
            'document_meta' key. Expected sub-keys include 'tower_code',
            'tower_name', and 'client_name'.

    Raises:
        KeyError: If required keys (e.g., 'document_meta', 'tower_code') are
            missing from `payload_dict`.
    """
    enable_update_fields_on_open(doc)

    heading_level_1 = {
        "Resumen ejecutivo de la torre",
        "Perfil de madurez por pilar",
        "AS-IS resumido",
        "Riesgos principales",
        "Estado objetivo y brechas",
        "Iniciativas prioritarias",
        "Conclusión",
        "Desarrollo ampliado",
        "1. AS-IS detallado",
        "2. Riesgos detallados",
        "3. Estado objetivo detallado",
        "4. Brechas detalladas",
        "5. Iniciativas priorizadas detalladas",
        "6. Cierre ampliado",
    }
    heading_level_2 = {
        "Fortaleza principal",
        "Brecha principal",
        "Cuello de botella",
        "Gráfico radial",
        "Detalle compacto por pilar",
        "Fortalezas clave",
        "Brechas clave",
        "Implicaciones operativas clave",
        "Principios / capacidades objetivo",
        "Tabla de brechas por pilar",
        "Fichas de iniciativas priorizadas",
        "Mensaje ejecutivo final",
        "Áreas prioritarias de actuación",
        "Fortalezas observadas",
        "Brechas observadas",
        "Implicaciones operativas",
        "Lectura transversal de gaps",
        "Principios de arquitectura y operación",
        "Implicaciones para el modelo operativo",
    }

    title_text = (
        f"ANEXO {payload_dict['document_meta']['tower_code']} – "
        f"{payload_dict['document_meta']['tower_name']}"
    )
    subtitle_prefix = payload_dict["document_meta"]["client_name"]
    subtitle_anchor = None

    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue

        if text == title_text:
            rewrite_paragraph_with_style(paragraph, text, "Title", "Ttulo")
            subtitle_anchor = paragraph
            continue

        if (
            text.startswith(subtitle_prefix)
            and "Fast Infrastructure Assessment" in text
        ):
            rewrite_paragraph_with_style(paragraph, text, "Subtitle", "Subttulo")
            subtitle_anchor = paragraph
            continue

        if text in heading_level_1:
            rewrite_paragraph_with_style(paragraph, text, "Heading 1", "Ttulo1")
            continue

        if text in heading_level_2:
            rewrite_paragraph_with_style(paragraph, text, "Heading 2", "Ttulo2")
            continue

        if text == "Update field to build contents.":
            style_existing_field_paragraph(paragraph, "toc 1", "TDC1")
            continue

        if text.startswith("• "):
            clear_paragraph(paragraph)
            clear_paragraph_properties(paragraph)
            apply_bullet_list_format(paragraph)
            run = paragraph.add_run(text[2:].strip())
            run.font.size = Pt(10)
            continue

        rewrite_body_paragraph(paragraph, text)

    if (
        subtitle_anchor is not None
        and "Update field to build contents."
        not in "\n".join(paragraph.text for paragraph in doc.paragraphs)
    ):
        insert_toc_after(subtitle_anchor)

    polish_semantic_tables(doc)


def _resolve_render_mode(args: list[str]) -> tuple[str, list[str]]:
    render_mode = (
        os.environ.get(WORD_RENDER_MODE_ENV, "semantic").strip().lower() or "semantic"
    )
    filtered_args: list[str] = []
    for arg in args:
        if arg == "--semantic-styles":
            render_mode = "semantic"
            continue
        if arg == "--legacy-styles":
            render_mode = "legacy"
            continue
        filtered_args.append(arg)
    return render_mode, filtered_args


def _ensure_radar_chart_path(payload_path: Path, profile: dict) -> str:
    radar_chart = clean_text(profile.get("radar_chart", ""))
    if resolve_radar_chart_image(radar_chart) is not None:
        return radar_chart

    pillars = profile.get("pillars", [])
    if not isinstance(pillars, list) or not pillars:
        return radar_chart

    generated_path = payload_path.with_name("pillar_radar_chart.generated.png")
    try:
        generate_radar_chart_from_pillars(pillars, generated_path)
    except ValueError:
        return radar_chart
    return str(generated_path)


def clean_brackets_and_consultant_notes(doc, payload: dict) -> Any:
    """Cleans a document by replacing placeholders and removing instructional content.

    This function sanitizes a document object in-place by performing a series
    of text replacements and structural modifications. It replaces template
    placeholders (e.g., '[Cliente]') with data from the payload, removes
    instructional paragraphs intended for consultants, and purges placeholder
    rows from tables.

    Args:
        doc: A `docx.Document` object to be cleaned. The object is modified
            directly.
        payload: A dictionary containing metadata and content used to populate the
            document. Expected to contain keys like 'document_meta' and
            'domain_introduction'.

    Returns:
        None. The input `doc` object is modified in-place.
    """
    meta = payload.get("document_meta", {})
    intro = payload.get("domain_introduction", {})

    client_name = meta.get("client_name", "CLIENTE")
    tower_name = meta.get("tower_name", "TORRE")
    tower_code = meta.get("tower_code", "TX")

    # Step 1: Construct the primary replacement map. This dictionary contains key-value pairs for direct, inline text substitution throughout the document.
    replacements = {
        "[Cliente]": client_name,
        "[cliente]": client_name,
        "[Nombre del Cliente]": client_name,
        "[número]": "los",
        "[número de pilares]": "varios",
        "-[Lista de torres evaluadas]": f"• {tower_name}",
        "[Lista de torres evaluadas]": tower_name,
        "T[X]": tower_code,
        "[Nombre de la Torre]": tower_name,
        "[Código]": tower_code,
        "[Nombre de la torre]": tower_name,
        "[descripción del dominio]": intro.get("domain_objective", ""),
        "descripción del dominio tecnológico": intro.get("technological_domain", ""),
        "Nivel de madurez obtenido: [Nivel de madurez de la torre]": "",
        "El resultado de [X.X] sitúa la torre en Nivel [X] – [Descripción].": "",
        "Nivel objetivo:\n[Nivel objetivo recomendado]": "",
        "[INSTRUCCIÓN PARA EL AGENTE:Hazlo extenso y bien estructurado mediante bullets]": "",
        "El análisis del estado actual del dominio Resilience & Continuity permite concluir que la infraestructura presenta un nivel de madurez [nivel de madurez], caracterizado por:": "",
        "[fortaleza identificada]": "",
        "[limitación identificada]": "",
        "[riesgo identificado]": "",
        "Pilar [Nombre del pilar]": "",
        "[Descripción del estado actual basada en las respuestas del cuestionario. Hazla extensa y estructurada por bullets]": "",
        "[capacidad objetivo]": "",
        "[línea de evolución]": "",
        "[Fecha]": meta.get("date", "2026"),
    }

    # Step 2: Sanitize the document by removing paragraphs that serve as instructional notes or contain non-content template markers. This ensures the final output is clean of authoring artifacts.
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        # Identifies and marks for deletion any paragraphs containing instructional prompts, such as "Teniendo en cuenta el resumen de...". This sanitization step removes non-content markers intended for document authors from the final output.
        if (
            "[Teniendo en cuenta el resumen de" in p.text
            or "[Teniendo en cuenta el Resumen ejecutivo" in p.text
        ):
            paragraphs_to_remove.append(p)
            if "Resumen ejecutivo del documento de contexto" in p.text:
                p.insert_paragraph_before(intro.get("introduction_paragraph", ""))
            continue

        if "[Generar un gráfico" in p.text:
            paragraphs_to_remove.append(p)
            continue

        if (
            "[capacidad tecnológica evaluada]" in p.text
            or "[infraestructura evaluada]" in p.text
            or "[plataformas tecnológicas incluidas]" in p.text
            or "[alcance específico del assessment]" in p.text
        ):
            paragraphs_to_remove.append(p)
            continue

        if (
            "[P1]" in p.text
            or "[P2]" in p.text
            or "[P3]" in p.text
            or "[P4]" in p.text
            or "[P5]" in p.text
        ):
            paragraphs_to_remove.append(p)
            continue

        if p.text.strip().startswith("| [Torre]"):
            paragraphs_to_remove.append(p)
            continue

        # Performs a comprehensive inline text substitution pass across the entire document body. Each key in the predefined replacement map is replaced with its corresponding value.
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)

    for p in paragraphs_to_remove:
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass

    # Step 3: Purge placeholder rows from all tables. Rows are identified for removal if they contain markers, typically denoted by bracket notation (e.g., "[EXAMPLE]"), which are used in the template for illustrative purposes only.
    for table in doc.tables:
        rows_to_remove = []
        for row in table.rows:
            row_has_bracket = False
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "[" in p.text and "]" in p.text:
                        # Conditionally processes table rows. A row is preserved and its date placeholder is substituted if the cell contains *only* the placeholder. Otherwise, the entire row is flagged for subsequent deletion to sanitize the output.
                        if "[Fecha]" in p.text:
                            p.text = p.text.replace("[Fecha]", replacements["[Fecha]"])
                        else:
                            row_has_bracket = True
                            break
                if row_has_bracket:
                    break
            if row_has_bracket:
                rows_to_remove.append(row)

        #
        for row in rows_to_remove:
            try:
                table._tbl.remove(row._tr)
            except Exception:
                pass


def render_extended_variant(doc, payload) -> Any:
    """Render the extended, detailed sections of the tower annex report into a document.

    Populates a document with a detailed analysis of the tower's current state
    (AS-IS), risks, target state (TO-BE), gaps, prioritized initiatives,
    and an extended conclusion. The function retrieves this information from the
    `extended_sections` key within the provided payload. If this key is
    missing or its value is empty, the function returns immediately without
    modifying the document.

    Args:
        doc (Any): The document object (e.g., a `docx.document.Document` instance)
            to which content will be added. This object is modified in place.
        payload (dict): A dictionary containing the source data for the report.
            Expects a nested structure under the 'extended_sections' key.

    Returns:
        None. The function modifies the `doc` object in place.
    """
    extended = payload.get("extended_sections") or {}
    if not extended:
        return

    doc.add_page_break()
    add_heading_paragraph(doc, "Desarrollo ampliado", level=1)
    add_body_paragraph(
        doc,
        "Esta versión desarrolla con más detalle el diagnóstico, el estado objetivo, las brechas y las iniciativas prioritarias de la torre.",
    )

    asis = extended.get("asis", {})
    add_heading_paragraph(doc, "1. AS-IS detallado", level=1)
    add_body_paragraph(doc, asis.get("executive_narrative", ""))
    add_label_value_paragraph(
        doc, "Madurez actual", asis.get("current_maturity_band", "")
    )
    add_label_value_paragraph(
        doc, "Referencia de score", asis.get("current_score_reference", "")
    )
    if asis.get("strengths"):
        add_heading_paragraph(doc, "Fortalezas observadas", level=2)
        add_bullet_list(doc, asis.get("strengths", []))
    if asis.get("gaps"):
        add_heading_paragraph(doc, "Brechas observadas", level=2)
        add_bullet_list(doc, asis.get("gaps", []))
    if asis.get("operational_impacts"):
        add_heading_paragraph(doc, "Implicaciones operativas", level=2)
        add_bullet_list(doc, asis.get("operational_impacts", []))

    risks = extended.get("risks", {})
    add_heading_paragraph(doc, "2. Riesgos detallados", level=1)
    add_body_paragraph(doc, risks.get("introduction", ""))
    for idx, item in enumerate(risks.get("risk_details", []), start=1):
        add_long_detail_table(
            doc,
            f"Riesgo {idx}. {item.get('risk', '')}",
            [
                ("Causa", item.get("cause", "")),
                ("Impacto", item.get("impact", "")),
                ("Probabilidad", item.get("probability", "")),
                ("Mitigación", item.get("mitigation_summary", "")),
                ("Pilares afectados", item.get("affected_pillars_display", "")),
            ],
        )
        add_body_paragraph(doc, "", space_after=2)
    add_body_paragraph(doc, risks.get("closing_summary", ""))

    tobe = extended.get("tobe", {})
    add_heading_paragraph(doc, "3. Estado objetivo detallado", level=1)
    add_body_paragraph(doc, tobe.get("introduction", ""))
    add_label_value_paragraph(
        doc, "Madurez objetivo recomendada", tobe.get("target_maturity_level", "")
    )
    add_label_value_paragraph(
        doc, "Referencia de score objetivo", tobe.get("target_score_reference", "")
    )
    add_body_paragraph(doc, tobe.get("target_maturity_justification", ""))
    for pillar in tobe.get("target_capabilities_by_pillar", []):
        add_heading_paragraph(doc, pillar.get("pillar", ""), level=2)
        add_bullet_list(doc, pillar.get("target_capabilities", []))
    if tobe.get("architecture_principles"):
        add_heading_paragraph(doc, "Principios de arquitectura y operación", level=2)
        add_bullet_list(doc, tobe.get("architecture_principles", []))
    if tobe.get("operating_model_implications"):
        add_heading_paragraph(doc, "Implicaciones para el modelo operativo", level=2)
        add_bullet_list(doc, tobe.get("operating_model_implications", []))

    gap = extended.get("gap", {})
    add_heading_paragraph(doc, "4. Brechas detalladas", level=1)
    add_body_paragraph(doc, gap.get("introduction", ""))
    if gap.get("cross_cutting_gap_summary"):
        add_heading_paragraph(doc, "Lectura transversal de gaps", level=2)
        add_bullet_list(doc, gap.get("cross_cutting_gap_summary", []))
    for idx, item in enumerate(gap.get("gap_items", []), start=1):
        add_long_detail_table(
            doc,
            f"Gap {idx}. {item.get('pillar', '')}",
            [
                ("Situación actual", item.get("as_is_summary", "")),
                ("Estado objetivo", item.get("target_state", "")),
                ("Brecha clave", item.get("key_gap", "")),
                ("Implicación operativa", item.get("operational_implication", "")),
            ],
        )
        add_body_paragraph(doc, "", space_after=2)

    todo = extended.get("todo", {})
    add_heading_paragraph(doc, "5. Iniciativas priorizadas detalladas", level=1)
    add_body_paragraph(doc, todo.get("introduction", ""))
    for item in todo.get("todo_items", []):
        add_long_detail_table(
            doc,
            f"{item.get('sequence', '')}. {item.get('initiative', '')}",
            [
                ("Objetivo", item.get("objective", "")),
                ("Prioridad", item.get("priority", "")),
                ("Pilares relacionados", item.get("related_pillars_display", "")),
                ("Resultado esperado", item.get("expected_outcome", "")),
                ("Dependencias", item.get("dependencies_display", "")),
            ],
        )
        add_body_paragraph(doc, "", space_after=2)
    add_body_paragraph(doc, todo.get("closing_summary", ""))

    conclusion = extended.get("conclusion", {})
    add_heading_paragraph(doc, "6. Cierre ampliado", level=1)
    add_body_paragraph(doc, conclusion.get("final_assessment", ""))
    add_body_paragraph(doc, conclusion.get("executive_message", ""))
    if conclusion.get("priority_focus_areas"):
        add_heading_paragraph(doc, "Áreas de foco prioritarias", level=2)
        add_bullet_list(doc, conclusion.get("priority_focus_areas", []))
    add_body_paragraph(doc, conclusion.get("closing_statement", ""))


def main(argv: list[str] | None = None) -> None:
    """Renders a Tower Annex DOCX document from a template and a JSON payload.

    This function serves as the main entry point for the script. It orchestrates the
    entire document generation process by parsing command-line arguments, loading
    and validating the JSON data payload, and then systematically populating a
    .docx template with the provided data. The process involves replacing simple
    text placeholders, rendering complex data structures like lists, tables,
    and charts, and applying specific formatting styles before saving the final
    document to the specified output path.

    Args:
        argv: A list of command-line arguments. If None, `sys.argv` is
            used. The script expects three positional arguments and an
            optional flag: `<payload_json> <template_docx> <output_docx>
            [--semantic-styles|--legacy-styles]`.

    Returns:
        None. The function writes the generated document to a file as a side
        effect.

    Raises:
        SystemExit: If the number of command-line arguments is incorrect.
        FileNotFoundError: If the specified payload JSON or template DOCX file
            does not exist.
        ValueError: If the payload JSON is malformed or fails data validation.
    """
    raw_args = list(argv if argv is not None else sys.argv)
    render_mode, parsed_args = _resolve_render_mode(raw_args[1:])
    if len(parsed_args) != 3:
        raise SystemExit(
            "Uso: python -m scripts.render_tower_annex_from_template <payload_json> <template_docx> <output_docx> [--semantic-styles|--legacy-styles]"
        )

    payload_path = Path(parsed_args[0]).resolve()
    template_path = Path(parsed_args[1]).resolve()
    output_path = Path(parsed_args[2]).resolve()

    # Instantiates and validates the contract data model. This critical step ensures the integrity and correctness of input data prior to commencing the document rendering process.
    payload = robust_load_payload(
        payload_path,
        AnnexPayload,
        "Annex",
        mode="strict",
    )
    payload_dict = normalize_annex_payload(payload.model_dump(by_alias=True))
    payload_dict = _replace_client_placeholders(
        payload_dict,
        str(payload_dict.get("document_meta", {}).get("client_name", "CLIENTE")),
    )

    doc = Document(str(template_path))
    remove_page_break_only_paragraphs(doc)
    clean_brackets_and_consultant_notes(doc, payload_dict)

    meta = payload_dict["document_meta"]
    exec_summary = payload_dict["executive_summary"]
    profile = payload_dict["pillar_score_profile"]
    profile["radar_chart"] = _ensure_radar_chart_path(payload_path, profile)
    sections = payload_dict["sections"]
    variant = clean_text(meta.get("report_variant", "short")).lower()
    subtitle_suffix = " · Versión Extendida" if variant == "long" else ""

    replace_simple_placeholder(
        doc,
        "{{ANNEX_TITLE}}",
        f"ANEXO {meta['tower_code']} – {meta['tower_name']}",
        font_size=15,
        bold=True,
    )
    replace_simple_placeholder(doc, "{{TOWER_CODE}}", meta["tower_code"], font_size=15)
    replace_simple_placeholder(doc, "{{TOWER_NAME}}", meta["tower_name"], font_size=15)
    replace_simple_placeholder(
        doc,
        "{{ANNEX_SUBTITLE}}",
        f"{meta['client_name']} · Fast Infrastructure Assessment{subtitle_suffix}",
        font_size=11,
    )
    replace_simple_placeholder(
        doc, "{{CLIENT_NAME}}", meta["client_name"], font_size=11
    )

    replace_simple_placeholder(
        doc,
        "{{GLOBAL_SCORE}}",
        exec_summary.get("global_score", ""),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=13,
    )
    replace_simple_placeholder(
        doc,
        "{{GLOBAL_BAND}}",
        exec_summary.get("global_band", ""),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=13,
    )
    replace_simple_placeholder(
        doc,
        "{{TARGET_MATURITY}}",
        exec_summary.get("target_maturity", ""),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=13,
    )

    render_multi_paragraph_block(
        doc, "{{EXEC_SUMMARY_BODY}}", exec_summary.get("summary_body", "")
    )
    replace_simple_placeholder(
        doc,
        "{{MSG_STRENGTH_VALUE}}",
        exec_summary.get("message_strength", ""),
        font_size=10.5,
    )
    replace_simple_placeholder(
        doc, "{{MSG_GAP_VALUE}}", exec_summary.get("message_gap", ""), font_size=10.5
    )
    replace_simple_placeholder(
        doc,
        "{{MSG_BOTTLENECK_VALUE}}",
        exec_summary.get("message_bottleneck", ""),
        font_size=10.5,
    )

    replace_simple_placeholder(
        doc,
        "{{PILLAR_PROFILE_INTRO}}",
        profile.get("profile_intro", ""),
        font_size=10.5,
    )
    render_note_box(
        doc, "{{SCORING_METHOD_NOTE}}", profile.get("scoring_method_note", "")
    )
    render_radar_chart(doc, "{{RADAR_CHART_BLOCK}}", profile.get("radar_chart", ""))
    render_pillar_score_table(doc, "{{PILLAR_SCORE_TABLE}}", profile.get("pillars", []))

    replace_simple_placeholder(
        doc,
        "{{ASIS_NARRATIVE}}",
        sections.get("asis", {}).get("narrative", ""),
        font_size=10.5,
    )
    render_list_at_placeholder(
        doc, "{{ASIS_STRENGTHS_LIST}}", sections.get("asis", {}).get("strengths", [])
    )
    render_list_at_placeholder(
        doc, "{{ASIS_GAPS_LIST}}", sections.get("asis", {}).get("gaps", [])
    )
    render_list_at_placeholder(
        doc,
        "{{ASIS_OPERATIONAL_IMPACTS_LIST}}",
        sections.get("asis", {}).get("operational_impacts", []),
    )

    replace_simple_placeholder(
        doc,
        "{{RISKS_INTRO}}",
        sections.get("risks", {}).get("introduction", ""),
        font_size=10.5,
    )
    render_risks_table(
        doc, "{{RISKS_TABLE}}", sections.get("risks", {}).get("risks", [])
    )
    replace_simple_placeholder(
        doc,
        "{{RISKS_CLOSING}}",
        sections.get("risks", {}).get("closing_summary", ""),
        font_size=10.5,
    )

    replace_simple_placeholder(
        doc,
        "{{TOBE_INTRO}}",
        sections.get("tobe", {}).get("vision", "")
        or sections.get("gap", {}).get("introduction", ""),
        font_size=10.5,
    )
    render_list_at_placeholder(
        doc,
        "{{TARGET_CAPABILITIES_LIST}}",
        sections.get("gap", {}).get("target_capabilities", [])
        or sections.get("tobe", {}).get("design_principles", []),
    )
    render_gap_table(doc, "{{GAP_TABLE}}", sections.get("gap", {}).get("gap_rows", []))

    replace_simple_placeholder(
        doc,
        "{{TODO_INTRO}}",
        sections.get("todo", {}).get("introduction", ""),
        font_size=10.5,
    )
    render_initiative_cards(
        doc,
        "{{PRIORITY_INITIATIVES_CARDS}}",
        sections.get("todo", {}).get("priority_initiatives", []),
    )

    replace_simple_placeholder(
        doc,
        "{{FINAL_ASSESSMENT}}",
        sections.get("conclusion", {}).get("final_assessment", ""),
        font_size=10.5,
    )
    replace_simple_placeholder(
        doc,
        "{{EXECUTIVE_MESSAGE}}",
        sections.get("conclusion", {}).get("executive_message", ""),
        font_size=10.5,
    )
    render_list_at_placeholder(
        doc,
        "{{PRIORITY_AREAS_LIST}}",
        sections.get("conclusion", {}).get("priority_focus_areas", []),
    )
    replace_simple_placeholder(
        doc,
        "{{CLOSING_STATEMENT}}",
        sections.get("conclusion", {}).get("closing_statement", ""),
        font_size=10.5,
    )

    if variant == "long":
        render_extended_variant(doc, payload_dict)

    if render_mode == "semantic":
        apply_semantic_annex_styles(doc, payload_dict)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print("Documento renderizado en:", output_path)


if __name__ == "__main__":
    main()
