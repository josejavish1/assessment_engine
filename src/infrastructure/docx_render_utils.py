from __future__ import annotations

from typing import Any

"""Shared helpers for DOCX-based renderers."""


import base64
import tempfile
from pathlib import Path

from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from infrastructure.text_utils import clean_text_for_word

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def find_style(doc, *candidates: Any) -> Any:
    """Find the first style in a document matching a given name or ID.

    Searches a document's styles collection for the first style object that
    matches any of the provided candidate identifiers. The search compares each
    candidate against both the style's `name` and `style_id` attributes.

    The comparison is performed case-insensitively after stripping leading and
    trailing whitespace from both the candidate and the style attributes. Any
    candidate that evaluates to False (e.g., `None` or an empty string) is
    ignored.

    Args:
        doc: An object representing a document. This object must possess an
            iterable `styles` attribute, where each element is an object with
            readable `name` and `style_id` attributes.
        *candidates: A variable-length argument list of style names or style
            IDs to search for. These values will be converted to strings.

    Returns:
        The first matching style object found. Returns `None` if no candidate
        matches any style, if the document lacks a `styles` attribute, or if
        no valid candidates are provided.

    Raises:
        TypeError: If `doc.styles` exists but is not iterable.
    """
    styles = getattr(doc, "styles", None)
    if styles is None:
        return None

    normalized_candidates = [
        str(candidate).strip().lower() for candidate in candidates if candidate
    ]
    for wanted in normalized_candidates:
        for style in styles:
            style_id = str(getattr(style, "style_id", "") or "").strip().lower()
            style_name = str(getattr(style, "name", "") or "").strip().lower()
            if style_id == wanted or style_name == wanted:
                return style
    return None


def apply_paragraph_style(paragraph: Any, *candidates: Any) -> Any:
    r"""{'docstring': "Applies the first available style from a list of candidate names to a paragraph.\n\n    This function searches for each style name in `candidates` within the\n    paragraph's parent document. If a valid style is found, it is assigned to\n    the paragraph object's `style` attribute.\n\n    Args:\n        paragraph (Any): The paragraph object to which the style will be\n            applied. This object must possess a `part.document` attribute\n            referencing its parent document and a settable `style` attribute.\n        *candidates (str): A variable-length sequence of style names to\n            attempt to apply, in order of preference.\n\n    Returns:\n        bool: True if a style from the candidates was found and applied,\n            False otherwise.\n\n    Raises:\n        AttributeError: If the `paragraph` object does not possess the required\n            `part.document` attribute."}."""
    style = find_style(paragraph.part.document, *candidates)
    if style is None:
        return False
    paragraph.style = style
    return True


def apply_table_style(table: Any, *candidates: Any) -> Any:
    r"""{'docstring': "Apply the first valid style from a sequence of candidates to a table.\n\n    Iterates through a sequence of style name candidates in order of preference.\n    The first candidate that corresponds to a defined style in the parent\n    document is applied to the table. If no candidate style is found, the\n    table's style remains unmodified.\n\n    Args:\n        table (docx.table.Table): The `python-docx` Table object to be styled.\n        *candidates (str): A variable-length sequence of style names to search\n            for, in order of preference.\n\n    Returns:\n        bool: True if a style was successfully found and applied, False otherwise.\n\n    Raises:\n        AttributeError: If `table` is not a valid `python-docx` Table\n            object and lacks the expected `part` or `style` attributes."}."""
    style = find_style(table.part.document, *candidates)
    if style is None:
        return False
    table.style = style.name
    return True


def get_style_numbering(doc, *candidates: Any) -> Any:
    """Traverse a style's inheritance hierarchy to find its numbering properties.

    Searches for a style from a list of candidate names or IDs and then
    traverses its 'basedOn' inheritance chain. The traversal identifies the
    first style in the hierarchy (including the starting style) that contains
    a numbering properties (`numPr`) definition. The search is protected against
    circular style dependencies.

    Args:
        doc (docx.document.Document): The `python-docx` Document object containing
            the style definitions.
        *candidates (str): One or more style names or style IDs to search for.
            The first valid style found is used as the starting point for the
            hierarchy traversal.

    Returns:
        Tuple[Optional[str], Optional[str]]: A 2-tuple `(num_id, ilvl)` where
        `num_id` is the numbering definition ID and `ilvl` is the indentation
        level. If `ilvl` is not explicitly defined in the style, it defaults
        to "0". Returns `(None, None)` if no numbering definition is found in
        the style's inheritance hierarchy.
    """
    style = find_style(doc, *candidates)
    if style is None:
        return None, None

    current = style.element
    visited = set()
    while current is not None:
        style_id = current.get(qn("w:styleId"))
        if style_id in visited:
            break
        visited.add(style_id)

        num_id = current.find("./w:pPr/w:numPr/w:numId", WORD_NS)
        ilvl = current.find("./w:pPr/w:numPr/w:ilvl", WORD_NS)
        if num_id is not None:
            return (
                num_id.get(qn("w:val")),
                ilvl.get(qn("w:val")) if ilvl is not None else "0",
            )

        based_on = current.find("./w:basedOn", WORD_NS)
        if based_on is None:
            break
        parent_style_id = based_on.get(qn("w:val"))
        current = find_style(doc, parent_style_id)
        current = current.element if current is not None else None

    return None, None


def set_paragraph_numbering(paragraph: Any, num_id: Any, ilvl: Any = 0) -> Any:
    r"""{'docstring': "Applies a specific numbering definition to a paragraph by modifying its OOXML.\n\nThis function directly manipulates the underlying Office Open XML (OOXML)\nof a paragraph to associate it with a numbering definition, effectively\nrendering the paragraph as an item in a numbered or bulleted list.\n\nAny pre-existing `w:numPr` (numbering properties) element within the\nparagraph's properties is removed before the new numbering is applied.\n\nArgs:\n    paragraph (docx.paragraph.Paragraph): The paragraph object to modify.\n    num_id (int): The identifier of the numbering definition to apply.\n        The function will not modify the paragraph if this value is `None`.\n    ilvl (int): The 0-indexed indentation level of the list item.\n        Defaults to 0.\n\nReturns:\n    bool: `True` if the numbering was successfully applied, or `False` if\n        `num_id` was `None`.\n\nRaises:\n    AttributeError: If the `paragraph` object does not have the expected\n        internal structure (e.g., a `_p` attribute containing the\n        paragraph's XML element)."}."""
    if num_id is None:
        return False

    pPr = paragraph._p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag.split("}")[-1] == "numPr":
            pPr.remove(child)

    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numPr.append(ilvl_el)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numPr.append(num_id_el)
    pPr.append(numPr)
    return True


def apply_bullet_list_format(paragraph: Any) -> Any:
    r"""{'docstring': "Reset and apply a standardized bullet list format to a Paragraph object.\n\n    This function modifies the Paragraph object in-place. It begins by removing any\n    pre-existing list-related XML properties (numbering and indentation) to\n    ensure a clean state. It then applies the first available style from a\n    predefined sequence of known bullet style names. Finally, it sets explicit\n    numbering, spacing, and left-alignment properties to enforce a consistent\n    visual representation.\n\n    Args:\n        paragraph (docx.text.paragraph.Paragraph): The Paragraph instance to format.\n\n    Raises:\n        AttributeError: If the input `paragraph` object does not conform to the\n            expected internal structure of a python-docx Paragraph.\n        ValueError: If none of the predefined bullet list styles can be found in\n            the document's style definitions."}."""
    pPr = paragraph._p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag.split("}")[-1] in {"numPr", "ind"}:
            pPr.remove(child)

    apply_paragraph_style(
        paragraph, "Bullet", "NTTFlushBullet1", "List Paragraph", "Prrafodelista"
    )
    num_id, ilvl = get_style_numbering(
        paragraph.part.document,
        "Bullet",
        "NTTFlushBullet1",
        "List Paragraph",
        "Prrafodelista",
    )
    set_paragraph_numbering(paragraph, num_id, ilvl or 0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def enable_update_fields_on_open(doc) -> Any:
    r"""{'docstring': 'Enable the \'update fields on open\' setting within a .docx file.\n\nModifies the provided `python-docx` Document object in-place by directly\nmanipulating the `word/settings.xml` part of the underlying OOXML. This\nfunction ensures a `w:updateFields` element exists and sets its `w:val`\nattribute to "true". This instructs client applications, such as Microsoft\nWord, to automatically update all fields (e.g., Table of Contents, page\nnumbers) when the document is next opened.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` Document object to be\n        modified.\n\nRaises:\n    AttributeError: If the `doc` object does not have a `settings` attribute.'}."""
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def clean_text(value: Any) -> Any:
    """Clean a value for safe rendering as text within a .docx file."""
    return clean_text_for_word(value)


def clean_list(value: Any) -> Any:
    """Cleans and normalizes an input value into a list of non-empty strings.

    This function processes an input value to produce a list of non-empty
    strings. If the input is `None`, it returns an empty list. If the input
    is a list, each element is individually cleaned using the `clean_text`
    helper, and any resulting empty strings are filtered out. A non-list,
    non-None input is treated as a single value to be cleaned and returned
    within a single-element list if the result is not an empty string.

    Args:
        value (Any): The input value to process. Accepts None, a list of
            items, or a single scalar item.

    Returns:
        List[str]: A list containing cleaned, non-empty strings. An empty list
            is returned if the input is None or all elements are empty after
            cleaning.
    """
    if value is None:
        return []
    if isinstance(value, list):
        return [clean_text(item) for item in value if clean_text(item)]
    text = clean_text(value)
    return [text] if text else []


def clean_paragraph_list(value: Any) -> Any:
    r"""{'docstring': 'Clean and split a value into a list of paragraphs.\n\n    This function standardizes an input, which can be None, a list of strings,\n    or a single string, into a list of non-empty paragraph strings.\n\n    If the input is a string, it is cleaned and then split by double newlines\n    (r"\\n\\n"). If no double newlines are present, the entire cleaned string is\n    treated as a single paragraph. If the input is a list, each item is\n    individually cleaned.\n\n    Empty strings resulting from cleaning or splitting are discarded. An empty\n    list is returned if the input is `None` or contains no processable content.\n\n    Args:\n        value (Any): The input to process. Expected types are `str`, `list[str]`,\n            or `None`.\n\n    Returns:\n        list[str]: A list of cleaned, non-empty paragraph strings.\n\n    Raises:\n        TypeError: If the input value, or an item within a list, cannot be\n            coerced to a string by the underlying `clean_text` utility.'}."""
    if value is None:
        return []
    if isinstance(value, list):
        return [clean_text(item) for item in value if clean_text(item)]
    text = clean_text(value)
    if not text:
        return []
    paragraphs = [part.strip() for part in text.split("\n\n") if part.strip()]
    return paragraphs or [text]


def resolve_radar_chart_image(image_path: Any) -> Any:
    """Resolve an image path or data URI to a local filesystem `pathlib.Path`.

    This function validates and resolves an image source to a concrete filesystem
    path. It handles two types of input for `image_path`:

    1.  A standard filesystem path: The function verifies that the path exists
        and points to a regular file.
    2.  A base64-encoded data URI (e.g., "data:image/png;base64,..."): The
        function decodes the base64 data and writes it to a new temporary
        file. The caller is responsible for the cleanup of this temporary file.

    Args:
        image_path: The image source, specified as a filesystem path or a
            base64-encoded data URI.

    Returns:
        A `pathlib.Path` object for the local image file, or `None` if the
        input is empty, the path does not exist, the data URI is malformed,
        or any other processing error occurs.
    """
    image_path = clean_text(image_path)
    if not image_path:
        return None
    if image_path.startswith("data:image/") and ";base64," in image_path:
        try:
            _, encoded = image_path.split(";base64,", 1)
            raw = base64.b64decode(encoded)
            if not raw or len(raw) < 10:
                return None
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            tmp.write(raw)
            tmp.close()
            return Path(tmp.name)
        except Exception:
            return None
    path = Path(image_path)
    if path.exists() and path.is_file():
        return path
    return None


def iter_paragraphs(parent: Any) -> Any:
    """Yield all Paragraph objects from a container, including those within tables.

    Traverses and yields all `docx.text.paragraph.Paragraph` objects from a
    parent container. The function first yields paragraphs directly attached to
    the parent, then iterates through all tables to yield paragraphs from each
    cell. The traversal order is deterministic.

    Args:
        parent (Any): A python-docx object that exposes `.paragraphs` and `.tables`
            attributes, such as a `docx.document.Document` or `docx.table._Cell`.

    Yields:
        docx.text.paragraph.Paragraph: The next paragraph object found during
            the traversal.

    Raises:
        AttributeError: If the `parent` object lacks `.paragraphs` or `.tables`
            attributes.
    """
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def find_first_paragraph(doc, placeholder: Any) -> Any:
    """Finds the first paragraph in a document containing a placeholder string.

    Args:
        doc (docx.document.Document): The `python-docx` document object to search.
        placeholder (str): The text substring to find within the paragraphs.

    Returns:
        Optional[docx.text.paragraph.Paragraph]: The first paragraph object
            containing the placeholder, or `None` if it is not found.

    Raises:
        TypeError: If `placeholder` is not a string.
    """
    for paragraph in iter_paragraphs(doc):
        if placeholder in paragraph.text:
            return paragraph
    return None


def clear_paragraph(paragraph: Any) -> Any:
    r"""{'docstring': 'Clears all content from a paragraph while preserving its formatting.\n\nModifies the paragraph object in-place by removing all child elements\nrepresenting content (e.g., text runs, images) from its underlying XML.\nThe paragraph properties element (`w:pPr`) is explicitly preserved, which\nretains paragraph-level formatting such as style, indentation, and spacing.\n\nArgs:\n    paragraph (docx.text.paragraph.Paragraph): The `python-docx` paragraph\n        object to clear.\n\nReturns:\n    None. The paragraph is modified in-place.\n\nRaises:\n    AttributeError: If `paragraph` lacks the internal `_element` attribute\n        required for XML manipulation.'}."""
    element = paragraph._element
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def clear_paragraph_properties(paragraph: Any) -> Any:
    """Removes all direct formatting properties from a paragraph.

    This function directly manipulates the underlying OOXML `w:pPr` (Paragraph
    Properties) element by removing all of its child elements. This action
    effectively resets any locally applied formatting such as spacing,
    indentation, and borders, causing the paragraph to inherit these
    properties exclusively from its designated style.

    Args:
        paragraph (docx.paragraph.Paragraph): The paragraph object to be modified
            in-place.

    Raises:
        AttributeError: If the provided object does not conform to the expected
            `python-docx` Paragraph internal structure.
    """
    pPr = paragraph._p.get_or_add_pPr()
    for child in list(pPr):
        pPr.remove(child)


def remove_paragraph(paragraph: Any) -> Any:
    """Removes a paragraph element from its parent container in a docx document.

    Directly manipulates the underlying lxml element of a paragraph to detach
    it from the document's XML tree. This is a low-level operation intended
    for scenarios where the standard python-docx API lacks a direct method
    for paragraph deletion.

    Args:
        paragraph (docx.paragraph.Paragraph): The paragraph object to remove.

    Returns:
        None

    Raises:
        AttributeError: If the `paragraph` object lacks the internal `_element`
            attribute, indicating an invalid or unexpected object type.
    """
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_after_block(block: Any) -> Any:
    r"""{'docstring': 'Insert a new, empty paragraph after a specified block-level element.\n\nManipulates the underlying Office Open XML (OXML) to insert a new paragraph\nelement (`<w:p>`) immediately after the XML element of a specified `python-docx`\n`Paragraph` or `Table` object.\n\nArgs:\n    block (Paragraph | Table): The block-level object from `python-docx` after\n        which the new paragraph will be inserted.\n\nReturns:\n    Paragraph: The newly created and inserted `Paragraph` object.\n\nRaises:\n    TypeError: If `block` is not an instance of `Paragraph` or `Table`.'}."""
    if isinstance(block, Paragraph):
        anchor = block._p
        parent = block._parent
    elif isinstance(block, Table):
        anchor = block._tbl
        parent = block._parent
    else:
        raise TypeError("Unsupported block type")

    new_p = OxmlElement("w:p")
    anchor.addnext(new_p)
    return Paragraph(new_p, parent)


def insert_field_paragraph_after_block(
    block: Any, field_code: Any, placeholder_text: Any = ""
) -> Any:
    r"""["Inserts a new paragraph containing a complex Word field after a block element.\n\nDirectly manipulates the underlying Office Open XML (OXML) to construct\nand insert a complex field, such as a MERGEFIELD. A new paragraph is\ncreated after the specified block element. Within this paragraph, a run\nis populated with the required OXML elements: a 'begin' field character\n(`w:fldChar`), the field instructions (`w:instrText`), a 'separate'\ncharacter, the optional placeholder text (`w:t`), and an 'end' character.\n\nThe 'begin' field character is marked as 'dirty' to prompt a field\nupdate when the document is opened in a compatible application.\n\nArgs:\n    block (Union[docx.text.paragraph.Paragraph, docx.table.Table]): The\n        block-level element (paragraph or table) after which the new\n        field paragraph will be inserted.\n    field_code (str): The instruction string for the Word field (e.g.,\n        'MERGEFIELD my_field_name').\n    placeholder_text (str): The static text to display as the field's\n        result before it is updated. Defaults to an empty string.\n\nReturns:\n    docx.text.paragraph.Paragraph: The newly created paragraph object\n        containing the complex field.\n\nRaises:\n    TypeError: If the `block` element is not a supported type for\n        insertion by the underlying helper functions."]."""
    paragraph = insert_paragraph_after_block(block)
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    run._r.append(instr_text)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_sep)

    if placeholder_text:
        text = OxmlElement("w:t")
        text.text = placeholder_text
        run._r.append(text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)
    return paragraph


def add_run(paragraph: Any, text: Any, bold: Any = False, font_size: Any = 10.5) -> Any:
    """Add a text run to a paragraph with specified bolding and font size."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    return run


def apply_body_format(
    paragraph: Any, justify: Any = True, space_after: Any = 6, font_size: Any = 10.5
) -> Any:
    r"""{'docstring': 'Apply standard body formatting to a `docx.text.paragraph.Paragraph` object.\n\nModifies the Paragraph object in-place. This function sets text alignment,\nvertical spacing, line spacing, and a uniform font size for all\nconstituent text runs. Specifically, it sets spacing-before to 0 points,\nline spacing to 1.05, and alignment to justified or left-aligned based\non the `justify` argument.\n\nArgs:\n    paragraph (docx.text.paragraph.Paragraph): The `python-docx` Paragraph\n        object to be formatted.\n    justify (bool): If True, justifies the paragraph text. Otherwise,\n        left-aligns the text. Defaults to True.\n    space_after (Union[int, float]): The space in points to add after the\n        paragraph. Defaults to 6.\n    font_size (Union[int, float]): The font size in points to apply to all\n        text runs in the paragraph. Defaults to 10.5.\n\nReturns:\n    None.\n\nRaises:\n    AttributeError: If the provided `paragraph` object does not have the\n        expected attributes of a `docx.text.paragraph.Paragraph` instance.\n    TypeError: If `space_after` or `font_size` are not numeric types\n        compatible with the `docx.shared.Pt` constructor.'}."""
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.05
    for run in paragraph.runs:
        run.font.size = Pt(font_size)


def replace_simple_placeholder(
    doc, placeholder, value, align=None, font_size=None, bold=None
):
    """Replaces the entire content of paragraphs containing a specific placeholder.

    Searches all paragraphs in a `python-docx` document for a given placeholder
    string. For each matching paragraph, the placeholder is substituted with the
    provided value, and this new string replaces the paragraph's entire original
    content. The replacement is performed by clearing the paragraph and inserting
    the new text into a single, new run.

    Optional formatting for alignment, font size, and bolding can be applied.
    Paragraph spacing is also standardized. This function modifies the `doc`
    object in-place.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to modify.
        placeholder (str): The placeholder string to find within a paragraph's text.
        value (str): The string to substitute for the placeholder. The original
            paragraph text with this substitution becomes the new content.
        align (Optional[WD_ALIGN_PARAGRAPH]): The alignment to apply to the
            modified paragraph. If None, the alignment is not changed.
        font_size (Optional[Union[int, float]]): The font size in points to apply
            to the new text run. If None, the style's default is used.
        bold (Optional[bool]): Specifies whether to apply bold formatting to the
            new text run. If None, the style's default is used.

    Returns:
        None

    Raises:
        AttributeError: If `doc` is not a valid Document object with iterable
            paragraphs.
        TypeError: If `font_size` is provided and is not a numeric type (int or
            float).
    """
    for paragraph in iter_paragraphs(doc):
        if placeholder in paragraph.text:
            text = paragraph.text.replace(placeholder, clean_text(value))
            clear_paragraph(paragraph)
            run = paragraph.add_run(text)
            if font_size:
                run.font.size = Pt(font_size)
            if bold is not None:
                run.bold = bold
            if align:
                paragraph.alignment = align
            paragraph.paragraph_format.space_before = Pt(0)
            if placeholder != "{{SCORING_METHOD_NOTE}}":
                paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.05


def render_multi_paragraph_block(doc, placeholder: Any, texts: Any) -> Any:
    r"""{'docstring': "Replaces a placeholder paragraph with a block of new paragraphs.\n\n    Finds a paragraph containing the specified placeholder text and substitutes it\n    with a series of new paragraphs, one for each string in the `texts` list. The\n    first text item populates the original placeholder's paragraph, and all\n    subsequent items are inserted into new paragraphs immediately following it.\n\n    If the `texts` list is empty after an internal cleaning step, the original\n    placeholder paragraph is removed from the document. All created or modified\n    paragraphs are formatted with a consistent body style. This function modifies\n    the document object in-place.\n\n    Args:\n        doc (docx.document.Document): The `python-docx` Document object to modify.\n        placeholder (str): The text used to find the paragraph that will be\n            replaced.\n        texts (List[str]): A list of strings, where each string is rendered into\n            its own paragraph.\n\n    Returns:\n        None."}."""
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return
    texts = clean_paragraph_list(texts)
    if not texts:
        remove_paragraph(paragraph)
        return

    clear_paragraph(paragraph)
    add_run(paragraph, texts[0], font_size=10.5)
    apply_body_format(paragraph, justify=True, space_after=6)
    anchor = paragraph

    for text in texts[1:]:
        new_paragraph = insert_paragraph_after_block(anchor)
        add_run(new_paragraph, text, font_size=10.5)
        apply_body_format(new_paragraph, justify=True, space_after=6)
        anchor = new_paragraph


def strip_numbering_and_indents(paragraph: Any) -> Any:
    """Modifies a docx paragraph in-place to remove list numbering and reset formatting.

    This function directly manipulates the paragraph's underlying Open XML element
    (`<w:pPr>`) to remove list-related formatting. It finds and removes the
    `<w:numPr>` (numbering properties) and `<w:ind>` (indentation) child elements.

    After removing the list-specific tags, it resets other paragraph format
    settings to a consistent default: indents (left and first-line) are set
    to 0 points, spacing is set to 0 points before and 6 points after, line
    spacing is set to 1.05, and alignment is set to left.

    Args:
        paragraph (docx.paragraph.Paragraph): The `python-docx` Paragraph object
            to modify.

    Returns:
        None. The paragraph object is modified in-place.

    Raises:
        AttributeError: If the provided object does not conform to the expected
            `python-docx` Paragraph object structure (e.g., lacks the `_p`
            attribute).
    """
    pPr = paragraph._p.get_or_add_pPr()
    for child in list(pPr):
        tag = child.tag.split("}")[-1]
        if tag in {"numPr", "ind"}:
            pPr.remove(child)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def render_list_at_placeholder(doc, placeholder: Any, items: Any) -> Any:
    r"""{'docstring': 'Replaces a placeholder paragraph in a DOCX document with a bulleted list.\n\n    Finds the first paragraph containing the specified placeholder text and\n    replaces it with a bulleted list generated from the provided items. The\n    document is modified in-place.\n\n    If the `items` iterable is empty or contains only empty/whitespace strings\n    after filtering, the entire placeholder paragraph is removed from the\n    document. Otherwise, the placeholder paragraph is cleared and repurposed\n    as the first list item. Subsequent items are added in new paragraphs below\n    it. All resulting list item paragraphs are styled as bullet points with a\n    10-point font size, and any pre-existing numbering or indentation is\n    removed. The vertical alignment of the parent container (e.g., a table\n    cell) is set to top-aligned.\n\n    Args:\n        doc (docx.document.Document): The `python-docx` Document object to modify.\n        placeholder (str): The text content to find. The first paragraph\n            containing this text will be targeted for replacement.\n        items (Iterable[str]): A sequence of strings to render as the bulleted\n            list. Empty or whitespace-only strings are filtered out.\n\n    Returns:\n        None. The document is modified in-place.'}."""
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return
    items = [clean_text(x) for x in items if clean_text(x)]
    if not items:
        remove_paragraph(paragraph)
        return

    clear_paragraph(paragraph)
    strip_numbering_and_indents(paragraph)
    if hasattr(paragraph._parent, "vertical_alignment"):
        paragraph._parent.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    anchor = paragraph

    for idx, item in enumerate(items):
        current = anchor if idx == 0 else insert_paragraph_after_block(anchor)
        clear_paragraph(current)
        strip_numbering_and_indents(current)
        if hasattr(current._parent, "vertical_alignment"):
            current._parent.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        apply_bullet_list_format(current)
        content = current.add_run(item)
        content.font.size = Pt(10)
        anchor = current


def move_table_after_paragraph(paragraph: Any, table: Any) -> Any:
    """Repositions a table to appear immediately after a specified paragraph.

    This function operates directly on the underlying `lxml` element tree of the
    Office Open XML document. It accesses the private `_p` (paragraph) and `_tbl`
    (table) attributes to perform an in-place modification of the document
    structure, effectively circumventing the standard `python-docx` public API.

    Args:
        paragraph (docx.paragraph.Paragraph): The paragraph object that will
            immediately precede the table after the move.
        table (docx.table.Table): The table object to be moved.

    Raises:
        AttributeError: If the provided objects do not have the expected
            private `_p` or `_tbl` attributes, indicating they are not valid
            `python-docx` Paragraph or Table objects.
    """
    table_element = table._tbl
    paragraph._p.addnext(table_element)


def shade_cell(cell: Any, fill: Any) -> Any:
    r"""{'docstring': 'Apply a background color shade to a python-docx table cell.\n\nThis function directly manipulates the cell\'s underlying Open XML (OXML)\nrepresentation. It accesses the table cell properties (`tcPr`) and appends\na `w:shd` (shading) element configured with a `w:fill` attribute set to\nthe specified color. The cell object is modified in-place.\n\nArgs:\n    cell (docx.table._Cell): The python-docx table cell object to modify.\n    fill (str): The hex RGB color string for the background (e.g., "AAAAAA"),\n        or "auto" to use the default background.\n\nReturns:\n    None. The cell object is modified in-place.\n\nRaises:\n    AttributeError: If the `cell` object lacks the expected internal OXML\n        structure, such as the `_tc` attribute.'}."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def clear_cell_shading(cell: Any) -> Any:
    r"""{'docstring': "Remove shading from a table cell.\n\n    Modifies the cell's underlying XML properties in-place by iterating through\n    the table cell properties (`<w:tcPr>`) and removing the shading element\n    (`<w:shd>`), if present. This effectively clears any background color from\n    the cell.\n\n    Args:\n        cell (docx.table._Cell): The python-docx table cell object to modify.\n\n    Returns:\n        None:\n\n    Raises:\n        AttributeError: If the `cell` object does not conform to the expected\n            internal XML structure (e.g., lacks a `_tc` attribute with a\n            `get_or_add_tcPr` method)."}."""
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.split("}")[-1] == "shd":
            tcPr.remove(child)


def set_cell_text(
    cell,
    text,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    font_size=10.5,
    vertical=WD_ALIGN_VERTICAL.CENTER,
    space_after=2,
):
    """Clears and populates a table cell with formatted text.

    This function first erases all existing content within the specified `cell`.
    It then inserts a single new paragraph containing the provided text, which is
    sanitized before insertion. The paragraph and cell are formatted according
    to the specified alignment, font, and spacing parameters. The `cell` object
    is modified in-place.

    Args:
        cell (docx.table._Cell): The python-docx cell object to modify.
        text (str): The string content to write into the cell.
        bold (bool, optional): If True, formats the text in bold. Defaults to
            False.
        align (docx.enum.text.WD_ALIGN_PARAGRAPH, optional): An enum member for
            horizontal paragraph alignment. Defaults to
            WD_ALIGN_PARAGRAPH.JUSTIFY.
        font_size (Union[int, float], optional): The font size in points (pt).
            Defaults to 10.5.
        vertical (docx.enum.table.WD_ALIGN_VERTICAL, optional): An enum member
            for vertical cell content alignment. Defaults to
            WD_ALIGN_VERTICAL.CENTER.
        space_after (Union[int, float], optional): The spacing in points (pt) to
            add after the paragraph. Defaults to 2.

    Returns:
        None

    Raises:
        AttributeError: If `cell` is not a valid `docx` cell object and lacks
            the required attributes (e.g., `paragraphs`, `vertical_alignment`).
        TypeError: If `font_size` or `space_after` are not numeric types
            compatible with `docx.shared.Pt`.
        ValueError: If `align` or `vertical` are not valid `docx` alignment
            enum members.
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(clean_text(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = vertical


def set_repeat_table_header(row: Any) -> Any:
    """Marks a table row to be repeated as a header on subsequent pages.

    This function directly manipulates the underlying Office Open XML (OOXML) of a
    `python-docx` table row. It accesses the row's `trPr` (table row properties)
    element and appends a `w:tblHeader` child element. This instructs compliant
    word processing applications to repeat this row at the top of each new page
    that the table spans. The modification is performed in-place on the `row` object.

    Args:
        row (docx.table._Row): The `python-docx` table row object to configure.

    Returns:
        None

    Raises:
        AttributeError: If the provided `row` object does not possess the expected
            internal `_tr` attribute required for manipulation.
    """
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def prevent_row_break(row: Any) -> Any:
    """Prevent a table row from splitting across a page break.

    Modifies the row's underlying OOXML representation in-place by adding the
    `w:cantSplit` property. This setting instructs rendering applications to
    keep the entire row on a single page, preventing it from being divided by
    an automatic page break.

    Args:
        row (docx.table._Row): The table row object to modify.

    Returns:
        None.

    Raises:
        AttributeError: If the `row` object lacks the internal `_tr` attribute,
            which represents the underlying `CT_Tr` XML element.
    """
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)


def finalize_table(table: Any) -> Any:
    """Set the 'Table Grid' style, left alignment, and autofit on a Table object."""
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True


def autofit_table_to_contents(table: Any) -> Any:
    r"""{'docstring': "Forces a `python-docx` table to autofit its contents by manipulating OXML.\n\nThis function directly modifies the underlying Office Open XML (OXML)\nrepresentation of a `python-docx` Table object to enforce a content-based\nautofit layout. It serves as a more robust alternative to the standard\n`table.autofit = True` property, which may not override pre-existing\nfixed-width settings. The provided `table` object is mutated in-place.\n\nThe modification process consists of the following steps:\n1.  Removes the `w:tblGrid` element, which pre-defines fixed column widths.\n2.  Clears any existing table-level width (`w:tblW`) and layout (`w:tblLayout`)\n    properties from the table's properties (`w:tblPr`) element.\n3.  Sets the overall table layout type to `autofit` and its width type to\n    `auto`.\n4.  Iterates through each cell (`w:tc`) in the table, removing any\n    pre-existing cell-level fixed-width property (`w:tcW`).\n5.  Sets each cell's width type to `auto`, allowing it to dynamically resize\n    based on its content.\n\nArgs:\n    table (docx.table.Table): The `python-docx` Table object to be modified.\n\nRaises:\n    AttributeError: If the `table` object lacks the expected internal OXML\n        structure (e.g., `_tbl`, `_tc` attributes), indicating it is not a\n        valid `python-docx` Table object."}."""
    table.autofit = True
    tbl = table._tbl

    for child in list(tbl):
        if child.tag.split("}")[-1] == "tblGrid":
            tbl.remove(child)

    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    for child in list(tblPr):
        tag = child.tag.split("}")[-1]
        if tag in {"tblW", "tblLayout"}:
            tblPr.remove(child)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "auto")
    tblW.set(qn("w:w"), "0")
    tblPr.append(tblW)

    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for child in list(tcPr):
                if child.tag.split("}")[-1] == "tcW":
                    tcPr.remove(child)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:type"), "auto")
            tcW.set(qn("w:w"), "0")
            tcPr.append(tcW)


def render_note_box(doc, placeholder: Any, text: Any) -> Any:
    r"""{'docstring': 'Replaces a placeholder paragraph in a DOCX document with a styled note box.\n\n    Locates the first paragraph containing the `placeholder` text and inserts a\n    formatted, single-cell table immediately after it. The provided `text` is\n    placed within this cell, and specific styling (background shade, font,\n    alignment) is applied. The original placeholder paragraph is then removed.\n\n    If the `placeholder` is not found, the function performs no action. If the\n    `text` argument is `None` or an empty string after cleaning, the placeholder\n    paragraph is removed, but no table is inserted.\n\n    Args:\n        doc (docx.document.Document): The `python-docx` Document object to modify.\n        placeholder (str): The text content that uniquely identifies the target\n            paragraph for replacement.\n        text (str): The content to insert into the styled note box table.\n\n    Returns:\n        None: The document object is modified in-place.\n\n    Raises:\n        AttributeError: If `doc` is not a valid `python-docx` Document object\n            and lacks the expected attributes or methods.\n        TypeError: If `placeholder` or `text` are not string-like objects and\n            cannot be processed by internal helpers.'}."""
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return
    text = clean_text(text)
    if not text:
        remove_paragraph(paragraph)
        return

    table = doc.add_table(rows=1, cols=1)
    finalize_table(table)
    move_table_after_paragraph(paragraph, table)
    cell = table.cell(0, 0)
    set_cell_text(
        cell,
        text,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=9.5,
        vertical=WD_ALIGN_VERTICAL.CENTER,
        space_after=0,
    )
    shade_cell(cell, "F2F6FA")
    autofit_table_to_contents(table)
    remove_paragraph(paragraph)


def render_pillar_score_table(doc, placeholder: Any, rows: Any) -> Any:
    """Replaces a placeholder paragraph in a .docx document with a formatted table.

    Locates the first paragraph containing the `placeholder` text and inserts a
    styled table in its place. The generated table features a four-column layout
    with fixed Spanish headers: "Pilar", "Score", "Nivel", and
    "Lectura ejecutiva". The header row is shaded, bolded, and configured to
    repeat across page breaks. Data rows are populated from the `rows`
    iterable, and column widths are automatically adjusted to fit content.
    The original placeholder paragraph is removed after table creation.

    Note:
        If the specified `placeholder` text is not found within the document,
        the function returns silently without making any modifications.

    Args:
        doc (docx.document.Document): The python-docx Document object to modify.
        placeholder (str): The text within a paragraph that serves as a marker
            for table insertion.
        rows (Iterable[Dict[str, str]]): An iterable of dictionaries, each
            representing a table row. Expected keys include 'pillar_label',
            'score_display', 'maturity_band', and 'executive_reading'.

    Returns:
        None. This function modifies the 'doc' object in-place.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return

    table = doc.add_table(rows=1, cols=4)
    finalize_table(table)
    move_table_after_paragraph(paragraph, table)

    headers = ["Pilar", "Score", "Nivel", "Lectura ejecutiva"]
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    set_repeat_table_header(table.rows[0])

    for item in rows:
        row = table.add_row()
        prevent_row_break(row)
        set_cell_text(
            row.cells[0],
            item.get("pillar_label", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10.2,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[1],
            item.get("score_display", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10.2,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[2],
            item.get("maturity_band", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10.2,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[3],
            item.get("executive_reading", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10.2,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )

    autofit_table_to_contents(table)
    remove_paragraph(paragraph)


def render_risks_table(doc, placeholder: Any, rows: Any) -> Any:
    """Render and insert a project risks and mitigations table into a document.

    Locates the first paragraph containing the `placeholder` text and inserts a
    formatted table immediately after it. The table is populated with data from
    `rows` and features fixed Spanish headers for 'Riesgo', 'Impacto',
    'Probabilidad', and 'Mitigación resumida'. The original paragraph
    containing the placeholder is subsequently removed.

    If the placeholder is not found, the document remains unmodified.

    Args:
        doc (docx.document.Document.Document): The `python-docx` Document object to
            be modified.
        placeholder (str): The text content of the paragraph that serves as the
            insertion point for the table.
        rows (List[Dict[str, str]]): A list of dictionaries, where each
            dictionary represents a single risk and populates a table row. Each
            dictionary is expected to contain the keys 'risk', 'impact',
            'probability', and 'mitigation_summary'. Missing keys will result in
            empty cells for that row.

    Returns:
        None: This function modifies the `doc` object in place and does not
            return a value.

    Raises:
        AttributeError: If `doc` is not a valid `python-docx` Document object
            and lacks the expected methods.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return

    table = doc.add_table(rows=1, cols=4)
    finalize_table(table)
    move_table_after_paragraph(paragraph, table)

    headers = ["Riesgo", "Impacto", "Probabilidad", "Mitigación resumida"]
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    set_repeat_table_header(table.rows[0])

    for item in rows:
        row = table.add_row()
        prevent_row_break(row)
        set_cell_text(
            row.cells[0],
            item.get("risk", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[1],
            item.get("impact", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[2],
            item.get("probability", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[3],
            item.get("mitigation_summary", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )

    autofit_table_to_contents(table)
    remove_paragraph(paragraph)


def render_gap_table(doc, placeholder: Any, rows: Any) -> Any:
    r"""{'docstring': 'Inserts a styled gap analysis table into a python-docx Document object.\n\nLocates a paragraph containing the specified placeholder text and inserts a\n4-column table immediately after it. The table is populated with the\nprovided row data and formatted with specific styles, including hardcoded\nSpanish headers, a shaded header row, and defined text alignments. The\nheader row is configured to repeat on subsequent pages.\n\nThe original placeholder paragraph is removed after the table is created.\nIf the placeholder is not found, the document is not modified. This\nfunction modifies the `doc` object in-place.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` document object to modify.\n    placeholder (str): The text within a paragraph that serves as the anchor\n        for table insertion.\n    rows (List[Dict[str, str]]): A list of dictionaries, where each\n        dictionary represents a data row. Required keys are "pillar",\n        "as_is_summary", "target_state", and "key_gap".'}."""
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return

    table = doc.add_table(rows=1, cols=4)
    finalize_table(table)
    move_table_after_paragraph(paragraph, table)

    headers = ["Pilar", "Situación actual", "Estado objetivo", "Brecha clave"]
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    set_repeat_table_header(table.rows[0])

    for item in rows:
        row = table.add_row()
        prevent_row_break(row)
        set_cell_text(
            row.cells[0],
            item.get("pillar", ""),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[1],
            item.get("as_is_summary", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[2],
            item.get("target_state", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[3],
            item.get("key_gap", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )

    autofit_table_to_contents(table)
    remove_paragraph(paragraph)


def render_initiative_cards(doc, placeholder: Any, cards: Any) -> Any:
    """Renders a list of initiative cards as styled tables in a .docx document.

    Finds a paragraph matching the `placeholder` text and replaces it with a
    sequence of styled tables, one for each card provided. The tables are
    inserted sequentially at the original location of the placeholder. If the
    `cards` list is empty, the placeholder paragraph is removed. The original
    placeholder paragraph is deleted after all tables have been rendered.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to be
            modified in-place.
        placeholder (str): The exact text of the paragraph that serves as the
            insertion point for the card tables.
        cards (list[dict[str, Any]]): A list of dictionaries, where each
            dictionary represents a single initiative card. Expected keys include
            'sequence', 'initiative', 'objective', 'priority',
            'expected_outcome', and 'dependencies_display'.

    Returns:
        None: The `doc` object is modified directly.

    Raises:
        AttributeError: If an element in the `cards` list is not a dictionary
            or lacks a `.get()` method.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return
    if not cards:
        remove_paragraph(paragraph)
        return

    anchor = paragraph
    for idx, item in enumerate(cards, start=1):
        table = doc.add_table(rows=5, cols=2)
        finalize_table(table)
        move_table_after_paragraph(anchor, table)

        title_row = table.rows[0]
        merged = title_row.cells[0].merge(title_row.cells[1])
        set_cell_text(
            merged,
            f"{item.get('sequence', idx)}. {item.get('initiative', '')}",
            bold=True,
            font_size=11,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(merged, "D9EAF7")
        prevent_row_break(title_row)

        labels = ["Objetivo", "Prioridad", "Resultado esperado", "Dependencias"]
        values = [
            item.get("objective", ""),
            item.get("priority", ""),
            item.get("expected_outcome", ""),
            item.get("dependencies_display", ""),
        ]
        for row_idx in range(4):
            row = table.rows[row_idx + 1]
            set_cell_text(
                row.cells[0],
                labels[row_idx],
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=10,
                vertical=WD_ALIGN_VERTICAL.CENTER,
                space_after=0,
            )
            shade_cell(row.cells[0], "D9EAF7")
            set_cell_text(
                row.cells[1],
                values[row_idx],
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=10,
                vertical=WD_ALIGN_VERTICAL.CENTER,
            )
            prevent_row_break(row)

        autofit_table_to_contents(table)
        spacer = insert_paragraph_after_block(table)
        spacer.paragraph_format.space_after = Pt(6)
        add_run(spacer, "", font_size=1)
        anchor = spacer

    remove_paragraph(paragraph)


def render_radar_chart(doc, placeholder: Any, image_path: Any) -> Any:
    """Embeds a radar chart image into a Word document at a placeholder.

    Searches for the first paragraph containing the specified placeholder text
    and replaces its content with the provided image. The image is centered,
    resized to a fixed width of 4.35 inches, and followed by 6 points of
    spacing.

    If the placeholder text is not found, the document remains unmodified.
    If the provided `image_path` is invalid or cannot be resolved, the entire
    paragraph containing the placeholder is removed from the document.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to be
            modified in-place.
        placeholder (str): The placeholder text to search for, e.g., '{{radar_chart}}'.
        image_path (Union[str, os.PathLike]): The file system path to the radar
            chart image.

    Returns:
        None. The Document object is modified directly.

    Raises:
        FileNotFoundError: If the file at `image_path` does not exist.
        ValueError: If the file at `image_path` is not a recognized or
            supported image format.
        TypeError: If `doc` is not a valid `python-docx` Document instance.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return

    path = resolve_radar_chart_image(image_path)
    if path is None:
        remove_paragraph(paragraph)
        return

    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(4.35))
    paragraph.paragraph_format.space_after = Pt(6)


def add_heading_paragraph(doc, text: Any, level: Any = 1) -> Any:
    """Adds and styles a heading paragraph within a `python-docx` document.

    Creates a new paragraph, sanitizes the input text via an internal
    `clean_text` function, and applies specific formatting based on the heading
    level. Levels 1 and 2 are styled with a 'Georgia' font and a specific
    blue color (RGB 0, 114, 188). Levels 3 and higher receive a default bold
    styling. Paragraph spacing and font size are also adjusted per level.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to which
            the heading will be added.
        text (Any): The raw text content for the heading. This value is coerced
            to a string.
        level (Any): The heading level used to determine styling. This value is
            coerced to an integer. Defaults to 1.

    Returns:
        docx.text.paragraph.Paragraph: The Paragraph instance that was created and
            added to the document.

    Raises:
        AttributeError: If the `doc` object lacks the required `add_paragraph`
            method.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(clean_text(text))

    if level == 1:
        run.bold = False
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(18)
        paragraph.paragraph_format.page_break_before = False
        run.font.name = "Georgia"
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 114, 188)
    elif level == 2:
        run.bold = False
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        run.font.name = "Georgia"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 114, 188)
    else:
        run.bold = True
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(11.5)

    return paragraph


def add_body_paragraph(
    doc, text: Any, justify: Any = True, space_after: Any = 6, color_rgb: Any = None
) -> Any:
    r"""{'docstring': "Adds a formatted body paragraph to a `docx.document.Document` object.\n\n    This function first cleans the input `text`. If the resulting string is\n    non-empty, a new paragraph is created and appended to the document. The\n    function then applies specified formatting for justification, spacing, and\n    text color.\n\n    Args:\n        doc (docx.document.Document): The document object to which the paragraph\n            will be added.\n        text (str): The raw text content for the paragraph.\n        justify (bool): If `True`, the paragraph text will be fully justified.\n            Defaults to `True`.\n        space_after (Union[int, float]): The space in points (Pt) to add after\n            the paragraph. Defaults to 6.\n        color_rgb (Optional[docx.shared.RGBColor]): An instance of\n            `docx.shared.RGBColor` to apply to the text. If `None`, the\n            document's default text color is used. Defaults to `None`.\n\n    Returns:\n        Optional[docx.text.paragraph.Paragraph]: The created `Paragraph` object,\n        or `None` if the input `text` is empty after cleaning.\n\n    Raises:\n        AttributeError: If `doc` is not a valid Document object (e.g., it lacks\n            the `add_paragraph` method).\n        ValueError: If an invalid value is provided for `color_rgb` that cannot\n            be processed by the underlying `python-docx` library."}."""
    text = clean_text(text)
    if not text:
        return None
    paragraph = doc.add_paragraph()
    add_run(paragraph, text, font_size=10.5)
    apply_body_format(paragraph, justify=justify, space_after=space_after)
    if color_rgb:
        for run in paragraph.runs:
            run.font.color.rgb = color_rgb
    return paragraph


def add_bullet_list(doc, items: Any) -> Any:
    r"""{'docstring': 'Appends a formatted bullet list to a `python-docx` document.\n\nSanitizes each string in the input iterable and filters out any resulting\nempty items. The remaining items are added to the document as individual\nbullet points, each formatted with a 10-point font size.\n\nArgs:\n    doc (docx.document.Document): The document object to modify in-place.\n    items (Iterable[str]): A collection of strings to add as bullet points.\n\nReturns:\n    None\n\nRaises:\n    AttributeError: If `doc` does not have an `add_paragraph` method.\n    TypeError: If `items` is not an iterable.'}."""
    items = [clean_text(item) for item in items if clean_text(item)]
    for item in items:
        paragraph = doc.add_paragraph()
        strip_numbering_and_indents(paragraph)
        apply_bullet_list_format(paragraph)
        text = paragraph.add_run(item)
        text.font.size = Pt(10)


def add_label_value_paragraph(doc, label: Any, value: Any) -> Any:
    """Adds a formatted paragraph with a label and a value to a document.

    The function first sanitizes the `value` by calling `clean_text`. If the
    sanitized value is a non-empty string, a new paragraph is appended to the
    `doc`. The paragraph is styled with the `label` rendered in bold, followed
    by the sanitized `value`. Both parts are set to a 10.5pt font size, and
    specific paragraph spacing is applied. No action is taken if the sanitized
    value is empty.

    Args:
        doc (docx.document.Document): The `Document` object to which the paragraph
            will be added.
        label (Any): The object to be used as the label. It will be coerced into
            a string.
        value (Any): The object to be used as the value. It will be passed through
            a cleaning function and then coerced into a string.

    Returns:
        Optional[docx.text.paragraph.Paragraph]: The newly created and appended
            `Paragraph` object if the sanitized value was not empty, otherwise
            `None`.

    Raises:
        AttributeError: If `doc` is not a valid `docx.Document` object and lacks
            the `add_paragraph` method.
    """
    value = clean_text(value)
    if not value:
        return None
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10.5)
    value_run = paragraph.add_run(value)
    value_run.font.size = Pt(10.5)
    apply_body_format(paragraph, justify=False, space_after=4)
    return paragraph


def add_long_detail_table(doc, title: Any, rows: Any) -> Any:
    """Adds a styled, two-column table with a merged header to a document.

    The function first filters the input `rows`, retaining only those where the
    value is non-empty after being processed by a text cleaning utility. If no
    valid rows remain, the function returns `None` and the document is not
    modified.

    Otherwise, a two-column table is created and appended to the document. The
    table structure and styling are as follows:
    - A header row with two cells merged, containing the `title` text in a bold,
      11pt font, shaded with color "D9EAF7".
    - Subsequent rows for each valid (label, value) pair.
    - The first (label) column is styled with a bold, 10pt font and shaded
      with color "EEF5FB".
    - The second (value) column uses a regular 10pt font.
    - All rows are configured to prevent splitting across page breaks.
    - Finally, column widths are automatically adjusted to fit their content.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to which
            the table will be appended.
        title (str): The text to display in the merged header row.
        rows (Iterable[Tuple[Any, Any]]): An iterable of two-element sequences
            representing (label, value) pairs. Values are coerced to strings.

    Returns:
        Optional[docx.table.Table]: The created `docx.table.Table` object if at
            least one data row was added, otherwise `None`.

    Raises:
        ValueError: If an element in `rows` cannot be unpacked into two values.
        AttributeError: If `doc` lacks an `add_table` method.
    """
    rows = [
        (clean_text(label), clean_text(value))
        for label, value in rows
        if clean_text(value)
    ]
    if not rows:
        return None

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    finalize_table(table)

    header = table.rows[0]
    merged = header.cells[0].merge(header.cells[1])
    set_cell_text(
        merged,
        title,
        bold=True,
        font_size=11,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        vertical=WD_ALIGN_VERTICAL.CENTER,
        space_after=0,
    )
    shade_cell(merged, "D9EAF7")
    prevent_row_break(header)

    for idx, (label, value) in enumerate(rows, start=1):
        row = table.rows[idx]
        set_cell_text(
            row.cells[0],
            label,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(row.cells[0], "EEF5FB")
        set_cell_text(
            row.cells[1],
            value,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        prevent_row_break(row)

    autofit_table_to_contents(table)
    return table


def remove_page_break_only_paragraphs(doc) -> Any:
    """Remove paragraphs containing only a manual page break and whitespace.

    Scans a document's paragraphs to identify and delete those that serve only
    as page break containers. A paragraph is targeted for removal if it meets
    two criteria: it contains a run with an underlying OOXML `<w:br>` element
    that has a `type` attribute of `page`, and its textual content is empty
    after stripping whitespace characters.

    This is an in-place operation that directly modifies the input `doc` object.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to
            modify in-place.

    Returns:
        None
    """
    to_remove = []
    for paragraph in doc.paragraphs:
        has_text = bool(clean_text(paragraph.text))
        has_page_break = False
        for run in paragraph.runs:
            for br in run._r.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
            ):
                if (
                    br.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                    )
                    == "page"
                ):
                    has_page_break = True
        if has_page_break and not has_text:
            to_remove.append(paragraph)
    for paragraph in to_remove:
        remove_paragraph(paragraph)
