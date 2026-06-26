import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

"""
Módulo render_commercial_report.py.
Contiene la lógica y utilidades principales para el pipeline de Assessment Engine.
"""

from assessment_engine.adapters.render_global_report_from_template import (
    add_smart_bullet_list,
    add_spacer,
    clean_t_codes,
    create_page_number_footer,
)
from assessment_engine.domain.schemas.commercial import (
    CommercialPayload,
    CommercialSummaryDraft,
    GtmStrategy,
    OpportunityPipelineItem,
    ProposalDraft,
)
from assessment_engine.infrastructure.contract_utils import robust_load_payload
from assessment_engine.infrastructure.docx_render_utils import (
    add_body_paragraph,
    add_heading_paragraph,
    autofit_table_to_contents,
    clear_paragraph,
    finalize_table,
    set_cell_text,
    shade_cell,
)

BASE_TEXT_COLOR = RGBColor(46, 64, 77)


def clean_commercial_text(text) -> Any:
    """Clean and standardize a string for commercial reporting.

    Applies a sequence of cleaning and normalization operations to an input string.

    The operations are executed in the following order:
    1.  Removes 't-codes' via the `clean_t_codes` helper function.
    2.  Strips unresolved reference tags matching the pattern `[[REF:...]]`.
    3.  Removes em-dash (—) and box-drawing horizontal (─) characters.
    4.  Normalizes semicolons to commas.
    5.  Removes leading and trailing whitespace.

    If the input is not a string, it is returned unmodified. This provides a
    safe pass-through for non-string data types.

    Args:
        text (Any): The input value to process.

    Returns:
        Any: The cleaned string, or the original input if it was not a string.
    """
    if not isinstance(text, str):
        return text
    cleaned = clean_t_codes(text)
    # Post-processes the document to remove unresolved reference tags. This is a safeguard against rendering artifacts that can result from upstream data inconsistencies or an incomplete template context.
    cleaned = re.sub(r"\[\[REF:[^\]]*\]\]", "", cleaned)
    return cleaned.replace("─", "").replace("—", "").replace(";", ",").strip()


def load_json(path) -> Any:
    """Deserialize a UTF-8 encoded JSON file from a path into a Python object."""
    return json.loads(path.read_text(encoding="utf-8-sig"))


def clear_document_body(doc) -> Any:
    """Removes all content from the document body while preserving section properties.

    Iterates through the top-level XML elements within the document's body and
    removes all content elements (e.g., paragraphs, tables). The final section
    properties element (`w:sectPr`) is explicitly preserved. This is primarily
    used to clear a template document while retaining its page layout settings
    (e.g., margins, orientation). The input `Document` object is modified
    in-place.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to modify.

    Returns:
        None
    """
    body = doc._body._element
    for child in list(body):
        if (
            child.tag
            != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
        ):
            body.remove(child)


def load_payload(payload_path: Path) -> CommercialPayload:
    """Load and strictly validate a 'Commercial Account Plan' payload from a file."""
    return robust_load_payload(
        payload_path,
        CommercialPayload,
        "Commercial Account Plan",
        mode="strict",
    )


def render_commercial_cover(doc, payload: CommercialPayload) -> Any:
    """Render a formatted cover page into a `docx.Document` object.

    Populates the provided document with a standardized commercial report cover page.
    The page includes a main title, client name, date, classification notice, and a
    detailed confidentiality disclaimer in Spanish, applying specific text and
    paragraph formatting. This function modifies the document object in-place and
    concludes by inserting a page break.

    Args:
        doc: The `python-docx.document.Document` object to which the cover page
            content will be added.
        payload: A `CommercialPayload` data object containing the metadata for the
            report, such as the client name and report date.

    Returns:
        None. The function modifies the `doc` object in-place.
    """
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run("Account Action Plan")
    run.font.size = Pt(34)
    run.font.name = "Georgia"
    run.font.color.rgb = RGBColor(0, 114, 188)
    run.bold = False

    add_spacer(doc, 30)
    client_p = doc.add_paragraph()
    client_run = client_p.add_run(payload.meta.client.upper())
    client_run.font.size = Pt(24)
    client_run.font.name = "Arial"
    client_run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(100)

    version_p = doc.add_paragraph()
    v_text = f"Fecha: {payload.meta.date}\nClasificación: CONFIDENCIAL - USO INTERNO"
    version_run = version_p.add_run(v_text)
    version_run.font.size = Pt(14)
    version_run.font.name = "Arial"
    version_run.font.color.rgb = RGBColor(46, 64, 77)
    version_run.bold = False

    add_spacer(doc, 150)

    disclaimer_title_p = doc.add_paragraph()
    disclaimer_title_run = disclaimer_title_p.add_run(
        "Advertencia de Confidencialidad:"
    )
    disclaimer_title_run.font.size = Pt(8)
    disclaimer_title_run.font.name = "Arial"
    disclaimer_title_run.font.color.rgb = RGBColor(127, 127, 127)
    disclaimer_title_run.bold = True
    disclaimer_title_p.paragraph_format.space_after = Pt(2)

    disclaimer_p = doc.add_paragraph()
    disclaimer_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p.paragraph_format.space_before = Pt(0)
    p1_text = "El presente documento es estrictamente confidencial y de uso exclusivo interno para los equipos comerciales, de preventa y de entrega de NTT DATA. Contiene estimaciones de negocio, argumentarios de venta competitivos, estrategias de cuenta (Go-To-Market) y evaluaciones de riesgo de clientes. Bajo ninguna circunstancia este documento, ni parcial ni totalmente, debe ser compartido con el cliente, partners, proveedores o cualquier tercero."
    r1 = disclaimer_p.add_run(p1_text)
    r1.font.size = Pt(8)
    r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor(127, 127, 127)

    doc.add_page_break()


def render_commercial_summary(doc, data: CommercialSummaryDraft, matrix: list) -> Any:
    """Renders a commercial summary and stakeholder matrix into a `python-docx` document.

    This function populates a Word document with a formatted executive summary
    and deal snapshot. It includes a contextual warning, a "flash table" with key
    deal metrics, bulleted lists for the problem statement and proposed strategy,
    and an optional stakeholder value matrix. The function modifies the input
    `doc` object in-place.

    Args:
        doc (docx.document.Document): The document object to be modified.
        data (CommercialSummaryDraft): A data object containing the summary content.
            Must expose `estimated_tam`, `deal_flash` (an object with
            `purchase_driver` and `ntt_win_theme` attributes),
            `why_now_bullets`, and `how_we_win_bullets`.
        matrix (List[object]): A list of stakeholder data objects, each with
            `role`, `focus`, and `message` attributes. If empty, the stakeholder
            matrix section is not rendered.

    Raises:
        AttributeError: If `data` or any object within the `matrix` list is
            missing required attributes.
    """
    add_heading_paragraph(doc, "1. Executive Summary & Deal Snapshot", level=1)

    #
    warn_table = doc.add_table(rows=1, cols=1)
    finalize_table(warn_table)
    w_cell = warn_table.rows[0].cells[0]
    shade_cell(w_cell, "F2F2F2")
    clear_paragraph(w_cell.paragraphs[0])
    wp = w_cell.paragraphs[0]
    wp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    wr1 = wp.add_run("⚠️ CONTEXTO DEL INFORME (FAST ASSESSMENT): ")
    wr1.bold = True
    wr1.font.size = Pt(9)
    wr1.font.color.rgb = RGBColor(127, 127, 127)
    wr2 = wp.add_run(
        "Este documento es un artefacto comercial interno. Las oportunidades, riesgos y valoraciones económicas nacen de un diagnóstico preliminar y entrevistas de alto nivel, no de una due-diligence técnica. Su propósito es definir la estrategia de entrada (Go-To-Market), no es una oferta vinculante."
    )
    wr2.font.size = Pt(9)
    wr2.font.color.rgb = RGBColor(127, 127, 127)
    add_spacer(doc, 15)

    #
    flash_table = doc.add_table(rows=3, cols=2)
    flash_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    finalize_table(flash_table)
    flash_table.columns[0].width = Inches(2.0)
    flash_table.columns[1].width = Inches(4.5)

    set_cell_text(
        flash_table.rows[0].cells[0], "TAM Estimado NTT DATA:", bold=True, font_size=11
    )
    shade_cell(flash_table.rows[0].cells[0], "0072BC")
    set_cell_text(
        flash_table.rows[0].cells[1],
        f"{clean_commercial_text(data.estimated_tam)} (18-24 meses)",
        bold=True,
        font_size=11,
    )
    shade_cell(flash_table.rows[0].cells[1], "0072BC")
    for r in flash_table.rows[0].cells[0].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
    for r in flash_table.rows[0].cells[1].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    deal_flash = data.deal_flash
    set_cell_text(
        flash_table.rows[1].cells[0],
        "Driver de Compra (Urgencia):",
        bold=True,
        font_size=10.5,
    )
    shade_cell(flash_table.rows[1].cells[0], "F2F2F2")
    set_cell_text(
        flash_table.rows[1].cells[1],
        clean_commercial_text(deal_flash.purchase_driver),
        font_size=10.5,
    )

    set_cell_text(
        flash_table.rows[2].cells[0],
        "Nuestra Ventaja (Win Theme):",
        bold=True,
        font_size=10.5,
    )
    shade_cell(flash_table.rows[2].cells[0], "F2F2F2")
    set_cell_text(
        flash_table.rows[2].cells[1],
        clean_commercial_text(deal_flash.ntt_win_theme),
        font_size=10.5,
    )
    autofit_table_to_contents(flash_table)
    add_spacer(doc, 15)

    #
    p_why = doc.add_paragraph()
    r_why = p_why.add_run("El Problema (Why Now):")
    r_why.bold = True
    r_why.font.color.rgb = BASE_TEXT_COLOR
    r_why.font.size = Pt(11)
    p_why.paragraph_format.space_after = Pt(4)
    add_smart_bullet_list(
        doc,
        data.why_now_bullets,
        color_rgb=BASE_TEXT_COLOR,
        bold_prefix=False,
    )
    add_spacer(doc, 10)

    p_how = doc.add_paragraph()
    r_how = p_how.add_run("La Estrategia de NTT DATA (How we Win):")
    r_how.bold = True
    r_how.font.color.rgb = BASE_TEXT_COLOR
    r_how.font.size = Pt(11)
    p_how.paragraph_format.space_after = Pt(4)
    add_smart_bullet_list(
        doc,
        data.how_we_win_bullets,
        color_rgb=BASE_TEXT_COLOR,
        bold_prefix=False,
    )
    add_spacer(doc, 20)

    if matrix:
        doc.add_page_break()
        add_heading_paragraph(
            doc, "Matriz de Valor por Interlocutor (Stakeholders)", level=2
        )
        add_body_paragraph(
            doc,
            "Alineación del discurso comercial según las prioridades de cada rol directivo.",
            color_rgb=BASE_TEXT_COLOR,
            space_after=12,
        )

        m_table = doc.add_table(rows=1, cols=3)
        finalize_table(m_table)
        headers = [
            "Rol Directivo",
            "Foco Estratégico",
            "Mensaje Clave (Elevator Pitch)",
        ]
        for i, h in enumerate(headers):
            set_cell_text(
                m_table.rows[0].cells[i],
                h,
                bold=True,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(m_table.rows[0].cells[i], "0072BC")
            for r in m_table.rows[0].cells[i].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(255, 255, 255)

        for row_data in matrix:
            row = m_table.add_row()
            set_cell_text(
                row.cells[0],
                clean_commercial_text(row_data.role),
                bold=True,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(row.cells[0], "F2F2F2")
            set_cell_text(
                row.cells[1],
                clean_commercial_text(row_data.focus),
                font_size=10.5,
            )
            set_cell_text(
                row.cells[2],
                clean_commercial_text(row_data.message),
                font_size=10.5,
            )
        autofit_table_to_contents(m_table)
        add_spacer(doc, 20)


def render_gtm_strategy(doc, gtm: GtmStrategy) -> Any:
    r"""{'docstring': "Adds a 'Go-To-Market' strategy section to a Word document object.\n\nThis function generates a level-1 heading and a formatted two-column table\nsummarizing Go-To-Market strategies. The table populates rows with strategy\nnames and their corresponding descriptions, which are sourced from the `gtm`\ndata object. The table's appearance is customized with specific colors,\nfonts, and alignments, and its columns are auto-fitted to the content.\n\nArgs:\n    doc: The `python-docx` Document object to which the section is added.\n        This object is modified in-place.\n    gtm: An object containing go-to-market strategy descriptions, required\n        to have `trojan_horse`, `self_funded_transformation`, and `lock_in`\n        attributes.\n\nReturns:\n    None. The `doc` object is modified by side effect.\n\nRaises:\n    AttributeError: If the `gtm` object lacks a required strategy attribute."}."""
    add_heading_paragraph(doc, "2. Go-To-Market", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    finalize_table(table)
    headers = ["Vector de Entrada", "Estrategia Recomendada"]
    for i, h in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            bold=True,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table.rows[0].cells[i], "0072BC")
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
    vectors = [
        ("El Caballo de Troya (Wedge)", gtm.trojan_horse),
        ("Transformación Autofinanciada", gtm.self_funded_transformation),
        ("Estrategia de Lock-in (MRR)", gtm.lock_in),
    ]
    for v_name, v_desc in vectors:
        row = table.add_row()
        set_cell_text(
            row.cells[0],
            v_name,
            bold=True,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(row.cells[0], "F2F2F2")
        set_cell_text(
            row.cells[1], clean_commercial_text(v_desc), bold=False, font_size=10.5
        )
    autofit_table_to_contents(table)
    add_spacer(doc, 20)


def render_opportunities_pipeline(doc, pipeline: list[OpportunityPipelineItem]) -> Any:
    """Populates a Word document with a section detailing commercial opportunities.

    This function generates the "Pipeline de Oportunidades" section by first
    adding a primary heading and a descriptive paragraph to the document. It then
    iterates through each `OpportunityPipelineItem` in the `pipeline` list,
    constructing a distinct, formatted table for each one.

    Each table summarizes key details of an opportunity, including the
    initiative name, co-sell vendor, revenue type, estimated Total Contract
    Value (TCV), the client's cost of inaction, and a "Red Team" analysis of
    potential client objections paired with corresponding mitigation strategies.

    Args:
        doc (docx.document.Document): An instance of a `python-docx` Document
            object to which the pipeline section will be appended.
        pipeline (list[OpportunityPipelineItem]): A list of data objects, where
            each object contains the details for a single commercial opportunity.

    Returns:
        None. The function modifies the input `doc` object in-place.
    """
    add_heading_paragraph(doc, "3. Pipeline de Oportunidades", level=1)
    add_body_paragraph(
        doc,
        "Estimaciones financieras y estrategias de mitigación de objeciones (Red Team).",
        color_rgb=BASE_TEXT_COLOR,
        space_after=12,
    )
    for i, opp in enumerate(pipeline, start=1):
        table = doc.add_table(rows=4, cols=2)
        finalize_table(table)
        h_cell = table.rows[0].cells[0]
        h_cell.merge(table.rows[0].cells[1])
        set_cell_text(
            h_cell,
            f"Oportunidad {i}: {clean_commercial_text(opp.initiative)}",
            bold=True,
            font_size=11,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(h_cell, "0072BC")
        for r in h_cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_text(
            table.rows[1].cells[0],
            f"Vendor Co-Sell: {opp.vendor_cosell}",
            bold=True,
            font_size=10.5,
        )
        set_cell_text(
            table.rows[1].cells[1],
            f"Tipo de Ingreso: {opp.revenue_type}",
            bold=True,
            font_size=10.5,
        )
        shade_cell(table.rows[1].cells[0], "F2F2F2")
        shade_cell(table.rows[1].cells[1], "F2F2F2")
        set_cell_text(
            table.rows[2].cells[0],
            f"TCV Estimado NTT DATA:\n{opp.estimated_tcv}",
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(
            table.rows[2].cells[1],
            f"Cost of Inaction (COI) para Cliente:\n{clean_commercial_text(opp.cost_of_inaction)}",
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        obj_cell = table.rows[3].cells[0]
        obj_cell.merge(table.rows[3].cells[1])
        clear_paragraph(obj_cell.paragraphs[0])
        p1 = obj_cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run("🛑 Objeción del Cliente (Red Team): ")
        r1.bold = False
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = RGBColor(204, 0, 0)
        r_obj = p1.add_run(clean_commercial_text(opp.client_objection))
        r_obj.bold = False
        r_obj.font.size = Pt(10.5)

        p2 = obj_cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run("✅ Respuesta NTT DATA (Objection Handling): ")
        r2.bold = False
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = RGBColor(39, 78, 19)
        r_ans = p2.add_run(clean_commercial_text(opp.objection_handling))
        r_ans.bold = False
        r_ans.font.size = Pt(10.5)
        add_spacer(doc, 15)


def render_proactive_proposals(doc, proposals: list[ProposalDraft]) -> Any:
    """Appends a formatted appendix section for each proactive proposal to a document.

    This function first adds a main heading for the "Proactive Proposals" appendix. It then iterates through each `ProposalDraft` object, inserting a page break before each new proposal to ensure clear separation. For each proposal, it constructs a detailed, multi-part section with specific formatting, including:

    - Numbered title for the proposal.
    - Executive Synthesis, AI Strategy, Context, Solution, and Scope.
    - Delivery model, team composition, and critical assumptions.
    - A formatted table for risk identification and mitigation strategies.
    - Company differentiators and investment/timeline details.

    All text content is cleaned and rendered using specific styles for headings, body text, bullet points, and table cells to maintain document consistency.

    Args:
        doc (docx.document.Document): The `python-docx` document object that will be modified in-place.
        proposals (list[ProposalDraft]): A sequence of data objects, with each object containing the structured content for a single proposal.

    Raises:
        AttributeError: If a `ProposalDraft` object is missing a required attribute, indicating malformed input data.
    """
    add_heading_paragraph(doc, "4. Anexos: Propuestas Proactivas", level=1)
    add_body_paragraph(
        doc,
        "Borradores de propuestas detalladas listos para ser adaptados y enviados al cliente tras incluir las credenciales correspondientes.",
        color_rgb=BASE_TEXT_COLOR,
        space_after=12,
    )
    for idx, prop in enumerate(proposals, 1):
        doc.add_page_break()
        p_title = add_heading_paragraph(
            doc,
            f"Propuesta {idx}: {clean_commercial_text(prop.initiative_name)}",
            level=2,
        )
        for r in p_title.runs:
            r.font.size = Pt(16)
        add_heading_paragraph(doc, "1. Executive Synthesis", level=3)
        p_syn = doc.add_paragraph()
        p_syn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_syn_text = p_syn.add_run(clean_commercial_text(prop.executive_synthesis))
        r_syn_text.bold = False
        r_syn_text.font.color.rgb = RGBColor(46, 64, 77)
        p_syn.paragraph_format.space_after = Pt(12)

        ai_strat = prop.ai_transformation_strategy
        if ai_strat:
            add_heading_paragraph(
                doc, "Estrategia de Transformación y Automatización", level=3
            )
            p_ai = doc.add_paragraph()
            p_ai.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_ai = p_ai.add_run(clean_commercial_text(ai_strat))
            r_ai.font.color.rgb = RGBColor(46, 64, 77)
            p_ai.paragraph_format.space_after = Pt(12)

        add_heading_paragraph(
            doc, "2. Contexto y Comprensión del Reto (The 'Why')", level=3
        )
        ctx = prop.context_and_why
        add_smart_bullet_list(
            doc,
            [
                f"Origen del Proyecto: {clean_commercial_text(ctx.origin)}",
                f"Coste de la Inacción (COI): {clean_commercial_text(ctx.cost_of_inaction)}",
            ],
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        add_spacer(doc, 10)
        add_heading_paragraph(
            doc, "3. La Solución y Resultados Estratégicos (The 'What')", level=3
        )
        sol = prop.solution_and_what
        add_smart_bullet_list(
            doc,
            [
                f"Estado Objetivo (TO-BE): {clean_commercial_text(sol.target_state)}",
                f"Métrica de Éxito (North Star Metric): {clean_commercial_text(sol.north_star_metric)}",
            ],
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        add_spacer(doc, 10)
        add_heading_paragraph(doc, "4. Alcance y Entregables (The 'How')", level=3)
        scope = prop.scope_and_how
        p_phases = doc.add_paragraph()
        r_ph = p_phases.add_run("Fases del Proyecto:")
        r_ph.bold = True
        r_ph.font.color.rgb = RGBColor(46, 64, 77)
        p_phases.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            scope.phases,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        p_del = doc.add_paragraph()
        r_del = p_del.add_run("Entregables Principales:")
        r_del.bold = True
        r_del.font.color.rgb = RGBColor(46, 64, 77)
        p_del.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            scope.deliverables,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        p_out = doc.add_paragraph()
        r_out = p_out.add_run("Fuera de Alcance (Out of Scope):")
        r_out.bold = True
        r_out.font.color.rgb = RGBColor(46, 64, 77)
        p_out.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            scope.out_of_scope,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        add_spacer(doc, 10)
        add_heading_paragraph(doc, "5. Modelo de Entrega y Asunciones", level=3)
        gov = prop.governance_and_assumptions
        team = prop.delivery_team
        add_body_paragraph(
            doc,
            f"Modelo de Gobierno: {clean_commercial_text(gov.governance_model)}",
            color_rgb=RGBColor(46, 64, 77),
            space_after=6,
        )
        p_team = doc.add_paragraph()
        r_team = p_team.add_run("Perfiles Clave:")
        r_team.bold = True
        r_team.font.color.rgb = RGBColor(46, 64, 77)
        p_team.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            team.team_roles,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        p_as = doc.add_paragraph()
        r_as = p_as.add_run("Asunciones Críticas (Assumptions):")
        r_as.bold = True
        r_as.font.color.rgb = RGBColor(46, 64, 77)
        p_as.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            gov.assumptions,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        add_spacer(doc, 10)
        add_heading_paragraph(doc, "6. Gestión de Riesgos del Proyecto", level=3)
        risks = prop.risk_management
        if risks:
            rtable = doc.add_table(rows=1, cols=2)
            finalize_table(rtable)
            set_cell_text(
                rtable.rows[0].cells[0],
                "Riesgo Identificado",
                bold=True,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text(
                rtable.rows[0].cells[1],
                "Estrategia de Mitigación NTT DATA",
                bold=True,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(rtable.rows[0].cells[0], "0072BC")
            shade_cell(rtable.rows[0].cells[1], "0072BC")
            for cell in rtable.rows[0].cells:
                for r in cell.paragraphs[0].runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
            for risk_item in risks:
                row = rtable.add_row()
                set_cell_text(
                    row.cells[0],
                    clean_commercial_text(risk_item.risk),
                    font_size=10.5,
                )
                set_cell_text(
                    row.cells[1],
                    clean_commercial_text(risk_item.mitigation),
                    font_size=10.5,
                )
                shade_cell(row.cells[0], "F2F2F2")
                shade_cell(row.cells[1], "F2F2F2")
            add_spacer(doc, 10)
        add_heading_paragraph(doc, "7. Por qué NTT DATA", level=3)
        why = prop.why_ntt_data
        add_smart_bullet_list(
            doc,
            why.accelerators,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        add_body_paragraph(
            doc,
            f"Partnerships Clave: {clean_commercial_text(why.partnerships)}",
            color_rgb=RGBColor(46, 64, 77),
        )
        p_cred = doc.add_paragraph()
        r_cred = p_cred.add_run(
            "[⚠️ ACCIÓN COMERCIAL REQUERIDA: INSERTE AQUÍ AL MENOS 2 CREDENCIALES DE PROYECTOS SIMILARES EN EL SECTOR DEL CLIENTE ANTES DE ENVIAR]"
        )
        r_cred.bold = True
        r_cred.font.color.rgb = RGBColor(238, 63, 50)
        p_cred.paragraph_format.space_before = Pt(6)
        p_cred.paragraph_format.space_after = Pt(10)
        add_heading_paragraph(
            doc, "8. Inversión y Plan de Activación (Next 14 Days)", level=3
        )
        inv = prop.investment_and_timeline
        add_smart_bullet_list(
            doc,
            [
                f"Duración Estimada: {clean_commercial_text(inv.estimated_duration)}",
                f"Rango de Inversión (TCV): {clean_commercial_text(inv.tcv_range)}",
            ],
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        p_act = doc.add_paragraph()
        r_act = p_act.add_run("Plan de Activación Inmediata:")
        r_act.bold = True
        r_act.font.color.rgb = RGBColor(46, 64, 77)
        p_act.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            prop.activation_plan,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        add_spacer(doc, 15)


def process_footnotes(doc, dossier) -> Any:
    """Removes unresolved `[[REF:...]]` placeholders from a document object.

    This function performs a final sanitization pass on a document, modifying it
    in-place. It iterates through all paragraphs, both at the top level and
    within table cells, to find and strip any remaining placeholder strings that
    match the pattern `[[REF:...]]`. This cleanup ensures the final output is free
    of internal syntax artifacts that might persist if the corresponding data
    was not found during the rendering process.

    Args:
        doc: The document object to be sanitized. Its paragraphs and tables are
            scanned and modified in-place.
        dossier: Conforms to the standard processing signature; not used in this
            function.

    Returns:
        None. The document object is modified in-place.
    """
    # Executes a final sanitization pass to strip any unresolved `[[REF:...]]` placeholders from the document. This cleanup step ensures the final output is free of internal syntax artifacts that might persist if the corresponding data was not available during rendering.
    ref_pattern = re.compile(r"\[\[REF:[^\]]*\]\]")

    def strip_refs_from_paragraph(p):
        """Removes reference tags (e.g., `[[ref]]`) from a `docx` paragraph in-place.

        Iterates through the `run` objects within the provided paragraph. For each run,
        this function uses a globally-defined `ref_pattern` regular expression to
        find and remove all matching reference tag substrings. The modification is
        applied directly to the `run.text` attribute of each run, effectively
        altering the paragraph object passed as an argument.

        Args:
            p (docx.text.paragraph.Paragraph): The paragraph object to modify in-place.

        Raises:
            NameError: If the global regular expression pattern `ref_pattern` is not
                defined in the calling scope.
        """
        if "[[" in p.text:
            for run in p.runs:
                if "[[" in run.text:
                    run.text = ref_pattern.sub("", run.text)

    for p in doc.paragraphs:
        strip_refs_from_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    strip_refs_from_paragraph(p)


def render_commercial_report(
    payload: CommercialPayload,
    template_path: Path,
    output_path: Path,
) -> Path:
    r"""{'docstring': 'Generate a commercial report as a Word document from a data payload.\n\nOrchestrates the report generation by loading a specified .docx template,\nclearing its body content, and populating it with data. The function\nconditionally renders sections for a cover page, summary, go-to-market\nstrategy, opportunities pipeline, and proactive proposals based on the\nprovided payload. Finally, it processes footnotes, adds page numbering,\nand saves the completed document to the specified output path.\n\nArgs:\n    payload (CommercialPayload): The data transfer object containing all\n        information to be rendered in the report.\n    template_path (pathlib.Path): The file system path to the .docx\n        template document which provides the base structure and styles.\n    output_path (pathlib.Path): The destination file system path where the\n        generated .docx report will be saved.\n\nReturns:\n    pathlib.Path: The path to the newly created report file, identical to the\n        `output_path` argument.\n\nRaises:\n    docx.opc.exceptions.PackageNotFoundError: If the file at `template_path`\n        does not exist or is not a valid Word document.\n    OSError: If an error occurs while writing the document to `output_path`,\n        such as a permissions issue or an invalid path.'}."""
    doc = Document(str(template_path))
    clear_document_body(doc)

    if doc.sections:
        create_page_number_footer(doc.sections[0])

    render_commercial_cover(doc, payload)

    if payload.commercial_summary:
        render_commercial_summary(
            doc, payload.commercial_summary, payload.stakeholder_matrix
        )
    if payload.gtm_strategy:
        render_gtm_strategy(doc, payload.gtm_strategy)
    if payload.opportunities_pipeline:
        render_opportunities_pipeline(doc, payload.opportunities_pipeline)
    if payload.proactive_proposals:
        render_proactive_proposals(doc, payload.proactive_proposals)

    process_footnotes(doc, payload.intelligence_dossier)
    doc.save(str(output_path))
    return output_path


def main(argv: list[str] | None = None) -> None:
    """Orchestrates the rendering of a commercial report from the command line.

    This function serves as the main entry point for a command-line utility. It
    parses three required arguments: a path to a JSON payload file, a path to a
    template file, and an output file path. It then loads the data and invokes
    the rendering process.

    Args:
        argv: An optional list of command-line arguments, including the script
            name. Defaults to `sys.argv` if None. Provided primarily for test
            isolation.

    Raises:
        SystemExit: If the number of provided arguments is not four (script name
            plus three required paths).
        FileNotFoundError: If the payload or template file does not exist.
        IOError: If the rendered report cannot be written to the output path.
        json.JSONDecodeError: If the payload file contains malformed JSON.
    """
    if len(argv if argv is not None else sys.argv) != 4:
        sys.exit(1)
    payload_path = Path((argv if argv is not None else sys.argv)[1])
    template_path = Path((argv if argv is not None else sys.argv)[2])
    output_path = Path((argv if argv is not None else sys.argv)[3])
    payload = load_payload(payload_path)

    render_commercial_report(
        payload=payload,
        template_path=template_path,
        output_path=output_path,
    )
    print(f"Account Action Plan renderizado: {output_path}")


if __name__ == "__main__":
    main()
