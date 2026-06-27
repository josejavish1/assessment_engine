import json
import os
import sys
import tempfile
import uuid
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from assessment_engine.infrastructure.text_utils import clean_text_for_word

# Defines global constants for branding and styling to centralize document format control.
COLOR_BLUE = (
    "0072BC"  # Provides a fallback value if the primary data retrieval operation fails.
)
COLOR_HEADER_BG = "D9EAF7"
COLOR_ROW_ALT = "F2F2F2"


def shade_cell(cell, color_hex):
    """Applies a solid background color fill to a table cell.

    This function directly manipulates the underlying OOXML of the cell's
    properties (`w:tcPr`) by adding a `w:shd` (shading) element. This
    provides a low-level mechanism for cell shading.

    Args:
        cell (docx.table._Cell): The table cell object to modify.
        color_hex (str): The background color specified as a six-character
            hexadecimal RGB string (e.g., 'FFC000').
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_text_custom(
    cell, text, bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=None
):
    """Set the text content and apply custom formatting to a table cell."""
    set_cell_text_custom_impl(cell, text, bold, font_size, align, color_rgb)


def set_cell_text_custom_impl(cell, text, bold, font_size, align, color_rgb):
    r"""{'docstring': "Sets the text and formatting for a `python-docx` table cell.\n\nThis function first erases all existing content within the target cell. It\nthen configures the cell's primary paragraph with specified alignment,\nspacing, and line height. A new run is added to this paragraph containing\nthe provided text, which is sanitized to ensure XML compatibility before\ninsertion. Finally, font properties such as bold style, size, and color\nare applied to the new run.\n\nArgs:\n    cell (docx.table._Cell): The `_Cell` object from a `python-docx` table\n        that will be modified.\n    text (str): The string content to insert into the cell.\n    bold (bool): A flag to determine if the text should be bold.\n    font_size (Union[int, float]): The desired font size in points.\n    align (int): A member of the `docx.enum.text.WD_ALIGN_PARAGRAPH`\n        enumeration used to set the paragraph's horizontal alignment.\n    color_rgb (Optional[docx.shared.RGBColor]): An `RGBColor` instance for the\n        font color. If None, the document's default font color is applied."}."""
    # Ensures content is XML-escaped for structural integrity within the document.
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(clean_text_for_word(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    if color_rgb:
        run.font.color.rgb = color_rgb


def autofit_table_to_contents(table):
    """Set a `python-docx` table to span 100% container width with autofitting.

    This function modifies a table's properties to ensure it occupies the full
    width of its container. It enables the high-level `autofit` and
    `allow_autofit` properties. Concurrently, it manipulates the underlying
    OOXML structure by injecting a `w:tblW` element into the table's properties
    (`tblPr`). This element explicitly sets the table width type to percentage
    (`pct`) and the value to '5000', which corresponds to 100% in the OOXML
    specification (5000 units of 1/50th of a percent).

    Args:
        table (docx.table.Table): The `python-docx` table object to modify.

    Returns:
        None. The table object is modified in-place.
    """
    table.autofit = True
    table.allow_autofit = True
    tblPr = table._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(
        qn("w:w"), "5000"
    )  # The OOXML table width unit is 1/50th of a percent. 5000 corresponds to 100%.
    tblPr.append(tblW)


def add_body_paragraph(
    doc,
    text,
    bold=False,
    italic=False,
    space_after=6,
    text_color_rgb=(46, 64, 77),
    style="Normal",
) -> Any:
    """Appends a paragraph with specialized formatting to a `docx.Document` object.

    This function creates a paragraph, applying justification, a fixed 10pt font
    size, and configurable text color and spacing. It supports conditional
    bolding logic that overrides the global `bold` argument:

    1.  **List Bullet Style**: For paragraphs assigned the 'List Bullet' style,
        text preceding the first colon (':') is automatically bolded.
    2.  **Markdown-style Bold**: Text enclosed in double asterisks (`**text**`)
        is rendered in bold, allowing for multiple bolded segments.

    If a specified paragraph style does not exist in the document template, the
    function gracefully falls back to the 'Normal' style.

    Args:
        doc (docx.document.Document): The document instance to which the
            paragraph will be added.
        text (str): The text content of the paragraph.
        bold (bool, optional): If True, applies bold formatting to the entire
            paragraph. This is ignored if 'List Bullet' or markdown-style
            bolding rules are triggered. Defaults to False.
        italic (bool, optional): If True, applies italic formatting. This is
            only applied when the text is added as a single run (i.e., no
            special bolding rules are triggered). Defaults to False.
        space_after (int, optional): The spacing, in points, after the
            paragraph. Defaults to 6.
        text_color_rgb (Tuple[int, int, int], optional): An RGB tuple for the
            font color. Defaults to (46, 64, 77).
        style (str, optional): The name of the paragraph style to apply.
            Defaults to 'Normal'.

    Returns:
        docx.text.paragraph.Paragraph: The paragraph object added to the document.

    Raises:
        ValueError: If `text_color_rgb` contains values outside the 0-255 range,
            as raised by the underlying `python-docx` library.
    """
    try:
        p = doc.add_paragraph(style=style)
        is_bullet = style == "List Bullet"
    except KeyError:
        p = doc.add_paragraph(style="Normal")
        is_bullet = style == "List Bullet"
        if is_bullet:
            run_b = p.add_run("• ")
            run_b.font.name = "Arial"
            run_b.font.size = Pt(10)
            run_b.font.color.rgb = RGBColor(*text_color_rgb)

    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Applies bold formatting to the first 3-4 words of bullet points to meet a specified document design standard.
    if is_bullet and ":" in text:
        parts = text.split(":", 1)
        run_bold = p.add_run(parts[0] + ":")
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_bold.font.color.rgb = RGBColor(*text_color_rgb)

        run_rest = p.add_run(parts[1])
        run_rest.font.size = Pt(10)
        run_rest.font.color.rgb = RGBColor(*text_color_rgb)
    elif "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 1:
                run = p.add_run(part)
                run.font.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*text_color_rgb)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*text_color_rgb)
    return p


def add_heading(doc, text, level, primary_color_rgb=(0, 114, 188)):
    """Adds a styled heading to a `python-docx` document object.

    Creates a heading paragraph with specified text and level, applying custom
    paragraph formatting (space before/after, keep with next) and run-level
    formatting (font color and size). Font size is determined by the heading
    level: 16pt for level 1, 14pt for 2, 12pt for 3, and 11pt otherwise.

    Args:
        doc (docx.document.Document): The document object to which the heading
            will be added.
        text (str): The text content for the heading.
        level (int): The heading level, where 0 is a title and 1-9 are
            standard heading levels.
        primary_color_rgb (tuple[int, int, int]): A tuple of three integers
            (0-255) for the RGB font color. Defaults to (0, 114, 188).

    Returns:
        docx.text.paragraph.Paragraph: The newly created and styled heading
            paragraph object.

    Raises:
        ValueError: If `level` is outside the valid range (0-9) or if any value
            in `primary_color_rgb` is outside the range (0-255), as enforced
            by the underlying `python-docx` library.
    """
    h = doc.add_heading(text, level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True

    # Applies specified color and size; strips numeric prefixes to use the renderer's native list numbering.
    for run in h.runs:
        run.font.color.rgb = RGBColor(*primary_color_rgb)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
    return h


def add_toc(doc):
    r"""{'docstring': 'Inserts an OOXML field code for a Table of Contents into a document.\n\nThis function directly manipulates the underlying Office Open XML (OOXML)\nby injecting the elements that constitute a Word field code. The specific\ninstruction text inserted is `TOC \\o "1-3" \\h \\z \\u`, which directs a\nword processor to generate a Table of Contents from headings of levels 1\nthrough 3, with entries formatted as hyperlinks.\n\nNote:\n    The TOC is not populated by this function. A word processing\n    application, such as Microsoft Word, must open the document and update\n    its fields to render the final Table of Contents.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` Document object to which\n        the TOC field will be added.\n\nReturns:\n    None: The function modifies the `doc` object in-place.\n\nRaises:\n    AttributeError: If `doc` is not a valid `python-docx` Document object\n        and lacks the required methods.'}."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)


def generate_radar_chart(
    labels: list, actual_scores: list, target_scores: list, title: str, output_path: str
):
    """Generate and save a radar chart visualizing actual vs. target scores.

    Constructs a polar plot (radar chart) using matplotlib to compare two sets
    of scores ('actual' and 'target') across multiple categorical dimensions.
    The function automatically wraps long labels for readability and closes the
    polygon shape by repeating the first data point. The final chart is saved as
    a high-resolution (300 DPI) PNG image.

    Args:
        labels (list[str]): The categorical labels for each axis of the chart.
        actual_scores (list[int | float]): Numerical scores representing the "AS-IS"
            state. Must have the same length as 'labels'.
        target_scores (list[int | float]): Numerical scores representing the "TO-BE"
            state. Must have the same length as 'labels'.
        title (str): The main title displayed above the chart.
        output_path (str): The full file path, including the '.png' extension,
            where the chart will be saved.

    Returns:
        bool: True if the chart was successfully generated and saved, or False if
            the 'labels' list is empty.

    Raises:
        ValueError: If the lengths of 'actual_scores' or 'target_scores' do not
            match the length of 'labels'.
        FileNotFoundError: If the directory specified in 'output_path' does not
            exist.
        PermissionError: If write permissions are denied for the specified
            'output_path'.
    """
    if not labels:
        return False
    num_vars = len(labels)
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()

    actual_scores += actual_scores[:1]
    target_scores += target_scores[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 5), subplot_kw=dict(polar=True))

    #
    ax.plot(
        angles,
        target_scores,
        color="#0072BC",
        linewidth=2,
        linestyle="solid",
        label="Objetivo (TO-BE)",
    )
    ax.fill(angles, target_scores, color="#0072BC", alpha=0.1)

    #
    ax.plot(
        angles,
        actual_scores,
        color="#C00000",
        linewidth=2,
        linestyle="solid",
        label="Actual (AS-IS)",
    )
    ax.fill(angles, actual_scores, color="#C00000", alpha=0.25)

    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)

    import textwrap

    wrapped_labels = [textwrap.fill(lbl, 15) for lbl in labels]
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(wrapped_labels, size=8)

    ax.set_ylim(0, 5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(["1", "2", "3", "4", "5"], color="grey", size=7)

    ax.set_title(title, size=11, weight="bold", pad=15)
    ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.1))

    plt.tight_layout()
    plt.savefig(output_path, format="png", dpi=300)
    plt.close()
    return True


def render_consolidated_tobe(ast_json_path: str, output_path: str):
    """Renders a consolidated TO-BE architecture report into a Microsoft Word document.

    This function orchestrates the creation of a multi-section Word document from
    a structured Abstract Syntax Tree (AST) defined in a JSON file. It enriches
    the core AST data by integrating external configurations for branding,
    internationalization (i18n), and client-specific strategic intelligence.
    The process involves loading a pre-styled Word template and programmatically
    populating it with content, including an executive summary, maturity model
    analysis with dynamically generated radar charts, detailed breakdowns by
    technology domain, a high-level transformation roadmap, and strategic
    conclusions.

    Args:
        ast_json_path (str): The file path to the input JSON file containing
            the document's Abstract Syntax Tree (AST). The file must exist
            and conform to the expected schema.
        output_path (str): The file path where the generated Word (.docx)
            document will be saved. The parent directory is created if it
            does not exist.

    Raises:
        json.JSONDecodeError: If the main AST file or any auxiliary JSON
            configuration files are malformed or cannot be parsed.
        KeyError: If the input AST or its nested data structures are missing
            mandatory keys that are essential for the rendering process (e.g.,
            'tower_name', 'score', 'pillars' in a domain definition).
    """
    print("🚀 Cargando AST de documento TO-BE para Renderizado...")

    if not os.path.exists(ast_json_path):
        print(f"❌ Error: {ast_json_path} no existe.")
        return

    with open(ast_json_path, "r", encoding="utf-8-sig") as f:
        ast = json.load(f)

    client_name = ast.get("metadata", {}).get("client_name", "Cliente")
    meta_lang = ast.get("metadata", {}).get("language", "es").lower()

    # Loads configurations for branding and i18n to decouple presentation from logic.
    from assessment_engine.infrastructure.config_loader import load_brand_profile

    brand = load_brand_profile()
    styling = brand.get("styling", {})
    color_blue = styling.get("primary_color_hex", "0072BC")

    from assessment_engine.infrastructure.config_loader import (
        resolve_localized_vocabulary,
    )

    vocab = resolve_localized_vocabulary(meta_lang)
    global_sum = ast.get("global_summary", {})
    towers = ast.get("towers", [])

    # Loads and integrates client-specific strategic intelligence data.
    working_dir = Path(ast_json_path).parent
    client_intel_path = working_dir / "client_intelligence.json"
    intel = {}
    if client_intel_path.exists():
        with open(client_intel_path, "r", encoding="utf-8-sig") as cif:
            try:
                intel = json.load(cif)
                print(
                    "   ├─ Carga exitosa de Inteligencia Estratégica del Cliente en Nivel Consolidado."
                )
            except Exception as e:
                print(f"   ⚠️ Error cargando inteligencia: {e}")

    business_context = intel.get("business_context", {})
    ceo_agenda_raw = business_context.get("ceo_agenda", {}).get(
        "summary", "No disponible."
    )

    # Ensure absolute path resolution to support execution from any directory context.
    from assessment_engine.infrastructure.runtime_paths import ROOT

    template_path = ROOT / "templates" / "docx" / "template_tobe_consolidated.docx"
    if template_path.exists():
        doc = Document(str(template_path))
        print("   ├─ Plantilla Word pre-estilizada cargada correctamente.")
    else:
        doc = Document()
        print("   ⚠️ Plantilla no encontrada. Generando con estilos por defecto.")

    # Post-processing step: Removes all placeholder elements from the document to produce the final output.
    for p in list(doc.paragraphs):
        p._p.getparent().remove(p._p)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    #
    #
    #
    doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Converts the hexadecimal brand color string to an RGB tuple required by the rendering library.
    p_color_hex = color_blue
    r_color = int(p_color_hex[0:2], 16)
    g_color = int(p_color_hex[2:4], 16)
    b_color = int(p_color_hex[4:6], 16)
    p_color_rgb = (r_color, g_color, b_color)

    for _ in range(5):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run(
        vocab.get("conclusions_title", "Informe Consolidado de Situación Futura")
    )
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(*p_color_rgb)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(40)
    run_sub = p_sub.add_run(
        "Tier-1 Infrastructure Modernisation & Resilience Assessment"
    )
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(128, 128, 128)

    p_client = doc.add_paragraph()
    run_client = p_client.add_run(f"Preparado para: {client_name}")
    run_client.font.size = Pt(12)
    run_client.font.bold = True

    # Applies internationalization (i18n) and branding transformations based on the loaded configuration profile.
    company = brand.get("company_name", "NTT DATA")
    classification = brand.get("default_classification", "Confidencial")

    # Retrieves the document version identifier from the main configuration.
    doc_version = ast.get("metadata", {}).get("version", "1.0")

    # Formats the current system date according to the configured i18n locale.
    doc_date = ast.get("metadata", {}).get("date")
    if not doc_date:
        import datetime

        now = datetime.datetime.now()
        months_es = [
            "Enero",
            "Febrero",
            "Marzo",
            "Abril",
            "Mayo",
            "Junio",
            "Julio",
            "Agosto",
            "Septiembre",
            "Octubre",
            "Noviembre",
            "Diciembre",
        ]
        months_en = [
            "January",
            "February",
            "March",
            "April",
            "May",
            "June",
            "July",
            "August",
            "September",
            "October",
            "November",
            "December",
        ]
        months_pt = [
            "Janeiro",
            "Fevereiro",
            "Março",
            "Abril",
            "Maio",
            "Junho",
            "Julho",
            "Agosto",
            "Setembro",
            "Outubro",
            "Novembro",
            "Dezembro",
        ]
        months_fr = [
            "Janvier",
            "Février",
            "Mars",
            "Avril",
            "Mai",
            "Juin",
            "Juillet",
            "Août",
            "Septembre",
            "Octobre",
            "Novembre",
            "Décembre",
        ]

        if meta_lang == "es":
            doc_date = f"{months_es[now.month - 1]} {now.year}"
        elif meta_lang == "pt":
            doc_date = f"{months_pt[now.month - 1]} {now.year}"
        elif meta_lang == "fr":
            doc_date = f"{months_fr[now.month - 1]} {now.year}"
        elif meta_lang == "ja":
            doc_date = f"{now.year}年{now.month}月"
        else:
            doc_date = f"{months_en[now.month - 1]} {now.year}"

    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run(
        f"{company} | {classification}\n{vocab.get('version_label', 'Versión')} {doc_version} | {doc_date}"
    )
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True

    doc.add_page_break()

    #
    add_heading(
        doc,
        vocab.get("table_of_contents_title", "Índice de Contenidos"),
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "Nota: Para actualizar el índice en Word, haga clic derecho sobre el texto inferior y seleccione 'Actualizar campos'.",
        italic=True,
    )
    add_toc(doc)
    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        vocab.get("asis_consolidated_title", "Resumen Ejecutivo"),
        level=1,
        primary_color_rgb=p_color_rgb,
    )

    add_heading(
        doc,
        vocab.get("asis_consolidated_title", "Visión de Estado Objetivo"),
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(doc, global_sum.get("executive_vision", "No definida."))

    add_heading(
        doc,
        "Estrategia de Transformación de la Plataforma",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(doc, global_sum.get("transformation_strategy", "No definida."))

    # Integrates content from the Redeia Leadership Agenda and Investment Plan for contextual enrichment.
    if ceo_agenda_raw and ceo_agenda_raw != "No disponible.":
        add_heading(
            doc,
            "Por qué importa al negocio (Agenda de Liderazgo)",
            level=2,
            primary_color_rgb=p_color_rgb,
        )
        for block in ceo_agenda_raw.split("\n\n"):
            if block.strip().startswith("###"):
                add_heading(
                    doc,
                    block.replace("###", "").strip(),
                    level=3,
                    primary_color_rgb=p_color_rgb,
                )
            elif block.strip().startswith("####"):
                add_heading(
                    doc,
                    block.replace("####", "").strip(),
                    level=4,
                    primary_color_rgb=p_color_rgb,
                )
            elif block.strip().startswith("1.") or block.strip().startswith("-"):
                add_body_paragraph(doc, block.strip(), style="List Bullet")
            else:
                add_body_paragraph(doc, block.strip())

    add_heading(
        doc, "Pilares Estratégicos de Evolución", level=2, primary_color_rgb=p_color_rgb
    )
    for lever in global_sum.get("strategic_levers", []):
        add_body_paragraph(doc, lever, style="List Bullet")

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        "Propósito, alcance y enfoque del documento",
        level=1,
        primary_color_rgb=p_color_rgb,
    )

    add_heading(doc, "Propósito del documento", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, global_sum.get("document_purpose", "No definido."))

    add_heading(doc, "Alcance", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, global_sum.get("document_scope", "No definido."))

    add_heading(doc, "Enfoque y criterios", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(doc, global_sum.get("maturity_approach", "No definido."))

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc, "Modelo de madurez objetivo", level=1, primary_color_rgb=p_color_rgb
    )

    add_heading(
        doc,
        "Principios de evolución de la madurez",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "La evolución tecnológica y operativa para alcanzar el nivel objetivo se rige por principios inquebrantables de resiliencia proactiva:",
    )
    add_body_paragraph(
        doc,
        "**Resiliencia y Ciber-Resiliencia Primero:** Las inversiones se priorizan para erradicar puntos de fallo únicos y asegurar la continuidad sobre la reducción de costes inmediatos.",
        style="List Bullet",
    )
    add_body_paragraph(
        doc,
        "**Automatización Extrema:** La dependencia del conocimiento táctico y de la operación manual debe eliminarse mediante Runbooks-as-Code.",
        style="List Bullet",
    )
    add_body_paragraph(
        doc,
        "**Soberanía y Evidencia:** Todo salto de madurez debe ser auditable y sustentado por pruebas continuas basadas en evidencias empíricas (Pruebas de Resiliencia).",
        style="List Bullet",
    )

    add_heading(
        doc,
        "Estrategia de la resiliencia geográfica",
        level=2,
        primary_color_rgb=p_color_rgb,
    )

    add_heading(
        doc,
        "Estrategia de emplazamiento de CPD",
        level=3,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(doc, global_sum.get("cpd_strategy_3_years", "No definida."))
    add_body_paragraph(doc, global_sum.get("cpd_strategy_5_years", "No definida."))

    add_heading(
        doc,
        "Modelo operativo dual-site avanzado",
        level=3,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(doc, global_sum.get("active_active_transition", "No definida."))

    add_heading(
        doc,
        "Madurez actual vs madurez objetivo",
        level=2,
        primary_color_rgb=p_color_rgb,
    )

    # Generates a UUID for the global radar chart image to ensure a unique filename and prevent filesystem collisions.
    global_labels = [t["tower_name"] for t in towers]
    global_actual = [t["score"] for t in towers]
    global_target = [t["target"] for t in towers]
    if global_labels:
        chart_path = (
            Path(tempfile.gettempdir()) / f"global_radar_tobe_{uuid.uuid4().hex}.png"
        )
        if generate_radar_chart(
            global_labels,
            global_actual,
            global_target,
            "Comparativa de Madurez Global (AS-IS vs TO-BE)",
            str(chart_path),
        ):
            doc.add_picture(str(chart_path), width=Inches(5.0))
            p_img = doc.paragraphs[-1]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                os.remove(chart_path)
            except OSError:
                pass

    add_heading(
        doc,
        "Matriz Resumen de Evolución por Dominio",
        level=3,
        primary_color_rgb=p_color_rgb,
    )
    mat_table = doc.add_table(rows=1, cols=4)
    mat_table.style = "Table Grid"

    headers = [
        "Dominio Tecnológico",
        "Situación Actual (AS-IS)",
        "Objetivo a 3 años",
        "Visión a 5 años",
    ]
    for i, header in enumerate(headers):
        set_cell_text_custom(
            mat_table.rows[0].cells[i],
            header,
            bold=True,
            font_size=9,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(mat_table.rows[0].cells[i], COLOR_BLUE)

    for r_idx, t in enumerate(towers):
        row = mat_table.add_row()
        score_text = f"Nivel {t['score']:.1f}"
        target_text = f"Nivel {t['target']:.1f}"

        # Aggregates data from multiple pillar-specific sources into a canonical structure for analysis and rendering.
        asis_bullets = []
        v3_bullets = []
        v5_bullets = []
        for p in t["pillars"]:
            asis_bullets.append(p["pilar_name"])
            v3_bullets.append(f"{p['pilar_name']}: {p['vision_3_years'][:250]}...")
            v5_bullets.append(f"{p['pilar_name']}: {p['vision_5_years'][:250]}...")

        set_cell_text_custom(
            row.cells[0],
            f"{t['tower_name']}\n({score_text} ➔ {target_text})",
            bold=True,
            font_size=8,
        )
        set_cell_text_custom(
            row.cells[1], "\n".join([f"• {b}" for b in asis_bullets]), font_size=8
        )
        set_cell_text_custom(
            row.cells[2], "\n".join([f"• {b}" for b in v3_bullets]), font_size=8
        )
        set_cell_text_custom(
            row.cells[3], "\n".join([f"• {b}" for b in v5_bullets]), font_size=8
        )

        # Applies alternating row colors (row banding) to improve table readability.
        if r_idx % 2 == 1:
            for cell in row.cells:
                shade_cell(cell, COLOR_ROW_ALT)

    autofit_table_to_contents(mat_table)
    doc.add_paragraph()  #

    add_heading(
        doc,
        "Relación con la modernización de aplicaciones",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc, global_sum.get("app_modernization_relation", "No definida.")
    )

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        "Modelo TO-BE por dominio tecnológico",
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "Este capítulo desglosa en detalle el estado objetivo, las palancas de cambio requeridas y los riesgos de inacción para cada dominio de infraestructura evaluado:",
    )

    for t_idx, t in enumerate(towers):
        add_heading(
            doc, f"Dominio: {t['tower_name']}", level=2, primary_color_rgb=p_color_rgb
        )

        # Generates a UUID for each radar chart image to ensure a unique filename and prevent filesystem collisions.
        p_labels = [p["pilar_name"] for p in t["pillars"]]
        p_actual = [p["score"] for p in t["pillars"]]
        p_target = [p["target_score"] for p in t["pillars"]]
        if p_labels:
            chart_path = (
                Path(tempfile.gettempdir())
                / f"radar_tobe_{t['tower_name'][:10].replace(' ', '_')}_{uuid.uuid4().hex}.png"
            )
            if generate_radar_chart(
                p_labels,
                p_actual,
                p_target,
                f"Evolución de Madurez: {t['tower_name']}",
                str(chart_path),
            ):
                doc.add_picture(str(chart_path), width=Inches(4.5))
                p_img = doc.paragraphs[-1]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    os.remove(chart_path)
                except OSError:
                    pass

        for p in t["pillars"]:
            add_heading(doc, p["pilar_name"], level=3, primary_color_rgb=p_color_rgb)

            add_heading(
                doc,
                "Situación Actual (AS-IS Resumido)",
                level=4,
                primary_color_rgb=p_color_rgb,
            )
            add_body_paragraph(doc, p["asis_description"])

            add_heading(
                doc,
                "Objetivo a 3 años (Nivel 5 Básico)",
                level=4,
                primary_color_rgb=p_color_rgb,
            )
            add_body_paragraph(doc, p["vision_3_years"])

            add_heading(
                doc,
                "Visión Aspiracional a 5 años",
                level=4,
                primary_color_rgb=p_color_rgb,
            )
            add_body_paragraph(doc, p["vision_5_years"])

            add_heading(
                doc,
                "Palancas Principales de Evolución",
                level=4,
                primary_color_rgb=p_color_rgb,
            )

            #
            doc.add_heading("Palancas Tecnológicas:", level=5)
            tech = p["levers_technology"]
            if not tech:
                add_body_paragraph(doc, "N/A")
            for item in tech:
                add_body_paragraph(doc, item, style="List Bullet")

            doc.add_heading("Palancas de Procesos:", level=5)
            proc = p["levers_process"]
            if not proc:
                add_body_paragraph(doc, "N/A")
            for item in proc:
                add_body_paragraph(doc, item, style="List Bullet")

            doc.add_heading("Palancas de Operación:", level=5)
            op = p["levers_operation"]
            if not op:
                add_body_paragraph(doc, "N/A")
            for item in op:
                add_body_paragraph(doc, item, style="List Bullet")

            add_heading(
                doc,
                "Beneficios Esperados en el Nuevo Modelo",
                level=4,
                primary_color_rgb=p_color_rgb,
            )
            bens = p["expected_benefits"]
            if not bens:
                add_body_paragraph(doc, "N/A")
            for item in bens:
                add_body_paragraph(doc, item, style="List Bullet")

            add_heading(
                doc,
                "Riesgos si no se actúa y dependencias clave (Cost of Inaction)",
                level=4,
                primary_color_rgb=p_color_rgb,
            )
            coi = p["cost_of_inaction_risks"]
            if not coi:
                add_body_paragraph(doc, "N/A")
            for item in coi:
                add_body_paragraph(doc, item, style="List Bullet")

            doc.add_paragraph()  #

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        "Roadmap de alto nivel y horizontes temporales",
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "La transformación tecnológica no se plantea como un macroproyecto único, sino como una trayectoria escalonada y secuencial basada en dependencias técnicas lógicas:",
    )

    roadmap = global_sum.get("roadmap", {})

    add_heading(
        doc,
        "Corto plazo (0-12 meses): Estabilización y reducción de riesgos críticos",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    for r in roadmap.get("phase_1_short_term", []):
        add_body_paragraph(doc, r, style="List Bullet")

    add_heading(
        doc,
        "Medio plazo (Años 2-3): Consolidación del nivel 5 básico",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    for r in roadmap.get("phase_2_mid_term", []):
        add_body_paragraph(doc, r, style="List Bullet")

    add_heading(
        doc,
        "Largo plazo (Años 4-5): Consolidación del nivel 5 y preparación de la evolución futura",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    for r in roadmap.get("phase_3_long_term", []):
        add_body_paragraph(doc, r, style="List Bullet")

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        "Beneficios esperados y posicionamiento futuro",
        level=1,
        primary_color_rgb=p_color_rgb,
    )

    add_heading(
        doc,
        "Beneficios globales y transversales",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    for b in global_sum.get("global_benefits", []):
        add_body_paragraph(doc, b, style="List Bullet")

    add_heading(
        doc,
        "Impacto sobre la misión estratégica y posicionamiento corporativo",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(doc, global_sum.get("public_mission_impact", "No definido."))

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        "Supuestos y riesgos del Modelo TO-BE",
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "El éxito de la arquitectura objetivo descansa sobre varios supuestos estructurales que requieren una participación activa de la dirección del cliente:",
    )

    add_heading(
        doc,
        "Supuestos estructurales y organizativos",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    for a in global_sum.get("assumptions_structural", []):
        add_body_paragraph(doc, a, style="List Bullet")

    add_heading(
        doc,
        "Supuestos tecnológicos y de compatibilidad de arquitectura",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    for a in global_sum.get("assumptions_technological", []):
        add_body_paragraph(doc, a, style="List Bullet")

    add_heading(
        doc,
        "Supuestos sobre el entorno regulatorio y de cumplimiento",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    for a in global_sum.get("assumptions_regulatory", []):
        add_body_paragraph(doc, a, style="List Bullet")

    add_heading(
        doc,
        "Riesgos si no se actúa (Cost of Inaction a nivel corporativo)",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    for r in global_sum.get("risks_if_not_acted", []):
        add_body_paragraph(doc, r, style="List Bullet")

    add_heading(
        doc,
        "Riesgos durante la implantación del plan",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    for r in global_sum.get("risks_of_implementation", []):
        add_body_paragraph(doc, r, style="List Bullet")

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc, "Conclusiones y Siguientes pasos", level=1, primary_color_rgb=p_color_rgb
    )
    add_body_paragraph(
        doc,
        "El paso hacia el modelo de resiliencia TO-BE es un imperativo estratégico. Se proponen las siguientes acciones inmediatas de gobernanza para iniciar el plan:",
    )

    for step in global_sum.get("next_steps", []):
        add_body_paragraph(doc, step, style="List Bullet")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"🎉 ¡Documento Global TO-BE COMPLETO generado con éxito en: {output_path}!")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(
            "Uso: python render_togaf_consolidated_tobe.py <working_dir/tobe_master_ast.json> <output_doc.docx>"
        )
        sys.exit(1)
    render_consolidated_tobe(sys.argv[1], sys.argv[2])
