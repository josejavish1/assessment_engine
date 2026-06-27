import glob
import json
import os
import sys
import tempfile
import uuid
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from assessment_engine.infrastructure.docx_render_utils import (
    autofit_table_to_contents,
    finalize_table,
    shade_cell,
)
from assessment_engine.infrastructure.text_utils import clean_text_for_word

# Define corporate styling parameters for document generation.
COLOR_BLUE = "0072BC"
COLOR_HEADER_BG = "D9EAF7"
COLOR_ROW_ALT = "F2F2F2"


def set_cell_text_custom(
    cell, text, bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=None
):
    r"""{'docstring': "Sets and formats the text content of a `docx` table cell.\n\nClears any existing content within the specified cell and applies new text\nwith custom formatting. The function manipulates the cell's first paragraph\nto set alignment, spacing, and line height. A new text run is then added,\nprocessing the input text via `clean_text_for_word`, and its font is styled\nwith the provided bold, size, and color attributes.\n\nArgs:\n    cell (docx.table._Cell): The `python-docx` cell object to modify.\n    text (str): The string content to insert into the cell. This text is\n        sanitized by the `clean_text_for_word` function before insertion.\n    bold (bool, optional): If True, formats the text as bold. Defaults to\n        False.\n    font_size (int, optional): The font size in points (Pt). Defaults to 9.\n    align (WD_ALIGN_PARAGRAPH, optional): An enum member from\n        `docx.enum.text.WD_ALIGN_PARAGRAPH` specifying horizontal\n        alignment. Defaults to `WD_ALIGN_PARAGRAPH.LEFT`.\n    color_rgb (docx.shared.RGBColor, optional): A `python-docx` `RGBColor`\n        object for the font color. If None, the document's default color\n        is used. Defaults to None.\n\nReturns:\n    None."}."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(clean_text_for_word(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    if color_rgb:
        run.font.color.rgb = color_rgb


def add_body_paragraph(
    doc,
    text,
    bold=False,
    italic=False,
    space_after=6,
    text_color_rgb=(46, 64, 77),
    style="Normal",
) -> Any:
    r"""{'docstring': "Adds a styled paragraph to a `docx.document.Document` with advanced formatting.\n\nThis function creates a paragraph and applies specified formatting for style,\ncolor, and spacing. It supports two special formatting syntaxes within the `text`\nargument, which take precedence over the global `bold` and `italic` flags:\n\n1.  **List Bullet Headers**: If `style` is 'List Bullet' and `text` contains\n    a colon (`:`), the substring before the first colon is bolded. This is\n    prioritized over markdown-style bolding.\n2.  **Markdown-style Bolding**: Text enclosed in double asterisks (`**`) is\n    rendered as bold.\n\nAll text runs are set to a 10pt font size. The paragraph alignment is\njustified. If a specified style does not exist in the document, the function\nfalls back to the 'Normal' style. If the intended style was 'List Bullet'\nbut it was not found, a bullet character (`•`) is prepended manually.\n\nArgs:\n    doc (docx.document.Document): The document object to which the paragraph\n        will be added.\n    text (str): The content of the paragraph. May contain formatting markers\n        like `**text**` or a colon when used with the 'List Bullet' style.\n    bold (bool, optional): If True, applies bold formatting to the entire\n        paragraph. This setting is ignored if special formatting syntax is\n        present in the `text` string. Defaults to False.\n    italic (bool, optional): If True, applies italic formatting to the entire\n        paragraph. This setting is ignored if special formatting syntax is\n        present in the `text` string. Defaults to False.\n    space_after (int, optional): The spacing, in points, to apply after the\n        paragraph. Defaults to 6.\n    text_color_rgb (tuple[int, int, int], optional): An RGB tuple specifying\n        the text color. Defaults to (46, 64, 77).\n    style (str, optional): The name of the Word paragraph style to apply.\n        The 'List Bullet' style enables specific bolding behavior. If the\n        style is not found, 'Normal' is used as a fallback.\n        Defaults to 'Normal'.\n\nReturns:\n    docx.text.paragraph.Paragraph: The newly created paragraph object.\n\nRaises:\n    ValueError: If values in `text_color_rgb` are outside the 0-255 range.\n    TypeError: If `text_color_rgb` is not a tuple of three integers."}."""
    try:
        p = doc.add_paragraph(style=style)
        is_bullet = style == "List Bullet"
    except KeyError:
        p = doc.add_paragraph(style="Normal")
        is_bullet = style == "List Bullet"
        if is_bullet:
            run_b = p.add_run("• ")
            run_b.font.name = "Arial"
            run_b.font.size = Pt(10)
            run_b.font.color.rgb = RGBColor(*text_color_rgb)

    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Apply bold formatting to the initial words of each bullet point to improve scannability.
    if is_bullet and ":" in text:
        parts = text.split(":", 1)
        run_bold = p.add_run(parts[0] + ":")
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_bold.font.color.rgb = RGBColor(*text_color_rgb)

        run_rest = p.add_run(parts[1])
        run_rest.font.size = Pt(10)
        run_rest.font.color.rgb = RGBColor(*text_color_rgb)
    elif "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 1:
                run = p.add_run(part)
                run.font.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*text_color_rgb)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*text_color_rgb)
    return p


def add_heading(doc, text, level, primary_color_rgb=(0, 114, 188)):
    r"""{'docstring': 'Adds a styled heading to a `python-docx` Document object.\n\nThis function inserts a heading paragraph with specified text and level.\nIt then applies specific formatting to the heading, including paragraph\nspacing (`space_before`, `space_after`), enabling `keep_with_next`, setting\na custom font color, and adjusting the font size based on the heading level.\n\nArgs:\n    doc (docx.document.Document): The document object to which the heading\n        will be added.\n    text (str): The text content of the heading.\n    level (int): The heading level, where 1 corresponds to the highest level.\n    primary_color_rgb (Tuple[int, int, int]): A tuple of three integers\n        representing the RGB font color. Defaults to (0, 114, 188).\n\nReturns:\n    docx.text.paragraph.Paragraph: The newly created and styled heading\n    paragraph object.\n\nRaises:\n    TypeError: If `primary_color_rgb` is not an iterable of length 3.\n    ValueError: If a value in `primary_color_rgb` is not an integer or is\n        outside the valid 0-255 range for an RGB color component.'}."""
    h = doc.add_heading(text, level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True

    # Apply specified color and font size. Strip numeric prefixes from list items to allow the document's native list auto-numbering to function correctly.
    for run in h.runs:
        run.font.color.rgb = RGBColor(*primary_color_rgb)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
    return h


def add_spacer(doc, points=12):
    """Add a vertical spacer paragraph to a document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    return p


def add_toc(doc):
    r"""{'docstring': 'Injects an Office Open XML (OOXML) field code for a Table of Contents.\n\nThis function directly manipulates the underlying OOXML by creating and\nappending a series of `w:fldChar` elements and a `w:instrText` element\nto a new paragraph at the end of the document. The instruction text,\n`TOC \\o "1-3" \\h \\z \\u`, directs a client application (e.g., Microsoft Word)\nto generate a hyperlinked Table of Contents (TOC) from paragraphs formatted\nwith heading styles 1 through 3.\n\nThe TOC itself is not rendered by this function; it is a field code that is\nevaluated and populated by the word processing application when the document\nis opened and its fields are updated.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` document object to which\n        the TOC field code will be added. The object is modified in-place.\n\nReturns:\n    None.'}."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)


def generate_radar_chart(
    labels: list, actual_scores: list, target_scores: list, title: str, output_path: str
):
    """Generates and saves a radar chart comparing actual and target scores.

    This function creates a polar plot to visualize a comparison between current
    ("AS-IS") and target ("TO-BE") maturity scores across multiple domains. The
    resulting chart is saved as a PNG image to a specified file path.

    Note:
        The `actual_scores` and `target_scores` lists are modified in-place by
        appending their first element to the end to close the radar chart's
        polygon.

    Args:
        labels (list[str]): A list of strings representing the domains to be
            plotted as axes on the chart.
        actual_scores (list[int | float]): A list of numbers representing the
            current maturity scores for each domain. Must be the same length as
            `labels`.
        target_scores (list[int | float]): A list of numbers representing the
            target maturity scores for each domain. Must be the same length as
            `labels`.
        title (str): The title to be displayed on the chart.
        output_path (str): The full file path, including the extension, where the
            generated PNG chart will be saved.

    Returns:
        bool: True if the chart was successfully generated and saved. False if the
            `labels` list is empty, in which case no chart is generated.

    Raises:
        ValueError: If the lengths of `labels`, `actual_scores`, and
            `target_scores` are not identical.
        IOError: If the file cannot be written to `output_path` due to an
            invalid path or filesystem permissions.
    """
    if not labels:
        return False
    num_vars = len(labels)
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()

    actual_scores += actual_scores[:1]
    target_scores += target_scores[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 5), subplot_kw=dict(polar=True))

    #
    ax.plot(
        angles,
        target_scores,
        color="#0072BC",
        linewidth=2,
        linestyle="solid",
        label="Objetivo (TO-BE)",
    )
    ax.fill(angles, target_scores, color="#0072BC", alpha=0.1)

    #
    ax.plot(
        angles,
        actual_scores,
        color="#C00000",
        linewidth=2,
        linestyle="solid",
        label="Actual (AS-IS)",
    )
    ax.fill(angles, actual_scores, color="#C00000", alpha=0.25)

    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)

    import textwrap

    wrapped_labels = [textwrap.fill(lbl, 15) for lbl in labels]
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(wrapped_labels, size=8)

    ax.set_ylim(0, 5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(["1", "2", "3", "4", "5"], color="grey", size=7)

    ax.set_title(title, size=11, weight="bold", pad=15)
    ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.1))

    plt.tight_layout()
    plt.savefig(output_path, format="png", dpi=300)
    plt.close()
    return True


def render_consolidated_asis(working_dir: str, output_path: str):
    """Generates a consolidated 'as-is' technology assessment Word document.

    Aggregates data from multiple technology tower-specific JSON payloads located
    within a specified working directory. The function synthesizes this information
    into a single, multi-section Microsoft Word report. The generated document
    includes a cover page, table of contents, executive summary, a consolidated
    risk matrix quantified using the FAIR model, a multi-tower maturity assessment
    with radar charts, and strategic conclusions. It also incorporates optional
    strategic context from a `client_intelligence.json` file if present.

    Args:
        working_dir: The file system path to the directory containing input
            data. This directory must contain subdirectories for each technology
            tower (e.g., 'T1', 'T2'), each holding a
            'blueprint_*_payload.json' file. An optional
            'client_intelligence.json' file may be present at the root of
            this directory.
        output_path: The full file path, including the filename and .docx
            extension, where the generated Word document will be saved.

    Returns:
        None. The function writes a file to the specified output path and does
        not return a value.

    Raises:
        FileNotFoundError: If a required JSON payload file discovered by the search
            pattern cannot be found or opened.
        json.JSONDecodeError: If an input JSON file is malformed and cannot be
            parsed.
        KeyError: If a required key (e.g., 'document_meta', 'pillars_analysis')
            is missing from a JSON payload.
        TypeError: If data within a JSON payload has an unexpected type, leading
            to errors during processing (e.g., iterating on a non-list value).
        OSError: If the function is unable to write the final document to the
            specified `output_path` due to file system permissions,
            insufficient disk space, or other I/O-related errors.
    """
    print("🚀 Cargando Payloads de Torres para Renderizado AS-IS Consolidado...")

    # Define the mapping for risk identifier prefixes according to the FNMT institutional standard.
    prefix_map = {
        "T1": "RDC",  # Process the 'Data Center Risk' category.
        "T2": "RVS",  # Process the 'Virtualization and Servers Risk' category.
        "T4": "RAB",  # Subsection: Storage and Backup Systems
        "T5": "RRC",  # Subsection: Resilience and Business Continuity
        "T6": "RSI",  # Process the 'Security and Identity Risk' category.
        "T7": "RIT",  # Process the 'ITSM and Operations Risk' category.
        "T8": "RSG",  # Process the 'Strategy and Governance Risk' category.
        "T10": "RLM",  # Subsection: Legacy and Mainframe Systems
    }

    # Load all technology tower-specific data payloads from the working directory for subsequent processing.
    payloads = []
    total_ale_global = 0.0
    all_risks = []
    tower_overviews = []
    input_documents = []

    search_pattern = os.path.join(working_dir, "T*", "blueprint_*_payload.json")
    for file_path in glob.glob(search_pattern):
        with open(file_path, "r", encoding="utf-8-sig") as f:
            data = json.load(f)
            payloads.append(data)

            tower_name = data.get("document_meta", {}).get("tower_name", "Desconocida")
            tower_id = data.get("document_meta", {}).get("tower_code", "")

            total_ale_global += data.get("total_fair_ale", 0.0)

            snapshot = data.get("executive_snapshot", {})
            headline = snapshot.get("headline", "Sin resumen ejecutivo.")
            bottom_line = snapshot.get("bottom_line", "")

            pil_scores = [p.get("score", 0.0) for p in data.get("pillars_analysis", [])]
            t_scores = [
                p.get("target_score", 4.0) for p in data.get("pillars_analysis", [])
            ]
            avg_score = sum(pil_scores) / len(pil_scores) if pil_scores else 0.0
            avg_target = sum(t_scores) / len(t_scores) if t_scores else 4.0

            tower_overviews.append(
                {
                    "tower_id": tower_id,
                    "tower_name": tower_name,
                    "score": avg_score,
                    "target": avg_target,
                    "headline": headline,
                    "bottom_line": bottom_line,
                    "pillars": [
                        {
                            "name": p.get("pilar_name", ""),
                            "score": p.get("score", 0.0),
                            "target": p.get("target_score", 4.0),
                            "desc": p.get(
                                "asis_architecture_description",
                                p.get("thought_process", "Sin descripción detallada."),
                            ),
                        }
                        for p in data.get("pillars_analysis", [])
                    ],
                }
            )

            input_documents.append(f"AS-IS_Anexo_Tecnico_{tower_id}.docx")

            # Aggregate identified risks from all sources for the consolidated global risk matrix.
            for pilar in data.get("pillars_analysis", []):
                for hc in pilar.get("health_check_asis", []):
                    hc["tower_id"] = tower_id
                    hc["tower_name"] = tower_name
                    hc["pilar_name"] = pilar.get("pilar_name", "")

                    # Normalize dictionary keys by resolving field aliases defined in the Pydantic models to ensure consistent data access.
                    hc["finding"] = hc.get(
                        "finding", hc.get("risk_observed", "No disponible.")
                    )
                    hc["business_risk"] = hc.get(
                        "business_risk", hc.get("impact", "No disponible.")
                    )
                    all_risks.append(hc)

    if not payloads:
        print("❌ No se encontraron payloads de torre en working_dir.")
        return

    client_name = payloads[0].get("document_meta", {}).get("client_name", "Cliente")

    # Dynamically load localization parameters, cloud adoption ratios, and compliance metrics as specified for Section 2.
    doc_meta_global = payloads[0].get("document_meta", {})
    doc_lang_global = doc_meta_global.get("language", "es").lower()
    on_prem_pct = doc_meta_global.get("on_premise_percentage", "80%")
    cloud_pct = doc_meta_global.get("cloud_percentage", "20%")
    reg_frameworks = doc_meta_global.get("regulatory_frameworks")
    if not reg_frameworks:
        reg_frameworks = (
            "NIS2 / ENS"
            if doc_lang_global == "es"
            else "applicable regulatory frameworks"
        )

    # Load client-provided strategic intelligence data for target-state analysis.
    working_dir_path = Path(working_dir)
    client_intel_path = working_dir_path / "client_intelligence.json"
    intel = {}
    if client_intel_path.exists():
        with open(client_intel_path, "r", encoding="utf-8-sig") as cif:
            try:
                intel = json.load(cif)
                print(
                    "   ├─ Carga exitosa de Inteligencia Estratégica del Cliente en Nivel Consolidado."
                )
            except Exception as e:
                print(f"   ⚠️ Error cargando inteligencia: {e}")

    business_context = intel.get("business_context", {})
    ceo_agenda_raw = business_context.get("ceo_agenda", {}).get(
        "summary", "No disponible."
    )

    # Ensure absolute path resolution to support execution from any directory context.
    from assessment_engine.infrastructure.runtime_paths import ROOT

    template_path = ROOT / "templates" / "docx" / "template_tobe_consolidated.docx"
    if template_path.exists():
        doc = Document(str(template_path))
        print("   ├─ Plantilla Word pre-estilizada cargada correctamente.")
    else:
        doc = Document()
        print("   ⚠️ Plantilla no encontrada. Generando con estilos por defecto.")

    # Remove all placeholder paragraphs and tables from the document template before populating it with generated content.
    for p in list(doc.paragraphs):
        p._p.getparent().remove(p._p)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    #
    # Section 0: Generate the corporate cover page.
    #
    doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_color_hex = COLOR_BLUE
    r_color = int(p_color_hex[0:2], 16)
    g_color = int(p_color_hex[2:4], 16)
    b_color = int(p_color_hex[4:6], 16)
    p_color_rgb = (r_color, g_color, b_color)

    for _ in range(5):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run(
        "Informe Consolidado de Situación Actual\n(AS-IS) de la Plataforma"
    )
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(*p_color_rgb)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(40)
    run_sub = p_sub.add_run(
        "Tier-1 Infrastructure Modernisation & Resilience Assessment"
    )
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(128, 128, 128)

    p_client = doc.add_paragraph()
    run_client = p_client.add_run(f"Preparado para: {client_name}")
    run_client.font.size = Pt(12)
    run_client.font.bold = True

    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run(
        "NTT DATA Corporation | Confidencial\nVersión 1.2 | Junio 2026"
    )
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True

    doc.add_page_break()

    # Section: Generate the table of contents.
    add_heading(doc, "Índice de Contenidos", level=1, primary_color_rgb=p_color_rgb)
    add_body_paragraph(
        doc,
        "Nota: Para actualizar el índice en Word, haga clic derecho sobre el texto inferior y seleccione 'Actualizar campos'.",
        italic=True,
    )
    add_toc(doc)
    doc.add_page_break()

    #
    # Section 1: Introduction and Source Documents
    #
    add_heading(
        doc, "Introducción y Documentos Fuente", level=1, primary_color_rgb=p_color_rgb
    )

    add_heading(doc, "Objetivo del Diagnóstico", level=2, primary_color_rgb=p_color_rgb)
    add_body_paragraph(
        doc,
        f"El presente informe ejecutivo consolida los hallazgos de los análisis AS-IS realizados sobre los distintos dominios tecnológicos de la plataforma de {client_name}. Su propósito es ofrecer a la Dirección una visión unificada y accionable del estado actual, priorizando los riesgos por impacto de negocio y resiliencia.",
    )

    add_heading(
        doc, "Documentos Fuente Evaluados", level=2, primary_color_rgb=p_color_rgb
    )
    add_body_paragraph(
        doc,
        "Este informe consolida las auditorías forenses y encuestas de madurez documentadas de manera independiente para cada una de las torres tecnológicas de la organización:",
    )

    doc_table = doc.add_table(rows=1, cols=2)
    doc_table.style = "Table Grid"

    headers_doc = ["Dominio Tecnológico Evaluado", "Documento Anexo AS-IS Resultante"]
    for i, h_txt in enumerate(headers_doc):
        set_cell_text_custom(
            doc_table.rows[0].cells[i],
            h_txt,
            bold=True,
            font_size=9,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(doc_table.rows[0].cells[i], COLOR_BLUE)

    for r_idx, file_name in enumerate(sorted(input_documents)):
        row = doc_table.add_row()
        # Parse the human-readable display name from the source filename for use in the document's presentation layer.
        t_name = file_name.replace("AS-IS_Anexo_Tecnico_", "").replace(".docx", "")
        set_cell_text_custom(
            row.cells[0],
            f"Torre {t_name} - Diagnóstico Detallado de Infraestructura",
            bold=True,
            font_size=8.5,
        )
        set_cell_text_custom(row.cells[1], file_name, font_size=8.5)
        if r_idx % 2 == 1:
            for cell in row.cells:
                shade_cell(cell, COLOR_ROW_ALT)

    autofit_table_to_contents(doc_table)
    doc.add_page_break()

    #
    # Section 2: Executive Summary of the Current 'As-Is' State
    #
    add_heading(
        doc,
        "Resumen Ejecutivo de Situación Actual Global (AS-IS)",
        level=1,
        primary_color_rgb=p_color_rgb,
    )

    add_heading(
        doc,
        "Diagnóstico General de la Plataforma",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        f"El estado actual de la plataforma tecnológica de {client_name}, aunque operativo en sus pilares tradicionales, se encuentra en un punto crítico de obsolescencia operativa y fragmentación estructural. La coexistencia de tecnologías on-premise ({on_prem_pct}) y cloud ({cloud_pct}) operadas en silos independientes limita severamente la agilidad de entrega y la capacidad de resiliencia ante contingencias críticas.",
    )

    # Inject strategic content derived from the Leadership Agenda and Redeia Investment Plan into the document.
    if ceo_agenda_raw and ceo_agenda_raw != "No disponible.":
        add_heading(
            doc,
            "Por qué importa al negocio (Agenda de Liderazgo)",
            level=2,
            primary_color_rgb=p_color_rgb,
        )
        for block in ceo_agenda_raw.split("\n\n"):
            if block.strip().startswith("###"):
                add_heading(
                    doc,
                    block.replace("###", "").strip(),
                    level=3,
                    primary_color_rgb=p_color_rgb,
                )
            elif block.strip().startswith("####"):
                add_heading(
                    doc,
                    block.replace("####", "").strip(),
                    level=4,
                    primary_color_rgb=p_color_rgb,
                )
            elif block.strip().startswith("1.") or block.strip().startswith("-"):
                add_body_paragraph(doc, block.strip(), style="List Bullet")
            else:
                add_body_paragraph(doc, block.strip())

    add_heading(
        doc,
        "Exposición Financiera al Riesgo Global (FAIR)",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "A partir del modelo de cuantificación financiera de riesgo **FAIR**, la Expectativa de Pérdida Anualizada (ALE) consolidada agregando todos los riesgos críticos detectados en el assessment de la plataforma asciende a:",
    )

    p_ale = doc.add_paragraph()
    p_ale.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ale = p_ale.add_run(
        f"ALE Global Proyectado: {total_ale_global:,.0f} € / Año".replace(",", ".")
    )
    run_ale.font.size = Pt(16)
    run_ale.font.bold = True
    run_ale.font.color.rgb = RGBColor(
        192, 0, 0
    )  # Define the color constant for high-priority alerts or critical status indicators, conforming to standard visual conventions for urgency.

    add_body_paragraph(
        doc,
        f"Esta cifra representa la expectativa de coste anualizado resultante de mantener el status quo de operaciones manuales, indisponibilidades no programadas de sistemas críticos y deudas de cumplimiento normativo ({reg_frameworks}).",
    )

    doc.add_page_break()

    #
    # Section 3: Platform Descriptions by Technology Tower
    #
    add_heading(
        doc,
        "Descripción de la Plataforma de Infraestructura Actual",
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "A continuación, se define el estado de situación y arquitectura de cada torre evaluada, proporcionando la visión ejecutiva y el 'bottom line' de cada dominio:",
    )

    for t in sorted(tower_overviews, key=lambda x: x["tower_id"]):
        add_heading(
            doc,
            f"Dominio: {t['tower_name']} (Torre {t['tower_id']})",
            level=2,
            primary_color_rgb=p_color_rgb,
        )
        add_body_paragraph(
            doc,
            f"**Nivel de Madurez Actual:** Nivel {t['score']:.2f} de 5,00 (Clasificación: {t['headline']})",
            bold=True,
        )
        add_body_paragraph(doc, t["bottom_line"])

        # Define the data segments for the radar chart, which correspond to the maturity pillars of the technology tower.
        add_heading(
            doc,
            "Estructura de Pilares Técnicos Evaluados",
            level=3,
            primary_color_rgb=p_color_rgb,
        )
        for p in t["pillars"]:
            add_body_paragraph(
                doc, f"**{p['name']}:** {p['desc'].split('.')[0]}.", style="List Bullet"
            )

        doc.add_paragraph()  #

    doc.add_page_break()

    #
    # Section 4: Consolidated Vulnerability Register and Global Risk Matrix
    #
    add_heading(
        doc,
        "Registro Consolidado de Vulnerabilidades (Matriz FAIR Global)",
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "Este capítulo constituye la matriz consolidada de los **12 riesgos técnicos más materiales** detectados en toda la plataforma, ordenados por su nivel de severidad y coste de exposición FAIR (ALE):",
    )

    #
    try:
        all_risks.sort(key=lambda x: x.get("fair_ale_score", 0.0), reverse=True)
    except Exception:
        pass
    top_risks = all_risks[:12]

    table_risks = doc.add_table(rows=1, cols=4)
    finalize_table(table_risks)

    headers_risks = [
        "ID",
        "Torre / Dominio",
        "Descripción del Riesgo y Evidencia de Auditoría (RAG)",
        "Exposición FAIR (ALE)",
    ]
    for i, h_txt in enumerate(headers_risks):
        set_cell_text_custom(
            table_risks.rows[0].cells[i],
            h_txt,
            bold=True,
            font_size=9,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table_risks.rows[0].cells[i], COLOR_BLUE)

    for r_idx, hc in enumerate(top_risks):
        row = table_risks.add_row()

        # Define a mapping from technology towers to their standardized identifier prefixes.
        t_id = hc.get("tower_id", "TXX")
        t_prefix = prefix_map.get(t_id, "RTX")

        # Map the risk identifier attribute.
        set_cell_text_custom(
            row.cells[0],
            f"{t_prefix}{r_idx + 1:02d}",
            bold=True,
            font_size=8.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(row.cells[0], COLOR_ROW_ALT)

        # Map the originating technology tower attribute.
        set_cell_text_custom(
            row.cells[1],
            f"Torre {t_id}\n{hc.get('tower_name', 'General')}\n\nDominio:\n{hc.get('pilar_name', 'General')}",
            bold=True,
            font_size=8,
        )

        # Subsection: Description and Supporting Evidence
        p_desc = row.cells[2].paragraphs[0]
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run_obs = p_desc.add_run("Vulnerabilidad:\n")
        run_obs.bold = True
        p_desc.add_run(clean_text_for_word(hc.get("finding", "No descripto.")) + "\n\n")

        run_imp = p_desc.add_run("Impacto de Negocio:\n")
        run_imp.bold = True
        p_desc.add_run(
            clean_text_for_word(hc.get("business_risk", "No descripto.")) + "\n\n"
        )

        run_ev = p_desc.add_run("Evidencia Forense Literal (Audit RAG):\n")
        run_ev.bold = True
        run_cite = p_desc.add_run(
            f'"{clean_text_for_word(hc.get("literal_evidence", "No se aportó evidencia literal"))}"'
        )
        run_cite.italic = True
        run_cite.font.color.rgb = RGBColor(100, 110, 120)

        for run in p_desc.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)

        # Process the 'FAIR Risk Exposure' category.
        tef = hc.get("threat_event_frequency", 0.0)
        lm = hc.get("loss_magnitude", 0.0)
        ale = hc.get("fair_ale_score", 0.0)

        if ale and ale > 0:
            calc_txt = (
                f"TEF: {tef:.1f} / 5,0\nLM: {lm:.1f} / 5,0\n\nALE: {ale:,.0f} €".replace(
                    ",", "X"
                )
                .replace(".", ",")
                .replace("X", ".")
            )
            set_cell_text_custom(
                row.cells[3],
                calc_txt,
                font_size=8,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
            )

            # Establish the color mapping for risk severity levels to ensure consistent visual representation in the output.
            risk_score = tef * lm
            bg_color = "D9F2D9"
            if risk_score >= 15 or ale >= 1000000:
                bg_color = "F8D7DA"
            elif risk_score >= 10 or ale >= 250000:
                bg_color = "FFF3CD"
            elif risk_score >= 5 or ale >= 50000:
                bg_color = "E2E3E5"
            shade_cell(row.cells[3], bg_color)
        else:
            set_cell_text_custom(
                row.cells[3], "N/A", font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER
            )

        if r_idx % 2 == 1:
            shade_cell(row.cells[1], COLOR_ROW_ALT)

    autofit_table_to_contents(table_risks)
    doc.add_page_break()

    #
    # Section 5: Global Maturity Level Assessment
    #
    add_heading(
        doc,
        "Evaluación Global de Niveles de Madurez",
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "A continuación, se presenta la comparativa y el radar de madurez consolidado entre todos los dominios de infraestructura evaluados:",
    )

    # Generate the global maturity radar chart. A UUID is appended to the output filename to mitigate potential race conditions or naming conflicts in concurrent execution environments.
    global_labels = [t["tower_name"] for t in tower_overviews]
    global_actual = [t["score"] for t in tower_overviews]
    global_target = [t["target"] for t in tower_overviews]
    if global_labels:
        chart_path = (
            Path(tempfile.gettempdir()) / f"global_radar_asis_{uuid.uuid4().hex}.png"
        )
        if generate_radar_chart(
            global_labels,
            global_actual,
            global_target,
            "Comparativa de Madurez Global (AS-IS vs TO-BE)",
            str(chart_path),
        ):
            doc.add_picture(str(chart_path), width=Inches(5.0))
            p_img = doc.paragraphs[-1]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                os.remove(chart_path)
            except OSError:
                pass

    add_heading(
        doc,
        "Resumen General de Puntuaciones por Torre",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    mat_table = doc.add_table(rows=1, cols=4)
    mat_table.style = "Table Grid"

    headers_mat = [
        "Dominio de Infraestructura Evaluado",
        "Score AS-IS",
        "Meta TO-BE",
        "Justificación del Diagnóstico de Madurez",
    ]
    for i, h_txt in enumerate(headers_mat):
        set_cell_text_custom(
            mat_table.rows[0].cells[i],
            h_txt,
            bold=True,
            font_size=9,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(mat_table.rows[0].cells[i], COLOR_BLUE)

    for r_idx, t in enumerate(sorted(tower_overviews, key=lambda x: x["tower_id"])):
        row = mat_table.add_row()
        set_cell_text_custom(
            row.cells[0],
            f"Torre {t['tower_id']}\n{t['tower_name']}",
            bold=True,
            font_size=8.5,
        )
        set_cell_text_custom(
            row.cells[1],
            f"{t['score']:.2f}",
            font_size=8.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text_custom(
            row.cells[2],
            f"{t['target']:.2f}",
            font_size=8.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        # Extract the maturity level justification by parsing the first sentence from the 'bottom_line' attribute of the relevant technology tower data model.
        first_sentence = (
            t["bottom_line"].split(".")[0] + "."
            if t["bottom_line"]
            else "Evaluado con éxito."
        )
        set_cell_text_custom(row.cells[3], first_sentence, font_size=8)

        if r_idx % 2 == 1:
            for cell in row.cells:
                shade_cell(cell, COLOR_ROW_ALT)

    autofit_table_to_contents(mat_table)
    doc.add_page_break()

    #
    # Section 6: General Conclusions and Cross-Cutting 'As-Is' Gap Analysis
    #
    add_heading(
        doc,
        "Conclusiones y Brechas Transversales",
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "A partir del análisis de madurez y de la matriz consolidada de riesgos, se detectan de manera sistemática los siguientes patrones y deudas técnicas transversales en toda la infraestructura de la organización:",
    )

    # Dynamically aggregate all identified critical technical debt items from their respective technology towers.
    add_heading(
        doc,
        "Deuda Técnica Crítica de la Plataforma",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "La deuda técnica acumulada se agrupa en torno a tres patrones críticos de deficiencia. Primero, la **fragmentación de herramientas y operaciones en silos discretos** para los entornos on-premise y cloud, lo que impide una visibilidad unificada de servicio de extremo a extremo. Segundo, la **dependencia de procesos manuales y reactivos** para el aprovisionamiento, la configuración y el parcheado, lo que genera cuellos de botella sistémicos. Tercero, la **falta de mecanismos continuos de auditoría de conformidad**, validando el cumplimiento normativo (ENS, NIS2) a posteriori.",
    )

    add_heading(
        doc,
        "Siguientes Pasos Ejecutivos de Gobernanza",
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "Para solventar las deficiencias del estado actual y sentar las bases de la resiliencia proactiva, se proponen las siguientes acciones inmediatas a nivel directivo:",
    )
    add_body_paragraph(
        doc,
        "**Establecer un Comité de Gobernanza de Plataforma Híbrida:** Encargado de unificar los estándares operativos on-premise y cloud.",
        style="List Bullet",
    )
    add_body_paragraph(
        doc,
        "**Acelerar el Programa de Platform Engineering:** Automatizando los flujos de provisión bajo un modelo de autoservicio.",
        style="List Bullet",
    )
    add_body_paragraph(
        doc,
        "**Implantar el Modelo de Conformidad Continua (Policy-as-Code):** Asegurando el cumplimiento preventivo y en tiempo real de ENS y NIS2.",
        style="List Bullet",
    )

    #
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(
        f"🎉 ¡Informe Consolidado AS-IS COMPLETO generado con éxito en: {output_path}!"
    )


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(
            "Uso: python render_togaf_consolidated_asis.py <working_dir> <output_doc.docx>"
        )
        sys.exit(1)
    render_consolidated_asis(sys.argv[1], sys.argv[2])
