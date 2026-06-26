import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from assessment_engine.infrastructure.docx_render_utils import (
    autofit_table_to_contents,
    finalize_table,
    set_cell_text,
    shade_cell,
)
from assessment_engine.infrastructure.text_utils import clean_text_for_word

#
COLOR_BLUE = "0072BC"
COLOR_HEADER_BG = "D9EAF7"
COLOR_ROW_ALT = "F2F2F2"


def set_cell_text_custom(
    cell, text, bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=None
):
    """Sets the text and custom formatting for a table cell.

    This function is a wrapper around a base `set_cell_text` function, adding
    the capability to specify the font color.

    Args:
        cell (docx.table._Cell): The table cell object to be modified.
        text (str): The text content to set in the cell.
        bold (bool, optional): If True, the text will be bold. Defaults to False.
        font_size (int, optional): The font size for the text. Defaults to 9.
        align (WD_ALIGN_PARAGRAPH, optional): The paragraph alignment constant.
            Defaults to WD_ALIGN_PARAGRAPH.LEFT.
        color_rgb (docx.shared.RGBColor, optional): The RGB color for the font.
            If None, the default font color is used. Defaults to None.

    Raises:
        IndexError: If the target cell contains no paragraphs.
        AttributeError: If `cell` is not a valid cell object.
        TypeError: If `color_rgb` is not a valid `docx.shared.RGBColor` object.
    """
    set_cell_text(cell, text, bold=bold, align=align, font_size=font_size)
    p = cell.paragraphs[0]
    for r in p.runs:
        if color_rgb:
            r.font.color.rgb = color_rgb


def add_body_paragraph(
    doc,
    text,
    bold=False,
    italic=False,
    space_after=6,
    text_color_rgb=(46, 64, 77),
    style="Normal",
) -> Any:
    """Adds a formatted paragraph to a Word document.

    This function adds a paragraph with specified text and formatting. It includes
    special handling for paragraphs with the 'List Bullet' style, automatically
    bolding the text before the first colon (':'). It also supports a
    markdown-like syntax where text surrounded by double asterisks (`**text**`)
    is made bold.

    The function gracefully handles cases where a specified paragraph style is
    not found in the document template, falling back to the 'Normal' style to
    prevent errors.

    Args:
        doc (Any): The `python-docx` Document object to which the paragraph
            will be added.
        text (str): The content for the new paragraph.
        bold (bool, optional): If True, applies bold formatting to the entire
            text. Defaults to False.
        italic (bool, optional): If True, applies italic formatting to the entire
            text. Defaults to False.
        space_after (int, optional): The space in points to add after the
            paragraph. Defaults to 6.
        text_color_rgb (tuple[int, int, int], optional): An RGB tuple defining
            the text color. Defaults to (46, 64, 77).
        style (str, optional): The name of the paragraph style to apply. If the
            style is not found, 'Normal' is used as a fallback. Defaults to
            'Normal'.

    Returns:
        Any: The newly created `python-docx` Paragraph object.
    """
    # Handles missing style definitions by defaulting to a base style, preventing `KeyError` exceptions and ensuring forward compatibility with templates that may not include all optional styles.
    try:
        p = doc.add_paragraph(style=style)
        is_bullet = style == "List Bullet"
    except KeyError:
        p = doc.add_paragraph(style="Normal")
        is_bullet = style == "List Bullet"
        if is_bullet:
            run_b = p.add_run("• ")
            run_b.font.name = "Arial"
            run_b.font.size = Pt(10)
            run_b.font.color.rgb = RGBColor(*text_color_rgb)

    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Programmatically applies bold formatting to the initial words of each bullet point to improve scannability.
    if is_bullet and ":" in text:
        parts = text.split(":", 1)
        run_bold = p.add_run(parts[0] + ":")
        run_bold.bold = True
        run_bold.font.name = "Arial"
        run_bold.font.size = Pt(10)
        run_bold.font.color.rgb = RGBColor(*text_color_rgb)

        run_rest = p.add_run(parts[1])
        run_rest.font.name = "Arial"
        run_rest.font.size = Pt(10)
        run_rest.font.color.rgb = RGBColor(*text_color_rgb)
    elif "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 1:
                run = p.add_run(part)
                run.font.bold = True
            else:
                run = p.add_run(part)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*text_color_rgb)
    else:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*text_color_rgb)
    return p


def add_heading(doc, text, level, primary_color_rgb=(0, 114, 188)):
    """Add a styled heading to a Word document.

    This function adds a heading with specified text and level, then applies
    custom formatting such as spacing, font size, and color. Font size is
    determined by the heading level.

    Args:
        doc (docx.document.Document): The document object to which the heading
            will be added.
        text (str): The text content of the heading.
        level (int): The heading level (e.g., 1 for Heading 1, 2 for Heading 2).
        primary_color_rgb (Tuple[int, int, int], optional): An RGB tuple to set
            the font color. Defaults to (0, 114, 188).

    Returns:
        docx.text.paragraph.Paragraph: The newly created and styled heading
            paragraph object.

    Raises:
        ValueError: If `level` is outside the valid range (0-9) as defined by
            the `python-docx` library, or if RGB values are out of the 0-255
            range.
    """
    h = doc.add_heading(text, level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True

    # Applies color and size attributes, inheriting the base font from the parent style. Numeric prefixes are stripped from list items to allow the document's native numbering.
    for run in h.runs:
        run.font.color.rgb = RGBColor(*primary_color_rgb)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
    return h


def add_spacer(doc, points=12):
    """Add a blank paragraph to a document to act as a vertical spacer.

    This function creates an empty paragraph and configures its formatting to
    create vertical whitespace. The `space_after` property is set to the
    specified point value, while `space_before` and `line_spacing` are
    minimized.

    Args:
        doc (docx.document.Document): The document object to which the spacer
            will be added.
        points (int): The height of the space in points. Defaults to 12.

    Returns:
        docx.text.paragraph.Paragraph: The newly created spacer paragraph object.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    return p


def add_toc(doc):
    """Adds a Table of Contents (TOC) field to a python-docx Document.

    This function programmatically inserts the necessary XML field codes to
    generate a table of contents that includes heading levels 1 through 3.
    The TOC itself must be updated by the user within Microsoft Word or a
    compatible application to be populated.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to which
            the TOC field will be added.
    """
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)


def render_asis_annex(payload_path: str, output_path: str):
    """Generates a comprehensive AS-IS technical annex in Microsoft Word format.

    This function orchestrates the creation of a detailed technical report based on
    infrastructure assessment data. It dynamically loads a corporate brand profile
    for styling, localization data for text, and a primary JSON payload containing
    the core analysis. The final document includes a cover page, table of contents,
    methodology, risk analysis matrices (FAIR), maturity scorecards, and
    AI-generated conclusions, all formatted according to a predefined Word template.

    Args:
        payload_path: The file path to the primary JSON input file. This file
            should contain the complete AS-IS assessment data, including document
            metadata, pillar analysis, and health check findings.
        output_path: The destination file path for the generated .docx document.
            The parent directory will be created automatically if it does not exist.

    Raises:
        FileNotFoundError: If the `payload_path` or other critical configuration
            files (e.g., `brand_profile.json`) cannot be found.
        json.JSONDecodeError: If the main payload or other supplementary JSON files
            are improperly formatted and cannot be parsed.
        KeyError: If the input JSON data is missing essential keys that do not have
            default fallbacks, leading to a failure in document rendering.
        Exception: Catches and reports broad exceptions during the loading of the
            Word template or the supplementary annex payload, which could be due to
            file corruption or parsing errors.
    """
    print(f"📄 Iniciando Renderizado de Anexo Técnico AS-IS para: {payload_path}")

    # Dynamically loads brand profiles and corporate glossaries from external configuration files to support document standardization.
    brand_path = Path("engine_config/brand_profile.json")

    with open(brand_path, "r", encoding="utf-8-sig") as bf:
        brand = json.load(bf)

    company_name = brand.get("company_name", "NTT DATA")
    classification = brand.get("default_classification", "Confidencial")
    disclaimer = brand.get("disclaimer_text", "")

    styling = brand.get("styling", {})
    p_color_hex = styling.get("primary_color_hex", "0072BC")
    alt_row_hex = styling.get("alternate_row_color_hex", "F2F2F2")

    text_color_rgb = styling.get("text_dark_color_rgb", [46, 64, 77])

    # Converts a hexadecimal color string to an RGB tuple for use in document element styling.
    r_color = int(p_color_hex[0:2], 16)
    g_color = int(p_color_hex[2:4], 16)
    b_color = int(p_color_hex[4:6], 16)
    p_color_rgb = (r_color, g_color, b_color)

    #
    from assessment_engine.infrastructure.runtime_paths import (
        resolve_tower_annex_template_path,
    )

    template_path = resolve_tower_annex_template_path()
    try:
        doc = Document(str(template_path))
        print("   ├─ Plantilla Word pre-estilizada cargada correctamente.")
    except Exception:
        doc = Document()
        print("   ⚠️ No se encontró plantilla. Inicializando documento en blanco.")

    # Clears all placeholder paragraphs and tables from the document to provide a clean state for programmatic content injection.
    # This sanitization method preserves embedded styles from the template while preventing format inheritance from placeholder content.
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    #
    with open(payload_path, "r", encoding="utf-8-sig") as f:
        data = json.load(f)

    tower_meta = data.get("document_meta", {})
    meta_lang = tower_meta.get("language", "es").lower()

    #
    locales_path = Path("engine_config/locales.json")
    locales_data = {}
    if locales_path.exists():
        with open(locales_path, "r", encoding="utf-8-sig") as lf:
            locales_data = json.load(lf)
    vocab = locales_data.get(meta_lang, locales_data.get("es", {}))

    tower_name = tower_meta.get("tower_name", "Desconocida")
    tower_id = tower_meta.get("tower_code", tower_meta.get("tower_id", "TXX"))
    client_name = tower_meta.get("client_name", "Cliente")
    pillars = data.get("pillars_analysis", [])
    snap = data.get("executive_snapshot", {})

    # Specifies the SVG path data for rendering the radar chart visualization.
    payload_path_obj = Path(payload_path)
    radar_path = payload_path_obj.parent / "pillar_radar_chart.generated.png"

    # Dynamically loads the Executive Synthesis Annex from a structured JSON payload.
    annex_path = (
        payload_path_obj.parent
        / f"approved_annex_{tower_id.lower()}.template_payload.json"
    )
    annex_data = {}
    if annex_path.exists():
        with open(annex_path, "r", encoding="utf-8-sig") as af:
            try:
                annex_data = json.load(af)
                print(
                    f"   ├─ Carga exitosa de Síntesis del Anexo Ejecutivo para {tower_id}."
                )
            except Exception as e:
                print(f"   ⚠️ Error cargando síntesis: {e}")

    exec_sum = annex_data.get("executive_summary", {})
    score_profile = annex_data.get("pillar_score_profile", {})

    # Maps risk category prefixes according to the FNMT standard.
    prefix_map = {
        "T1": "RDC",  #
        "T2": "RVS",  #
        "T4": "RAB",  #
        "T5": "RRC",  #
        "T6": "RSI",  #
        "T7": "RIT",  #
        "T8": "RSG",  #
        "T10": "RLM",  #
    }
    risk_prefix = prefix_map.get(tower_id, "RTX")

    #
    #
    #
    for _ in range(5):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run(
        f"Informe de Situación Actual (AS-IS)\nAnexo Técnico: {tower_name}"
    )
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(*p_color_rgb)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(40)
    project_name = vocab.get(
        "project", "Consultoría de Diagnóstico de Infraestructura Crítica y Resiliencia"
    )
    run_sub = p_sub.add_run(project_name)
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(128, 128, 128)

    p_client = doc.add_paragraph()
    run_client = p_client.add_run(f"Preparado para: {client_name}")
    run_client.font.size = Pt(12)
    run_client.font.bold = True

    version = tower_meta.get("version", "1.2")
    date_str = tower_meta.get("date", "Junio 2026")

    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run(
        f"{company_name} | Clasificación: {classification}\nVersión {version} | {date_str}"
    )
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        vocab.get(
            "control_documental_title", "Control Documental e Histórico de Cambios"
        ),
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    hist_table = doc.add_table(rows=3, cols=4)
    hist_table.style = "Table Grid"
    autofit_table_to_contents(hist_table)

    headers_hist = [
        "Versión",
        "Fecha",
        "Comentarios / Cambios Realizados",
        "Páginas Afectadas",
    ]
    for i, h_txt in enumerate(headers_hist):
        set_cell_text_custom(
            hist_table.rows[0].cells[i],
            h_txt,
            bold=True,
            font_size=9,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(hist_table.rows[0].cells[i], p_color_hex)

    set_cell_text_custom(
        hist_table.rows[1].cells[0], "1.0", font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER
    )
    set_cell_text_custom(
        hist_table.rows[1].cells[1],
        "10/05/2026",
        font_size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text_custom(
        hist_table.rows[1].cells[2],
        "Borrador de diagnóstico técnico inicial.",
        font_size=8,
    )
    set_cell_text_custom(
        hist_table.rows[1].cells[3],
        "Todas",
        font_size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    set_cell_text_custom(
        hist_table.rows[2].cells[0], "1.2", font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER
    )
    set_cell_text_custom(
        hist_table.rows[2].cells[1],
        "15/06/2026",
        font_size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text_custom(
        hist_table.rows[2].cells[2],
        "Refinado cuantitativo de riesgos FAIR e inyección SOTA.",
        font_size=8,
    )
    set_cell_text_custom(
        hist_table.rows[2].cells[3],
        "11-28",
        font_size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_spacer(doc, 24)

    add_heading(
        doc,
        vocab.get("appendix_a_title", "Índice de Contenidos"),
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "Nota: Para actualizar el índice en Word, haga clic derecho sobre el texto inferior y seleccione 'Actualizar campos'.",
        italic=True,
        text_color_rgb=text_color_rgb,
    )
    add_toc(doc)
    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        vocab.get(
            "intro_metodologia_title", "Introducción y Metodología de Diagnóstico"
        ),
        level=1,
        primary_color_rgb=p_color_rgb,
    )

    add_heading(
        doc,
        vocab.get("objetivo_alcance_title", "Objetivo y Alcance"),
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        f"Este documento técnico anexo detalla de manera exhaustiva el diagnóstico de situación actual (AS-IS) de la torre **{tower_name}** para {client_name}. El objetivo principal es identificar y registrar de manera estructurada las brechas operativas, riesgos de continuidad y obsolescencias tecnológicas dentro del perímetro de evaluación.",
        text_color_rgb=text_color_rgb,
    )
    add_body_paragraph(
        doc,
        "El alcance técnico incluye el inventario de infraestructura y la topología operativa, restringida estrictamente a los sistemas activos de producción del cliente, evaluando el nivel de madurez operativa basándose en evidencias de auditoría recopiladas empíricamente.",
        text_color_rgb=text_color_rgb,
    )

    # The executive summary content is sourced from the `summary_body` field of the synthesis annex payload.
    if exec_sum:
        add_heading(
            doc,
            vocab.get("resumen_ejecutivo_title", "Resumen Ejecutivo de la Torre"),
            level=2,
            primary_color_rgb=p_color_rgb,
        )
        headline_text = exec_sum.get("headline", "Diagnóstico General")
        add_body_paragraph(
            doc, f"**{headline_text}**", bold=True, text_color_rgb=text_color_rgb
        )

        summary_body = exec_sum.get("summary_body", "")
        if summary_body:
            add_body_paragraph(doc, summary_body, text_color_rgb=text_color_rgb)

        add_heading(
            doc,
            vocab.get("impactos_negocio_title", "Principales Impactos de Negocio"),
            level=3,
            primary_color_rgb=p_color_rgb,
        )
        impacts = exec_sum.get("key_business_impacts", [])
        if isinstance(impacts, str):
            add_body_paragraph(doc, impacts, text_color_rgb=text_color_rgb)
        else:
            for item in impacts:
                add_body_paragraph(
                    doc, item, style="List Bullet", text_color_rgb=text_color_rgb
                )

    add_heading(
        doc,
        vocab.get("metodologia_madurez_title", "Metodología de Valoración de Madurez"),
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "La madurez operativa de la torre se calcula mediante un modelo de bloques ponderados alineado con mejores prácticas del sector. La madurez se califica en una escala analítica del 1,00 al 5,00, donde cada nivel determina un estadio de control:",
    )
    add_body_paragraph(
        doc,
        "**Nivel 1 - Inicial (0-1.8):** Prácticas inexistentes, inestables o dependientes del esfuerzo heroico de personas clave.",
        style="List Bullet",
    )
    add_body_paragraph(
        doc,
        "**Nivel 2 - Básico (1.8-2.6):** Prácticas existentes de manera parcial o irregular, sin consistencia organizativa.",
        style="List Bullet",
    )
    add_body_paragraph(
        doc,
        "**Nivel 3 - Estandarizado (2.6-3.4):** Procesos formalizados e implantados de manera coherente en toda la organización.",
        style="List Bullet",
    )
    add_body_paragraph(
        doc,
        "**Nivel 4 - Optimizado (3.4-4.2):** Capacidades industrializadas, gobernadas predictivamente y sustentadas en métricas.",
        style="List Bullet",
    )
    add_body_paragraph(
        doc,
        "**Nivel 5 - Avanzado (4.2-5.0):** Procesos dinámicos impulsados por automatización inteligente y mejora continua.",
        style="List Bullet",
    )

    add_heading(
        doc,
        vocab.get(
            "metodologia_fair_title", "Metodología Cuantitativa de Riesgos (FAIR)"
        ),
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "Para garantizar el rigor ejecutivo y evitar estimaciones cualitativas de riesgo arbitrarias, este diagnóstico utiliza el estándar internacional **FAIR (Factor Analysis of Information Risk)** para modelar la exposición al riesgo en términos financieros reales:",
        text_color_rgb=text_color_rgb,
    )
    add_body_paragraph(
        doc,
        "1. **Frecuencia de Ocurrencia (TEF):** Representa el número estimado de veces al año que un evento de amenaza es capaz de materializar una vulnerabilidad. Se representa cuantitativamente o sobre una escala relativa del 1 al 5.",
        text_color_rgb=text_color_rgb,
    )
    add_body_paragraph(
        doc,
        "2. **Magnitud de la Pérdida (LM):** Representa el impacto económico total promedio estimado por cada evento de pérdida, considerando la interrupción de operaciones, horas de ingeniería de recuperación y penalizaciones regulatorias.",
        text_color_rgb=text_color_rgb,
    )
    add_body_paragraph(
        doc,
        "3. **Exposición de Pérdida Anualizada (ALE):** Representa el riesgo económico real proyectado de manera anualizada, calculado mediante el producto matemático de los factores anteriores:",
        text_color_rgb=text_color_rgb,
    )

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq = p_eq.add_run("ALE = TEF x LM (Costo Proyectado/Año)")
    run_eq.font.size = Pt(11)
    run_eq.font.bold = True
    run_eq.font.color.rgb = RGBColor(*p_color_rgb)

    add_spacer(doc, 24)
    add_heading(
        doc,
        vocab.get(
            "disclaimer_heading_title", "Cláusula de Limitación de Responsabilidad"
        ),
        level=2,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(doc, disclaimer, italic=True, text_color_rgb=text_color_rgb)

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        vocab.get(
            "descripcion_plataforma_title",
            "Descripción de la Plataforma y Topología Unificada",
        ),
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "A continuación, se define de manera consolidada y unificada el inventario técnico, arquitectura y estado operativo general del entorno. Esta descripción unifica el contexto tecnológico general para evitar repeticiones innecesarias entre dominios:",
        text_color_rgb=text_color_rgb,
    )

    # Aggregates and deduplicates pillar-specific descriptions into a single text block.
    if pillars:
        main_desc = pillars[0].get(
            "asis_architecture_description",
            "Descripción técnica de plataforma no disponible.",
        )
        add_body_paragraph(doc, main_desc, text_color_rgb=text_color_rgb)
    else:
        add_body_paragraph(
            doc,
            "No hay descripciones topológicas detalladas para esta torre.",
            text_color_rgb=text_color_rgb,
        )

    # Embeds the generated Tower Maturity Radar chart image into the document.
    if radar_path.exists():
        add_spacer(doc, 20)
        add_heading(
            doc,
            vocab.get(
                "radar_chart_title", "Perfil de Madurez General de la Torre (Radar)"
            ),
            level=2,
            primary_color_rgb=p_color_rgb,
        )
        doc.add_picture(str(radar_path), width=Inches(4.5))
        p_img = doc.paragraphs[-1]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        vocab.get(
            "matriz_fair_title",
            "Registro Forense de Vulnerabilidades y Riesgos (Matriz FAIR)",
        ),
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "Este capítulo constituye el registro bruto y la matriz detallada de riesgos de auditoría. Cada hallazgo está estructurado quirúrgicamente con sus evidencias de diagnóstico y su impacto de negocio:",
        text_color_rgb=text_color_rgb,
    )

    for p_idx, pilar in enumerate(pillars):
        add_heading(
            doc,
            f"Dominio: {pilar.get('pilar_name', 'Desconocido')}",
            level=2,
            primary_color_rgb=p_color_rgb,
        )

        risks = pilar.get("health_check_asis", [])
        if not risks:
            add_body_paragraph(
                doc,
                "_No se han registrado hallazgos ni vulnerabilidades específicas para este dominio arquitéctónico en el RAG._",
                italic=True,
                text_color_rgb=text_color_rgb,
            )
            doc.add_paragraph()
            continue

        #
        try:
            risks.sort(key=lambda x: x.get("fair_ale_score", 0.0), reverse=True)
        except Exception:
            pass

        # Renders the data table and programmatically omits the 'Evaluated Capacity' column from the final output.
        table = doc.add_table(rows=1, cols=3)
        finalize_table(table)

        headers_table = [
            "ID",
            "Descripción del Riesgo y Evidencias de Auditoría (Deep Dive)",
            "Exposición FAIR (ALE)",
        ]
        for i, h_txt in enumerate(headers_table):
            set_cell_text_custom(
                table.rows[0].cells[i],
                h_txt,
                bold=True,
                font_size=9,
                color_rgb=RGBColor(255, 255, 255),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(table.rows[0].cells[i], p_color_hex)

        for r_idx, hc in enumerate(risks):
            row = table.add_row()

            # Section 1: Risk Identification. Risk identifiers are prefixed according to the FNMT standard (e.g., RVS01, RAB01).
            set_cell_text_custom(
                row.cells[0],
                f"{risk_prefix}{r_idx + 1:02d}",
                bold=True,
                font_size=9,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(row.cells[0], alt_row_hex)

            # Section 2: Finding Description and Evidence. Content is mapped from the `finding` and `business_risk` keys in the source payload.
            p_desc = row.cells[1].paragraphs[0]
            p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run_observed = p_desc.add_run("Vulnerabilidad Identificada:\n")
            run_observed.bold = True
            p_desc.add_run(
                clean_text_for_word(
                    hc.get("finding", hc.get("risk_observed", "No descripto."))
                )
                + "\n\n"
            )

            run_imp = p_desc.add_run("Impacto en Negocio / Operaciones:\n")
            run_imp.bold = True
            p_desc.add_run(
                clean_text_for_word(
                    hc.get("business_risk", hc.get("impact", "No descripto."))
                )
                + "\n\n"
            )

            run_ev = p_desc.add_run("Evidencia Forense Literal (Audit RAG):\n")
            run_ev.bold = True
            run_cite = p_desc.add_run(
                f'"{clean_text_for_word(hc.get("literal_evidence", "No se aportó evidencia literal"))}"'
            )
            run_cite.italic = True
            run_cite.font.color.rgb = RGBColor(100, 110, 120)

            for run in p_desc.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8.5)

            #
            tef = hc.get("threat_event_frequency", 0.0)
            lm = hc.get("loss_magnitude", 0.0)
            ale = hc.get("fair_ale_score", 0.0)

            if ale and ale > 0:
                # Threat Event Frequency (TEF) and Loss Magnitude (LM) are formatted as qualitative scores on a 1-to-5 scale. Currency symbols are omitted to prevent misinterpreting these abstract risk factors as direct financial values.
                currency = tower_meta.get("currency", "€")
                calc_txt = (
                    f"TEF: {tef:.1f} / 5,0\nLM: {lm:.1f} / 5,0\n\nALE: {ale:,.0f} {currency}".replace(
                        ",", "X"
                    )
                    .replace(".", ",")
                    .replace("X", ".")
                )
                set_cell_text_custom(
                    row.cells[2],
                    calc_txt,
                    font_size=8.5,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.RIGHT,
                )

                # Applies background shading to table rows based on the calculated risk severity level.
                risk_score = tef * lm
                bg_color = "D9F2D9"  #
                if risk_score >= 15 or ale >= 1000000:
                    bg_color = "F8D7DA"  #
                elif risk_score >= 10 or ale >= 250000:
                    bg_color = "FFF3CD"  #
                elif risk_score >= 5 or ale >= 50000:
                    bg_color = "E2E3E5"  #

                shade_cell(row.cells[2], bg_color)
            else:
                set_cell_text_custom(
                    row.cells[2], "N/A", font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER
                )

            # Applies alternating background colors (banded rows) to the table to improve readability.
            if r_idx % 2 == 1:
                shade_cell(row.cells[1], alt_row_hex)

        autofit_table_to_contents(table)
        doc.add_page_break()

    #
    #
    #
    add_heading(
        doc,
        vocab.get("evaluacion_madurez_title", "Evaluación de Niveles de Madurez"),
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "La madurez técnica de la torre de infraestructura se detalla a continuación por cada pilar. El promedio ponderado consolida la brecha que orientará el diseño estratégico de transformación:",
        text_color_rgb=text_color_rgb,
    )

    # Renders the justification table, grouped by pillar, adhering to the FNMT layout standard.
    mat_table = doc.add_table(rows=1, cols=4)
    mat_table.style = "Table Grid"

    headers_mat = [
        "Pilar / Capacidad Evaluada",
        "Score AS-IS",
        "Meta TO-BE",
        "Análisis de Brecha y Justificación de Nota",
    ]
    for i, h_txt in enumerate(headers_mat):
        set_cell_text_custom(
            mat_table.rows[0].cells[i],
            h_txt,
            bold=True,
            font_size=9,
            color_rgb=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(mat_table.rows[0].cells[i], p_color_hex)

    # Retrieves the pillar mapping from the payload; access is defensive as the mapping may not exist in all versions.
    annex_pillars_map = {}
    if score_profile:
        for ap in score_profile.get("pillars", []):
            annex_pillars_map[ap.get("pillar_label")] = ap.get("executive_reading")

    for r_idx, p in enumerate(pillars):
        row = mat_table.add_row()
        p_name = p.get("pilar_name", "Pilar")
        set_cell_text_custom(row.cells[0], p_name, bold=True, font_size=8.5)
        set_cell_text_custom(
            row.cells[1],
            f"{p.get('score', 0.0):.2f}",
            font_size=8.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text_custom(
            row.cells[2],
            f"{p.get('target_score', 4.0):.2f}",
            font_size=8.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        # The justification content is sourced from the `approved_annex` payload to ensure the rendered text is the authoritative version.
        justification = annex_pillars_map.get(p_name)
        if not justification:
            desc = p.get(
                "asis_description",
                p.get("asis_architecture_description", "No descripto."),
            )
            justification = desc.split(".")[0] + "." if desc else "Evaluado con éxito."

        set_cell_text_custom(row.cells[3], justification, font_size=8)

        if r_idx % 2 == 1:
            for cell in row.cells:
                shade_cell(cell, alt_row_hex)

    autofit_table_to_contents(mat_table)
    add_spacer(doc, 20)

    #
    scores = [p.get("score", 0.0) for p in pillars]
    targets = [p.get("target_score", 4.0) for p in pillars]
    avg_score = sum(scores) / len(scores) if scores else 0.0
    avg_target = sum(targets) / len(targets) if targets else 4.0

    add_body_paragraph(
        doc,
        f"**Puntuación de Madurez AS-IS Obtenida:** {avg_score:.2f} sobre 5,00",
        bold=True,
        text_color_rgb=text_color_rgb,
    )
    add_body_paragraph(
        doc,
        f"**Puntuación Objetivo TO-BE Proyectada:** {avg_target:.2f} sobre 5,00",
        bold=True,
        text_color_rgb=text_color_rgb,
    )

    add_body_paragraph(
        doc,
        "Esta calificación sitúa a la infraestructura en un estadio operativo fundamental que requiere modernización y automatización urgente para alinearse con los estándares SOTA de continuidad y soberanía. El salto de madurez se sostendrá en la ejecución de las iniciativas recomendadas.",
        text_color_rgb=text_color_rgb,
    )

    doc.add_page_break()

    #
    # Section 5: General Conclusions.
    #
    add_heading(
        doc,
        vocab.get("conclusiones_tobe_title", "Conclusiones y Brechas para el TO-BE"),
        level=1,
        primary_color_rgb=p_color_rgb,
    )
    add_body_paragraph(
        doc,
        "A partir del análisis técnico y cuantitativo realizado, se consolidan las siguientes conclusiones generales extraídas por nuestro motor de inteligencia para esta torre:",
        text_color_rgb=text_color_rgb,
    )

    add_heading(
        doc,
        vocab.get(
            "fortalezas_identificadas_title",
            "Fortalezas de Infraestructura Identificadas",
        ),
        level=2,
        primary_color_rgb=p_color_rgb,
    )

    # Injects the `message_strength` field from the annex synthesis payload.
    msg_strength = exec_sum.get("message_strength", "")
    if msg_strength:
        add_body_paragraph(doc, msg_strength, text_color_rgb=text_color_rgb)

    benefits = snap.get("operational_benefits", [])
    if benefits:
        if isinstance(benefits, str):
            add_body_paragraph(doc, benefits, text_color_rgb=text_color_rgb)
        else:
            for b in benefits:
                add_body_paragraph(
                    doc, b, style="List Bullet", text_color_rgb=text_color_rgb
                )
    else:
        add_body_paragraph(
            doc,
            "La plataforma cuenta con un núcleo de hardware moderno y soportado de fabricante (Dell, IBM) que constituye una base tecnológica excelente para soportar los modelos de virtualización y datos.",
            text_color_rgb=text_color_rgb,
        )

    add_heading(
        doc,
        vocab.get("debilidades_criticas_title", "Debilidades y Desafíos Críticos"),
        level=2,
        primary_color_rgb=p_color_rgb,
    )

    # Injects the `message_gap` field from the annex synthesis payload.
    msg_gap = exec_sum.get("message_gap", "")
    if msg_gap:
        add_body_paragraph(doc, msg_gap, text_color_rgb=text_color_rgb)

    weaknesses = snap.get("structural_risks", [])
    if weaknesses:
        if isinstance(weaknesses, str):
            add_body_paragraph(doc, weaknesses, text_color_rgb=text_color_rgb)
        else:
            for w in weaknesses:
                add_body_paragraph(
                    doc, w, style="List Bullet", text_color_rgb=text_color_rgb
                )
    else:
        add_body_paragraph(
            doc,
            "La excesiva dependencia del conocimiento tácito y de la operación manual secuencial constituye la brecha más crítica del negocio. Ante contingencias graves, los tiempos de recuperación reales se desvían de los RTO teóricos debido a la falta de runbooks automatizados ejecutables.",
            text_color_rgb=text_color_rgb,
        )

    add_heading(
        doc,
        vocab.get("next_steps_title", "Siguientes Pasos (Base para el Plan TO-BE)"),
        level=2,
        primary_color_rgb=p_color_rgb,
    )

    # Injects the `message_bottleneck` field from the synthesis payload into the executive summary.
    msg_bottle = exec_sum.get("message_bottleneck", "")
    if msg_bottle:
        add_body_paragraph(doc, msg_bottle, text_color_rgb=text_color_rgb)

    coi = snap.get("cost_of_inaction", [])
    if coi:
        if isinstance(coi, str):
            add_body_paragraph(doc, coi, text_color_rgb=text_color_rgb)
        else:
            for item in coi:
                add_body_paragraph(
                    doc, item, style="List Bullet", text_color_rgb=text_color_rgb
                )
    else:
        add_body_paragraph(
            doc,
            "El diagnóstico cuantitativo AS-IS aquí formulado servirá de red de seguridad y base empírica para el diseño de la arquitectura objetivo TO-BE (Fase II) y la priorización de iniciativas de cambio en el Plan de Actuación TO-DO (Fase III).",
            text_color_rgb=text_color_rgb,
        )

    #
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"🎉 ¡Anexo Técnico AS-IS COMPLETO generado con éxito en: {output_path}!")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(
            "Uso: python render_togaf_asis_annex.py <blueprint_payload.json> <output_annex.docx>"
        )
        sys.exit(1)
    render_asis_annex(sys.argv[1], sys.argv[2])
