import os
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from assessment_engine.domain.schemas.ast import (
    CellNode,
    DocumentAST,
    HeadingNode,
    PageBreakNode,
    ParagraphNode,
    PictureNode,
    SpacerNode,
    TableNode,
)
from assessment_engine.infrastructure.docx_render_utils import (
    add_heading_paragraph,
    apply_bullet_list_format,
    finalize_table,
)
from assessment_engine.ports.document_compiler import DocumentCompiler


def safe_shade_cell(cell: Any, fill: str) -> None:
    """Apply a solid color shading to a table cell, ensuring XML schema validity.

    Adds a `w:shd` (shading) element to the cell's `w:tcPr` (table cell
    properties). Any pre-existing `w:shd` element is removed before the new
    shading is applied.

    The function inserts the new `w:shd` element at the correct position as
    defined by the Office Open XML schema for the `CT_TcPr` complex type.
    This prevents document corruption that can result from improperly ordered
    child elements within the cell's properties.

    Args:
        cell (Any): The `python-docx` cell object, which must expose the
            underlying `lxml` element through a `_tc` attribute.
        fill (str): The RRGGBB hex string for the fill color (e.g., "FFA500").

    Raises:
        AttributeError: If `cell` does not have the expected `_tc` attribute.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith("shd"):
            tcPr.remove(child)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)

    # This sequence is dictated by the official w:tcPr XSD schema definition.
    order = [
        "cnfStyle",
        "tcW",
        "tcBorders",
        "shd",
        "noWrap",
        "tcMar",
        "textDirection",
        "tcFitText",
        "vAlign",
        "wHeader",
        "vMerge",
    ]

    # Finds the first element that must succeed w:shd according to the OpenXML schema to determine the correct insertion point.
    insert_before = None
    for child in list(tcPr):
        tag_name = child.tag.split("}")[-1]
        if tag_name in order:
            if order.index(tag_name) > order.index("shd"):
                insert_before = child
                break

    if insert_before is not None:
        insert_before.addprevious(shd)
    else:
        tcPr.append(shd)


def safe_autofit_table_to_contents(table: Any) -> None:
    r"""{'docstring': "Manipulates a table's OOXML to reliably autofit column widths.\n\nThis function directly modifies the Office Open XML (OOXML) of a table to\nenforce a robust column autofit behavior, bypassing potential inconsistencies\nwith higher-level library properties. It operates by systematically resetting\ntable- and cell-level width properties to ensure document validity and\npredictable rendering.\n\nFirst, at the table level, it removes any pre-existing width (`w:tblW`) and\nlayout (`w:tblLayout`) elements from the table properties (`w:tblPr`). It then\nadds a `w:tblLayout` element with its type set to 'autofit' and a `w:tblW`\nelement with its type set to 'auto'.\n\nSecond, it iterates through every cell in the table. For each cell, it\nremoves any existing cell width element (`w:tcW`) from the cell's properties\n(`w:tcPr`). A new `w:tcW` element with type 'auto' is then programmatically\ninserted into the correct position within the `w:tcPr` children, adhering to\nthe strict element sequence defined by the OOXML schema.\n\nArgs:\n    table (Any): The `python-docx` Table object to be modified in-place.\n        Expected to be an instance of `docx.table.Table`.\n\nReturns:\n    None\n\nRaises:\n    AttributeError: If the `table` object lacks the expected attributes and\n        methods of a `python-docx` Table instance."}."""
    table.autofit = True
    tbl = table._tbl

    #
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    for child in list(tblPr):
        tag = child.tag.split("}")[-1]
        if tag in {"tblW", "tblLayout"}:
            tblPr.remove(child)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "auto")
    tblW.set(qn("w:w"), "0")
    tblPr.append(tblW)

    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)

    # Forces automatic width (type='auto', w='0') on each cell, ensuring the property is inserted in the correct sequence within the XML schema.
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()

            # Removes any pre-existing tcW (cell width) element to prevent conflicts with the new auto-fit properties.
            for child in list(tcPr):
                if child.tag.endswith("tcW"):
                    tcPr.remove(child)

            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:type"), "auto")
            tcW.set(qn("w:w"), "0")

            # The sequence of child elements within w:tcPr is strictly defined by the OpenXML standard.
            order = [
                "cnfStyle",
                "tcW",
                "tcBorders",
                "shd",
                "noWrap",
                "tcMar",
                "textDirection",
                "tcFitText",
                "vAlign",
                "wHeader",
                "vMerge",
            ]

            insert_before = None
            for child in list(tcPr):
                tag_name = child.tag.split("}")[-1]
                if tag_name in order:
                    if order.index(tag_name) > order.index("tcW"):
                        insert_before = child
                        break

            if insert_before is not None:
                insert_before.addprevious(tcW)
            else:
                tcPr.append(tcW)


def add_markdown_runs(
    paragraph: Any,
    text: str,
    default_bold: bool = False,
    default_italic: bool = False,
    font_size: float = 10.5,
    font_name: str = "Arial",
    color_rgb: Any = None,
) -> None:
    """Parses a string with Markdown-style bold syntax into formatted docx runs.

    This function processes an input string to identify and format text enclosed
    in double asterisks (`**`) or double underscores (`__`), treating both as
    indicators for bold formatting. The function segments the string into a
    series of runs and adds them to a provided `docx.paragraph.Paragraph`
    object.

    Each generated run is configured with the specified font name, size, and color.
    Text not enclosed in bold markers is formatted according to the `default_bold`
    flag, while the `default_italic` flag is applied uniformly to all runs.

    Args:
        paragraph ('docx.paragraph.Paragraph'): The paragraph object to which
            the formatted runs will be added.
        text (str): The input string containing text with potential Markdown-style
            bold formatting.
        default_bold (bool): If True, non-bolded segments of text are also
            rendered in bold. Defaults to False.
        default_italic (bool): If True, all generated runs (both bold and regular)
            are rendered in italics. Defaults to False.
        font_size (float): The font size in points for all runs. Defaults to 10.5.
        font_name (str): The font face name for all runs. Defaults to "Arial".
        color_rgb (Optional[Tuple[int, int, int]]): A 3-tuple of integers (0-255)
            representing the RGB color for the font. If None, the default document
            font color is used. Defaults to None.

    Returns:
        None: This function modifies the `paragraph` object in place.

    Raises:
        AttributeError: If the `paragraph` object lacks the required `add_run`
            method or its subsequent font properties.
        TypeError: If `color_rgb` is not a 3-tuple of integers compatible with
            `docx.shared.RGBColor`.
    """
    sanitized = str(text or "").replace("__", "**")
    parts = re.split(r"(\*\*.*?\*\*)", sanitized)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            #
            clean_part = part[2:-2]
            run = paragraph.add_run(clean_part)
            run.bold = True
            run.italic = default_italic
        else:
            #
            run = paragraph.add_run(part)
            run.bold = default_bold
            run.italic = default_italic

        run.font.name = font_name
        run.font.size = Pt(font_size)
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)


def safe_set_cell_text(
    cell: Any,
    text: str,
    bold: bool = False,
    align: Any = WD_ALIGN_PARAGRAPH.LEFT,
    font_size: float = 9.0,
    color_rgb: Any = None,
) -> None:
    r"""Replaces the content of a `python-docx` table cell with formatted text.

    This function clears any pre-existing content in the cell and populates it
    based on the input text, applying several formatting heuristics. It handles
    multi-line input by creating distinct paragraphs for each line. The function
    interprets simple list structures (e.g., lines starting with '•', '-', or
    whitespace indentation) and "Label: Description" patterns, where the label
    is rendered in bold. It also parses inline Markdown-style bold syntax
    (e.g., "**text**").

    A critical behavior is ensuring the cell contains at least one paragraph,
    even if empty. This is a safeguard to prevent the generation of a corrupted
    document that cannot be opened by Microsoft Word.

    Args:
        cell (Any): The `python-docx` `_Cell` object to populate. Any existing
            content within this cell will be replaced.
        text (str): The text content to add. Both standard newlines (`\n`) and
            literal escaped newlines (`\\n`) are treated as paragraph breaks.
        bold (bool, optional): If True, applies bold formatting to text that is
            not otherwise styled by other heuristics. Defaults to False.
        align (Any, optional): A `WD_ALIGN_PARAGRAPH` enum member for text
            alignment. Defaults to `WD_ALIGN_PARAGRAPH.LEFT`.
        font_size (float, optional): The font size in points. Defaults to 9.0.
        color_rgb (Any, optional): A tuple or list of three integers (0-255)
            representing an RGB color for the font. Defaults to None.

    Returns:
        None.

    Raises:
        ValueError: If a value in `color_rgb` is outside the valid 0-255 range.
        TypeError: If `color_rgb` is provided with an incompatible type or
            structure (e.g., not a sequence of three numbers).
    """
    clean_txt = str(text or "")

    # If the text contains newlines, it is split into distinct paragraphs. This ensures hard line breaks rather than soft line breaks within a single run.
    if "\n" in clean_txt or "\\n" in clean_txt:
        # Removes the default empty paragraph automatically added by the python-docx library upon cell creation.
        for p in list(cell.paragraphs):
            cell._tc.remove(p._element)

        raw_lines = clean_txt.replace("\\n", "\n").split("\n")
        for line in raw_lines:
            # Leading whitespace is preserved initially as it indicates hierarchical depth; stripping is deferred.
            is_subpoint = line.startswith("   ") or line.strip().startswith("-")

            clean_line = line.strip()
            if not clean_line:
                continue

            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.alignment = align

            # Removes any existing bullet formatting to ensure a consistent, uniform reconstruction.
            if clean_line.startswith("•") or clean_line.startswith("-"):
                clean_line = clean_line[1:].strip()

            if is_subpoint:
                p.paragraph_format.left_indent = Inches(0.25)
                bullet_prefix = "- "
            else:
                p.paragraph_format.left_indent = Inches(0.0)
                bullet_prefix = (
                    "• "
                    if clean_txt.startswith(("-", "•"))
                    or line.strip().startswith(("-", "•"))
                    else ""
                )

            # Heuristic to detect and structure list items that follow a '• Title: Description' pattern.
            if ":" in clean_line:
                label, desc = clean_line.split(":", 1)

                #
                run_lbl = p.add_run(f"{bullet_prefix}{label.strip()}: ")
                run_lbl.bold = True
                run_lbl.font.name = "Arial"
                run_lbl.font.size = Pt(font_size)
                if color_rgb:
                    run_lbl.font.color.rgb = RGBColor(*color_rgb)

                #
                add_markdown_runs(
                    p, desc, default_bold=bold, font_size=font_size, color_rgb=color_rgb
                )
            else:
                p.add_run(bullet_prefix)
                add_markdown_runs(
                    p,
                    clean_line,
                    default_bold=bold,
                    font_size=font_size,
                    color_rgb=color_rgb,
                )
    else:
        #
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0

        # Clears any pre-existing text runs from the paragraph to prevent content duplication.
        p.text = ""

        # Heuristic to detect and structure paragraphs that follow a 'Title: Description' pattern.
        if ":" in clean_txt and not clean_txt.startswith("http"):
            label, desc = clean_txt.split(":", 1)
            run_lbl = p.add_run(f"{label.strip()}: ")
            run_lbl.bold = True
            run_lbl.font.name = "Arial"
            run_lbl.font.size = Pt(font_size)
            if color_rgb:
                run_lbl.font.color.rgb = RGBColor(*color_rgb)

            add_markdown_runs(
                p, desc, default_bold=bold, font_size=font_size, color_rgb=color_rgb
            )
        else:
            add_markdown_runs(
                p,
                clean_txt,
                default_bold=bold,
                font_size=font_size,
                color_rgb=color_rgb,
            )

    # CRITICAL: A table cell must contain at least one paragraph to prevent document corruption, which can cause a fatal error in Microsoft Word.
    if not cell.paragraphs:
        cell.add_paragraph()


class DocxCompiler(DocumentCompiler):
    r"""{'DocxCompiler': 'Compiles a `DocumentAST` into a Microsoft Word (.docx) document.\n\n    This class implements the visitor pattern to traverse a `DocumentAST`,\n    dispatching each node to a dedicated compilation method that translates it\n    into the corresponding `python-docx` object. It manages document\n    structure, text formatting, styles, and embedded objects like tables and\n    images. The process is designed to be robust, with fallbacks for missing\n    assets.', 'compile': "Translates a `DocumentAST` and saves it as a .docx file.\n\n        This method orchestrates the compilation pipeline. It first attempts to\n        load a DOCX template; if unavailable, it proceeds with a blank document.\n        The template's existing body content is cleared, and a default font is\n        applied to the 'Normal' style. It then iterates through the AST,\n        delegating each node to the `_compile_node` dispatcher. Finally, it saves\n        the generated document to the specified path, creating parent directories\n        if necessary.\n\n        Args:\n            ast: The abstract syntax tree of the document.\n            output_path: The file system path for the output .docx file.\n\n        Raises:\n            IOError: If the output file cannot be written due to permissions,\n                disk space, or other file system issues.", '_compile_node': 'Dispatches an AST node to its corresponding compiler method.\n\n        This method acts as the central router in the visitor pattern. It\n        inspects the `type` attribute or class of the node and calls the\n        appropriate `_compile_*` handler. Unrecognized node types are silently\n        ignored to maintain forward compatibility with evolving AST schemas.\n\n        Args:\n            doc: The `python-docx` Document object being constructed.\n            node: The `DocumentAST` node to compile.', '_compile_paragraph': "Compiles a `ParagraphNode` into a DOCX paragraph.\n\n        This method creates a new paragraph, applying specified alignment and\n        spacing. It handles styled text by processing markdown-like syntax for\n        bold and italic formatting. If the node's style indicates a bullet list,\n        the corresponding paragraph format is applied.\n\n        Args:\n            doc: The `python-docx` Document object being constructed.\n            node: The `ParagraphNode` containing text and formatting details.", '_compile_heading': 'Compiles a `HeadingNode` into a styled DOCX heading.\n\n        A heading paragraph is added to the document with the specified level and\n        text. If a `primary_color_rgb` is defined in the node, the font color\n        of all runs within the heading is set accordingly.\n\n        Args:\n            doc: The `python-docx` Document object being constructed.\n            node: The `HeadingNode` containing the heading text, level, and\n                color information.', '_compile_spacer': 'Compiles a `SpacerNode` by inserting vertical whitespace.\n\n        This is achieved by adding an empty paragraph and setting its\n        `space_after`, `space_before`, and `line_spacing` format properties.\n\n        Args:\n            doc: The `python-docx` Document object being constructed.\n            node: The `SpacerNode` specifying the height of the space in points.', '_compile_picture': 'Compiles a `PictureNode` into an embedded, centered image.\n\n        The method first verifies the existence of the image file at the\n        specified path. If found, the image is added to the document and\n        centered. If the file is not found, a descriptive, italicized\n        placeholder text is inserted in its place to prevent compilation\n        failure.\n\n        Args:\n            doc: The `python-docx` Document object being constructed.\n            node: The `PictureNode` containing the image path and width.', '_compile_table': "Compiles a `TableNode` into a DOCX table.\n\n        The method constructs a table with the specified dimensions and populates\n        each cell based on the corresponding `CellNode` data. It applies\n        cell-level formatting, including text content, alignment, bolding,\n        font size, text color, and background shading. If the `autofit`\n        attribute is true, column widths are adjusted to fit their content after\n        population.\n\n        Args:\n            doc: The `python-docx` Document object being constructed.\n            node: The `TableNode` containing the table's structure and data."}."""

    def compile(self, ast: DocumentAST, output_path: str) -> None:
        """Compiles a Document Abstract Syntax Tree (AST) into a .docx file.

        This function performs a single-pass compilation of a `DocumentAST` object
        into a Microsoft Word (.docx) document. It initializes a document object,
        attempting to load from a standard template and falling back to a default
        empty document if the template is not found. Any pre-existing content in the
        document body is cleared, and the base "Normal" style is configured with a
        default font.

        The function then traverses the input AST, delegating the compilation of
        each node to a corresponding handler method (visitor pattern). The resulting
        document is saved to the specified output path, creating parent directories
        as needed.

        Args:
            ast (DocumentAST): The Abstract Syntax Tree representing the document's
                structure and content.
            output_path (str): The destination file path for the generated .docx file.
                Parent directories will be created if they do not exist.

        Raises:
            PermissionError: If the process lacks sufficient permissions to create the
                output directory or write the output file.
            ValueError: If an unsupported or malformed node type is encountered
                during the AST traversal.
        """
        #
        from assessment_engine.infrastructure.runtime_paths import (
            resolve_tower_annex_template_path,
        )

        template_path = resolve_tower_annex_template_path()
        try:
            doc = Document(str(template_path))
        except Exception:
            doc = Document()

        # Clears all pre-existing paragraphs and tables from the template body.
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(doc.tables):
            t._element.getparent().remove(t._element)

        #
        style_normal = doc.styles["Normal"]
        font = style_normal.font
        font.name = "Arial"
        font.size = Pt(10.5)

        # Sequentially compiles all AST nodes using a Visitor design pattern.
        for node in ast.nodes:
            self._compile_node(doc, node)

        #
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))

    def _compile_node(self, doc: Document, node: Any) -> None:
        """Dispatches the node compilation to the appropriate handler based on its concrete type."""
        node_type = getattr(node, "type", None)

        if node_type == "paragraph" or isinstance(node, ParagraphNode):
            self._compile_paragraph(doc, node)
        elif node_type == "heading" or isinstance(node, HeadingNode):
            self._compile_heading(doc, node)
        elif node_type == "table" or isinstance(node, TableNode):
            self._compile_table(doc, node)
        elif node_type == "spacer" or isinstance(node, SpacerNode):
            self._compile_spacer(doc, node)
        elif node_type == "page_break" or isinstance(node, PageBreakNode):
            doc.add_page_break()
        elif node_type == "picture" or isinstance(node, PictureNode):
            self._compile_picture(doc, node)
        else:
            # Unidentified nodes are safely ignored to ensure forward compatibility and operational robustness.
            pass

    def _compile_paragraph(self, doc: Document, node: ParagraphNode) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(node.space_after)

        # Map string alignment
        align_const = WD_ALIGN_PARAGRAPH.JUSTIFY
        align_str = node.align.upper()
        if align_str == "CENTER":
            align_const = WD_ALIGN_PARAGRAPH.CENTER
        elif align_str == "LEFT":
            align_const = WD_ALIGN_PARAGRAPH.LEFT
        elif align_str == "RIGHT":
            align_const = WD_ALIGN_PARAGRAPH.RIGHT

        p.paragraph_format.alignment = align_const

        #
        is_bullet = node.style == "List Bullet" or node.style == "Bullet"
        if is_bullet:
            apply_bullet_list_format(p)

        #
        add_markdown_runs(
            p,
            node.text,
            default_bold=node.bold,
            default_italic=node.italic,
            font_size=node.font_size,
            color_rgb=node.text_color_rgb,
        )

    def _compile_heading(self, doc: Document, node: HeadingNode) -> None:
        is_appendix = node.text.startswith("Apéndice")
        p = add_heading_paragraph(doc, node.text, level=node.level)

        if is_appendix and p:
            try:
                p.style = "Appendix Heading 1"
            except Exception:
                pass

        #
        if node.primary_color_rgb and p:
            for run in p.runs:
                run.font.color.rgb = RGBColor(*node.primary_color_rgb)

    def _compile_spacer(self, doc: Document, node: SpacerNode) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(node.points)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(1)

    def _compile_picture(self, doc: Document, node: PictureNode) -> None:
        if os.path.exists(node.path):
            doc.add_picture(node.path, width=Inches(node.width_inches))
            p_img = doc.paragraphs[-1]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Defensive fallback to prevent production failures if the image asset failed to generate.
            p = doc.add_paragraph()
            run = p.add_run(f"[Imagen no disponible: {Path(node.path).name}]")
            run.italic = True
            run.font.color.rgb = RGBColor(120, 120, 120)

    def _compile_table(self, doc: Document, node: TableNode) -> None:
        #
        num_rows = len(node.rows)
        if num_rows == 0:
            return

        num_cols = len(node.rows[0].cells)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        finalize_table(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  #

        #
        for r_idx, row_node in enumerate(node.rows):
            for c_idx in range(num_cols):
                cell = table.rows[r_idx].cells[c_idx]

                #
                if c_idx < len(row_node.cells):
                    cell_node = row_node.cells[c_idx]
                else:
                    cell_node = CellNode(text="", bold=False, align="LEFT")

                #
                align_const = WD_ALIGN_PARAGRAPH.LEFT
                align_str = cell_node.align.upper()
                if align_str == "CENTER":
                    align_const = WD_ALIGN_PARAGRAPH.CENTER
                elif align_str == "RIGHT":
                    align_const = WD_ALIGN_PARAGRAPH.RIGHT
                elif align_str == "JUSTIFY":
                    align_const = WD_ALIGN_PARAGRAPH.JUSTIFY

                #
                safe_set_cell_text(
                    cell,
                    cell_node.text,
                    bold=cell_node.bold,
                    align=align_const,
                    font_size=cell_node.font_size,
                    color_rgb=cell_node.text_color_rgb,
                )

                #
                if cell_node.bg_color:
                    safe_shade_cell(cell, cell_node.bg_color)

        # Applies a safe, schema-compliant method for auto-fitting column widths.
        if node.autofit:
            safe_autofit_table_to_contents(table)
