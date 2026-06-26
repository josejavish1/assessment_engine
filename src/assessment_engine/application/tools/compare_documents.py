from docx import Document


def inspect_doc(path: str) -> dict:
    """Analyze a .docx file and extract key structural statistics.

    This function processes a Microsoft Word (.docx) document to quantify its
    structure. It iterates through paragraphs and tables to count elements
    such as non-empty paragraphs, tables, rows, and cells. An approximate
    word count is calculated by splitting text content on whitespace.

    Headings are identified using a heuristic that checks for style names
    beginning with "Heading", specific numeric prefixes, or the presence of
    pre-defined Spanish keywords (e.g., "INFORME", "ANEXO").

    Args:
        path (str): The file system path to the .docx document.

    Returns:
        dict: A dictionary containing the extracted statistics:
            "headings" (list[tuple[str, str]]): A list of up to the first 15
                detected headings. Each heading is a tuple containing the
                paragraph's style name and its text, truncated to 60
                characters.
            "num_paragraphs" (int): The total count of non-empty paragraphs.
            "num_tables" (int): The total count of tables in the document.
            "num_rows" (int): The aggregate count of rows across all tables.
            "num_cells" (int): The aggregate count of cells across all tables.
            "word_count" (int): An approximate total word count derived from
                whitespace-splitting text in paragraphs and table cells.

    Raises:
        FileNotFoundError: If the document at the specified path does not exist.
        docx.opc.exceptions.PackageNotFoundError: If the file is not a valid
            or uncorrupted .docx package.
    """
    doc = Document(path)
    headings = []
    num_paras = 0
    num_tables = 0
    num_rows = 0
    num_cells = 0
    word_count = 0

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            num_paras += 1
            word_count += len(text.split())
            # Evaluates if a given paragraph qualifies as a heading by analyzing its associated style properties or matching its text against a heading-specific regular expression.
            if (
                p.style.name.startswith("Heading")
                or (
                    p.text.startswith(
                        (
                            "1.",
                            "2.",
                            "3.",
                            "4.",
                            "5.",
                            "6.",
                            "7.",
                            "8.",
                            "9.",
                            "Dominio:",
                        )
                    )
                )
                or (
                    any(
                        keyword in text.upper()
                        for keyword in [
                            "INFORME",
                            "PORTADA",
                            "ANEXO",
                            "CONTROL DOCUMENTAL",
                            "INDICE",
                            "RESUMEN EJECUTIVO",
                            "REGISTRO FORENSE",
                            "SIGUIENTES PASOS",
                        ]
                    )
                )
            ):
                headings.append((p.style.name, text[:60]))

    for t in doc.tables:
        num_tables += 1
        for row in t.rows:
            num_rows += 1
            for cell in row.cells:
                num_cells += 1
                word_count += len(cell.text.split())

    return {
        "headings": headings[:15],  #
        "num_paragraphs": num_paras,
        "num_tables": num_tables,
        "num_rows": num_rows,
        "num_cells": num_cells,
        "word_count": word_count,
    }


def compare():
    """Orchestrates the comparison of two hardcoded DOCX documents and prints a report.

    This function compares a hardcoded legacy DOCX document against its
    corresponding 'shadow' version. It utilizes the `inspect_doc` helper to
    extract structural statistics (e.g., paragraph, table, and word counts) and
    headings from each file. The collected data is then formatted and printed to
    standard output as a comparative report, allowing for a direct assessment of
    structural and content differences.

    Raises:
        FileNotFoundError: If either the legacy or shadow document file is not
            found at its hardcoded path.
    """
    legacy_path = "working/redeia_v3/T2/AS-IS_Anexo_Tecnico_T2.docx"
    shadow_path = "working/redeia_v3/T2/AS-IS_Anexo_Tecnico_T2.shadow.docx"

    print("🔍 Inspeccionando reporte original (Legacy V3)...")
    legacy_stats = inspect_doc(legacy_path)

    print("\n🔍 Inspeccionando reporte nuevo (Shadow V4 AST)...")
    shadow_stats = inspect_doc(shadow_path)

    print("\n=======================================================")
    print("      INFORME DE COMPARACIÓN DE CALIDAD Y CONTENIDO")
    print("=======================================================")
    print(f"{'Métrica':<30} | {'Legacy V3 (Modules)':<20} | {'Shadow V4 (AST)':<20}")
    print("-" * 75)
    print(
        f"{'Número de Párrafos':<30} | {legacy_stats['num_paragraphs']:<20} | {shadow_stats['num_paragraphs']:<20}"
    )
    print(
        f"{'Número de Tablas':<30} | {legacy_stats['num_tables']:<20} | {shadow_stats['num_tables']:<20}"
    )
    print(
        f"{'Número Total de Filas':<30} | {legacy_stats['num_rows']:<20} | {shadow_stats['num_rows']:<20}"
    )
    print(
        f"{'Número Total de Celdas':<30} | {legacy_stats['num_cells']:<20} | {shadow_stats['num_cells']:<20}"
    )
    print(
        f"{'Recuento Estimado de Palabras':<30} | {legacy_stats['word_count']:<20} | {shadow_stats['word_count']:<20}"
    )
    print("=======================================================\n")

    print("📋 Encabezados del Documento Original (Legacy V3):")
    for style, txt in legacy_stats["headings"]:
        print(f"   - [{style}] {txt}")

    print("\n📋 Encabezados del Documento Nuevo (Shadow V4):")
    for style, txt in shadow_stats["headings"]:
        print(f"   - [{style}] {txt}")


if __name__ == "__main__":
    compare()
