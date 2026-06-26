"""Defines the core logic and utilities for the unresolved DOCX placeholder validation stage within the Assessment Engine pipeline."""

import logging
import re
import sys
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")


def main(argv: list[str] | None = None) -> None:
    """Scans a Microsoft Word (.docx) document for unresolved placeholders.

    This function serves as the main entry point for a command-line script that
    validates a DOCX file. It systematically inspects all paragraphs in the
    document's main body and within all table cells, searching for text that
    matches the `PLACEHOLDER_RE` regular expression.

    If any unresolved placeholders are found, the script logs each unique
    instance, prints a machine-readable marker "UNRESOLVED_PLACEHOLDERS=YES"
    to standard output, and terminates the process with a non-zero exit code.
    If the document is clean, it prints "UNRESOLVED_PLACEHOLDERS=NO" and exits
    with a status code of 0.

    This behavior is designed for integration into automated workflows, such as
    CI/CD pipelines, to ensure document templating processes have completed
    successfully.

    Args:
        argv: A list of command-line arguments. If `None`, `sys.argv` is used.
            The list is expected to contain two elements: the script name and
            the path to the target .docx file.

    Raises:
        SystemExit:
            - If the number of command-line arguments is not two.
            - If the specified .docx file does not exist at the given path.
            - If one or more unresolved placeholders are found (exit code 1).
    """
    if len(argv if argv is not None else sys.argv) != 2:
        raise SystemExit(
            "Uso: python -m scripts.tools.check_docx_unresolved_placeholders <docx_path>"
        )

    docx_path = Path((argv if argv is not None else sys.argv)[1]).resolve()
    if not docx_path.exists():
        raise SystemExit(f"No existe el archivo: {docx_path}")

    doc = Document(str(docx_path))
    found = []

    for p in doc.paragraphs:
        matches = PLACEHOLDER_RE.findall(p.text or "")
        for m in matches:
            found.append(("paragraph", m, p.text.strip()))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    matches = PLACEHOLDER_RE.findall(p.text or "")
                    for m in matches:
                        found.append(("table", m, p.text.strip()))

    unique = []
    seen = set()
    for item in found:
        key = (item[0], item[1], item[2])
        if key not in seen:
            seen.add(key)
            unique.append(item)

    if unique:
        logger.info("UNRESOLVED_PLACEHOLDERS=YES")
        for kind, placeholder, context in unique:
            logger.info(f"{kind}: {placeholder} :: {context}")
        raise SystemExit(1)

    logger.info("UNRESOLVED_PLACEHOLDERS=NO")


if __name__ == "__main__":
    main()
