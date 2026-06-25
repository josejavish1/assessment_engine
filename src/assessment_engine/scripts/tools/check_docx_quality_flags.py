"""Implements the core logic and utility functions for evaluating DOCX quality flags within the Assessment Engine pipeline."""

import logging
import re
import sys
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

BAD_PATTERNS = [
    r"\{\{[^{}]+\}\}",
    r"\{'statement':",
    r'"statement":',
]


def collect_text(doc):
    """Extracts and concatenates all textual content from a docx document.

    This function iterates through all paragraphs and tables within the provided
    document object. The text from each non-empty paragraph and table cell is
    collected and joined into a single string.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to process.

    Returns:
        str: A single string containing the concatenated text from all non-empty
            paragraphs and table cells, with each part separated by a newline
            character.

    Raises:
        AttributeError: If the input object is not a valid `python-docx`
            Document and lacks the expected `paragraphs` or `tables` attributes.
    """
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


def main(argv: list[str] | None = None) -> None:
    """Check a .docx file for predefined quality issues and excessive ellipses.

    This function serves as the main entry point for a command-line script that
    validates the quality of a Microsoft Word document. It parses command-line
    arguments to obtain the path to a .docx file, extracts its text content,
    and then performs two primary checks:
    1.  Scans the text for regular expression patterns defined in the global
        `BAD_PATTERNS` list.
    2.  Counts the occurrences of ellipses ("...") to ensure they do not
        exceed a threshold of 3.

    If any patterns are matched or the ellipsis count is excessive, the function
    logs the specific findings and exits with a status code of 1. If no issues
    are detected, it logs a success message and completes execution, resulting
    in a status code of 0.

    Args:
        argv: An optional list of command-line arguments. If None, `sys.argv`
            is used. The script expects exactly one argument following the
            script name: the path to the .docx file.

    Raises:
        SystemExit: Exits the program under the following conditions:
            - With an error message if the number of command-line arguments is
              not equal to two (script name and file path).
            - With an error message if the specified .docx file does not exist.
            - With a status code of 1 if any quality checks fail.
    """
    if len(argv if argv is not None else sys.argv) != 2:
        raise SystemExit(
            "Uso: python -m scripts.tools.check_docx_quality_flags <docx_path>"
        )

    path = Path((argv if argv is not None else sys.argv)[1]).resolve()
    if not path.exists():
        raise SystemExit(f"No existe el archivo: {path}")

    doc = Document(str(path))
    text = collect_text(doc)

    flags = []
    for pattern in BAD_PATTERNS:
        matches = re.findall(pattern, text)
        if matches:
            flags.append((pattern, len(matches)))

    excessive_ellipsis = text.count("...")

    if flags or excessive_ellipsis > 3:
        logger.info("DOCX_QUALITY_FLAGS=YES")
        for pattern, count in flags:
            logger.info(f"pattern={pattern} count={count}")
        if excessive_ellipsis > 3:
            logger.info(f"ellipsis_count={excessive_ellipsis}")
        raise SystemExit(1)

    logger.info("DOCX_QUALITY_FLAGS=NO")


if __name__ == "__main__":
    main()
