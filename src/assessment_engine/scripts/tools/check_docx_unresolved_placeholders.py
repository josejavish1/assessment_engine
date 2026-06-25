"""Verifies that all placeholders in a DOCX document have been resolved.

This script functions as a quality gate within the Assessment Engine pipeline to ensure the integrity and completeness of generated document artifacts.
"""

import logging
import re
import sys
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")


def main(argv: list[str] | None = None) -> None:
    """Scans a .docx document for unresolved template placeholders.

    This function serves as the main entry point for a command-line script. It
    parses a Microsoft Word (.docx) file, iterating through all paragraphs and
    table cells to find text matching a predefined placeholder regular expression.
    If any placeholders are found, their details are logged, and the program
    terminates with an exit code of 1. If the document contains no
    placeholders, a success message is logged, and the program exits cleanly.

    Args:
        argv: A list of command-line arguments. If `None`, `sys.argv` is used.
            The list is expected to contain exactly two elements: the script name
            and the path to the .docx file to be inspected.

    Returns:
        None.

    Raises:
        SystemExit: If the number of provided arguments is not two, if the
            specified .docx file does not exist, or if unresolved placeholders
            are found (exit code 1).
    """
    if len(argv if argv is not None else sys.argv) != 2:
        raise SystemExit(
            "Uso: python -m scripts.tools.check_docx_unresolved_placeholders <docx_path>"
        )

    docx_path = Path((argv if argv is not None else sys.argv)[1]).resolve()
    if not docx_path.exists():
        raise SystemExit(f"No existe el archivo: {docx_path}")

    doc = Document(str(docx_path))
    found = []

    for p in doc.paragraphs:
        matches = PLACEHOLDER_RE.findall(p.text or "")
        for m in matches:
            found.append(("paragraph", m, p.text.strip()))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    matches = PLACEHOLDER_RE.findall(p.text or "")
                    for m in matches:
                        found.append(("table", m, p.text.strip()))

    unique = []
    seen = set()
    for item in found:
        key = (item[0], item[1], item[2])
        if key not in seen:
            seen.add(key)
            unique.append(item)

    if unique:
        logger.info("UNRESOLVED_PLACEHOLDERS=YES")
        for kind, placeholder, context in unique:
            logger.info(f"{kind}: {placeholder} :: {context}")
        raise SystemExit(1)

    logger.info("UNRESOLVED_PLACEHOLDERS=NO")


if __name__ == "__main__":
    main()
