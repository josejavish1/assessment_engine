"""Implements the core logic and utilities for the Assessment Engine's commercial report generation pipeline."""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from assessment_engine.schemas.commercial import (
    CommercialPayload,
    CommercialSummaryDraft,
    GtmStrategy,
    OpportunityPipelineItem,
    ProposalDraft,
)
from assessment_engine.scripts.lib.contract_utils import robust_load_payload
from assessment_engine.scripts.lib.docx_render_utils import (
    add_body_paragraph,
    add_heading_paragraph,
    autofit_table_to_contents,
    clear_paragraph,
    finalize_table,
    set_cell_text,
    shade_cell,
)
from assessment_engine.scripts.render_global_report_from_template import (
    add_smart_bullet_list,
    add_spacer,
    clean_t_codes,
    create_page_number_footer,
)

BASE_TEXT_COLOR = RGBColor(46, 64, 77)


def clean_commercial_text(text):
    r"""{'docstring': "Sanitizes a string for inclusion in a commercial report.\n\n    Performs a series of cleaning operations on an input string. This function\n    removes T-codes by calling `clean_t_codes`, strips unresolved `[[REF:...]]`\n    tags via a regular expression, eliminates specific dash characters ('─', '—'),\n    replaces semicolons with commas, and removes leading and trailing whitespace.\n\n    If the provided input is not a string, it is returned unmodified.\n\n    Args:\n        text (Any): The input value to sanitize.\n\n    Returns:\n        str | Any: The sanitized string, or the original input `text` if it was\n            not of type `str`."}."""
    if not isinstance(text, str):
        return text
    cleaned = clean_t_codes(text)
    # Executes a final sanitization pass to remove unresolved reference tags that may have persisted through the generation pipeline. This operation guarantees a clean, artifact-free document output.
    cleaned = re.sub(r"\[\[REF:[^\]]*\]\]", "", cleaned)
    return cleaned.replace("─", "").replace("—", "").replace(";", ",").strip()


def load_json(path):
    """Load and parse a UTF-8 encoded JSON file from the specified path."""
    return json.loads(path.read_text(encoding="utf-8-sig"))


def clear_document_body(doc):
    r"""{'docstring': "Clears all block-level content from the document body, preserving section properties.\n\nThis function directly manipulates the `lxml` element representing the\ndocument's body (`w:body`). It iterates through all direct child elements,\nsuch as paragraphs (`w:p`) and tables (`w:tbl`), and removes them. The\nfinal `w:sectPr` (section properties) element is explicitly preserved.\n\nPreserving the section properties is crucial for maintaining the document's\npage layout, including margins, orientation, and page size, when the document\nis cleared for subsequent repopulation.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` Document object to modify\n        in-place."}."""
    body = doc._body._element
    for child in list(body):
        if (
            child.tag
            != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
        ):
            body.remove(child)


def load_payload(payload_path: Path) -> CommercialPayload:
    """Load and strictly parse a CommercialPayload from a file."""
    return robust_load_payload(
        payload_path,
        CommercialPayload,
        "Commercial Account Plan",
        mode="strict",
    )


def render_commercial_cover(doc, payload: CommercialPayload):
    r"""{'docstring': "Constructs and appends the cover page for an Account Action Plan report.\n\nPopulates the provided `docx.document.Document` object with a formatted\ncover page. The page includes a main title ('Account Action Plan'), the\nclient's name, the report date, a classification notice, and a detailed\nconfidentiality disclaimer. The content is styled with specific fonts,\nsizes, colors, and spacing. The function concludes by inserting a page break\nto ensure subsequent content begins on a new page.\n\nArgs:\n    doc: An active `docx.document.Document` object to which the cover page\n        content will be written. This object is modified in-place.\n    payload: A `CommercialPayload` data object containing the report's\n        metadata, specifically the client name (`payload.meta.client`) and\n        date (`payload.meta.date`).\n\nRaises:\n    AttributeError: If the `payload` object lacks the required `meta.client`\n        or `meta.date` attributes, or if `doc` does not conform to the\n        `docx.document.Document` interface."}."""
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run("Account Action Plan")
    run.font.size = Pt(34)
    run.font.name = "Georgia"
    run.font.color.rgb = RGBColor(0, 114, 188)
    run.bold = False

    add_spacer(doc, 30)
    client_p = doc.add_paragraph()
    client_run = client_p.add_run(payload.meta.client.upper())
    client_run.font.size = Pt(24)
    client_run.font.name = "Arial"
    client_run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(100)

    version_p = doc.add_paragraph()
    v_text = f"Fecha: {payload.meta.date}\nClasificación: CONFIDENCIAL - USO INTERNO"
    version_run = version_p.add_run(v_text)
    version_run.font.size = Pt(14)
    version_run.font.name = "Arial"
    version_run.font.color.rgb = RGBColor(46, 64, 77)
    version_run.bold = False

    add_spacer(doc, 150)

    disclaimer_title_p = doc.add_paragraph()
    disclaimer_title_run = disclaimer_title_p.add_run(
        "Advertencia de Confidencialidad:"
    )
    disclaimer_title_run.font.size = Pt(8)
    disclaimer_title_run.font.name = "Arial"
    disclaimer_title_run.font.color.rgb = RGBColor(127, 127, 127)
    disclaimer_title_run.bold = True
    disclaimer_title_p.paragraph_format.space_after = Pt(2)

    disclaimer_p = doc.add_paragraph()
    disclaimer_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p.paragraph_format.space_before = Pt(0)
    p1_text = "El presente documento es estrictamente confidencial y de uso exclusivo interno para los equipos comerciales, de preventa y de entrega de NTT DATA. Contiene estimaciones de negocio, argumentarios de venta competitivos, estrategias de cuenta (Go-To-Market) y evaluaciones de riesgo de clientes. Bajo ninguna circunstancia este documento, ni parcial ni totalmente, debe ser compartido con el cliente, partners, proveedores o cualquier tercero."
    r1 = disclaimer_p.add_run(p1_text)
    r1.font.size = Pt(8)
    r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor(127, 127, 127)

    doc.add_page_break()


def render_commercial_summary(doc, data: CommercialSummaryDraft, matrix: list):
    r"""{'docstring': 'Render the executive summary and deal snapshot into a python-docx Document.\n\n    Constructs the initial section of a commercial report within a given\n    `docx.document.Document` object. This function adds a static context\n    disclaimer, a deal snapshot table detailing Total Addressable Market (TAM),\n    purchase drivers, and win themes, followed by bulleted lists for the\n    problem statement ("Why Now") and the proposed strategy ("How We Win").\n    If stakeholder data is provided, a stakeholder value matrix is appended.\n\n    Args:\n        doc (docx.document.Document): The document object to be modified.\n        data (CommercialSummaryDraft): A data container with summary content.\n            Expected attributes include `estimated_tam`, `deal_flash` (an object\n            with `purchase_driver` and `ntt_win_theme` attributes),\n            `why_now_bullets` (list[str]), and `how_we_win_bullets` (list[str]).\n        matrix (list): A list of objects representing stakeholders. Each object\n            must have `role`, `focus`, and `message` attributes. If the list\n            is empty, the stakeholder matrix section is omitted.\n\n    Returns:\n        None. The `doc` object is modified in place.\n\n    Raises:\n        AttributeError: If `data` or an element within `matrix` is missing an\n            expected attribute.'}."""
    add_heading_paragraph(doc, "1. Executive Summary & Deal Snapshot", level=1)

    #
    warn_table = doc.add_table(rows=1, cols=1)
    finalize_table(warn_table)
    w_cell = warn_table.rows[0].cells[0]
    shade_cell(w_cell, "F2F2F2")
    clear_paragraph(w_cell.paragraphs[0])
    wp = w_cell.paragraphs[0]
    wp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    wr1 = wp.add_run("⚠️ CONTEXTO DEL INFORME (FAST ASSESSMENT): ")
    wr1.bold = True
    wr1.font.size = Pt(9)
    wr1.font.color.rgb = RGBColor(127, 127, 127)
    wr2 = wp.add_run(
        "Este documento es un artefacto comercial interno. Las oportunidades, riesgos y valoraciones económicas nacen de un diagnóstico preliminar y entrevistas de alto nivel, no de una due-diligence técnica. Su propósito es definir la estrategia de entrada (Go-To-Market), no es una oferta vinculante."
    )
    wr2.font.size = Pt(9)
    wr2.font.color.rgb = RGBColor(127, 127, 127)
    add_spacer(doc, 15)

    #
    flash_table = doc.add_table(rows=3, cols=2)
    flash_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    finalize_table(flash_table)
    flash_table.columns[0].width = Inches(2.0)
    flash_table.columns[1].width = Inches(4.5)

    set_cell_text(
        flash_table.rows[0].cells[0], "TAM Estimado NTT DATA:", bold=True, font_size=11
    )
    shade_cell(flash_table.rows[0].cells[0], "0072BC")
    set_cell_text(
        flash_table.rows[0].cells[1],
        f"{clean_commercial_text(data.estimated_tam)} (18-24 meses)",
        bold=True,
        font_size=11,
    )
    shade_cell(flash_table.rows[0].cells[1], "0072BC")
    for r in flash_table.rows[0].cells[0].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
    for r in flash_table.rows[0].cells[1].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    deal_flash = data.deal_flash
    set_cell_text(
        flash_table.rows[1].cells[0],
        "Driver de Compra (Urgencia):",
        bold=True,
        font_size=10.5,
    )
    shade_cell(flash_table.rows[1].cells[0], "F2F2F2")
    set_cell_text(
        flash_table.rows[1].cells[1],
        clean_commercial_text(deal_flash.purchase_driver),
        font_size=10.5,
    )

    set_cell_text(
        flash_table.rows[2].cells[0],
        "Nuestra Ventaja (Win Theme):",
        bold=True,
        font_size=10.5,
    )
    shade_cell(flash_table.rows[2].cells[0], "F2F2F2")
    set_cell_text(
        flash_table.rows[2].cells[1],
        clean_commercial_text(deal_flash.ntt_win_theme),
        font_size=10.5,
    )
    autofit_table_to_contents(flash_table)
    add_spacer(doc, 15)

    #
    p_why = doc.add_paragraph()
    r_why = p_why.add_run("El Problema (Why Now):")
    r_why.bold = True
    r_why.font.color.rgb = BASE_TEXT_COLOR
    r_why.font.size = Pt(11)
    p_why.paragraph_format.space_after = Pt(4)
    add_smart_bullet_list(
        doc,
        data.why_now_bullets,
        color_rgb=BASE_TEXT_COLOR,
        bold_prefix=False,
    )
    add_spacer(doc, 10)

    p_how = doc.add_paragraph()
    r_how = p_how.add_run("La Estrategia de NTT DATA (How we Win):")
    r_how.bold = True
    r_how.font.color.rgb = BASE_TEXT_COLOR
    r_how.font.size = Pt(11)
    p_how.paragraph_format.space_after = Pt(4)
    add_smart_bullet_list(
        doc,
        data.how_we_win_bullets,
        color_rgb=BASE_TEXT_COLOR,
        bold_prefix=False,
    )
    add_spacer(doc, 20)

    if matrix:
        doc.add_page_break()
        add_heading_paragraph(
            doc, "Matriz de Valor por Interlocutor (Stakeholders)", level=2
        )
        add_body_paragraph(
            doc,
            "Alineación del discurso comercial según las prioridades de cada rol directivo.",
            color_rgb=BASE_TEXT_COLOR,
            space_after=12,
        )

        m_table = doc.add_table(rows=1, cols=3)
        finalize_table(m_table)
        headers = [
            "Rol Directivo",
            "Foco Estratégico",
            "Mensaje Clave (Elevator Pitch)",
        ]
        for i, h in enumerate(headers):
            set_cell_text(
                m_table.rows[0].cells[i],
                h,
                bold=True,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(m_table.rows[0].cells[i], "0072BC")
            for r in m_table.rows[0].cells[i].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(255, 255, 255)

        for row_data in matrix:
            row = m_table.add_row()
            set_cell_text(
                row.cells[0],
                clean_commercial_text(row_data.role),
                bold=True,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(row.cells[0], "F2F2F2")
            set_cell_text(
                row.cells[1],
                clean_commercial_text(row_data.focus),
                font_size=10.5,
            )
            set_cell_text(
                row.cells[2],
                clean_commercial_text(row_data.message),
                font_size=10.5,
            )
        autofit_table_to_contents(m_table)
        add_spacer(doc, 20)


def render_gtm_strategy(doc, gtm: GtmStrategy):
    r"""{'docstring': 'Renders and appends a go-to-market strategy section to a Word document.\n\nThis function adds a level 1 heading, "2. Go-To-Market", to the specified\ndocument object. It then constructs and styles a two-column table detailing\npredefined go-to-market vectors and their corresponding strategies. The\nstrategy descriptions are sourced from the provided `gtm` object\'s attributes.\nThe `doc` object is modified in-place.\n\nArgs:\n    doc (docx.document.Document): The document object to which the section\n        will be appended. This object is modified in-place.\n    gtm (GtmStrategy): An object possessing `trojan_horse`,\n        `self_funded_transformation`, and `lock_in` string attributes\n        that detail the respective strategies.\n\nReturns:\n    None.\n\nRaises:\n    AttributeError: If the `gtm` object lacks one of the required strategy\n        attributes.'}."""
    add_heading_paragraph(doc, "2. Go-To-Market", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    finalize_table(table)
    headers = ["Vector de Entrada", "Estrategia Recomendada"]
    for i, h in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            bold=True,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table.rows[0].cells[i], "0072BC")
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
    vectors = [
        ("El Caballo de Troya (Wedge)", gtm.trojan_horse),
        ("Transformación Autofinanciada", gtm.self_funded_transformation),
        ("Estrategia de Lock-in (MRR)", gtm.lock_in),
    ]
    for v_name, v_desc in vectors:
        row = table.add_row()
        set_cell_text(
            row.cells[0],
            v_name,
            bold=True,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(row.cells[0], "F2F2F2")
        set_cell_text(
            row.cells[1], clean_commercial_text(v_desc), bold=False, font_size=10.5
        )
    autofit_table_to_contents(table)
    add_spacer(doc, 20)


def render_opportunities_pipeline(doc, pipeline: list[OpportunityPipelineItem]):
    r"""{'docstring': 'Appends a formatted commercial opportunities pipeline section to a Word document.\n\nThis function adds a section titled "3. Pipeline de Oportunidades". For each\nopportunity in the provided pipeline, it generates and appends a formatted\ntable. Each table details financial estimates (Total Contract Value),\nco-sell vendor, revenue type, the client\'s cost of inaction (COI), and a\n"Red Team" analysis. The analysis consists of a potential client objection\nand the corresponding mitigation strategy.\n\nArgs:\n    doc (docx.document.Document): The Word document object to which the\n        pipeline section will be appended.\n    pipeline (list[OpportunityPipelineItem]): A list of objects, where each\n        represents a single commercial opportunity with its associated\n        details.\n\nReturns:\n    None. The function modifies the input `doc` object in-place.\n\nRaises:\n    AttributeError: If an object within the `pipeline` list is missing an\n        expected attribute (e.g., \'initiative\', \'estimated_tcv\',\n        \'client_objection\').'}."""
    add_heading_paragraph(doc, "3. Pipeline de Oportunidades", level=1)
    add_body_paragraph(
        doc,
        "Estimaciones financieras y estrategias de mitigación de objeciones (Red Team).",
        color_rgb=BASE_TEXT_COLOR,
        space_after=12,
    )
    for i, opp in enumerate(pipeline, start=1):
        table = doc.add_table(rows=4, cols=2)
        finalize_table(table)
        h_cell = table.rows[0].cells[0]
        h_cell.merge(table.rows[0].cells[1])
        set_cell_text(
            h_cell,
            f"Oportunidad {i}: {clean_commercial_text(opp.initiative)}",
            bold=True,
            font_size=11,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(h_cell, "0072BC")
        for r in h_cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_text(
            table.rows[1].cells[0],
            f"Vendor Co-Sell: {opp.vendor_cosell}",
            bold=True,
            font_size=10.5,
        )
        set_cell_text(
            table.rows[1].cells[1],
            f"Tipo de Ingreso: {opp.revenue_type}",
            bold=True,
            font_size=10.5,
        )
        shade_cell(table.rows[1].cells[0], "F2F2F2")
        shade_cell(table.rows[1].cells[1], "F2F2F2")
        set_cell_text(
            table.rows[2].cells[0],
            f"TCV Estimado NTT DATA:\n{opp.estimated_tcv}",
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(
            table.rows[2].cells[1],
            f"Cost of Inaction (COI) para Cliente:\n{clean_commercial_text(opp.cost_of_inaction)}",
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        obj_cell = table.rows[3].cells[0]
        obj_cell.merge(table.rows[3].cells[1])
        clear_paragraph(obj_cell.paragraphs[0])
        p1 = obj_cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run("🛑 Objeción del Cliente (Red Team): ")
        r1.bold = False
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = RGBColor(204, 0, 0)
        r_obj = p1.add_run(clean_commercial_text(opp.client_objection))
        r_obj.bold = False
        r_obj.font.size = Pt(10.5)

        p2 = obj_cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run("✅ Respuesta NTT DATA (Objection Handling): ")
        r2.bold = False
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = RGBColor(39, 78, 19)
        r_ans = p2.add_run(clean_commercial_text(opp.objection_handling))
        r_ans.bold = False
        r_ans.font.size = Pt(10.5)
        add_spacer(doc, 15)


def render_proactive_proposals(doc, proposals: list[ProposalDraft]):
    r"""{'docstring': 'Renders and appends detailed proactive proposal annexes to a Word document.\n\nIterates through a list of `ProposalDraft` objects and appends a formatted,\nmulti-section annex for each to the provided document. A page break precedes\neach proposal annex. The content for each proposal is structured into\ndistinct sections, including executive synthesis, strategy, context,\nsolution, scope, governance, risks (as a table), and investment.\n\nArgs:\n    doc (docx.document.Document): The document object to which the rendered\n        proposal content will be appended. This object is modified in-place.\n    proposals (list[ProposalDraft]): A list of data objects, where each\n        object contains the structured content for a single proposal.\n\nReturns:\n    None. This function modifies the `doc` object directly.\n\nRaises:\n    AttributeError: If a `ProposalDraft` object or its nested structures\n        lack a required attribute that the rendering logic attempts to access.\n    TypeError: If `proposals` is not a list or another iterable type.'}."""
    add_heading_paragraph(doc, "4. Anexos: Propuestas Proactivas", level=1)
    add_body_paragraph(
        doc,
        "Borradores de propuestas detalladas listos para ser adaptados y enviados al cliente tras incluir las credenciales correspondientes.",
        color_rgb=BASE_TEXT_COLOR,
        space_after=12,
    )
    for idx, prop in enumerate(proposals, 1):
        doc.add_page_break()
        p_title = add_heading_paragraph(
            doc,
            f"Propuesta {idx}: {clean_commercial_text(prop.initiative_name)}",
            level=2,
        )
        for r in p_title.runs:
            r.font.size = Pt(16)
        add_heading_paragraph(doc, "1. Executive Synthesis", level=3)
        p_syn = doc.add_paragraph()
        p_syn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_syn_text = p_syn.add_run(clean_commercial_text(prop.executive_synthesis))
        r_syn_text.bold = False
        r_syn_text.font.color.rgb = RGBColor(46, 64, 77)
        p_syn.paragraph_format.space_after = Pt(12)

        ai_strat = prop.ai_transformation_strategy
        if ai_strat:
            add_heading_paragraph(
                doc, "Estrategia de Transformación y Automatización", level=3
            )
            p_ai = doc.add_paragraph()
            p_ai.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_ai = p_ai.add_run(clean_commercial_text(ai_strat))
            r_ai.font.color.rgb = RGBColor(46, 64, 77)
            p_ai.paragraph_format.space_after = Pt(12)

        add_heading_paragraph(
            doc, "2. Contexto y Comprensión del Reto (The 'Why')", level=3
        )
        ctx = prop.context_and_why
        add_smart_bullet_list(
            doc,
            [
                f"Origen del Proyecto: {clean_commercial_text(ctx.origin)}",
                f"Coste de la Inacción (COI): {clean_commercial_text(ctx.cost_of_inaction)}",
            ],
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        add_spacer(doc, 10)
        add_heading_paragraph(
            doc, "3. La Solución y Resultados Estratégicos (The 'What')", level=3
        )
        sol = prop.solution_and_what
        add_smart_bullet_list(
            doc,
            [
                f"Estado Objetivo (TO-BE): {clean_commercial_text(sol.target_state)}",
                f"Métrica de Éxito (North Star Metric): {clean_commercial_text(sol.north_star_metric)}",
            ],
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        add_spacer(doc, 10)
        add_heading_paragraph(doc, "4. Alcance y Entregables (The 'How')", level=3)
        scope = prop.scope_and_how
        p_phases = doc.add_paragraph()
        r_ph = p_phases.add_run("Fases del Proyecto:")
        r_ph.bold = True
        r_ph.font.color.rgb = RGBColor(46, 64, 77)
        p_phases.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            scope.phases,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        p_del = doc.add_paragraph()
        r_del = p_del.add_run("Entregables Principales:")
        r_del.bold = True
        r_del.font.color.rgb = RGBColor(46, 64, 77)
        p_del.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            scope.deliverables,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        p_out = doc.add_paragraph()
        r_out = p_out.add_run("Fuera de Alcance (Out of Scope):")
        r_out.bold = True
        r_out.font.color.rgb = RGBColor(46, 64, 77)
        p_out.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            scope.out_of_scope,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        add_spacer(doc, 10)
        add_heading_paragraph(doc, "5. Modelo de Entrega y Asunciones", level=3)
        gov = prop.governance_and_assumptions
        team = prop.delivery_team
        add_body_paragraph(
            doc,
            f"Modelo de Gobierno: {clean_commercial_text(gov.governance_model)}",
            color_rgb=RGBColor(46, 64, 77),
            space_after=6,
        )
        p_team = doc.add_paragraph()
        r_team = p_team.add_run("Perfiles Clave:")
        r_team.bold = True
        r_team.font.color.rgb = RGBColor(46, 64, 77)
        p_team.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            team.team_roles,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        p_as = doc.add_paragraph()
        r_as = p_as.add_run("Asunciones Críticas (Assumptions):")
        r_as.bold = True
        r_as.font.color.rgb = RGBColor(46, 64, 77)
        p_as.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            gov.assumptions,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        add_spacer(doc, 10)
        add_heading_paragraph(doc, "6. Gestión de Riesgos del Proyecto", level=3)
        risks = prop.risk_management
        if risks:
            rtable = doc.add_table(rows=1, cols=2)
            finalize_table(rtable)
            set_cell_text(
                rtable.rows[0].cells[0],
                "Riesgo Identificado",
                bold=True,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text(
                rtable.rows[0].cells[1],
                "Estrategia de Mitigación NTT DATA",
                bold=True,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(rtable.rows[0].cells[0], "0072BC")
            shade_cell(rtable.rows[0].cells[1], "0072BC")
            for cell in rtable.rows[0].cells:
                for r in cell.paragraphs[0].runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
            for risk_item in risks:
                row = rtable.add_row()
                set_cell_text(
                    row.cells[0],
                    clean_commercial_text(risk_item.risk),
                    font_size=10.5,
                )
                set_cell_text(
                    row.cells[1],
                    clean_commercial_text(risk_item.mitigation),
                    font_size=10.5,
                )
                shade_cell(row.cells[0], "F2F2F2")
                shade_cell(row.cells[1], "F2F2F2")
            add_spacer(doc, 10)
        add_heading_paragraph(doc, "7. Por qué NTT DATA", level=3)
        why = prop.why_ntt_data
        add_smart_bullet_list(
            doc,
            why.accelerators,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=False,
        )
        add_body_paragraph(
            doc,
            f"Partnerships Clave: {clean_commercial_text(why.partnerships)}",
            color_rgb=RGBColor(46, 64, 77),
        )
        p_cred = doc.add_paragraph()
        r_cred = p_cred.add_run(
            "[⚠️ ACCIÓN COMERCIAL REQUERIDA: INSERTE AQUÍ AL MENOS 2 CREDENCIALES DE PROYECTOS SIMILARES EN EL SECTOR DEL CLIENTE ANTES DE ENVIAR]"
        )
        r_cred.bold = True
        r_cred.font.color.rgb = RGBColor(238, 63, 50)
        p_cred.paragraph_format.space_before = Pt(6)
        p_cred.paragraph_format.space_after = Pt(10)
        add_heading_paragraph(
            doc, "8. Inversión y Plan de Activación (Next 14 Days)", level=3
        )
        inv = prop.investment_and_timeline
        add_smart_bullet_list(
            doc,
            [
                f"Duración Estimada: {clean_commercial_text(inv.estimated_duration)}",
                f"Rango de Inversión (TCV): {clean_commercial_text(inv.tcv_range)}",
            ],
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        p_act = doc.add_paragraph()
        r_act = p_act.add_run("Plan de Activación Inmediata:")
        r_act.bold = True
        r_act.font.color.rgb = RGBColor(46, 64, 77)
        p_act.paragraph_format.space_after = Pt(4)
        add_smart_bullet_list(
            doc,
            prop.activation_plan,
            color_rgb=RGBColor(46, 64, 77),
            bold_prefix=True,
        )
        add_spacer(doc, 15)


def process_footnotes(doc, dossier):
    """Strip residual placeholder reference tags from a `python-docx` document.

    Performs a final sanitization pass on a document object, modifying it in-place
    to remove placeholder tags before serialization. The function iterates through
    all paragraphs in the main body and within all table cells, removing any text
    matching the pattern ``[[REF:...]]``. This prevents rendering artifacts from
    appearing in the final output.

    Args:
        doc (docx.document.Document): The `python-docx` document to process.
            This object is modified in-place.
        dossier (Any): Unused argument included for API consistency; reserved for
            future use.

    Returns:
        None.

    Raises:
        AttributeError: If `doc` is not a `python-docx` Document object and lacks
            expected attributes such as `.paragraphs` or `.tables`.
    """
    # Prior to serialization, executes a final recursive scan to strip any residual `[[REF:...]]` placeholder tags from all document objects. This sanitization step prevents rendering artifacts from appearing in the final output.
    ref_pattern = re.compile(r"\[\[REF:[^\]]*\]\]")

    def strip_refs_from_paragraph(p):
        """Strip internal reference markers from a paragraph object in-place."""
        if "[[" in p.text:
            for run in p.runs:
                if "[[" in run.text:
                    run.text = ref_pattern.sub("", run.text)

    for p in doc.paragraphs:
        strip_refs_from_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    strip_refs_from_paragraph(p)


def render_commercial_report(
    payload: CommercialPayload,
    template_path: Path,
    output_path: Path,
) -> Path:
    """Generates a commercial report by populating a Word document template.

    Orchestrates the creation of a complete commercial report. The function loads
    a specified `.docx` template, clears its existing body content, and adds
    page numbering to the footer. It then sequentially renders various sections
    of the report—such as the cover, summary, and opportunity pipeline—based
    on the presence of corresponding data in the `payload` object. Finally,
    it processes and inserts footnotes before saving the completed document to
    the specified output path.

    Args:
        payload (CommercialPayload): A data object containing all structured
            information required for the report sections.
        template_path (Path): The file system path to the source `.docx`
            template document.
        output_path (Path): The file system path where the final rendered `.docx`
            report will be saved.

    Returns:
        Path: The file system path of the generated report, identical to
            `output_path`.

    Raises:
        FileNotFoundError: If the document template at `template_path` does not
            exist.
        IOError: If the rendered document cannot be saved to `output_path` due
            to file system or permission issues.
    """
    doc = Document(str(template_path))
    clear_document_body(doc)

    if doc.sections:
        create_page_number_footer(doc.sections[0])

    render_commercial_cover(doc, payload)

    if payload.commercial_summary:
        render_commercial_summary(
            doc, payload.commercial_summary, payload.stakeholder_matrix
        )
    if payload.gtm_strategy:
        render_gtm_strategy(doc, payload.gtm_strategy)
    if payload.opportunities_pipeline:
        render_opportunities_pipeline(doc, payload.opportunities_pipeline)
    if payload.proactive_proposals:
        render_proactive_proposals(doc, payload.proactive_proposals)

    process_footnotes(doc, payload.intelligence_dossier)
    doc.save(str(output_path))
    return output_path


def main(argv: list[str] | None = None) -> None:
    """Renders a commercial report from the command line.

    This function serves as the primary entry point for script execution. It
    parses command-line arguments specifying a payload file, a template file,
    and an output path. It then orchestrates the data loading, report
    rendering, and result serialization processes.

    The script exits with a non-zero status code if an incorrect number of
    arguments is provided.

    Command-line usage:
        python render_commercial_report.py <payload_path> <template_path> <output_path>

    Args:
        argv: A list of command-line arguments, defaulting to `sys.argv`. The
            list must contain exactly four string elements in the following
            order: script name, path to the JSON payload file, path to the
            report template file, and the destination path for the output file.

    Returns:
        None.

    Raises:
        SystemExit: If the number of command-line arguments is not equal to 4.
        FileNotFoundError: If the specified payload or template file path does
            not exist.
    """
    if len(argv if argv is not None else sys.argv) != 4:
        sys.exit(1)
    payload_path = Path((argv if argv is not None else sys.argv)[1])
    template_path = Path((argv if argv is not None else sys.argv)[2])
    output_path = Path((argv if argv is not None else sys.argv)[3])
    payload = load_payload(payload_path)

    render_commercial_report(
        payload=payload,
        template_path=template_path,
        output_path=output_path,
    )
    print(f"Account Action Plan renderizado: {output_path}")


if __name__ == "__main__":
    main()
