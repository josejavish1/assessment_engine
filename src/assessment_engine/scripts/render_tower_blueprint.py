"""Renders a DOCX tower blueprint document from a corresponding JSON definition."""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt, RGBColor

from assessment_engine.schemas.blueprint import BlueprintPayload, PillarBlueprintDraft
from assessment_engine.scripts.lib.client_intelligence import (
    load_client_intelligence_legacy_view,
)
from assessment_engine.scripts.lib.docx_render_utils import (
    add_body_paragraph as _orig_add_body_paragraph,
)
from assessment_engine.scripts.lib.docx_render_utils import (
    add_heading_paragraph,
    autofit_table_to_contents,
    clear_paragraph,
    finalize_table,
    set_cell_text,
    shade_cell,
)
from assessment_engine.scripts.lib.maturity_band import (
    ANNEX_MATURITY_BANDS,
    resolve_maturity_band,
)
from assessment_engine.scripts.lib.runtime_paths import (
    resolve_tower_annex_template_path,
)
from assessment_engine.scripts.lib.text_utils import clean_text_for_word

BASE_TEXT_COLOR = RGBColor(46, 64, 77)
ROOT = Path(__file__).resolve().parents[3]
DEFAULT_TEMPLATE_PATH = resolve_tower_annex_template_path()


def load_json(path: Path) -> dict:
    """Return a dictionary by parsing a JSON file with UTF-8-SIG encoding."""
    return json.loads(path.read_text(encoding="utf-8-sig"))


def add_spacer(doc, points=12):
    """Add a vertical spacer of a specified point size to a document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)


def clean_text_for_render(text):
    """Formats an object into a human-readable string.

    This function applies conditional formatting based on the input's type and
    structure. If the input is a dictionary containing both "name" and
    "description" keys, it is formatted as "name: description". Other
    dictionaries are formatted by joining their string-converted values with
    " - ". All other data types are passed to the `clean_text_for_word`
    helper function for processing.

    Args:
        text (typing.Any): The input object to format.

    Returns:
        str: A string representation of the object, formatted for display.
    """
    if isinstance(text, dict):
        if "name" in text and "description" in text:
            return f"{text['name']}: {text['description']}"
        return " - ".join(str(v) for v in text.values())
    return clean_text_for_word(text)


def _safe_float(value: object) -> float | None:
    if value in (None, ""):
        return None
    try:
        return float(str(value).split()[0].replace(",", "."))
    except (TypeError, ValueError):
        return None


def _resolve_annex_band(value: object) -> str:
    score = _safe_float(value)
    if score is None:
        return ""
    return resolve_maturity_band(score, ANNEX_MATURITY_BANDS)["label"]


def add_body_paragraph(doc, text, color_rgb=BASE_TEXT_COLOR):
    """Add a cleaned and justified body paragraph with 12 points of trailing space."""
    text = clean_text_for_render(text)
    return _orig_add_body_paragraph(
        doc,
        text,
        space_after=12,
        color_rgb=color_rgb,
        justify=True,
    )


def add_bullet_p(container, text, color_rgb=BASE_TEXT_COLOR):
    """Adds a paragraph formatted as a bullet point to a python-docx container.

    The function sanitizes the input text and adds a new, justified paragraph to the specified container. It attempts to apply the 'List Bullet' style. If this style is not available in the document (a `KeyError`), it manually simulates a bullet point by prepending a '•' character run and setting a hanging indent.

    If the input text contains a colon (':'), the portion of the string preceding the colon is rendered in bold.

    Args:
        container (Union[docx.document.Document, docx.table._Cell]): The `python-docx` object to which the new paragraph will be added, such as a document body or a table cell.
        text (str): The string content for the bullet point.
        color_rgb (docx.shared.RGBColor): The RGB color to apply to the text. Defaults to the `BASE_TEXT_COLOR` constant.

    Returns:
        docx.text.paragraph.Paragraph: The newly created and appended paragraph object.
    """
    text = clean_text_for_render(text)
    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)

    try:
        p.style = "List Bullet"
    except KeyError:
        p.paragraph_format.left_indent = Pt(20)
        p.paragraph_format.first_line_indent = Pt(-15)
        bullet = p.add_run("• ")
        bullet.font.size = Pt(10.5)
        if color_rgb:
            bullet.font.color.rgb = color_rgb

    if ":" in text:
        prefix, rest = text.split(":", 1)
        prefix_run = p.add_run(prefix + ":")
        prefix_run.bold = True
        prefix_run.font.size = Pt(10.5)
        if color_rgb:
            prefix_run.font.color.rgb = color_rgb

        rest_run = p.add_run(rest)
        rest_run.font.size = Pt(10.5)
        if color_rgb:
            rest_run.font.color.rgb = color_rgb
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        if color_rgb:
            run.font.color.rgb = color_rgb
    return p


def create_page_number_footer(section):
    """Configures a document section footer to display 'Page X of Y' numbering.

    This function modifies the footer of a given `docx.section.Section` object to
    include dynamic page numbering in the format "Página {current_page} de
    {total_pages}". It also configures the section to use a distinct footer for
    the first page, a common requirement for title pages.

    The function first clears the content of the footer's default paragraph.
    It then constructs the page number fields by directly manipulating the
    underlying Office Open XML (OXML). This low-level approach is employed as
    the `python-docx` library lacks a high-level API for inserting the `PAGE` and
    `NUMPAGES` fields required for this functionality.

    Args:
        section (docx.section.Section): The document section object to modify.

    Returns:
        None.

    Raises:
        IndexError: If the section's footer does not contain any paragraphs to
            modify, which can occur with newly created or malformed sections.
    """
    section.different_first_page_header_footer = True
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph(paragraph)

    def add_field(p, field_code):
        r"""{'docstring': "Inserts a Word field into a paragraph via direct OpenXML manipulation.\n\n    Directly manipulates the underlying OpenXML tree to insert a complex field\n    into the specified paragraph. This is achieved by creating a new run (`w:r`)\n    and appending the required sequence of field character (`w:fldChar`) and\n    instruction text (`w:instrText`) elements.\n\n    The entire field is contained within this single run, which is styled with a\n    9pt Arial font and grey color (RGB 127, 127, 127). A default result text\n    of '0' is included as a placeholder, which is displayed until the\n    document's fields are updated by a word processor.\n\n    Args:\n        p (docx.text.paragraph.Paragraph): The target paragraph object for field\n            insertion.\n        field_code (str): The instruction string that defines the field's\n            behavior (e.g., 'PAGE \\* MERGEFORMAT').\n\n    Returns:\n        None"}."""
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(127, 127, 127)
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(ns.qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)
        instr_text = OxmlElement("w:instrText")
        instr_text.set(ns.qn("xml:space"), "preserve")
        instr_text.text = field_code
        run._r.append(instr_text)
        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(ns.qn("w:fldCharType"), "separate")
        run._r.append(fld_char_sep)
        text = OxmlElement("w:t")
        text.text = "0"
        run._r.append(text)
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(ns.qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)

    prefix = paragraph.add_run("Página ")
    prefix.font.size = Pt(9)
    prefix.font.name = "Arial"
    prefix.font.color.rgb = RGBColor(127, 127, 127)
    add_field(paragraph, "PAGE")
    middle = paragraph.add_run(" de ")
    middle.font.size = Pt(9)
    middle.font.name = "Arial"
    middle.font.color.rgb = RGBColor(127, 127, 127)
    add_field(paragraph, "NUMPAGES")


def clear_document_body(doc) -> None:
    r"""{'docstring': 'Clear all block-level content from the document body while preserving section properties.\n\nThis function performs an in-place modification of the document by directly\nmanipulating its underlying OOXML lxml element tree. It iterates through all\ndirect children of the `<w:body>` element and removes all block-level content,\nsuch as paragraphs (`<w:p>`) and tables (`<w:tbl>`).\n\nThe final `<w:sectPr>` (section properties) child element of the body is\nexplicitly preserved. This is critical for maintaining page layout settings\nlike margins, page size, and orientation when new content is subsequently\nadded to the document.\n\nNote: This function relies on the private `_body._element` attribute of the\n`python-docx` Document object, which is not a public API and may change in\nfuture versions of the library.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` Document object to modify\n        in-place.\n\nRaises:\n    AttributeError: If the provided `doc` object does not have the expected\n        internal `_body._element` structure, potentially due to a change in\n        the `python-docx` library version.'}."""
    body = doc._body._element
    for child in list(body):
        if (
            child.tag
            != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
        ):
            body.remove(child)


def resolve_client_dir(payload_path: Path, payload_data: dict) -> Path:
    r"""{'docstring': 'Derives a client\'s working directory from its payload file path.\n\nThe function first attempts to infer the directory by inspecting the\ngrandparent directory of `payload_path`. It performs a case-insensitive\ncomparison between this directory\'s name and the `client_name` from\n`payload_data` after normalizing the name by lowercasing and replacing\nspaces with underscores.\n\nIf the names match, the grandparent directory is returned. Otherwise,\na fallback path is constructed by joining a predefined `ROOT` directory,\na "working" subdirectory, and the normalized client name.\n\nArgs:\n    payload_path: The `pathlib.Path` object for the payload file.\n    payload_data: A dictionary containing payload data, expected to have the\n        structure `{\'document_meta\': {\'client_name\': \'...\'}}`.\n\nReturns:\n    A `pathlib.Path` object representing the resolved client directory.'}."""
    client_name = payload_data.get("document_meta", {}).get("client_name", "")
    candidate = payload_path.parents[1]
    if candidate.name.lower() == str(client_name).lower().replace(" ", "_"):
        return candidate
    return ROOT / "working" / str(client_name).lower().replace(" ", "_")


def load_client_intelligence(client_dir: Path) -> dict:
    """Load client intelligence data from `client_intelligence.json` within a directory."""
    path = client_dir / "client_intelligence.json"
    return load_client_intelligence_legacy_view(path)


def load_annex_data(client_dir: Path, tower_code: str) -> dict:
    """Load the approved annex template payload for a tower, returning an empty dict if the file does not exist."""
    tower_dir = client_dir / tower_code.upper()
    path = tower_dir / f"approved_annex_{tower_code.lower()}.template_payload.json"
    if path.exists():
        return load_json(path)
    return {}


def _list_or_default(value, default=None):
    if isinstance(value, list):
        return value
    return default or []


def _string_or_default(value, default=""):
    if isinstance(value, str):
        return value
    return default


def _derive_executive_snapshot(data: dict, annex_data: dict) -> dict:
    snapshot = data.get("executive_snapshot") or {}
    annex_summary = annex_data.get("executive_summary", {})
    annex_sections = annex_data.get("sections", {})
    risks = annex_sections.get("risks", {}).get("risks", [])
    todo_items = annex_sections.get("todo", {}).get("priority_initiatives", [])
    target_capabilities = annex_sections.get("tobe", {}).get("target_capabilities", [])

    if isinstance(annex_summary.get("summary_body"), list):
        summary_body = " ".join(str(x) for x in annex_summary.get("summary_body", []))
    else:
        summary_body = _string_or_default(annex_summary.get("summary_body"))

    return {
        "bottom_line": _string_or_default(snapshot.get("bottom_line"))
        or summary_body
        or _string_or_default(annex_summary.get("headline"))
        or "La torre requiere una transformación priorizada para reducir riesgo y habilitar el negocio.",
        "decisions": _list_or_default(snapshot.get("decisions"))
        or [
            clean_text_for_render(item.get("initiative", ""))
            for item in todo_items[:4]
            if item.get("initiative")
        ]
        or ["Validar el backlog priorizado y su secuencia de ejecución."],
        "cost_of_inaction": _string_or_default(snapshot.get("cost_of_inaction"))
        or "Mantener el estado actual prolonga el riesgo operativo, la deuda técnica y la incapacidad de escalar con control.",
        "structural_risks": _list_or_default(snapshot.get("structural_risks"))
        or [
            clean_text_for_render(item.get("risk", ""))
            for item in risks[:4]
            if item.get("risk")
        ],
        "business_impact": _string_or_default(snapshot.get("business_impact"))
        or "La modernización de la torre reduce exposición al riesgo y mejora la capacidad de ejecución del negocio.",
        "operational_benefits": _list_or_default(snapshot.get("operational_benefits"))
        or [clean_text_for_render(item) for item in target_capabilities[:4]],
        "transformation_complexity": _string_or_default(
            snapshot.get("transformation_complexity")
        )
        or "La transformación exige coordinación de arquitectura, operación y gobierno, pero es abordable por fases.",
    }


def _derive_cross_capabilities_analysis(data: dict) -> dict:
    cca = data.get("cross_capabilities_analysis") or {}
    pillars = _list_or_default(data.get("pillars_analysis"))
    low_score_pillars = [
        p.get("pilar_name", "") for p in pillars if float(p.get("score", 0) or 0) < 3.0
    ]

    if cca:
        return {
            "common_deficiency_patterns": _list_or_default(
                cca.get("common_deficiency_patterns")
            ),
            "transformation_paradigm": _string_or_default(
                cca.get("transformation_paradigm")
            )
            or "La torre requiere una evolución por dominios, priorizando estabilización, gobierno y posterior industrialización.",
            "critical_technical_debt": _string_or_default(
                cca.get("critical_technical_debt")
            )
            or "La deuda técnica limita la resiliencia y la capacidad de operar con consistencia.",
        }

    deficiency_patterns = []
    if low_score_pillars:
        deficiency_patterns.append(
            "Las mayores brechas se concentran en: "
            + ", ".join(low_score_pillars)
            + "."
        )

    return {
        "common_deficiency_patterns": deficiency_patterns
        or [
            "Persisten carencias repetidas en estandarización, automatización y gobierno técnico."
        ],
        "transformation_paradigm": "La transformación debe ejecutarse de forma incremental, combinando quick wins con capacidades fundacionales de largo recorrido.",
        "critical_technical_debt": "La deuda técnica acumulada incrementa el riesgo operativo y reduce la velocidad de adopción de nuevos servicios.",
    }


def _derive_roadmap(data: dict) -> list[dict]:
    roadmap = _list_or_default(data.get("roadmap"))
    if roadmap:
        return roadmap

    initiatives = []
    for pillar in _list_or_default(data.get("pillars_analysis")):
        for project in _list_or_default(pillar.get("projects_todo")):
            name = clean_text_for_render(project.get("name", ""))
            sizing = str(project.get("sizing", "")).lower()
            initiatives.append((name, sizing))

    quick_wins = [name for name, sizing in initiatives if name and "s" in sizing][:4]
    medium_term = [name for name, sizing in initiatives if name and "m" in sizing][:4]
    strategic = [name for name, sizing in initiatives if name and "l" in sizing][:4]

    fallback = [name for name, _ in initiatives[:4] if name]

    waves = [
        {"wave": "Wave 1", "projects": quick_wins or fallback[:2]},
        {"wave": "Wave 2", "projects": medium_term or fallback[2:4]},
        {"wave": "Wave 3", "projects": strategic or fallback[:2]},
    ]
    return [wave for wave in waves if wave["projects"]]


def normalize_blueprint_payload_dict(data: dict, annex_data: dict) -> dict:
    """Derives and defaults key sections to normalize a blueprint payload.

    Creates a normalized version of a raw blueprint payload. This function
    operates on a shallow copy of the input `data`, leaving the original
    dictionary unmodified. It enriches the payload by deriving values for the
    'executive_snapshot', 'cross_capabilities_analysis', and 'roadmap' sections.
    It also ensures structural consistency by defaulting 'external_dependencies'
    and 'pillars_analysis' to an empty list if they are absent from the input.

    Args:
        data: The source blueprint payload dictionary.
        annex_data: A dictionary containing auxiliary data required for
            derivation logic (e.g., for the executive snapshot).

    Returns:
        A new, normalized blueprint dictionary with all required sections
        populated.
    """
    normalized = dict(data)
    normalized["executive_snapshot"] = _derive_executive_snapshot(data, annex_data)
    normalized["cross_capabilities_analysis"] = _derive_cross_capabilities_analysis(
        data
    )
    normalized["roadmap"] = _derive_roadmap(data)
    normalized["external_dependencies"] = _list_or_default(
        data.get("external_dependencies")
    )
    normalized["pillars_analysis"] = _list_or_default(data.get("pillars_analysis"))
    return normalized


def load_payload(
    payload_path: Path, annex_data: dict | None = None
) -> BlueprintPayload:
    """Load, normalize, and validate a blueprint payload from a JSON file.

    Reads a JSON file from a given path, normalizes its content by merging it
    with optional supplementary data, and validates the result against the
    `BlueprintPayload` Pydantic model.

    Args:
        payload_path: The file system path to the blueprint's JSON payload.
        annex_data: An optional dictionary of supplementary data to merge into
            the main payload before validation.

    Returns:
        A validated `BlueprintPayload` instance.

    Raises:
        FileNotFoundError: If `payload_path` does not point to an existing file.
        json.JSONDecodeError: If the file content is not valid JSON.
        pydantic.ValidationError: If the normalized data does not conform to the
            `BlueprintPayload` schema.
    """
    raw_data = load_json(payload_path)
    normalized_data = normalize_blueprint_payload_dict(raw_data, annex_data or {})
    return BlueprintPayload.model_validate(normalized_data)


def render_cover(doc, payload: BlueprintPayload):
    """Constructs and appends a styled cover page to a document object.

    This function generates a cover page containing a report title, client name,
    technical tower details, and a confidentiality disclaimer. All content is
    styled with specific fonts, sizes, colors, and layout spacing. The
    provided document object is modified in place, and a page break is added
    at the end.

    Args:
        doc: The `python-docx` Document object to which the cover page will
            be appended.
        payload: A data structure containing the necessary metadata, specifically
            from its `document_meta` attribute.

    Raises:
        AttributeError: If `payload.document_meta` or its required attributes
            (e.g., `tower_name`, `client_name`) do not exist.
    """
    meta = payload.document_meta
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    title_p = doc.add_paragraph()
    run = title_p.add_run(f"{meta.tower_name}\nInforme de Madurez Tecnológica")
    run.font.size = Pt(34)
    run.font.name = "Georgia"
    run.font.color.rgb = RGBColor(0, 114, 188)
    add_spacer(doc, 30)

    client_p = doc.add_paragraph()
    client_run = client_p.add_run(meta.client_name.upper())
    client_run.font.size = Pt(24)
    client_run.font.name = "Arial"
    client_run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(100)
    version_p = doc.add_paragraph()
    version_text = (
        f"Torre Técnica: {meta.tower_code}\nHorizonte: {meta.transformation_horizon}"
    )
    version_run = version_p.add_run(version_text)
    version_run.font.size = Pt(14)
    version_run.font.name = "Arial"
    version_run.font.color.rgb = RGBColor(127, 127, 127)

    add_spacer(doc, 150)
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    title_run = disclaimer.add_run("Confidencialidad: ")
    title_run.bold = True
    title_run.font.size = Pt(8)
    title_run.font.color.rgb = RGBColor(127, 127, 127)
    text_run = disclaimer.add_run(
        "Este documento técnico es propiedad de NTT DATA y el cliente. Contiene "
        "información estratégica y de arquitectura sujeta a acuerdos de confidencialidad."
    )
    text_run.font.size = Pt(8)
    text_run.font.color.rgb = RGBColor(127, 127, 127)
    doc.add_page_break()


def render_snapshot_page(
    doc,
    payload: BlueprintPayload,
    client_intelligence: dict,
    annex_data: dict,
):
    """Constructs and appends the executive snapshot section to a document object.

    This function generates the 'Executive Snapshot' chapter of a report by
    populating a `docx.Document` object. It constructs a summary table with
    current, target, and maturity scores, followed by several analytical
    sections: business context (derived from client intelligence),
    structural business risks, cost of inaction, expected business impact,
    operational benefits, transformation complexity, and key decisions.

    Args:
        doc (docx.document.Document): The document object to be populated. This
            object is modified in-place.
        payload (BlueprintPayload): The main data transfer object containing the
            `executive_snapshot` attribute, which holds the core content for the
            section.
        client_intelligence (dict): A dictionary containing client-specific
            business intelligence. Expected keys include `ceo_agenda`,
            `regulatory_frameworks`, and `technological_drivers`.
        annex_data (dict): A dictionary containing supplemental data. The function
            extracts the `executive_summary` key, which is expected to contain
            `global_score`, `global_band`, and `target_maturity` values.

    Raises:
        AttributeError: If `payload.executive_snapshot` or any of its required
            nested fields (e.g., `bottom_line`, `decisions`,
            `structural_risks`) are missing from the payload object.
    """
    snap = payload.executive_snapshot
    exec_sum = annex_data.get("executive_summary", {})
    global_score_val = exec_sum.get("global_score", "N/A")
    global_band_val = exec_sum.get("global_band") or _resolve_annex_band(
        global_score_val
    )
    target_score_val = exec_sum.get("target_maturity", "N/A")

    add_heading_paragraph(doc, "1. Executive Snapshot (Resumen ejecutivo)", level=1)

    table = doc.add_table(rows=2, cols=3)
    finalize_table(table)
    headers = ["SCORE ACTUAL", "NIVEL DE MADUREZ", "MADUREZ OBJETIVO"]
    values = [global_score_val, global_band_val or "N/A", target_score_val]

    for i, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            header,
            bold=True,
            font_size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table.rows[0].cells[i], "0072BC")
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_text(
            table.rows[1].cells[i],
            values[i],
            font_size=14,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    if tbl_pr is not None:
        widths = tbl_pr.xpath("w:tblW")
        tbl_w = widths[0] if widths else OxmlElement("w:tblW")
        if not widths:
            tbl_pr.append(tbl_w)
        tbl_w.set(ns.qn("w:type"), "pct")
        tbl_w.set(ns.qn("w:w"), "5000")

    add_spacer(doc, 15)
    add_body_paragraph(doc, snap.bottom_line, color_rgb=BASE_TEXT_COLOR)

    business_angles = []
    ceo_agenda = clean_text_for_render(client_intelligence.get("ceo_agenda", ""))
    if ceo_agenda:
        business_angles.append("Agenda de negocio prioritaria: " + ceo_agenda)
    regulatory = client_intelligence.get("regulatory_frameworks", []) or []
    if regulatory:
        business_angles.append(
            "Presión regulatoria material: " + ", ".join(str(x) for x in regulatory)
        )
    drivers = client_intelligence.get("technological_drivers", []) or []
    for text in drivers:
        lowered = str(text).lower()
        if any(
            token in lowered
            for token in ["adquis", "m&a", "expansi", "ia", "pacient", "eficien"]
        ):
            business_angles.append(clean_text_for_render(text))

    if not business_angles:
        for candidate in [
            snap.business_impact,
            snap.cost_of_inaction,
            snap.bottom_line,
        ]:
            cleaned = clean_text_for_render(candidate)
            if cleaned and cleaned not in business_angles:
                business_angles.append(cleaned)
            if len(business_angles) >= 3:
                break

    add_heading_paragraph(doc, "Por qué importa al negocio", level=2)
    for item in business_angles[:4]:
        add_bullet_p(doc, item)

    if snap.structural_risks:
        add_heading_paragraph(doc, "Riesgos de negocio más materiales", level=2)
        for risk in snap.structural_risks:
            add_bullet_p(doc, risk)

    add_heading_paragraph(doc, "Coste de Inacción (Do Nothing)", level=2)
    add_body_paragraph(doc, snap.cost_of_inaction, color_rgb=BASE_TEXT_COLOR)

    if snap.business_impact:
        add_heading_paragraph(doc, "Impacto Esperado en Negocio", level=2)
        add_body_paragraph(doc, snap.business_impact, color_rgb=BASE_TEXT_COLOR)

    if snap.operational_benefits:
        add_heading_paragraph(doc, "Beneficios Operativos Target", level=2)
        for benefit in snap.operational_benefits:
            add_bullet_p(doc, benefit)

    if snap.transformation_complexity:
        add_heading_paragraph(doc, "Complejidad de la Transformación", level=2)
        add_body_paragraph(
            doc,
            snap.transformation_complexity,
            color_rgb=BASE_TEXT_COLOR,
        )

    add_heading_paragraph(doc, "Decisiones prioritarias", level=2)
    for decision in snap.decisions:
        add_bullet_p(doc, decision)


def render_cross_capabilities_analysis(doc, payload: BlueprintPayload):
    """Populates a document object with the Cross-Capabilities Analysis section.

    This function renders a structured section into the provided document. It adds
    headings for "El Paradigma de Transformación," "Deuda Técnica Crítica," and
    "Patrones Comunes de Deficiencia," populating them with content from the
    payload. The deficiency patterns are rendered as a bulleted list.

    The function performs no action if the `cross_capabilities_analysis` attribute
    of the payload is falsy.

    Args:
        doc: The document object (e.g., `docx.document.Document`) to which the
            section will be added. This object is modified in-place.
        payload: A `BlueprintPayload` data object containing the analysis
            information.

    Returns:
        None
    """
    cca = payload.cross_capabilities_analysis
    if not cca:
        return

    add_heading_paragraph(doc, "Análisis Transversal de Capacidades", level=1)
    add_heading_paragraph(doc, "El Paradigma de Transformación", level=2)
    add_body_paragraph(doc, cca.transformation_paradigm, color_rgb=BASE_TEXT_COLOR)

    add_heading_paragraph(doc, "Deuda Técnica Crítica", level=2)
    add_body_paragraph(doc, cca.critical_technical_debt, color_rgb=BASE_TEXT_COLOR)

    add_heading_paragraph(doc, "Patrones Comunes de Deficiencia", level=2)
    for item in cca.common_deficiency_patterns:
        add_bullet_p(doc, item)


def render_pilar_detail(doc, pilar: PillarBlueprintDraft):
    r"""{'docstring': "Renders a comprehensive pillar section into a `python-docx` document.\n\n    Populates the document with three primary subsections for a given pillar:\n    an 'AS-IS' technical health check table, a 'TO-BE' target architecture\n    description, and a 'TO-DO' transformation backlog. The function directly\n    manipulates the document object, adding formatted headings, tables, and\n    paragraphs.\n\n    Args:\n        doc: The `python-docx` `Document` object to be modified in-place.\n        pilar: A data object encapsulating the pillar's blueprint. It is\n            expected to have attributes such as `pilar_name`, `health_check_asis`\n            (an iterable of health check findings), `target_architecture_tobe`\n            (an object with `vision` and `design_principles`), and `projects_todo`\n            (an iterable of project details).\n\n    Returns:\n        None. The `doc` object is modified in-place.\n\n    Raises:\n        AttributeError: If `pilar` or its nested objects lack required attributes.\n        TypeError: If an attribute on `pilar` expected to be iterable (e.g.,\n            `projects_todo`) is not."}."""
    add_heading_paragraph(doc, f"Capacidad: {pilar.pilar_name}", level=1)

    add_heading_paragraph(doc, "A. Health Check Técnico (AS-IS)", level=2)
    table = doc.add_table(rows=1, cols=3)
    finalize_table(table)
    headers = ["Capacidad Técnica", "Hallazgo / Evidencia", "Riesgo de Negocio"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=10)
        shade_cell(table.rows[0].cells[i], "D9EAF7")

    for row_data in pilar.health_check_asis:
        row = table.add_row()
        set_cell_text(row.cells[0], row_data.target_state, bold=True, font_size=10)
        set_cell_text(row.cells[1], row_data.risk_observed, font_size=10)
        set_cell_text(row.cells[2], row_data.impact, font_size=10)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = BASE_TEXT_COLOR
    autofit_table_to_contents(table)
    add_spacer(doc, 10)

    add_heading_paragraph(doc, "B. Arquitectura Objetivo (TO-BE)", level=2)
    add_body_paragraph(
        doc, pilar.target_architecture_tobe.vision, color_rgb=BASE_TEXT_COLOR
    )
    principles = doc.add_paragraph()
    principles_run = principles.add_run("Principios de Diseño:")
    principles_run.bold = True
    principles_run.font.size = Pt(10)
    principles_run.font.color.rgb = BASE_TEXT_COLOR
    for principle in pilar.target_architecture_tobe.design_principles:
        add_bullet_p(doc, principle)

    add_heading_paragraph(doc, "C. Transformation Backlog (Iniciativas TO-DO)", level=2)
    for project in pilar.projects_todo:
        project_table = doc.add_table(rows=5, cols=2)
        finalize_table(project_table)
        merged = project_table.rows[0].cells[0].merge(project_table.rows[0].cells[1])
        set_cell_text(merged, project.initiative.upper(), bold=True, font_size=11)
        shade_cell(merged, "0072BC")
        for run in merged.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

        rows = [
            ("Business Rationale", project.expected_outcome),
            ("Objetivo Técnico", project.objective),
            (
                "Entregables (DoD)",
                "\n".join([f"• {item}" for item in project.deliverables]),
            ),
            (
                "Sizing & Duración",
                f"Complejidad: {project.sizing} | Estimación: {project.duration}",
            ),
        ]
        for idx, (label, value) in enumerate(rows, 1):
            set_cell_text(
                project_table.rows[idx].cells[0], label, bold=True, font_size=9
            )
            shade_cell(project_table.rows[idx].cells[0], "F2F2F2")
            set_cell_text(project_table.rows[idx].cells[1], value, font_size=9.5)
            for run in project_table.rows[idx].cells[1].paragraphs[0].runs:
                run.font.color.rgb = BASE_TEXT_COLOR
        autofit_table_to_contents(project_table)
        add_spacer(doc, 10)


def render_roadmap_page(doc, payload: BlueprintPayload):
    """Renders the strategic roadmap and dependency matrix into a document object.

    This function populates a document with a section titled "4. Strategic
    Roadmap & Dependencies". This section contains two parts:

    1.  A hierarchical list of projects, where each wave from the payload is
        rendered as a level-2 heading followed by its constituent projects as
        bullet points.
    2.  A dependency matrix table titled "Matriz de Sinergias Cruzadas", which
        lists external dependencies with columns for the initiative, its
        dependency, and the technical reason.

    The function directly modifies the passed `doc` object.

    Args:
        doc: The `python-docx` `Document` object to be modified.
        payload: A `BlueprintPayload` data object containing the roadmap waves
            and external dependency data.

    Returns:
        None.

    Raises:
        AttributeError: If the `payload` object lacks the required `roadmap` or
            `external_dependencies` attributes or if their nested objects do not
            conform to the expected structure.
    """
    add_heading_paragraph(doc, "4. Strategic Roadmap & Dependencies", level=1)

    for wave in payload.roadmap:
        add_heading_paragraph(doc, wave.wave, level=2)
        for project in wave.projects:
            add_bullet_p(doc, project)

    add_heading_paragraph(doc, "Matriz de Sinergias Cruzadas", level=2)
    table = doc.add_table(rows=1, cols=3)
    finalize_table(table)
    headers = ["Iniciativa", "Dependencia / Habilita a", "Razón Técnica"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=10)
        shade_cell(table.rows[0].cells[i], "D9EAF7")

    for dependency in payload.external_dependencies:
        row = table.add_row()
        set_cell_text(row.cells[0], dependency.project, bold=True, font_size=9.5)
        set_cell_text(row.cells[1], dependency.depends_on, font_size=9.5)
        set_cell_text(row.cells[2], dependency.reason, font_size=9.5)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = BASE_TEXT_COLOR
    autofit_table_to_contents(table)


def render_maturity_profile(doc, annex_data: dict):
    r"""{'docstring': 'Renders the complete maturity profile section into a `python-docx` document.\n\nThis function constructs and appends a multi-part maturity profile section based on the provided data. The generated section includes a main heading, an introductory narrative, an optional radar chart visualization, a detailed table of scores for each pillar, an "AS-IS" summary, a table contrasting key strengths and gaps, and a bulleted list of key operational impacts.\n\nIf the `annex_data` dictionary is empty or evaluates to False, a placeholder paragraph is added instead and the function returns early. The function is designed to handle missing keys within the data structure gracefully by using default empty values, preventing crashes from incomplete data.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` document object to be modified in-place.\n    annex_data (dict): A dictionary containing all data required for rendering the profile. The expected structure is:\n        {\n            "pillar_score_profile": {\n                "profile_intro": str,\n                "scoring_method_note": str,\n                "radar_chart": str,  # Path to the radar chart image file.\n                "pillars": [\n                    {\n                        "pillar_label": str,\n                        "score_display": str,\n                        "maturity_band": str,  # Optional.\n                        "executive_reading": str,\n                    }, ...\n                ]\n            },\n            "sections": {\n                "asis": {\n                    "narrative": str,\n                    "strengths": [str, ...],\n                    "gaps": [str, ...],\n                    "operational_impacts": [str, ...]\n                }\n            }\n        }'}."""
    add_heading_paragraph(doc, "Perfil de madurez por pilar", level=1)

    if not annex_data:
        add_body_paragraph(
            doc,
            "No se ha encontrado el perfil de madurez detallado para esta torre.",
            color_rgb=BASE_TEXT_COLOR,
        )
        return

    profile = annex_data.get("pillar_score_profile", {})
    asis = annex_data.get("sections", {}).get("asis", {})

    add_body_paragraph(doc, profile.get("profile_intro", ""), color_rgb=BASE_TEXT_COLOR)
    note = add_body_paragraph(
        doc,
        profile.get("scoring_method_note", ""),
        color_rgb=BASE_TEXT_COLOR,
    )
    if note:
        for run in note.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(127, 127, 127)

    radar_path = profile.get("radar_chart", "")
    if radar_path and Path(radar_path).exists():
        add_heading_paragraph(doc, "Gráfico radial", level=2)
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(radar_path, width=Inches(4.35))
        add_spacer(doc, 10)

    pillars = profile.get("pillars", [])
    add_heading_paragraph(doc, "Detalle compacto por pilar", level=2)
    if pillars:
        table = doc.add_table(rows=1, cols=4)
        finalize_table(table)
        headers = ["Pilar", "Score", "Nivel", "Lectura ejecutiva"]
        for i, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=10)
            shade_cell(table.rows[0].cells[i], "D9EAF7")

        for pillar in pillars:
            row = table.add_row()
            maturity_band = pillar.get("maturity_band") or _resolve_annex_band(
                pillar.get("score_display")
            )
            set_cell_text(
                row.cells[0], pillar.get("pillar_label", ""), bold=True, font_size=9.5
            )
            set_cell_text(
                row.cells[1],
                pillar.get("score_display", ""),
                font_size=9.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text(row.cells[2], maturity_band, font_size=9.5)
            set_cell_text(
                row.cells[3], pillar.get("executive_reading", ""), font_size=9.5
            )
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BASE_TEXT_COLOR
        autofit_table_to_contents(table)
        add_spacer(doc, 10)

    add_heading_paragraph(doc, "AS-IS resumido", level=2)
    add_body_paragraph(doc, asis.get("narrative", ""), color_rgb=BASE_TEXT_COLOR)

    strengths = asis.get("strengths", [])
    gaps = asis.get("gaps", [])
    if strengths or gaps:
        sg_table = doc.add_table(rows=2, cols=2)
        finalize_table(sg_table)
        set_cell_text(
            sg_table.rows[0].cells[0], "Fortalezas clave", bold=True, font_size=10
        )
        shade_cell(sg_table.rows[0].cells[0], "D9EAF7")
        set_cell_text(
            sg_table.rows[0].cells[1], "Brechas clave", bold=True, font_size=10
        )
        shade_cell(sg_table.rows[0].cells[1], "D9EAF7")

        strengths_cell = sg_table.rows[1].cells[0]
        strengths_cell.text = ""
        for item in strengths:
            p = strengths_cell.add_paragraph(f"• {item}")
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.first_line_indent = Pt(-10)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = BASE_TEXT_COLOR

        gaps_cell = sg_table.rows[1].cells[1]
        gaps_cell.text = ""
        for item in gaps:
            p = gaps_cell.add_paragraph(f"• {item}")
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.first_line_indent = Pt(-10)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = BASE_TEXT_COLOR

        autofit_table_to_contents(sg_table)
        add_spacer(doc, 10)

    impacts = asis.get("operational_impacts", [])
    if impacts:
        add_heading_paragraph(doc, "Implicaciones operativas clave", level=2)
        for impact in impacts:
            p = doc.add_paragraph(f"• {impact}")
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.first_line_indent = Pt(-10)
            for run in p.runs:
                run.font.color.rgb = BASE_TEXT_COLOR


def render_conclusion(doc, annex_data: dict):
    r"""{'docstring': 'Renders the conclusion section of a report into a docx document.\n\n    This function populates a document with the conclusion chapter (Chapter 5),\n    including subsections for a final assessment, a message for the technical\n    manager, priority focus areas formatted as a bulleted list, and a closing\n    statement. If conclusion data is not present within `annex_data`, a\n    placeholder message is added instead.\n\n    Args:\n        doc (docx.document.Document): The document object to be modified.\n        annex_data (dict): A dictionary containing the report data. The function\n            expects a key path `sections.conclusion`. The `conclusion`\n            dictionary may contain the following optional keys:\n            - `final_assessment` (str): The text for the final assessment.\n            - `executive_message` (str): The message for the technical manager.\n            - `priority_focus_areas` (list[str]): A list of priority areas.\n            - `closing_statement` (str): The text for the next steps.\n\n    Returns:\n        None. The `doc` object is modified in-place.\n\n    Raises:\n        AttributeError: If `annex_data` is not a dictionary or a dictionary-like\n            object that supports the `.get()` method.'}."""
    add_heading_paragraph(doc, "5. Conclusión", level=1)

    if not annex_data:
        add_body_paragraph(
            doc,
            "No se ha encontrado información de conclusión para esta torre.",
            color_rgb=BASE_TEXT_COLOR,
        )
        return

    conclusion = annex_data.get("sections", {}).get("conclusion", {})

    if conclusion.get("final_assessment"):
        add_heading_paragraph(doc, "Evaluación final", level=2)
        add_body_paragraph(
            doc, conclusion.get("final_assessment"), color_rgb=BASE_TEXT_COLOR
        )

    if conclusion.get("executive_message"):
        add_heading_paragraph(doc, "Mensaje para el responsable técnico", level=2)
        add_body_paragraph(
            doc, conclusion.get("executive_message"), color_rgb=BASE_TEXT_COLOR
        )

    priority_areas = conclusion.get("priority_focus_areas", [])
    if priority_areas:
        add_heading_paragraph(doc, "Áreas de foco prioritarias", level=2)
        for area in priority_areas:
            p = doc.add_paragraph(f"• {area}")
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.first_line_indent = Pt(-10)
            for run in p.runs:
                run.font.color.rgb = BASE_TEXT_COLOR

    if conclusion.get("closing_statement"):
        add_heading_paragraph(doc, "Próximos pasos", level=2)
        add_body_paragraph(
            doc, conclusion.get("closing_statement"), color_rgb=BASE_TEXT_COLOR
        )


def render_blueprint(
    payload: BlueprintPayload,
    output_path: Path,
    client_dir: Path,
    template_path: Path = DEFAULT_TEMPLATE_PATH,
) -> Path:
    """Generate and save a complete tower blueprint document from a data payload.

    Orchestrates the generation of a Microsoft Word (.docx) document from a
    structured data payload. The function loads a base template and supplementary
    client data. After clearing the template's body, it sequentially calls
    rendering functions to populate the document with a cover page, analysis
    sections, pillar details, a roadmap, and a conclusion. The resulting
    document is then saved to the specified output path, creating parent
    directories if they do not exist.

    Args:
        payload (BlueprintPayload): A data object containing analysis results and
            metadata required for document population.
        output_path (Path): The filesystem path where the generated .docx file
            will be saved.
        client_dir (Path): The directory path containing supplementary
            client-specific data, such as annex and intelligence files.
        template_path (Path): The path to the base .docx template. Defaults to
            a pre-configured system path if not provided.

    Returns:
        Path: The path to the newly created document, identical to `output_path`.

    Raises:
        FileNotFoundError: If the specified `template_path` or required data
            files within `client_dir` cannot be found.
        PermissionError: If the process lacks permissions to create the output
            directory or write to the `output_path`.
        KeyError: If the `payload` object is missing required data attributes
            that are expected by the rendering subroutines.
    """
    annex_data = load_annex_data(client_dir, payload.document_meta.tower_code)
    client_intelligence = load_client_intelligence(client_dir)

    doc = Document(str(template_path))
    clear_document_body(doc)
    if doc.sections:
        create_page_number_footer(doc.sections[0])

    render_cover(doc, payload)
    render_snapshot_page(doc, payload, client_intelligence, annex_data)
    render_maturity_profile(doc, annex_data)
    render_cross_capabilities_analysis(doc, payload)

    for pillar in payload.pillars_analysis:
        render_pilar_detail(doc, pillar)

    render_roadmap_page(doc, payload)
    render_conclusion(doc, annex_data)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main(argv: list[str] | None = None) -> None:
    r"""{'docstring': 'Orchestrate the rendering of a Tower Blueprint from a command-line interface.\n\nServes as the main entry point for the command-line script. This function\nvalidates and parses command-line arguments to obtain input and output file\npaths. It then coordinates the loading of a primary JSON payload, resolves\nthe associated client and annex data, and invokes the rendering process to\ngenerate the final blueprint file.\n\nThe script terminates with a non-zero exit code if the command-line\narguments are not provided correctly.\n\nArgs:\n    argv: A list of command-line arguments. If `None`, `sys.argv` is used.\n        The list is expected to contain exactly three elements: the script\n        name, the path to the input JSON payload, and the path for the\n        output blueprint file.\n\nRaises:\n    FileNotFoundError: If the input payload file or other required data files,\n        such as annex data, cannot be found at their expected locations.\n    json.JSONDecodeError: If the input payload file is not a valid JSON document.\n    ValueError: If the content of the JSON payload is invalid or if a\n        required client directory cannot be resolved from the given paths.\n    IOError: If an error occurs while writing the rendered blueprint to the\n        specified output path.'}."""
    if len(argv if argv is not None else sys.argv) != 3:
        sys.exit(1)

    payload_path = Path((argv if argv is not None else sys.argv)[1])
    output_path = Path((argv if argv is not None else sys.argv)[2])
    client_dir = resolve_client_dir(payload_path, load_json(payload_path))
    annex_data = load_annex_data(client_dir, payload_path.parent.name)
    payload = load_payload(payload_path, annex_data=annex_data)

    render_blueprint(
        payload=payload,
        output_path=output_path,
        client_dir=client_dir,
    )
    print(f"Blueprint de Torre renderizado: {output_path}")


if __name__ == "__main__":
    main()
