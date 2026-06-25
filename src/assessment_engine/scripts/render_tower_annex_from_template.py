"""Generates a DOCX tower annex document by populating a template with data from a validated payload."""

import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from assessment_engine.schemas.annex_synthesis import AnnexPayload
from assessment_engine.scripts.generate_tower_radar_chart import (
    generate_radar_chart_from_pillars,
)
from assessment_engine.scripts.lib.contract_utils import robust_load_payload
from assessment_engine.scripts.lib.docx_render_utils import (
    add_body_paragraph,
    add_bullet_list,
    add_heading_paragraph,
    add_label_value_paragraph,
    add_long_detail_table,
    apply_bullet_list_format,
    apply_paragraph_style,
    apply_table_style,
    clean_text,
    clear_cell_shading,
    clear_paragraph,
    clear_paragraph_properties,
    enable_update_fields_on_open,
    insert_field_paragraph_after_block,
    insert_paragraph_after_block,
    remove_page_break_only_paragraphs,
    render_gap_table,
    render_initiative_cards,
    render_list_at_placeholder,
    render_multi_paragraph_block,
    render_note_box,
    render_pillar_score_table,
    render_radar_chart,
    render_risks_table,
    replace_simple_placeholder,
    resolve_radar_chart_image,
)

WORD_RENDER_MODE_ENV = "AE_WORD_RENDER_MODE"


def normalize_annex_payload(payload_dict: dict) -> dict:
    """Ensures a tower annex payload dictionary conforms to a canonical structure.

    This function takes a raw payload dictionary and returns a new, normalized
    dictionary. It enforces a standard data schema by initializing missing keys,
    populating default values, and transforming legacy structures for backward
    compatibility. The original input dictionary is not modified.

    Key normalization steps include:
      - Schema Enforcement: Guarantees the existence of top-level keys
        (`document_meta`, `executive_summary`, `sections`, and
        `domain_introduction`), initializing each as an empty dictionary if absent.
      - Executive Summary Population: Populates default values for `headline` and
        `key_business_impacts`. It also coalesces a list-based `summary_body`
        into a single multi-paragraph string.
      - Legacy Migration: If a `tobe_gap` section exists and `tobe` does not,
        it transforms the former into the canonical `tobe` and `gap` sections.
      - Default Content Generation: Constructs a `domain_introduction` section
        from other payload data (e.g., metadata and summary content) if it is
        not explicitly provided.

    Args:
        payload_dict: A dictionary representing the raw annex payload. It may
            be missing keys or contain legacy structures that require
            transformation.

    Returns:
        A new dictionary instance conforming to the canonical annex payload
        structure.
    """
    normalized = dict(payload_dict)
    meta = dict(normalized.get("document_meta") or {})
    executive_summary = dict(normalized.get("executive_summary") or {})
    sections = dict(normalized.get("sections") or {})

    summary_body = executive_summary.get("summary_body", "")
    if isinstance(summary_body, list):
        executive_summary["summary_body"] = "\n\n".join(
            part for part in summary_body if clean_text(part)
        )
    executive_summary.setdefault(
        "headline", f"Anexo ejecutivo de la torre {meta.get('tower_code', 'TX')}"
    )
    executive_summary.setdefault("key_business_impacts", [])

    if "tobe" not in sections and "tobe_gap" in sections:
        tobe_gap = dict(sections.get("tobe_gap") or {})
        sections["tobe"] = {
            "vision": tobe_gap.get("introduction", ""),
            "design_principles": tobe_gap.get("target_capabilities", []),
        }
        sections["gap"] = {
            "introduction": tobe_gap.get("introduction", ""),
            "target_capabilities": tobe_gap.get("target_capabilities", []),
            "gap_rows": tobe_gap.get("gap_rows", []),
            "closing_summary": tobe_gap.get("closing_summary", ""),
        }

    domain_introduction = dict(normalized.get("domain_introduction") or {})
    if not domain_introduction:
        domain_introduction = {
            "introduction_paragraph": executive_summary.get("summary_body", ""),
            "technological_domain": meta.get("tower_name", ""),
            "domain_objective": meta.get("tower_name", ""),
            "evaluated_capabilities": [],
            "included_components": [],
        }

    normalized["document_meta"] = meta
    normalized["executive_summary"] = executive_summary
    normalized["sections"] = sections
    normalized["domain_introduction"] = domain_introduction
    return normalized


def _replace_client_placeholders(value, client_name: str):
    if isinstance(value, dict):
        return {
            key: _replace_client_placeholders(item, client_name)
            for key, item in value.items()
        }
    if isinstance(value, list):
        return [_replace_client_placeholders(item, client_name) for item in value]
    if isinstance(value, str):
        return value.replace("[Cliente]", client_name).replace("[cliente]", client_name)
    return value


def insert_toc_after(paragraph) -> None:
    r"""{'docstring': 'Insert a styled Table of Contents (TOC) heading and field after a paragraph.\n\n    This function first inserts a new paragraph styled as a \'TOC Heading\' with\n    the text "Contents". It then inserts a second paragraph immediately after the\n    heading, which contains a standard Word TOC field code:\n    `TOC \\o "1-2" \\h \\z \\u`.\n\n    The field code is configured to build a table from heading levels 1 and 2\n    (`\\o "1-2"`), create hyperlinks to the entries (`\\h`), suppress tab leaders\n    for web layout (`\\z`), and use the applied paragraph style\'s outline level (`\\u`).\n\n    Note:\n        The generated TOC field must be updated within the Microsoft Word\n        application (e.g., by right-clicking and selecting "Update Field") to\n        populate the table of contents.\n\n    Args:\n        paragraph (docx.text.paragraph.Paragraph): The paragraph object after\n            which the new TOC section will be inserted.'}."""
    heading = insert_paragraph_after_block(paragraph)
    heading.add_run("Contents")
    apply_paragraph_style(heading, "TOC Heading", "TtuloTDC")
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    toc_paragraph = insert_field_paragraph_after_block(
        heading,
        'TOC \\o "1-2" \\h \\z \\u',
        placeholder_text="Update field to build contents.",
    )
    style_existing_field_paragraph(toc_paragraph, "toc 1", "TDC1")


def rewrite_paragraph_with_style(paragraph, text: str, *styles: str) -> None:
    r"""{'docstring': 'Clears, restyles, and sets the text content for a paragraph.\n\nModifies a paragraph object in-place by first removing all existing runs\nand formatting properties. It then applies the specified styles before adding\nthe cleaned input text as a new, single run. If the text is empty or\nbecomes empty after cleaning, the paragraph will be styled but have no\ncontent.\n\nArgs:\n    paragraph: The paragraph object to modify.\n    text: The new string content for the paragraph.\n    *styles: A variable number of string style names to apply.'}."""
    clean_value = clean_text(text)
    clear_paragraph(paragraph)
    clear_paragraph_properties(paragraph)
    apply_paragraph_style(paragraph, *styles)
    if clean_value:
        paragraph.add_run(clean_value)


def rewrite_body_paragraph(paragraph, text: str) -> None:
    """Rewrite a `docx.paragraph.Paragraph` object with new text and body styling.

    The function first sanitizes the input text. If the resulting string is
    empty, no action is taken. Otherwise, the specified paragraph is cleared of
    all existing content and formatting properties.

    A primary style (e.g., "NTT Body Text") is attempted first. If it fails,
    a fallback style ("Normal") is applied. The paragraph's format is then explicitly
    set to be justified, with 0pt space before, 8pt space after, and a 1.05
    line spacing. The new text is added as a run with a 10pt font size.

    Args:
        paragraph (docx.paragraph.Paragraph): The `python-docx` Paragraph object to
            be modified in-place.
        text (str): The new string content to insert into the paragraph.
    """
    clean_value = clean_text(text)
    if not clean_value:
        return
    clear_paragraph(paragraph)
    clear_paragraph_properties(paragraph)
    if not apply_paragraph_style(
        paragraph,
        "NTT Body Text",
        "NTTBodyText",
        "Body Text",
        "Textoindependiente",
    ):
        apply_paragraph_style(paragraph, "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(clean_value)
    run.font.size = Pt(10)


def style_existing_field_paragraph(paragraph, *styles: str) -> None:
    """Resets and applies specified styles and spacing to a paragraph object.

    This function first clears all existing paragraph-level formatting from the
    provided paragraph object. It then applies the specified styles and sets a
    fixed spacing of 0pt before, 3pt after, and single line spacing.

    Args:
        paragraph: The paragraph object to be modified in-place.
        *styles: The string names of one or more styles to apply.

    Returns:
        None.
    """
    clear_paragraph_properties(paragraph)
    apply_paragraph_style(paragraph, *styles)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0


def polish_semantic_tables(doc) -> None:
    """Apply a standardized visual style to all tables in a DOCX document.

    This function iterates through each table within the provided `python-docx`
    document and applies a uniform set of formatting rules. It first attempts to
    set a base table style from a predefined list of known style names. It then
    iterates through every cell to enforce specific formatting, overriding any
    conflicting style defaults.

    Cell-level formatting includes clearing any background shading and setting
    vertical alignment to the top. Paragraph-level formatting within cells is
    normalized to have 0pt of space before, 2pt after, a line spacing of 1.0,
    and left alignment. All text runs are set to a 10pt font size.

    The first row of each table is treated as a header, and its font color is
    set to white.

    Args:
        doc (docx.document.Document): The document object whose tables will be
            restyled in-place.

    Returns:
        None. The document object is modified directly.
    """
    for table in doc.tables:
        apply_table_style(
            table,
            "Tabla con cuadrícula 1 clara - Énfasis 11",
            "Tablaconcuadrcula1clara-nfasis11",
            "Smart Navy table",
            "SmartNavytable",
            "NTT Future Blue table",
            "NTTFutureBluetable",
            "Table Grid",
        )
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                clear_cell_shading(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.0
                    if row_idx == 0:
                        if paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        if row_idx == 0:
                            run.font.color.rgb = RGBColor(255, 255, 255)


def apply_semantic_annex_styles(doc, payload_dict: dict) -> None:
    """Iteratively applies paragraph and table styles to a `docx.Document` annex.

    This function processes a `docx.Document` object, applying specific Microsoft
    Word styles to its contents to conform to a standard annex format. It
    iterates through each paragraph, identifying structural elements such as
    titles, subtitles, headings, and bullet points based on their text content.

    - Title and subtitle text are dynamically generated using metadata from the
      `payload_dict`.
    - Headings (Level 1 and 2) are identified by matching paragraph text
      against predefined sets of Spanish-language strings.
    - Bullet points, identified by a '•' prefix, are converted into a
      standard list format.
    - A Table of Contents (TOC) is inserted after the subtitle if one does
      not already exist.
    - All tables within the document are styled using a dedicated helper.

    The modification is performed in-place on the input `doc` object.

    Args:
        doc: The `docx.Document` object to be styled.
        payload_dict: A dictionary containing document metadata. It must
            contain a `document_meta` key, which in turn holds `tower_code`,
            `tower_name`, and `client_name` keys.

    Returns:
        None. The input `doc` object is modified directly.

    Raises:
        KeyError: If `payload_dict` is missing the `document_meta` key, or
            if the `document_meta` dictionary is missing any of its required
            nested keys (`tower_code`, `tower_name`, `client_name`).
    """
    enable_update_fields_on_open(doc)

    heading_level_1 = {
        "Resumen ejecutivo de la torre",
        "Perfil de madurez por pilar",
        "AS-IS resumido",
        "Riesgos principales",
        "Estado objetivo y brechas",
        "Iniciativas prioritarias",
        "Conclusión",
        "Desarrollo ampliado",
        "1. AS-IS detallado",
        "2. Riesgos detallados",
        "3. Estado objetivo detallado",
        "4. Brechas detalladas",
        "5. Iniciativas priorizadas detalladas",
        "6. Cierre ampliado",
    }
    heading_level_2 = {
        "Fortaleza principal",
        "Brecha principal",
        "Cuello de botella",
        "Gráfico radial",
        "Detalle compacto por pilar",
        "Fortalezas clave",
        "Brechas clave",
        "Implicaciones operativas clave",
        "Principios / capacidades objetivo",
        "Tabla de brechas por pilar",
        "Fichas de iniciativas priorizadas",
        "Mensaje ejecutivo final",
        "Áreas prioritarias de actuación",
        "Fortalezas observadas",
        "Brechas observadas",
        "Implicaciones operativas",
        "Lectura transversal de gaps",
        "Principios de arquitectura y operación",
        "Implicaciones para el modelo operativo",
    }

    title_text = (
        f"ANEXO {payload_dict['document_meta']['tower_code']} – "
        f"{payload_dict['document_meta']['tower_name']}"
    )
    subtitle_prefix = payload_dict["document_meta"]["client_name"]
    subtitle_anchor = None

    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue

        if text == title_text:
            rewrite_paragraph_with_style(paragraph, text, "Title", "Ttulo")
            subtitle_anchor = paragraph
            continue

        if (
            text.startswith(subtitle_prefix)
            and "Fast Infrastructure Assessment" in text
        ):
            rewrite_paragraph_with_style(paragraph, text, "Subtitle", "Subttulo")
            subtitle_anchor = paragraph
            continue

        if text in heading_level_1:
            rewrite_paragraph_with_style(paragraph, text, "Heading 1", "Ttulo1")
            continue

        if text in heading_level_2:
            rewrite_paragraph_with_style(paragraph, text, "Heading 2", "Ttulo2")
            continue

        if text == "Update field to build contents.":
            style_existing_field_paragraph(paragraph, "toc 1", "TDC1")
            continue

        if text.startswith("• "):
            clear_paragraph(paragraph)
            clear_paragraph_properties(paragraph)
            apply_bullet_list_format(paragraph)
            run = paragraph.add_run(text[2:].strip())
            run.font.size = Pt(10)
            continue

        rewrite_body_paragraph(paragraph, text)

    if (
        subtitle_anchor is not None
        and "Update field to build contents."
        not in "\n".join(paragraph.text for paragraph in doc.paragraphs)
    ):
        insert_toc_after(subtitle_anchor)

    polish_semantic_tables(doc)


def _resolve_render_mode(args: list[str]) -> tuple[str, list[str]]:
    render_mode = (
        os.environ.get(WORD_RENDER_MODE_ENV, "semantic").strip().lower() or "semantic"
    )
    filtered_args: list[str] = []
    for arg in args:
        if arg == "--semantic-styles":
            render_mode = "semantic"
            continue
        if arg == "--legacy-styles":
            render_mode = "legacy"
            continue
        filtered_args.append(arg)
    return render_mode, filtered_args


def _ensure_radar_chart_path(payload_path: Path, profile: dict) -> str:
    radar_chart = clean_text(profile.get("radar_chart", ""))
    if resolve_radar_chart_image(radar_chart) is not None:
        return radar_chart

    pillars = profile.get("pillars", [])
    if not isinstance(pillars, list) or not pillars:
        return radar_chart

    generated_path = payload_path.with_name("pillar_radar_chart.generated.png")
    try:
        generate_radar_chart_from_pillars(pillars, generated_path)
    except ValueError:
        return radar_chart
    return str(generated_path)


def clean_brackets_and_consultant_notes(doc, payload: dict):
    """Substitute template placeholders and remove instructional elements in a document.

    Modifies a `docx.Document` object in-place by performing substitutions and
    purging content intended for internal use. The function iterates through all
    paragraphs and tables, replacing placeholder strings (e.g., `[Cliente]`) with
    values from the `payload`. Concurrently, it identifies and removes entire
    paragraphs and table rows that contain predefined instructional text or
    generic placeholders enclosed in square brackets (`[]`), producing a clean,
    final document.

    Args:
        doc: The `python-docx` Document object to be processed. This object is
            modified in-place.
        payload: A dictionary containing data for template substitution. Expected
            nested keys include 'document_meta' (with 'client_name', 'tower_name',
            'tower_code', 'date') and 'domain_introduction' (with
            'domain_objective', 'technological_domain', 'introduction_paragraph').
    """
    meta = payload.get("document_meta", {})
    intro = payload.get("domain_introduction", {})

    client_name = meta.get("client_name", "CLIENTE")
    tower_name = meta.get("tower_name", "TORRE")
    tower_code = meta.get("tower_code", "TX")

    # Define a mapping between template placeholder keys and corresponding data values from the validated contract payload for direct text substitution.
    replacements = {
        "[Cliente]": client_name,
        "[cliente]": client_name,
        "[Nombre del Cliente]": client_name,
        "[número]": "los",
        "[número de pilares]": "varios",
        "-[Lista de torres evaluadas]": f"• {tower_name}",
        "[Lista de torres evaluadas]": tower_name,
        "T[X]": tower_code,
        "[Nombre de la Torre]": tower_name,
        "[Código]": tower_code,
        "[Nombre de la torre]": tower_name,
        "[descripción del dominio]": intro.get("domain_objective", ""),
        "descripción del dominio tecnológico": intro.get("technological_domain", ""),
        "Nivel de madurez obtenido: [Nivel de madurez de la torre]": "",
        "El resultado de [X.X] sitúa la torre en Nivel [X] – [Descripción].": "",
        "Nivel objetivo:\n[Nivel objetivo recomendado]": "",
        "[INSTRUCCIÓN PARA EL AGENTE:Hazlo extenso y bien estructurado mediante bullets]": "",
        "El análisis del estado actual del dominio Resilience & Continuity permite concluir que la infraestructura presenta un nivel de madurez [nivel de madurez], caracterizado por:": "",
        "[fortaleza identificada]": "",
        "[limitación identificada]": "",
        "[riesgo identificado]": "",
        "Pilar [Nombre del pilar]": "",
        "[Descripción del estado actual basada en las respuestas del cuestionario. Hazla extensa y estructurada por bullets]": "",
        "[capacidad objetivo]": "",
        "[línea de evolución]": "",
        "[Fecha]": meta.get("date", "2026"),
    }

    # Purge paragraphs that serve as instructional notes for consultants or contain residual, un-replaced templating markers to ensure a clean final document.
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        # Iterate through table rows to identify and mark for deletion any rows containing instructional text intended for consultants, typically enclosed in brackets.
        if (
            "[Teniendo en cuenta el resumen de" in p.text
            or "[Teniendo en cuenta el Resumen ejecutivo" in p.text
        ):
            paragraphs_to_remove.append(p)
            if "Resumen ejecutivo del documento de contexto" in p.text:
                p.insert_paragraph_before(intro.get("introduction_paragraph", ""))
            continue

        if "[Generar un gráfico" in p.text:
            paragraphs_to_remove.append(p)
            continue

        if (
            "[capacidad tecnológica evaluada]" in p.text
            or "[infraestructura evaluada]" in p.text
            or "[plataformas tecnológicas incluidas]" in p.text
            or "[alcance específico del assessment]" in p.text
        ):
            paragraphs_to_remove.append(p)
            continue

        if (
            "[P1]" in p.text
            or "[P2]" in p.text
            or "[P3]" in p.text
            or "[P4]" in p.text
            or "[P5]" in p.text
        ):
            paragraphs_to_remove.append(p)
            continue

        if p.text.strip().startswith("| [Torre]"):
            paragraphs_to_remove.append(p)
            continue

        # Execute a global search-and-replace operation across all document paragraphs and tables using the predefined key-value mapping.
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)

    for p in paragraphs_to_remove:
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass

    # Sanitize tables by removing placeholder rows, which are identified by the presence of bracketed example text, ensuring that only populated data rows are retained.
    for table in doc.tables:
        rows_to_remove = []
        for row in table.rows:
            row_has_bracket = False
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "[" in p.text and "]" in p.text:
                        # A special case is handled for date placeholders: if a cell exclusively contains the date placeholder, the value is substituted directly. Otherwise, if the placeholder is part of a larger instructional text, the entire row is flagged for deletion.
                        if "[Fecha]" in p.text:
                            p.text = p.text.replace("[Fecha]", replacements["[Fecha]"])
                        else:
                            row_has_bracket = True
                            break
                if row_has_bracket:
                    break
            if row_has_bracket:
                rows_to_remove.append(row)

        #
        for row in rows_to_remove:
            try:
                table._tbl.remove(row._tr)
            except Exception:
                pass


def render_extended_variant(doc, payload):
    """Renders detailed tower annex sections from a payload into a document.

    This function populates a document object with a comprehensive, multi-part
    analysis based on the `extended_sections` key within the input payload.
    It begins by adding a page break and an introductory title.

    The function then systematically constructs the following sections in order:
    1. Detailed AS-IS analysis, including maturity levels and strengths.
    2. Detailed risks, including cause, impact, and mitigation.
    3. Detailed TO-BE target state, with capabilities and principles.
    4. Detailed gap analysis between the AS-IS and TO-BE states.
    5. Detailed prioritized initiatives to bridge the gaps.
    6. An extended concluding summary with focus areas.

    If the `extended_sections` key is missing from the payload or its value is
    falsey (e.g., an empty dictionary), the function returns immediately without
    modifying the document.

    Args:
        doc: A document object (e.g., `docx.document.Document`) to which the
            content will be appended. This object is modified in place.
        payload (dict): The input data dictionary. The function expects this
            dictionary to contain an `extended_sections` key, which in turn
            holds nested dictionaries for 'asis', 'risks', 'tobe', 'gap', 'todo',
            and 'conclusion'.

    Returns:
        None. The `doc` object is modified directly.
    """
    extended = payload.get("extended_sections") or {}
    if not extended:
        return

    doc.add_page_break()
    add_heading_paragraph(doc, "Desarrollo ampliado", level=1)
    add_body_paragraph(
        doc,
        "Esta versión desarrolla con más detalle el diagnóstico, el estado objetivo, las brechas y las iniciativas prioritarias de la torre.",
    )

    asis = extended.get("asis", {})
    add_heading_paragraph(doc, "1. AS-IS detallado", level=1)
    add_body_paragraph(doc, asis.get("executive_narrative", ""))
    add_label_value_paragraph(
        doc, "Madurez actual", asis.get("current_maturity_band", "")
    )
    add_label_value_paragraph(
        doc, "Referencia de score", asis.get("current_score_reference", "")
    )
    if asis.get("strengths"):
        add_heading_paragraph(doc, "Fortalezas observadas", level=2)
        add_bullet_list(doc, asis.get("strengths", []))
    if asis.get("gaps"):
        add_heading_paragraph(doc, "Brechas observadas", level=2)
        add_bullet_list(doc, asis.get("gaps", []))
    if asis.get("operational_impacts"):
        add_heading_paragraph(doc, "Implicaciones operativas", level=2)
        add_bullet_list(doc, asis.get("operational_impacts", []))

    risks = extended.get("risks", {})
    add_heading_paragraph(doc, "2. Riesgos detallados", level=1)
    add_body_paragraph(doc, risks.get("introduction", ""))
    for idx, item in enumerate(risks.get("risk_details", []), start=1):
        add_long_detail_table(
            doc,
            f"Riesgo {idx}. {item.get('risk', '')}",
            [
                ("Causa", item.get("cause", "")),
                ("Impacto", item.get("impact", "")),
                ("Probabilidad", item.get("probability", "")),
                ("Mitigación", item.get("mitigation_summary", "")),
                ("Pilares afectados", item.get("affected_pillars_display", "")),
            ],
        )
        add_body_paragraph(doc, "", space_after=2)
    add_body_paragraph(doc, risks.get("closing_summary", ""))

    tobe = extended.get("tobe", {})
    add_heading_paragraph(doc, "3. Estado objetivo detallado", level=1)
    add_body_paragraph(doc, tobe.get("introduction", ""))
    add_label_value_paragraph(
        doc, "Madurez objetivo recomendada", tobe.get("target_maturity_level", "")
    )
    add_label_value_paragraph(
        doc, "Referencia de score objetivo", tobe.get("target_score_reference", "")
    )
    add_body_paragraph(doc, tobe.get("target_maturity_justification", ""))
    for pillar in tobe.get("target_capabilities_by_pillar", []):
        add_heading_paragraph(doc, pillar.get("pillar", ""), level=2)
        add_bullet_list(doc, pillar.get("target_capabilities", []))
    if tobe.get("architecture_principles"):
        add_heading_paragraph(doc, "Principios de arquitectura y operación", level=2)
        add_bullet_list(doc, tobe.get("architecture_principles", []))
    if tobe.get("operating_model_implications"):
        add_heading_paragraph(doc, "Implicaciones para el modelo operativo", level=2)
        add_bullet_list(doc, tobe.get("operating_model_implications", []))

    gap = extended.get("gap", {})
    add_heading_paragraph(doc, "4. Brechas detalladas", level=1)
    add_body_paragraph(doc, gap.get("introduction", ""))
    if gap.get("cross_cutting_gap_summary"):
        add_heading_paragraph(doc, "Lectura transversal de gaps", level=2)
        add_bullet_list(doc, gap.get("cross_cutting_gap_summary", []))
    for idx, item in enumerate(gap.get("gap_items", []), start=1):
        add_long_detail_table(
            doc,
            f"Gap {idx}. {item.get('pillar', '')}",
            [
                ("Situación actual", item.get("as_is_summary", "")),
                ("Estado objetivo", item.get("target_state", "")),
                ("Brecha clave", item.get("key_gap", "")),
                ("Implicación operativa", item.get("operational_implication", "")),
            ],
        )
        add_body_paragraph(doc, "", space_after=2)

    todo = extended.get("todo", {})
    add_heading_paragraph(doc, "5. Iniciativas priorizadas detalladas", level=1)
    add_body_paragraph(doc, todo.get("introduction", ""))
    for item in todo.get("todo_items", []):
        add_long_detail_table(
            doc,
            f"{item.get('sequence', '')}. {item.get('initiative', '')}",
            [
                ("Objetivo", item.get("objective", "")),
                ("Prioridad", item.get("priority", "")),
                ("Pilares relacionados", item.get("related_pillars_display", "")),
                ("Resultado esperado", item.get("expected_outcome", "")),
                ("Dependencias", item.get("dependencies_display", "")),
            ],
        )
        add_body_paragraph(doc, "", space_after=2)
    add_body_paragraph(doc, todo.get("closing_summary", ""))

    conclusion = extended.get("conclusion", {})
    add_heading_paragraph(doc, "6. Cierre ampliado", level=1)
    add_body_paragraph(doc, conclusion.get("final_assessment", ""))
    add_body_paragraph(doc, conclusion.get("executive_message", ""))
    if conclusion.get("priority_focus_areas"):
        add_heading_paragraph(doc, "Áreas de foco prioritarias", level=2)
        add_bullet_list(doc, conclusion.get("priority_focus_areas", []))
    add_body_paragraph(doc, conclusion.get("closing_statement", ""))


def main(argv: list[str] | None = None) -> None:
    r"""{'docstring': 'Render a Tower Annex DOCX document from a template and a JSON payload.\n\nServes as the main entry point for the document rendering command-line script.\n\nThis function orchestrates the document generation process. It parses command-line\narguments to determine the input JSON payload, the DOCX template, the output\nfile path, and the rendering mode. It then loads and validates the payload\ndata against the `AnnexPayload` schema, normalizes the data structure, and\nsystematically populates the template by replacing placeholders with content\nand rendering complex components such as tables, lists, and charts. The final\npopulated document is then saved to the specified output path.\n\nArgs:\n    argv: A list of command-line arguments. If None, `sys.argv` is used.\n        The expected format is `<payload_json> <template_docx> <output_docx>\n        [--semantic-styles|--legacy-styles]`, where the final flag\n        controls the styling mode applied to the document.\n\nRaises:\n    SystemExit: If the number of required positional command-line arguments\n        is not exactly three.\n    FileNotFoundError: If the specified payload JSON file or template\n        DOCX file cannot be found at the provided paths.\n    pydantic.ValidationError: If the payload data fails validation against the\n        `AnnexPayload` data model schema.\n    json.JSONDecodeError: If the payload file is not a valid JSON document.'}."""
    raw_args = list(argv if argv is not None else sys.argv)
    render_mode, parsed_args = _resolve_render_mode(raw_args[1:])
    if len(parsed_args) != 3:
        raise SystemExit(
            "Uso: python -m scripts.render_tower_annex_from_template <payload_json> <template_docx> <output_docx> [--semantic-styles|--legacy-styles]"
        )

    payload_path = Path(parsed_args[0]).resolve()
    template_path = Path(parsed_args[1]).resolve()
    output_path = Path(parsed_args[2]).resolve()

    # Load the contract data from the specified source and perform robust validation against the authoritative Pydantic schema to ensure data integrity.
    payload = robust_load_payload(
        payload_path,
        AnnexPayload,
        "Annex",
        mode="strict",
    )
    payload_dict = normalize_annex_payload(payload.model_dump(by_alias=True))
    payload_dict = _replace_client_placeholders(
        payload_dict,
        str(payload_dict.get("document_meta", {}).get("client_name", "CLIENTE")),
    )

    doc = Document(str(template_path))
    remove_page_break_only_paragraphs(doc)
    clean_brackets_and_consultant_notes(doc, payload_dict)

    meta = payload_dict["document_meta"]
    exec_summary = payload_dict["executive_summary"]
    profile = payload_dict["pillar_score_profile"]
    profile["radar_chart"] = _ensure_radar_chart_path(payload_path, profile)
    sections = payload_dict["sections"]
    variant = clean_text(meta.get("report_variant", "short")).lower()
    subtitle_suffix = " · Versión Extendida" if variant == "long" else ""

    replace_simple_placeholder(
        doc,
        "{{ANNEX_TITLE}}",
        f"ANEXO {meta['tower_code']} – {meta['tower_name']}",
        font_size=15,
        bold=True,
    )
    replace_simple_placeholder(doc, "{{TOWER_CODE}}", meta["tower_code"], font_size=15)
    replace_simple_placeholder(doc, "{{TOWER_NAME}}", meta["tower_name"], font_size=15)
    replace_simple_placeholder(
        doc,
        "{{ANNEX_SUBTITLE}}",
        f"{meta['client_name']} · Fast Infrastructure Assessment{subtitle_suffix}",
        font_size=11,
    )
    replace_simple_placeholder(
        doc, "{{CLIENT_NAME}}", meta["client_name"], font_size=11
    )

    replace_simple_placeholder(
        doc,
        "{{GLOBAL_SCORE}}",
        exec_summary.get("global_score", ""),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=13,
    )
    replace_simple_placeholder(
        doc,
        "{{GLOBAL_BAND}}",
        exec_summary.get("global_band", ""),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=13,
    )
    replace_simple_placeholder(
        doc,
        "{{TARGET_MATURITY}}",
        exec_summary.get("target_maturity", ""),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=13,
    )

    render_multi_paragraph_block(
        doc, "{{EXEC_SUMMARY_BODY}}", exec_summary.get("summary_body", "")
    )
    replace_simple_placeholder(
        doc,
        "{{MSG_STRENGTH_VALUE}}",
        exec_summary.get("message_strength", ""),
        font_size=10.5,
    )
    replace_simple_placeholder(
        doc, "{{MSG_GAP_VALUE}}", exec_summary.get("message_gap", ""), font_size=10.5
    )
    replace_simple_placeholder(
        doc,
        "{{MSG_BOTTLENECK_VALUE}}",
        exec_summary.get("message_bottleneck", ""),
        font_size=10.5,
    )

    replace_simple_placeholder(
        doc,
        "{{PILLAR_PROFILE_INTRO}}",
        profile.get("profile_intro", ""),
        font_size=10.5,
    )
    render_note_box(
        doc, "{{SCORING_METHOD_NOTE}}", profile.get("scoring_method_note", "")
    )
    render_radar_chart(doc, "{{RADAR_CHART_BLOCK}}", profile.get("radar_chart", ""))
    render_pillar_score_table(doc, "{{PILLAR_SCORE_TABLE}}", profile.get("pillars", []))

    replace_simple_placeholder(
        doc,
        "{{ASIS_NARRATIVE}}",
        sections.get("asis", {}).get("narrative", ""),
        font_size=10.5,
    )
    render_list_at_placeholder(
        doc, "{{ASIS_STRENGTHS_LIST}}", sections.get("asis", {}).get("strengths", [])
    )
    render_list_at_placeholder(
        doc, "{{ASIS_GAPS_LIST}}", sections.get("asis", {}).get("gaps", [])
    )
    render_list_at_placeholder(
        doc,
        "{{ASIS_OPERATIONAL_IMPACTS_LIST}}",
        sections.get("asis", {}).get("operational_impacts", []),
    )

    replace_simple_placeholder(
        doc,
        "{{RISKS_INTRO}}",
        sections.get("risks", {}).get("introduction", ""),
        font_size=10.5,
    )
    render_risks_table(
        doc, "{{RISKS_TABLE}}", sections.get("risks", {}).get("risks", [])
    )
    replace_simple_placeholder(
        doc,
        "{{RISKS_CLOSING}}",
        sections.get("risks", {}).get("closing_summary", ""),
        font_size=10.5,
    )

    replace_simple_placeholder(
        doc,
        "{{TOBE_INTRO}}",
        sections.get("tobe", {}).get("vision", "")
        or sections.get("gap", {}).get("introduction", ""),
        font_size=10.5,
    )
    render_list_at_placeholder(
        doc,
        "{{TARGET_CAPABILITIES_LIST}}",
        sections.get("gap", {}).get("target_capabilities", [])
        or sections.get("tobe", {}).get("design_principles", []),
    )
    render_gap_table(doc, "{{GAP_TABLE}}", sections.get("gap", {}).get("gap_rows", []))

    replace_simple_placeholder(
        doc,
        "{{TODO_INTRO}}",
        sections.get("todo", {}).get("introduction", ""),
        font_size=10.5,
    )
    render_initiative_cards(
        doc,
        "{{PRIORITY_INITIATIVES_CARDS}}",
        sections.get("todo", {}).get("priority_initiatives", []),
    )

    replace_simple_placeholder(
        doc,
        "{{FINAL_ASSESSMENT}}",
        sections.get("conclusion", {}).get("final_assessment", ""),
        font_size=10.5,
    )
    replace_simple_placeholder(
        doc,
        "{{EXECUTIVE_MESSAGE}}",
        sections.get("conclusion", {}).get("executive_message", ""),
        font_size=10.5,
    )
    render_list_at_placeholder(
        doc,
        "{{PRIORITY_AREAS_LIST}}",
        sections.get("conclusion", {}).get("priority_focus_areas", []),
    )
    replace_simple_placeholder(
        doc,
        "{{CLOSING_STATEMENT}}",
        sections.get("conclusion", {}).get("closing_statement", ""),
        font_size=10.5,
    )

    if variant == "long":
        render_extended_variant(doc, payload_dict)

    if render_mode == "semantic":
        apply_semantic_annex_styles(doc, payload_dict)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print("Documento renderizado en:", output_path)


if __name__ == "__main__":
    main()
