"""Provides core logic and utilities for the Assessment Engine's global report generation pipeline."""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt, RGBColor
from pydantic import ValidationError

from assessment_engine.schemas.global_report import (
    BurningPlatformItem,
    ExecutionRoadmapDraft,
    ExecutiveDecisionsDraft,
    ExecutiveSummaryDraft,
    GlobalReportPayload,
    TargetVisionDraft,
    TowerBottomLineItem,
)
from assessment_engine.scripts.lib.docx_render_utils import (
    add_body_paragraph,
    add_heading_paragraph,
    autofit_table_to_contents,
    clear_paragraph,
    finalize_table,
    set_cell_text,
    shade_cell,
)
from assessment_engine.scripts.lib.global_maturity_policy import (
    safe_float,
    status_color_for_score,
)


def load_json(path):
    """Load a JSON object from a file, decoding its content as 'utf-8-sig'."""
    return json.loads(path.read_text(encoding="utf-8-sig"))


def clean_t_codes(text):
    """Removes specific code patterns and internal reference tags from a string.

    This function applies a series of regular expression substitutions to remove
    predefined patterns, including T-codes and internal reference markers. Leading
    and trailing whitespace is also stripped from the result.

    The following patterns are removed:
    - Parenthesized T-codes (e.g., `(T12)`, `(T1 other text)`).
    - Standalone, word-bounded T-codes (e.g., `T1`, `T12`).
    - Internal reference tags (e.g., `[[REF:id]]`, `[REF:id]`).

    Args:
        text (Any): The input value to process. If not a `str`, it is
            returned unmodified.

    Returns:
        Any: The cleaned string or the original non-string input.
    """
    if not isinstance(text, str):
        return text
    text = re.sub(r"\(T\d{1,2}[^\)]*\)", "", text)
    text = re.sub(r"\bT\d{1,2}\b", "", text)
    # Remove internal reference tags (e.g., [[REF:...]]) from the text prior to rendering, as they are not intended for the final report.
    text = re.sub(r"\[\[?REF:[^\]]*\]\]?", "", text)
    return text.strip()


def sanitize_client_name(text, client_name):
    """Sanitizes a string by replacing client-specific names with a generic term.

    Performs a case-insensitive search and replace to anonymize text. This
    function replaces all occurrences of `client_name` with the generic Spanish
    phrase 'la organización'. It also removes preceding Spanish possessive
    prepositions (i.e., ' de ' or ' del '). The function handles `client_name`
    variations where spaces are substituted with underscores.

    Args:
        text (str): The input string to sanitize.
        client_name (str): The client name to be replaced.

    Returns:
        str: The sanitized text. If `text` is not a string or `client_name` is
            falsy, the original `text` is returned unmodified.
    """
    if not isinstance(text, str) or not client_name:
        return text

    # Sanitize the client name by replacing spaces and special characters with underscores to generate a valid, filesystem-safe identifier for use in filenames.
    clean_name = client_name.replace("_", " ")

    # Step 1: Normalize text by removing client-specific possessive prepositions (e.g., 'de [Client]', 'del [Client]').
    text = re.sub(
        rf"\s+del?\s+{re.escape(client_name)}\b", "", text, flags=re.IGNORECASE
    )
    text = re.sub(
        rf"\s+del?\s+{re.escape(clean_name)}\b", "", text, flags=re.IGNORECASE
    )

    # Step 2: Generalize the narrative by replacing direct client mentions with the term 'the organization' for standardization.
    text = re.sub(
        rf"\b{re.escape(client_name)}\b", "la organización", text, flags=re.IGNORECASE
    )
    text = re.sub(
        rf"\b{re.escape(clean_name)}\b", "la organización", text, flags=re.IGNORECASE
    )

    return text


def clear_document_body(doc):
    """Clears all content from the body of a `python-docx` document, preserving section properties.

    This function directly manipulates the underlying XML of the document. It iterates
    through all top-level child elements within the `<w:body>` tag and removes them,
    with the sole exception of the `<w:sectPr>` (section properties) element.

    This procedure is typically used to reset a template document, removing all
    placeholder text and objects while retaining the page layout configuration
    (e.g., margins, page orientation, headers, footers) defined in the section
    properties. The modification is performed in-place.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to modify.

    Returns:
        None.

    Raises:
        AttributeError: If the provided `doc` object does not have the expected
            internal `_body._element` structure.
    """
    body = doc._body._element
    for child in list(body):
        if (
            child.tag
            != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
        ):
            body.remove(child)


def add_spacer(doc, points=12):
    """{'docstring': 'Add a vertical spacer to a document object.'}."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)


def add_smart_bullet_list(container, items, color_rgb=None, bold_prefix=True):
    """Adds a formatted bulleted list to a python-docx container.

    This function iterates through a list of items, adding each as a paragraph
    styled as a bullet point. It first attempts to apply the 'List Bullet'
    style from the document's template. If this style is not available (a
    KeyError is caught internally), it falls back to manually formatting the
    paragraph with a '•' character and appropriate indentation.

    The function provides conditional formatting based on content. If an item
    string contains a `": "` or `" - "` separator, the text is split. The
    portion of the string preceding the first separator can be bolded based on
    the `bold_prefix` argument. The item text is sanitized using an internal
    `clean_t_codes` function before rendering.

    Args:
        container (Union[docx.document.Document, docx.table._Cell]): The
            python-docx object to which the list will be added (e.g., a
            Document or a table Cell).
        items (Union[str, List[str]]): A list of strings for the bullet points.
            A single string will be treated as a single-item list.
        color_rgb (Optional[docx.shared.RGBColor]): An RGBColor object to apply
            to the text. Defaults to None, using the style's default color.
        bold_prefix (bool): If True, bolds the text preceding a separator.
            If no separator is present, the entire item is bolded. Defaults to
            True.
    """
    # Robustness: Coerce input to a list if it is a string to prevent unintended character-wise iteration where a list of strings is expected.
    if isinstance(items, str):
        items = [items]
    if not items:
        return

    for item in items:
        p = container.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        try:
            p.style = "List Bullet"
        except KeyError:
            # Provide a fallback mechanism to ensure robust rendering if the specified style is not found within the document template.
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.first_line_indent = Pt(-15)
            b = p.add_run("• ")
            b.font.size = Pt(10.5)
            if color_rgb:
                b.font.color.rgb = color_rgb

        p.paragraph_format.space_after = Pt(6)
        text = clean_t_codes(str(item).strip())

        sep = ": " if ": " in text else (" - " if " - " in text else None)
        if sep:
            parts = text.split(sep, 1)
            r1 = p.add_run(parts[0] + sep)
            r1.bold = bold_prefix
            r1.font.size = Pt(10.5)
            r2 = p.add_run(parts[1])
            r2.font.size = Pt(10.5)
            if color_rgb:
                r1.font.color.rgb = color_rgb
                r2.font.color.rgb = color_rgb
        else:
            r = p.add_run(text)
            r.bold = bold_prefix
            r.font.size = Pt(10.5)
            if color_rgb:
                r.font.color.rgb = color_rgb


def render_cover(doc, payload: GlobalReportPayload):
    """Renders a formatted cover page into a `python-docx` document object.

    This function populates the first page of the provided document with a
    title, client name, report date, version reference, and a multi-paragraph
    legal disclaimer. It applies specific typographic and layout styles,
    including fonts, sizes, colors, alignment, and vertical spacing to
    structure the content. The operation concludes by inserting a page break,
    ensuring subsequent content begins on a new page.

    Args:
        doc (docx.document.Document): The document object to which the cover page
            content will be added. This object is modified in-place.
        payload (GlobalReportPayload): A data object containing the report's
            metadata, specifically the `client`, `date`, and `version`
            attributes required for rendering the cover page.

    Returns:
        None.

    Raises:
        AttributeError: If `payload.meta` or its required attributes (`client`,
            `date`, `version`) do not exist.
    """
    # Slightly reduce the document's initial top margin to optimize vertical space on the first page.
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Apply Title Case formatting, ensuring minor words such as articles and prepositions remain lowercase.
    run = title_p.add_run("Informe Estratégico de\nMadurez Tecnológica")
    run.font.size = Pt(34)
    run.font.name = "Georgia"
    run.font.color.rgb = RGBColor(0, 114, 188)
    run.bold = False

    # Reduce vertical spacing between the report title and client name to achieve a more compact header layout.
    add_spacer(doc, 30)

    client_p = doc.add_paragraph()
    client_run = client_p.add_run(payload.meta.client.upper())
    client_run.font.size = Pt(24)
    client_run.font.name = "Arial"
    client_run.bold = True

    # This block's vertical spacing is substantially reduced to correct a layout artifact where the date element was rendered with excessive top margin.
    doc.add_paragraph().paragraph_format.space_after = Pt(100)

    version_p = doc.add_paragraph()
    v_text = f"Fecha: {payload.meta.date}\nReferencia: {payload.meta.version}"
    version_run = version_p.add_run(v_text)
    version_run.font.size = Pt(14)
    version_run.font.name = "Arial"
    version_run.font.color.rgb = RGBColor(127, 127, 127)

    # A 150-point spacer is used to position the legal disclaimer near the page bottom while preventing overflow onto a subsequent page.
    add_spacer(doc, 150)

    disclaimer_p1 = doc.add_paragraph()
    disclaimer_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p1.paragraph_format.space_after = Pt(
        2
    )  # Defines the minimum required spacing between paragraphs within the legal disclaimer section.
    disclaimer_title_run = disclaimer_p1.add_run("Aviso Legal y Confidencialidad: ")
    disclaimer_title_run.font.size = Pt(8)
    disclaimer_title_run.font.name = "Arial"
    disclaimer_title_run.font.color.rgb = RGBColor(127, 127, 127)
    disclaimer_title_run.bold = True

    p1_text = "El presente documento constituye un informe de evaluación rápida (Fast Assessment) elaborado por NTT DATA basándose exclusivamente en la información, entrevistas y datos proporcionados por la organización, los cuales no han sido sometidos a una auditoría o verificación independiente por nuestra parte. Por tanto, NTT DATA no otorga ninguna garantía, expresa o implícita, sobre la exactitud, integridad o fiabilidad absoluta de dicha información base."
    r1 = disclaimer_p1.add_run(p1_text)
    r1.font.size = Pt(8)
    r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor(127, 127, 127)

    disclaimer_p2 = doc.add_paragraph()
    disclaimer_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p2.paragraph_format.space_after = Pt(2)
    p2_text = 'Las conclusiones, proyecciones (incluyendo el estado objetivo "TO-BE", los horizontes del roadmap y las estimaciones de impacto de negocio) y recomendaciones tienen un carácter puramente orientativo. Estas proyecciones constituyen estimaciones sujetas a incertidumbres operativas y de mercado, por lo que los resultados reales futuros podrían diferir de los planteados.'
    r2 = disclaimer_p2.add_run(p2_text)
    r2.font.size = Pt(8)
    r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor(127, 127, 127)

    disclaimer_p3 = doc.add_paragraph()
    disclaimer_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p3.paragraph_format.space_after = Pt(2)
    p3_text = "Este informe tiene un propósito estrictamente técnico y estratégico. Toda decisión ejecutiva, operativa o de inversión que la organización tome basándose en este informe es de su exclusiva responsabilidad. NTT DATA y sus representantes no asumen obligación legal ni responsabilidad alguna por posibles daños directos, indirectos, especiales o consecuentes derivados del uso o dependencia de la información aquí contenida sin una validación posterior exhaustiva."
    r3 = disclaimer_p3.add_run(p3_text)
    r3.font.size = Pt(8)
    r3.font.name = "Arial"
    r3.font.color.rgb = RGBColor(127, 127, 127)

    disclaimer_p4 = doc.add_paragraph()
    disclaimer_p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer_p4.paragraph_format.space_after = Pt(0)
    p4_text = "Finalmente, este documento y la información que contiene son confidenciales. Este material no podrá ser distribuido, reproducido, divulgado ni utilizado para ningún otro propósito, total o parcialmente, sin el consentimiento previo por escrito de NTT DATA."
    r4 = disclaimer_p4.add_run(p4_text)
    r4.font.size = Pt(8)
    r4.font.name = "Arial"
    r4.font.color.rgb = RGBColor(127, 127, 127)

    doc.add_page_break()


def render_executive_summary(
    doc,
    data: ExecutiveSummaryDraft,
    heatmap: list,
    visuals: dict,
    client_dir,
    client_name="",
):
    """Populates the executive summary section of a Word document object.

    Constructs a formatted summary by calculating an average score from a
    heatmap, adding a narrative headline and body, listing key business
    impacts in a table, and embedding a radar chart visualization. The function
    directly modifies the input document object, applying specific branding
    styles (e.g., fonts, colors, table shading) to all generated content.
    Textual content is sanitized to replace client-specific placeholders.

    Args:
        doc (docx.document.Document): The document object to be modified in-place.
        data (ExecutiveSummaryDraft): A data transfer object containing the textual
            content, requiring `headline`, `narrative`, and
            `key_business_impacts` attributes.
        heatmap (list[dict]): A list of dictionaries where each dictionary
            represents a topic and is expected to contain a 'score' key with a
            numerical value. Used to calculate the average score.
        visuals (dict[str, str]): A dictionary mapping visual element names to their
            corresponding filenames. A 'radar_chart' key is expected.
        client_dir (pathlib.Path): The filesystem path to the directory containing
            client-specific assets, such as the radar chart image.
        client_name (str): The client's name, used to replace placeholders in the
            text. Defaults to an empty string.

    Returns:
        None: The function modifies the `doc` object in place and does not return
            a value.

    Raises:
        AttributeError: If the `data` object is missing a required attribute
            (`headline`, `narrative`, or `key_business_impacts`).
        IOError: If the radar chart image file specified in `visuals` cannot be
            opened or read due to permissions or other filesystem issues.
        ValueError: If the radar chart image file is in an unsupported format
            or is otherwise invalid.
    """
    BASE_TEXT_COLOR = RGBColor(
        46, 64, 77
    )  # Define the primary text color (#2E404D) to maintain consistency with corporate branding guidelines.
    add_heading_paragraph(doc, "1. Resumen Ejecutivo", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Set column widths to an approximate 35/65 ratio (2.1" and 3.9" respectively). This distribution is optimized for readability on a standard A4 page layout.
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(3.9)

    score_cell = table.rows[0].cells[0]
    score_cell.width = Inches(2.1)
    score_val = "N/A"
    if heatmap:
        try:
            score_val = str(
                round(
                    sum(float(t.get("score", 0)) for t in heatmap) / len(heatmap),
                    1,
                )
            )
        except Exception:
            pass

    # Define styling for score rendering: format as '[X] / 5', apply a #0072BC background, set font size to 36pt, and set font color to white.
    set_cell_text(
        score_cell,
        f"{score_val} / 5",
        bold=True,
        font_size=36,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade_cell(score_cell, "0072BC")
    for r in score_cell.paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    desc_cell = table.rows[0].cells[1]
    desc_cell.width = Inches(3.9)
    desc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(desc_cell, "F2F2F2")

    headline = sanitize_client_name(clean_t_codes(data.headline), client_name)

    # Apply bold formatting exclusively to the title's substring preceding the colon, in accordance with style guide requirements.
    clear_paragraph(desc_cell.paragraphs[0])
    p = desc_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if ":" in headline:
        parts = headline.split(":", 1)
        r1 = p.add_run(parts[0] + ":")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.name = "Arial"
        r2 = p.add_run(parts[1])
        r2.bold = False
        r2.font.size = Pt(12)
        r2.font.name = "Arial"
    else:
        r = p.add_run(headline)
        r.bold = False
        r.font.size = Pt(12)
        r.font.name = "Arial"

    add_spacer(doc, 15)

    # Architectural decision: Narrative text and bullet points are rendered outside of any table structure to improve content flow and readability.
    narrative = sanitize_client_name(clean_t_codes(data.narrative), client_name)
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚ])", narrative.strip())

    if sentences:
        # The first sentence of the narrative is semantically designated as the 'Bottom Line' or executive summary statement.
        p_intro = doc.add_paragraph()
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_intro = p_intro.add_run(sentences[0])
        r_intro.font.size = Pt(10.5)
        r_intro.font.color.rgb = BASE_TEXT_COLOR
        p_intro.paragraph_format.space_before = Pt(12)
        p_intro.paragraph_format.space_after = Pt(6)

        # All subsequent sentences in the collection are formatted as distinct bullet points to enhance readability.
        if len(sentences) > 1:
            add_smart_bullet_list(
                doc, sentences[1:], color_rgb=BASE_TEXT_COLOR, bold_prefix=False
            )

    add_spacer(doc, 15)

    #
    impact_table = doc.add_table(rows=1, cols=1)
    finalize_table(impact_table)

    header_cell = impact_table.rows[0].cells[0]
    set_cell_text(
        header_cell,
        "Principales impactos de negocio",
        bold=True,
        font_size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade_cell(header_cell, "0072BC")
    for r in header_cell.paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    impacts = [sanitize_client_name(i, client_name) for i in data.key_business_impacts]
    for impact in impacts:
        row = impact_table.add_row()
        body_cell = row.cells[0]
        shade_cell(body_cell, "F2F2F2")

        # Inject content directly into the cell's first paragraph (p[0]) to preempt the library's default behavior of inserting an unwanted leading paragraph.
        p_bull = body_cell.paragraphs[0]
        clear_paragraph(p_bull)
        p_bull.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        try:
            p_bull.style = "List Bullet"
        except KeyError:
            p_bull.paragraph_format.left_indent = Pt(20)
            p_bull.paragraph_format.first_line_indent = Pt(-15)
            b_run = p_bull.add_run("• ")
            b_run.font.size = Pt(10.5)
            b_run.font.color.rgb = BASE_TEXT_COLOR

        p_bull.paragraph_format.space_before = Pt(4)
        p_bull.paragraph_format.space_after = Pt(4)

        text = clean_t_codes(str(impact).strip())
        r_text = p_bull.add_run(text)
        r_text.bold = False
        r_text.font.size = Pt(10.5)
        r_text.font.color.rgb = BASE_TEXT_COLOR

    radar_path = client_dir / visuals.get("radar_chart", "")
    if radar_path.exists():
        add_spacer(doc, 15)
        doc.add_picture(str(radar_path), width=Inches(6.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def render_burning_platform(
    doc, platform_risks: list[BurningPlatformItem], client_name=""
):
    """Renders the systemic threats ('Burning Platform') section into a document.

    Generates a 'Principales Amenazas Sistémicas' section within the provided
    `python-docx` document object. This function first adds a main heading and a
    descriptive paragraph. It then iterates through each systemic risk provided
    in `platform_risks`, creating a distinct, styled table for each one.

    Each table consists of a header cell for the threat's theme, a cell
    detailing the business risk, and a final cell containing a bulleted list
    of root causes. Text content is sanitized to replace generic placeholders
    with the specified `client_name`.

    Args:
        doc (docx.document.Document): An active `python-docx` Document object to
            which the content will be appended.
        platform_risks (list[BurningPlatformItem]): A sequence of data objects,
            where each object must contain `theme`, `business_risk`, and
            `root_causes` attributes.
        client_name (str): The client's name, used to replace placeholders in the
            risk descriptions. Defaults to an empty string.

    Returns:
        None: The function modifies the `doc` object in place.

    Raises:
        AttributeError: If an object within `platform_risks` lacks one of the
            required attributes (`theme`, `business_risk`, `root_causes`).
    """
    BASE_TEXT_COLOR = RGBColor(
        46, 64, 77
    )  # Define the primary text color (#2E404D) to maintain consistency with corporate branding guidelines.
    add_heading_paragraph(doc, "2. Principales Amenazas Sistémicas", level=1)
    add_body_paragraph(
        doc,
        "Identificación de riesgos críticos que comprometen la viabilidad operativa y la agilidad de la organización.",
        color_rgb=BASE_TEXT_COLOR,
    )
    for i, risk in enumerate(platform_risks, start=1):
        table = doc.add_table(rows=3, cols=1)
        finalize_table(table)

        #
        h_cell = table.rows[0].cells[0]
        theme = sanitize_client_name(clean_t_codes(risk.theme), client_name)
        theme = theme[0].upper() + theme[1:] if theme else ""
        set_cell_text(
            h_cell,
            f"Amenaza {i}: {theme}",
            bold=True,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(h_cell, "0072BC")
        for r in h_cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

        #
        r_cell = table.rows[1].cells[0]
        shade_cell(r_cell, "F2F2F2")
        clear_paragraph(r_cell.paragraphs[0])
        p1 = r_cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1_label = p1.add_run("Riesgo de Negocio: ")
        r1_label.bold = True
        r1_label.font.color.rgb = BASE_TEXT_COLOR
        r1_label.font.size = Pt(10.5)
        r1_text = p1.add_run(
            sanitize_client_name(clean_t_codes(risk.business_risk), client_name)
        )
        r1_text.bold = False
        r1_text.font.color.rgb = BASE_TEXT_COLOR
        r1_text.font.size = Pt(10.5)

        #
        c_cell = table.rows[2].cells[0]
        shade_cell(c_cell, "F2F2F2")
        clear_paragraph(c_cell.paragraphs[0])
        p2 = c_cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2_label = p2.add_run("Causas Raíz: ")
        r2_label.bold = True
        r2_label.font.color.rgb = BASE_TEXT_COLOR
        r2_label.font.size = Pt(10.5)
        causes_list = risk.root_causes
        sanitized_causes = [
            sanitize_client_name(clean_t_codes(c), client_name) for c in causes_list
        ]
        r2_text = p2.add_run("\n" + "\n".join([f"  • {c}" for c in sanitized_causes]))
        r2_text.bold = False
        r2_text.font.color.rgb = BASE_TEXT_COLOR
        r2_text.font.size = Pt(10.5)

        add_spacer(doc, 10)


def render_tower_bottom_lines(
    doc, heatmap: list, tower_texts: list[TowerBottomLineItem], client_name=""
):
    """Generates and appends a technology area diagnosis table to a Word document.

    Constructs a three-column table summarizing the maturity and executive
    diagnosis for each technology area (tower). The table is populated by merging
    data from a heatmap structure with corresponding executive summary texts.
    The second column, representing maturity, is color-coded based on the
    numerical score of the technology area. The diagnosis text is sourced from
    `tower_texts` by matching the tower 'id'; if no match is found, the function
    falls back to an 'executive_message' field within the `heatmap` data.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to which
            the table will be appended.
        heatmap (list[dict]): A list of dictionaries, where each dictionary
            represents a technology area. Each dictionary is expected to contain
            'id', 'name', 'score', 'band', and an optional 'executive_message' key.
        tower_texts (list[TowerBottomLineItem]): A list of data objects, each
            providing the executive diagnosis. Each object must have 'id' and
            'bottom_line' attributes to be correlated with an item in `heatmap`.
        client_name (str): The client name used to replace placeholders within the
            final diagnosis text. Defaults to an empty string.

    Returns:
        None.

    Raises:
        AttributeError: If an object in `tower_texts` lacks the required 'id' or
            'bottom_line' attributes.
    """
    BASE_TEXT_COLOR = RGBColor(
        46, 64, 77
    )  # Define the primary text color (#2E404D) to maintain consistency with corporate branding guidelines.
    add_heading_paragraph(doc, "3. Diagnóstico por Área Tecnológica", level=1)
    table = doc.add_table(rows=1, cols=3)
    finalize_table(table)
    for i, h in enumerate(["Área", "Madurez", "Diagnóstico Ejecutivo"]):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            bold=False,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table.rows[0].cells[i], "0072BC")
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
    color_map = {"E06666": "F4CCCC", "FFD966": "FFF2CC", "93C47D": "D9EAD3"}
    for t in heatmap:
        # The color value is explicitly recalculated here to enforce strict consistency and prevent potential state-related discrepancies across document sections.
        strict_color = status_color_for_score(safe_float(t.get("score"), 0.0))

        row = table.add_row()
        set_cell_text(
            row.cells[0],
            f"{t.get('id', '')}\n{t.get('name', '')}",
            bold=False,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(
            row.cells[1],
            f"{t.get('score', '')}\n({t.get('band', '')})",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10.5,
            bold=False,
        )
        shade_cell(row.cells[1], color_map.get(strict_color, "FFFFFF"))

        #
        tower_id = t.get("id", "")
        text_val = ""
        for bottom_line_item in tower_texts:
            if bottom_line_item.id == tower_id:
                text_val = bottom_line_item.bottom_line
                break

        if not text_val:
            text_val = t.get("executive_message", "")

        set_cell_text(
            row.cells[2],
            sanitize_client_name(clean_t_codes(str(text_val)), client_name),
            font_size=10.5,
        )
        #
        for p in row.cells[2].paragraphs:
            for r in p.runs:
                r.font.color.rgb = BASE_TEXT_COLOR
    autofit_table_to_contents(table)
    add_spacer(doc, 20)


def render_target_vision(doc, vision: TargetVisionDraft, client_name=""):
    """Renders the target state vision section into a Word document.

    Populates a `docx.document.Document` object with a complete section for the
    "Visión de Estado Objetivo (To-Be)". This section includes a main heading,
    a formatted table for the strategic value proposition, and bulleted lists
    detailing the maturity evolution principles and enabling strategic pillars.
    Placeholders in the source text are replaced with the provided client name.

    Args:
        doc (docx.document.Document): The document object to be mutated.
        vision (TargetVisionDraft): A data object containing the vision's textual
            components. Must expose `value_proposition` (str), `evolution_principles`
            (list), and `strategic_pillars` (list) attributes.
        client_name (str): The client's name, used to replace placeholders.

    Returns:
        None. The `doc` object is modified in place.

    Raises:
        AttributeError: If the `vision` object or its nested elements lack the
            required attributes (e.g., `value_proposition`, `pillar`, `description`).
    """
    BASE_TEXT_COLOR = RGBColor(
        46, 64, 77
    )  # Define the primary text color (#2E404D) to maintain consistency with corporate branding guidelines.
    add_heading_paragraph(doc, "4. Visión de Estado Objetivo (To-Be)", level=1)

    table = doc.add_table(rows=2, cols=1)
    finalize_table(table)

    #
    v_cell_header = table.rows[0].cells[0]
    set_cell_text(
        v_cell_header,
        "Propuesta de valor estratégica",
        bold=True,
        font_size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade_cell(v_cell_header, "0072BC")
    for r in v_cell_header.paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    #
    v_cell_body = table.rows[1].cells[0]
    shade_cell(v_cell_body, "F2F2F2")
    clear_paragraph(v_cell_body.paragraphs[0])
    p_body = v_cell_body.paragraphs[0]
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    prop = sanitize_client_name(clean_t_codes(vision.value_proposition), client_name)
    prop = prop[0].upper() + prop[1:] if prop else ""
    prop = prop.replace("─", "").replace("—", "").replace(";", ",")
    r_body = p_body.add_run(prop)
    r_body.font.color.rgb = BASE_TEXT_COLOR
    r_body.font.size = Pt(10.5)
    r_body.bold = False

    add_spacer(doc, 10)
    if vision.evolution_principles:
        add_heading_paragraph(doc, "Principios de Evolución de Madurez", level=2)
        p_list = [
            f"{p.principle}: {sanitize_client_name(p.description, client_name)}"
            for p in vision.evolution_principles
        ]
        add_smart_bullet_list(doc, p_list, color_rgb=BASE_TEXT_COLOR)
    add_heading_paragraph(doc, "Pilares Estratégicos Habilitadores", level=2)
    pillars = vision.strategic_pillars
    pillar_texts = []
    for p in pillars:
        p_name = sanitize_client_name(p.pillar, client_name)
        p_desc = sanitize_client_name(p.description, client_name)
        pillar_texts.append(f"{p_name}: {p_desc}")

    add_smart_bullet_list(
        doc, pillar_texts, color_rgb=BASE_TEXT_COLOR, bold_prefix=True
    )
    add_spacer(doc, 20)


def render_execution_roadmap(
    doc, roadmap: ExecutionRoadmapDraft, visuals: dict, client_dir, client_name=""
):
    """Generates and appends the execution roadmap section to a `docx` document.

    This function constructs the 'Implementation Plan and Time Horizons' section of a
    report. It first generates a table detailing the transversal programs and their
    descriptions. Subsequently, it creates distinct tables for each implementation
    phase (Quick Wins, Year 1, Year 2, and Year 3), itemizing the initiatives,
    their business cases, start months, and durations. All textual content is
    rendered in Spanish.

    Args:
        doc (docx.document.Document): The document object to be modified in-place.
        roadmap (ExecutionRoadmapDraft): A data object containing a list of programs
            and initiatives organized by time horizon.
        visuals (dict): A dictionary intended to hold visual assets. This parameter
            is currently unused by the function.
        client_dir (str): The path to the client's output directory. This parameter
            is currently unused by the function.
        client_name (str, optional): The client's name, used for placeholder
            substitution in text content. Defaults to an empty string.

    Returns:
        None: The function modifies the `doc` object directly.

    Raises:
        AttributeError: If the `roadmap` object or its nested structures do not
            conform to the expected schema (e.g., missing `programs`,
            `horizons`, or initiative attributes).
    """
    BASE_TEXT_COLOR = RGBColor(
        46, 64, 77
    )  # Define the primary text color (#2E404D) to maintain consistency with corporate branding guidelines.
    add_heading_paragraph(
        doc, "5. Plan de Implementación y Horizontes Temporales", level=1
    )
    add_body_paragraph(
        doc,
        "La transformación tecnológica se articula en programas transversales ejecutados en cuatro horizontes temporales para asegurar la resiliencia y escalabilidad del negocio.",
        color_rgb=BASE_TEXT_COLOR,
    )
    table_p = doc.add_table(rows=1, cols=1)
    finalize_table(table_p)

    #
    p_cell_header = table_p.rows[0].cells[0]
    set_cell_text(
        p_cell_header,
        "Programas transversales definidos",
        bold=True,
        font_size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade_cell(p_cell_header, "0072BC")
    for r in p_cell_header.paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    #
    programs = roadmap.programs
    for p in programs:
        row = table_p.add_row()
        body_cell = row.cells[0]
        shade_cell(body_cell, "F2F2F2")

        p_bull = body_cell.paragraphs[0]
        clear_paragraph(p_bull)
        p_bull.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        try:
            p_bull.style = "List Bullet"
        except KeyError:
            p_bull.paragraph_format.left_indent = Pt(20)
            p_bull.paragraph_format.first_line_indent = Pt(-15)
            b_run = p_bull.add_run("• ")
            b_run.font.size = Pt(10.5)
            b_run.font.color.rgb = BASE_TEXT_COLOR

        p_bull.paragraph_format.space_before = Pt(6)
        p_bull.paragraph_format.space_after = Pt(6)

        program_name = clean_t_codes(p.name)
        program_desc = sanitize_client_name(clean_t_codes(p.description), client_name)

        r_name = p_bull.add_run(f"{program_name}: ")
        r_name.bold = True
        r_name.font.size = Pt(10.5)
        r_name.font.color.rgb = BASE_TEXT_COLOR

        r_desc = p_bull.add_run(program_desc)
        r_desc.bold = False
        r_desc.font.size = Pt(10.5)
        r_desc.font.color.rgb = BASE_TEXT_COLOR

    add_spacer(doc, 15)
    horizons = roadmap.horizons
    mapping = [
        ("Quick Wins (0-3 meses)", horizons.quick_wins_0_3_months),
        ("Año 1 (4-12 meses)", horizons.year_1_3_12_months),
        ("Año 2 (13-24 meses)", horizons.year_2_12_24_months),
        ("Año 3 (25-36 meses)", horizons.year_3_24_36_months),
    ]
    for title, initiatives in mapping:
        if not initiatives:
            continue
        add_heading_paragraph(doc, title, level=3)
        table = doc.add_table(rows=1, cols=5)
        finalize_table(table)
        headers = [
            "Programa",
            "Iniciativa",
            "Impacto / Business Case",
            "Inicio",
            "Dur.",
        ]
        for i, h in enumerate(headers):
            set_cell_text(
                table.rows[0].cells[i],
                h,
                bold=False,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            shade_cell(table.rows[0].cells[i], "0072BC")
            for r in table.rows[0].cells[i].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
        for init in initiatives:
            row = table.add_row()
            set_cell_text(
                row.cells[0],
                clean_t_codes(init.program),
                bold=False,
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text(
                row.cells[1],
                sanitize_client_name(clean_t_codes(init.title), client_name),
                bold=False,
                font_size=10.5,
            )
            set_cell_text(
                row.cells[2],
                sanitize_client_name(clean_t_codes(init.business_case), client_name),
                font_size=10.5,
            )
            set_cell_text(
                row.cells[3],
                f"M{init.start_month}",
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text(
                row.cells[4],
                f"{init.duration_months}m",
                font_size=10.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            #
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = BASE_TEXT_COLOR
        autofit_table_to_contents(table)
        add_spacer(doc, 15)


def render_executive_decisions(doc, decisions: ExecutiveDecisionsDraft, client_name=""):
    r"""{'docstring': "Appends a formatted 'Executive Decisions' section to a Word document.\n\n    This function constructs and adds a section detailing priority executive\n    decisions. The section comprises a level-1 heading, an introductory\n    paragraph, and a three-column table. The table is populated with data from\n    the `decisions` object, listing the decision scope, required action, and\n    the impact of any delay for each item. All content is styled according to\n    predefined corporate branding guidelines, and textual data is sanitized\n    prior to rendering.\n\n    Args:\n        doc (docx.document.Document): The `python-docx` Document object to which\n            the executive decisions section will be appended.\n        decisions (ExecutiveDecisionsDraft): A data object encapsulating decision\n            details. This object must possess an `immediate_decisions` attribute,\n            which is an iterable of objects. Each of these inner objects must,\n            in turn, provide `decision_type`, `action_required`, and\n            `impact_if_delayed` attributes.\n        client_name (str): The name of the client, used to replace placeholder\n            text in the rendered content. Defaults to an empty string.\n\n    Returns:\n        None. The function modifies the `doc` object in-place.\n\n    Raises:\n        AttributeError: If `decisions` or its contained elements lack the\n            required attributes (e.g., `immediate_decisions`, `decision_type`)."}."""
    BASE_TEXT_COLOR = RGBColor(
        46, 64, 77
    )  # Define the primary text color (#2E404D) to maintain consistency with corporate branding guidelines.
    add_heading_paragraph(doc, "6. Decisiones Ejecutivas Prioritarias", level=1)
    add_body_paragraph(
        doc,
        "Líneas de decisión requeridas por la Dirección para desbloquear la ejecución del programa y mitigar los riesgos identificados.",
        color_rgb=BASE_TEXT_COLOR,
    )
    table = doc.add_table(rows=1, cols=3)
    finalize_table(table)
    headers = ["Ámbito de Decisión", "Acción Requerida", "Impacto de Demora"]
    for i, h in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            bold=True,
            font_size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table.rows[0].cells[i], "0072BC")
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
    for d in decisions.immediate_decisions:
        row = table.add_row()
        set_cell_text(
            row.cells[0],
            sanitize_client_name(clean_t_codes(d.decision_type), client_name),
            bold=False,
            font_size=10.5,
        )
        set_cell_text(
            row.cells[1],
            sanitize_client_name(clean_t_codes(d.action_required), client_name),
            font_size=10.5,
        )
        set_cell_text(
            row.cells[2],
            sanitize_client_name(clean_t_codes(d.impact_if_delayed), client_name),
            font_size=10.5,
        )
        shade_cell(row.cells[2], "FFFFFF")
        #
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = BASE_TEXT_COLOR
    autofit_table_to_contents(table)


def create_page_number_footer(section):
    """Adds a 'Página X de Y' page number to a document section's footer.

    Configures the section footer to display a dynamic page count on all pages
    except the first, which is reserved for a title page. The function enables
    the 'different_first_page' property and clears the first page's footer
    content. It then modifies the primary footer by directly manipulating the
    underlying OOXML to insert 'PAGE' and 'NUMPAGES' fields. The resulting
    footer text is centered and styled with a 9pt gray (RGB 127, 127, 127)
    Arial font.

    Args:
        section (docx.section.Section): The document section object to be
            modified in-place.

    Raises:
        IndexError: If the primary footer (`section.footer`) does not contain at
            least one paragraph to modify.
        AttributeError: If the `section` object does not have the expected
            attributes of a `docx.section.Section` instance.
    """
    # Enable the 'different first page' header/footer setting to accommodate a distinct title page layout that omits standard footer content.
    section.different_first_page_header_footer = True

    # Explicitly clear the first-page footer content to ensure the cover page remains devoid of footers, adhering to the specified report format.
    first_page_footer = section.first_page_footer
    for p in first_page_footer.paragraphs:
        clear_paragraph(p)

    # Configure the primary footer for application to all pages subsequent to the title page.
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #
    clear_paragraph(paragraph)

    def add_field(p, field_code):
        r"""Inserts a complex Word field into a specified paragraph.

        This function directly manipulates the underlying Office Open XML (OOXML)
        structure of a `docx.paragraph.Paragraph` object to insert a complex field.
        It constructs the required sequence of `w:fldChar` elements (with `begin`,
        `separate`, and `end` types) and a `w:instrText` element containing the
        field instructions. A placeholder text of "0" is added as the default
        field result, which is intended to be updated by a Word processing
        application upon opening the document. The added run is styled with a
        9pt gray Arial font.

        Args:
            p (docx.paragraph.Paragraph): The paragraph object into which the
                field will be inserted. This object is modified in place.
            field_code (str): The instruction string for the Word field, such as
                'PAGE \\* MERGEFORMAT' or 'NUMPAGES'.
        """
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(127, 127, 127)

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)

        instrText = OxmlElement("w:instrText")
        instrText.set(ns.qn("xml:space"), "preserve")
        instrText.text = field_code
        run._r.append(instrText)

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "separate")
        run._r.append(fldChar2)

        #
        t = OxmlElement("w:t")
        t.text = "0"
        run._r.append(t)

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "end")
        run._r.append(fldChar3)

    r = paragraph.add_run("Página ")
    r.font.size = Pt(9)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(127, 127, 127)

    add_field(paragraph, "PAGE")

    r2 = paragraph.add_run(" de ")
    r2.font.size = Pt(9)
    r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor(127, 127, 127)

    add_field(paragraph, "NUMPAGES")


def load_payload(payload_path: Path) -> GlobalReportPayload:
    """Load and validate a global report payload from a JSON file.

    Reads and parses a JSON file from the given path, then validates its
    contents against the `GlobalReportPayload` Pydantic model. If the data
    does not conform to the model's schema, this function prints a detailed
    validation error to standard output and terminates the process with a
    non-zero exit code.

    Args:
        payload_path (pathlib.Path): The file system path to the JSON payload.

    Returns:
        GlobalReportPayload: An instance of `GlobalReportPayload` populated
            with the validated data from the file.

    Raises:
        FileNotFoundError: If the file at `payload_path` does not exist.
        json.JSONDecodeError: If the file content is not valid JSON.
    """
    payload_dict = load_json(payload_path)
    try:
        return GlobalReportPayload.model_validate(payload_dict)
    except ValidationError as e:
        print(f"❌ Error de validación Type-Safety en {payload_path}:")
        print(e)
        sys.exit(1)


def render_global_report(
    payload: GlobalReportPayload,
    template_path: Path,
    output_path: Path,
    client_dir: Path,
) -> Path:
    r"""{'docstring': 'Populates a Word document template to generate a global report.\n\nThis function orchestrates the rendering of a multi-section report. It first\nclears the body of the provided `.docx` template, then conditionally populates\nit with content from the data payload. It processes each available section,\nsuch as the executive summary or roadmap, by calling dedicated rendering\nsubroutines. The final document is saved to the specified output path.\n\nArgs:\n    payload (GlobalReportPayload): A structured data object containing the\n        metadata and content for all potential report sections.\n    template_path (Path): Filesystem path to the source `.docx` template file.\n    output_path (Path): Filesystem path where the generated `.docx` report\n        will be saved. Parent directories are created if they do not exist.\n    client_dir (Path): Filesystem path to the client-specific asset\n        directory, used for resolving relative paths to resources like images.\n\nReturns:\n    Path: The filesystem path of the successfully generated report, identical\n    to the `output_path` argument.\n\nRaises:\n    docx.opc.exceptions.PackageNotFoundError: If the file at `template_path`\n        does not exist, is not a valid ZIP archive, or is not a valid\n        Office Open XML file.\n    OSError: If the output directory cannot be created or the output file\n        cannot be written due to filesystem permissions or other I/O errors.'}."""
    doc = Document(str(template_path))
    clear_document_body(doc)

    if doc.sections:
        create_page_number_footer(doc.sections[0])

    render_cover(doc, payload)
    client_name = payload.meta.client
    if payload.executive_summary:
        render_executive_summary(
            doc,
            payload.executive_summary,
            payload.heatmap,
            payload.visuals,
            client_dir,
            client_name=client_name,
        )
    if payload.burning_platform:
        render_burning_platform(doc, payload.burning_platform, client_name=client_name)
    if payload.tower_bottom_lines and payload.heatmap:
        render_tower_bottom_lines(
            doc, payload.heatmap, payload.tower_bottom_lines, client_name=client_name
        )
    if payload.target_vision:
        render_target_vision(doc, payload.target_vision, client_name=client_name)
    if payload.execution_roadmap:
        render_execution_roadmap(
            doc,
            payload.execution_roadmap,
            payload.visuals,
            client_dir,
            client_name=client_name,
        )
    if payload.executive_decisions:
        render_executive_decisions(
            doc, payload.executive_decisions, client_name=client_name
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Informe Global CIO-READY renderizado: {output_path}")
    return output_path


def main(argv: list[str] | None = None) -> None:
    """Parses command-line arguments to render a global report from a template.

    This function serves as the main entry point for the command-line report
    generation script. It orchestrates the loading of a JSON payload and the
    rendering of a template to produce a final report file.

    It requires exactly three command-line arguments in the following order:
    1. Path to the JSON payload file.
    2. Path to the template file.
    3. Path for the output report file.

    Args:
        argv: An optional list of command-line arguments. If `None`, `sys.argv`
            is used. The list is expected to contain the script name followed
            by the payload, template, and output paths.

    Raises:
        FileNotFoundError: If the file specified by the payload or template
            path does not exist. This exception is propagated from downstream
            functions.

    Side Effects:
        Writes the rendered report to the file specified by the output path.
        Exits the program via `sys.exit(1)` if the number of command-line
        arguments is not exactly four (script name plus three required arguments).
    """
    if len(argv if argv is not None else sys.argv) != 4:
        sys.exit(1)
    payload_path = Path((argv if argv is not None else sys.argv)[1])
    template_path = Path((argv if argv is not None else sys.argv)[2])
    output_path = Path((argv if argv is not None else sys.argv)[3])
    client_dir = payload_path.parent
    payload = load_payload(payload_path)
    render_global_report(payload, template_path, output_path, client_dir)


if __name__ == "__main__":
    main()
