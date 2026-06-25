"""Contains shared utility functions and constants utilized by DOCX document rendering modules."""

from __future__ import annotations

import base64
import tempfile
from pathlib import Path

from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from assessment_engine.scripts.lib.text_utils import clean_text_for_word

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def find_style(doc, *candidates):
    """Locates the first style in a document matching a candidate name or ID.

    Searches the styles collection of a document object for a style whose name
    or ID matches one of the provided candidates. The comparison is performed
    after normalizing both the candidate strings and the style's `name` and
    `style_id` attributes by stripping leading/trailing whitespace and
    converting to lowercase. The first matching style object found is returned.

    Args:
        doc: An object possessing a `.styles` attribute, which is an iterable
            of style objects. Each style object is expected to have `name`
            and `style_id` attributes.
        *candidates: A variable number of style name or style ID strings to
            search for. Falsy candidates are ignored.

    Returns:
        The first style object that matches a candidate, or `None` if the
        document lacks a `.styles` attribute or if no match is found.
    """
    styles = getattr(doc, "styles", None)
    if styles is None:
        return None

    normalized_candidates = [
        str(candidate).strip().lower() for candidate in candidates if candidate
    ]
    for wanted in normalized_candidates:
        for style in styles:
            style_id = str(getattr(style, "style_id", "") or "").strip().lower()
            style_name = str(getattr(style, "name", "") or "").strip().lower()
            if style_id == wanted or style_name == wanted:
                return style
    return None


def apply_paragraph_style(paragraph, *candidates):
    """Applies the first valid style from a sequence of candidate names to a paragraph.

    Args:
        paragraph (docx.paragraph.Paragraph): The paragraph object to modify.
        *candidates (str): An ordered sequence of style names to search for within
            the document. The first name that corresponds to a defined style is
            applied.

    Returns:
        bool: `True` if a style was successfully found and applied; `False` if no
            candidate style was found in the document.
    """
    style = find_style(paragraph.part.document, *candidates)
    if style is None:
        return False
    paragraph.style = style
    return True


def apply_table_style(table, *candidates):
    """Applies the first available style from a sequence of candidate names to a table.

    This function searches the document's style definitions for the provided
    candidate style names in their given order. The first style that is found
    is applied to the table.

    Args:
        table (docx.table.Table): The `Table` object to which the style will be
            applied.
        *candidates (str): A variable-length sequence of style names to attempt
            to apply, in order of preference.

    Returns:
        bool: True if a style was successfully found and applied, False otherwise.
    """
    style = find_style(table.part.document, *candidates)
    if style is None:
        return False
    table.style = style.name
    return True


def get_style_numbering(doc, *candidates):
    r"""{'docstring': "Retrieves the numbering ID and indentation level from a style hierarchy.\n\n    This function searches for a style from an ordered list of candidate names.\n    Once a style is found, it traverses its inheritance chain (via the 'basedOn'\n    property) to locate the first defined numbering properties: the numbering\n    definition ID (`numId`) and its corresponding indentation level (`ilvl`).\n\n    A cycle detection mechanism is included to handle circular style dependencies.\n\n    Args:\n        doc (docx.document.Document): The `python-docx` Document object containing\n            the style definitions.\n        *candidates (str): A variable-length sequence of style names. The search is\n            performed in order, and the first matching style found is used as\n            the starting point for the hierarchy traversal.\n\n    Returns:\n        A tuple `(num_id, ilvl)` where `num_id` is the numbering definition ID\n        string and `ilvl` is the indentation level string. The indentation\n        level defaults to '0' if not explicitly defined in the same style\n        element where the numbering ID is found. Returns `(None, None)` if no\n        candidate style is found or if no numbering information is present\n        in the entire style inheritance chain."}."""
    style = find_style(doc, *candidates)
    if style is None:
        return None, None

    current = style.element
    visited = set()
    while current is not None:
        style_id = current.get(qn("w:styleId"))
        if style_id in visited:
            break
        visited.add(style_id)

        num_id = current.find("./w:pPr/w:numPr/w:numId", WORD_NS)
        ilvl = current.find("./w:pPr/w:numPr/w:ilvl", WORD_NS)
        if num_id is not None:
            return (
                num_id.get(qn("w:val")),
                ilvl.get(qn("w:val")) if ilvl is not None else "0",
            )

        based_on = current.find("./w:basedOn", WORD_NS)
        if based_on is None:
            break
        parent_style_id = based_on.get(qn("w:val"))
        current = find_style(doc, parent_style_id)
        current = current.element if current is not None else None

    return None, None


def set_paragraph_numbering(paragraph, num_id, ilvl=0):
    """Apply a numbering definition to a paragraph by manipulating its OOXML.

    This function directly modifies the underlying Office Open XML (OOXML) of a
    `Paragraph` object to associate it with a specific numbering definition.
    It first removes any pre-existing `w:numPr` (numbering properties) element
    from the paragraph's properties before constructing and appending a new one.
    This ensures the specified numbering is applied atomically.

    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph instance to modify.
        num_id (Optional[int]): The identifier (`w:numId`) for a numbering
            definition that must exist within the document's numbering part.
            If `None`, no operation is performed.
        ilvl (int): The zero-based indentation level (`w:ilvl`) to apply.

    Returns:
        bool: `True` if the numbering properties were successfully applied, or
            `False` if `num_id` was `None`.
    """
    if num_id is None:
        return False

    pPr = paragraph._p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag.split("}")[-1] == "numPr":
            pPr.remove(child)

    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numPr.append(ilvl_el)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numPr.append(num_id_el)
    pPr.append(numPr)
    return True


def apply_bullet_list_format(paragraph):
    """Apply bullet list formatting to a paragraph object in-place.

    This function first removes any pre-existing numbering (`numPr`) and
    indentation (`ind`) properties from the paragraph's underlying XML element.
    It then attempts to apply a standard bullet list style from a predefined
    sequence of style names. After a style is applied, it retrieves the
    corresponding numbering definition ID and indentation level from the document's
    style definitions and assigns them to the paragraph. Finally, it standardizes
    the paragraph's spacing, line height, and sets the alignment to left.

    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph object to modify.

    Returns:
        None
    """
    pPr = paragraph._p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag.split("}")[-1] in {"numPr", "ind"}:
            pPr.remove(child)

    apply_paragraph_style(
        paragraph, "Bullet", "NTTFlushBullet1", "List Paragraph", "Prrafodelista"
    )
    num_id, ilvl = get_style_numbering(
        paragraph.part.document,
        "Bullet",
        "NTTFlushBullet1",
        "List Paragraph",
        "Prrafodelista",
    )
    set_paragraph_numbering(paragraph, num_id, ilvl or 0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def enable_update_fields_on_open(doc):
    """Sets the 'update fields on open' property within a .docx file's settings.

    This function modifies the document's underlying OOXML to instruct compatible
    word processors to automatically update all fields (e.g., Table of
    Contents) when the document is opened. It achieves this by locating or
    creating the `w:updateFields` element within the document's settings part
    and setting its `w:val` attribute to `true`.

    The modification is performed in-place on the provided Document object.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to be
            modified.
    """
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def clean_text(value):
    """Clean a text string for safe rendering in a Word document."""
    return clean_text_for_word(value)


def clean_list(value):
    """Normalizes an input value into a list of cleaned, non-empty strings.

    This function processes a value that may be a single string, a list of
    strings, or `None`. It uses an internal `clean_text` utility to process
    each string element. Scalar values are wrapped in a list after cleaning. Any
    elements that are empty strings after cleaning are discarded.

    Args:
        value (str | list[str] | None): The value to process. Can be a single
            string, a list of strings, or `None`.

    Returns:
        list[str]: A list containing the cleaned, non-empty string(s). An empty
            list is returned if the input is `None` or results in no non-empty
            strings after cleaning.
    """
    if value is None:
        return []
    if isinstance(value, list):
        return [clean_text(item) for item in value if clean_text(item)]
    text = clean_text(value)
    return [text] if text else []


def clean_paragraph_list(value):
    r"""Normalizes a string or list of strings into a list of cleaned paragraphs.

    This function processes an input (`None`, a single string, or a list of
    strings) into a standardized list of non-empty paragraph strings. All
    string content is cleaned using an internal utility to remove extraneous
    whitespace and control characters.

    The normalization logic is as follows:
    - If the input is `None`, an empty list is returned.
    - If the input is a list, each string element is individually cleaned, and
      any resulting empty strings are omitted from the output.
    - If the input is a single string, it is first cleaned. The result is then
      split into paragraphs using a double newline (`\n\n`) as a delimiter.
      Each resulting part is stripped of leading/trailing whitespace, and
      empty parts are discarded. If the cleaned input string was non-empty but
      resulted in an empty list after this process (e.g., the input was `"\n\n"`),
      the function returns a list containing the original cleaned string.

    Args:
        value (Optional[Union[str, List[str]]]): The value to normalize. Can be a
            single string, a list of strings, or `None`.

    Returns:
        List[str]: A list of cleaned, non-empty paragraph strings. Returns an
            empty list if the input is `None` or yields no content after cleaning.
    """
    if value is None:
        return []
    if isinstance(value, list):
        return [clean_text(item) for item in value if clean_text(item)]
    text = clean_text(value)
    if not text:
        return []
    paragraphs = [part.strip() for part in text.split("\n\n") if part.strip()]
    return paragraphs or [text]


def resolve_radar_chart_image(image_path):
    r"""["Resolve an image source string to a local `pathlib.Path` object.\n\n    This function translates a string identifier for an image into a validated\n    `pathlib.Path` object on the local filesystem. It handles two formats:\n    base64-encoded data URIs and standard file paths.\n\n    If the input string is a data URI (e.g., 'data:image/png;base64,...'),\n    its base64 content is decoded and written to a new temporary file. The\n    path to this file is returned. The temporary file persists after this\n    function returns and its lifecycle must be managed by the caller.\n\n    If the input is treated as a standard file path, the function verifies\n    that the path exists and points to a regular file.\n\n    Args:\n        image_path (str): The image source, specified as either a local\n            filesystem path or a base64-encoded data URI.\n\n    Returns:\n        Optional[pathlib.Path]: A `Path` object pointing to the resolved image\n            file on the local filesystem. Returns `None` if the input is an\n            empty string, the file path is invalid, or if an error occurs\n            while processing a data URI."]."""
    image_path = clean_text(image_path)
    if not image_path:
        return None
    if image_path.startswith("data:image/") and ";base64," in image_path:
        try:
            _, encoded = image_path.split(";base64,", 1)
            raw = base64.b64decode(encoded)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            tmp.write(raw)
            tmp.close()
            return Path(tmp.name)
        except Exception:
            return None
    path = Path(image_path)
    if path.exists() and path.is_file():
        return path
    return None


def iter_paragraphs(doc):
    """Iterate over all paragraphs in a document, including those within tables.

    Creates a flattened generator that yields paragraphs from the main document
    body, followed by paragraphs from all table cells, in document order.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to process.

    Yields:
        docx.text.paragraph.Paragraph: Successive paragraph objects from the
            document.

    Raises:
        AttributeError: If the `doc` object does not have `paragraphs` or
            `tables` attributes.
    """
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def find_first_paragraph(doc, placeholder):
    """Return the first paragraph object containing a given placeholder text."""
    for paragraph in iter_paragraphs(doc):
        if placeholder in paragraph.text:
            return paragraph
    return None


def clear_paragraph(paragraph):
    """Clears all content from a paragraph while preserving its formatting.

    This function directly manipulates the underlying Office Open XML (OOXML)
    structure of the paragraph. It iterates through and removes all child elements
    of the paragraph's XML representation, except for the paragraph properties
    element (`w:pPr`). This operation deletes all text runs, images, and other
    inline content, while retaining paragraph-level formatting such as style,
    indentation, and spacing. The `paragraph` object is modified in-place.

    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph object to clear.

    Returns:
        None
    """
    element = paragraph._element
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def clear_paragraph_properties(paragraph):
    """Remove all direct formatting properties from a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    for child in list(pPr):
        pPr.remove(child)


def remove_paragraph(paragraph):
    r"""{'docstring': "Detaches a paragraph from its parent element in the document's XML tree.\n\n    This function accesses the underlying `lxml` element of the paragraph\n    object and removes it from its parent container. This effectively deletes\n    the paragraph from the document structure. If the paragraph's element has\n    no parent (i.e., it is already detached), the function performs no\n    operation.\n\n    Args:\n        paragraph (docx.text.paragraph.Paragraph): The paragraph object to be\n            removed from the document.\n\n    Raises:\n        AttributeError: If the `paragraph` object lacks the expected `_element`\n            attribute, suggesting it is not a valid `python-docx` Paragraph."}."""
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_after_block(block):
    r"""{'docstring': 'Inserts a new, empty paragraph immediately after a specified block element.\n\nThis function directly manipulates the underlying Office Open XML (OXML) tree\nof the document. It identifies the XML element for the provided `block` (either\na <w:p> for a Paragraph or a <w:tbl> for a Table) and inserts a new <w:p>\nelement as its next sibling within the parent container. This operation is\nuseful for programmatically adding content between existing block-level items.\n\nArgs:\n    block (Union[docx.text.paragraph.Paragraph, docx.table.Table]): The\n        reference block element after which the new paragraph will be inserted.\n\nReturns:\n    docx.text.paragraph.Paragraph: A `Paragraph` object representing the newly\n    created and inserted element in the document tree.\n\nRaises:\n    TypeError: If the `block` argument is not an instance of `Paragraph` or\n        `Table`.'}."""
    if isinstance(block, Paragraph):
        anchor = block._p
        parent = block._parent
    elif isinstance(block, Table):
        anchor = block._tbl
        parent = block._parent
    else:
        raise TypeError("Unsupported block type")

    new_p = OxmlElement("w:p")
    anchor.addnext(new_p)
    return Paragraph(new_p, parent)


def insert_field_paragraph_after_block(block, field_code, placeholder_text=""):
    r"""Inserts a paragraph containing a complex field after a block element.

    Constructs the underlying OOXML for a complex field, which is composed of
    `w:fldChar` (begin, separate, end) and `w:instrText` elements. The field
    is marked as 'dirty', which instructs a client application like Microsoft
    Word to re-evaluate its contents when the document is opened. This is
    commonly used for features like a Table of Contents (TOC).

    Args:
        block (docx.block.BlockItem): The block element (Paragraph or Table)
            after which the new field paragraph will be inserted.
        field_code (str): The instruction string for the field, such as
            'TOC \\o "1-3" \\h \\z \\u'.
        placeholder_text (str): The text to be displayed as the field's
            result until the field is updated. If an empty string is provided,
            no result text is inserted between the 'separate' and 'end'
            field markers.

    Returns:
        docx.text.paragraph.Paragraph: The new Paragraph object containing the
            constructed complex field.
    """
    paragraph = insert_paragraph_after_block(block)
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    run._r.append(instr_text)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_sep)

    if placeholder_text:
        text = OxmlElement("w:t")
        text.text = placeholder_text
        run._r.append(text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)
    return paragraph


def add_run(paragraph, text, bold=False, font_size=10.5):
    r"""{'docstring': "Adds and formats a new text run within a `docx` paragraph.\n\nArgs:\n    paragraph (docx.text.paragraph.Paragraph): The paragraph object to which\n        the new run will be added.\n    text (str): The string content for the new run.\n    bold (bool): If True, the run's text is formatted as bold.\n    font_size (float): The font size for the run, specified in points.\n\nReturns:\n    docx.text.run.Run: The newly created and formatted run object."}."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    return run


def apply_body_format(paragraph, justify=True, space_after=6, font_size=10.5):
    """Applies a standard set of body text formatting to a Paragraph object.

    Modifies the provided `python-docx` Paragraph object in-place. This utility
    configures alignment, spacing, and font properties. Specifically, it sets
    the space before to 0, the line spacing to 1.05, and applies the
    specified alignment, space after, and font size to all text runs.

    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph object to format.
        justify (bool, optional): If True, justifies the paragraph text;
            otherwise, the text is left-aligned. Defaults to True.
        space_after (float, optional): The spacing in points to apply after the
            paragraph. Defaults to 6.
        font_size (float, optional): The font size in points for all text runs
            within the paragraph. Defaults to 10.5.

    Raises:
        AttributeError: If `paragraph` lacks the expected attributes of a
            `python-docx` Paragraph object.
    """
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.05
    for run in paragraph.runs:
        run.font.size = Pt(font_size)


def replace_simple_placeholder(
    doc, placeholder, value, align=None, font_size=None, bold=None
):
    """Replaces a paragraph containing a text placeholder with a new value and format.

    Searches all paragraphs in a `docx.document.Document` for a given
    placeholder string. If a paragraph's text contains the placeholder, the
    entire paragraph is cleared. A new run is then added containing the original
    text with the placeholder substituted by the new value.

    This function modifies the document in-place. It also applies specific
    paragraph formatting: `space_before` is set to 0 Pt, `line_spacing` to
    1.05, and `space_after` is set to 6 Pt (unless the placeholder is
    "{{SCORING_METHOD_NOTE}}").

    Args:
        doc (docx.document.Document): The `python-docx` Document object to modify.
        placeholder (str): The placeholder text to find (e.g., "{{CLIENT_NAME}}").
        value (Any): The value to substitute for the placeholder. The value is
            converted to a string before insertion.
        align (Optional[WD_ALIGN_PARAGRAPH]): The alignment to apply to the
            paragraph, typically from `docx.enum.text.WD_ALIGN_PARAGRAPH`.
            Defaults to None.
        font_size (Optional[int]): The font size in points. Defaults to None.
        bold (Optional[bool]): Specifies whether to apply bold formatting.
            Defaults to None.
    """
    for paragraph in iter_paragraphs(doc):
        if placeholder in paragraph.text:
            text = paragraph.text.replace(placeholder, clean_text(value))
            clear_paragraph(paragraph)
            run = paragraph.add_run(text)
            if font_size:
                run.font.size = Pt(font_size)
            if bold is not None:
                run.bold = bold
            if align:
                paragraph.alignment = align
            paragraph.paragraph_format.space_before = Pt(0)
            if placeholder != "{{SCORING_METHOD_NOTE}}":
                paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.05


def render_multi_paragraph_block(doc, placeholder, texts):
    """Replaces a text placeholder with a multi-paragraph block in a document.

    This function locates the first paragraph containing the `placeholder` text.
    The content of this paragraph is cleared and replaced with the first string
    from the `texts` list. Each subsequent string in the list is then rendered
    in a new paragraph, inserted sequentially after the previous one. All
    generated paragraphs are formatted with a consistent body style (justified
    alignment, 10.5pt font, 6pt spacing after).

    If the `texts` list is empty (or becomes empty after internal cleaning of
    empty strings), the entire paragraph containing the placeholder is removed
    from the document. If the placeholder is not found, the document remains
    unmodified.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to be
            modified.
        placeholder (str): The placeholder text to find within a paragraph.
        texts (List[str]): A list of strings, with each string forming the content
            of a new paragraph.

    Returns:
        None. The `doc` object is modified in place.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return
    texts = clean_paragraph_list(texts)
    if not texts:
        remove_paragraph(paragraph)
        return

    clear_paragraph(paragraph)
    add_run(paragraph, texts[0], font_size=10.5)
    apply_body_format(paragraph, justify=True, space_after=6)
    anchor = paragraph

    for text in texts[1:]:
        new_paragraph = insert_paragraph_after_block(anchor)
        add_run(new_paragraph, text, font_size=10.5)
        apply_body_format(new_paragraph, justify=True, space_after=6)
        anchor = new_paragraph


def strip_numbering_and_indents(paragraph):
    """Remove numbering and reset formatting for a python-docx Paragraph object.

    This function modifies the paragraph object in-place. It directly manipulates
    the underlying OOXML by removing the `<w:numPr>` (numbering properties) and
    `<w:ind>` (indentation) child elements from the paragraph's properties
    (`<w:pPr>`).

    Following the XML manipulation, it resets other visual formatting attributes
    to a neutral baseline state:
      * Left indent: 0 points
      * First-line indent: 0 points
      * Space before: 0 points
      * Space after: 6 points
      * Line spacing: 1.05
      * Alignment: Left

    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph object to modify.
    """
    pPr = paragraph._p.get_or_add_pPr()
    for child in list(pPr):
        tag = child.tag.split("}")[-1]
        if tag in {"numPr", "ind"}:
            pPr.remove(child)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def render_list_at_placeholder(doc, placeholder, items):
    """Replaces a placeholder paragraph in a .docx document with a bulleted list.

    Locates the first paragraph containing the `placeholder` text and replaces its
    content with a bulleted list generated from `items`. The modification occurs
    in-place.

    If the `items` iterable is empty or contains only whitespace strings after
    cleaning, the entire placeholder paragraph is removed from the document.
    Otherwise, the placeholder paragraph is cleared, repurposed as the first list
    item, and subsequent items are added in new paragraphs below it. Each list
    item is formatted with a standard bullet point style and a 10pt font size.
    Any pre-existing numbering or indentation on the affected paragraphs is
    removed. Additionally, the vertical alignment of the parent container (e.g.,
    a table cell) is set to top-aligned (`WD_ALIGN_VERTICAL.TOP`).

    Args:
        doc (docx.document.Document): The python-docx Document object to modify.
        placeholder (str): The unique text string to find. The paragraph
            containing this text serves as the anchor for replacement.
        items (Iterable[str]): A collection of strings to render as bullet
            points. Empty or whitespace-only strings are ignored.

    Returns:
        None.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return
    items = [clean_text(x) for x in items if clean_text(x)]
    if not items:
        remove_paragraph(paragraph)
        return

    clear_paragraph(paragraph)
    strip_numbering_and_indents(paragraph)
    if hasattr(paragraph._parent, "vertical_alignment"):
        paragraph._parent.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    anchor = paragraph

    for idx, item in enumerate(items):
        current = anchor if idx == 0 else insert_paragraph_after_block(anchor)
        clear_paragraph(current)
        strip_numbering_and_indents(current)
        if hasattr(current._parent, "vertical_alignment"):
            current._parent.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        apply_bullet_list_format(current)
        content = current.add_run(item)
        content.font.size = Pt(10)
        anchor = current


def move_table_after_paragraph(paragraph, table):
    r"""{'docstring': "Move a table to the position immediately after a specified paragraph.\n\nThis function directly manipulates the document's underlying lxml object\ntree. It accesses the internal XML elements of the paragraph (`_p`) and\ntable (`_tbl`) objects, then repositions the table element to be the\nimmediate subsequent sibling of the paragraph element. The change is made\nin-place.\n\nArgs:\n    paragraph (docx.paragraph.Paragraph): The paragraph object after which the\n        table should be placed.\n    table (docx.table.Table): The table object to move.\n\nReturns:\n    None"}."""
    table_element = table._tbl
    paragraph._p.addnext(table_element)


def shade_cell(cell, fill):
    r"""{'docstring': "Sets the background shading for a table cell.\n\nThis function directly manipulates the underlying Open XML of the cell's\n`tcPr` (table cell properties) element by appending a `w:shd` (shading)\nchild.\n\nArgs:\n    cell (docx.table._Cell): The `python-docx` cell object to modify.\n    fill (str): The color value for the background fill. This value is\n        assigned to the `w:fill` attribute and is typically a hex-encoded\n        RGB color string (e.g., 'FF0000').\n\nReturns:\n    None. The cell object is modified in-place."}."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def clear_cell_shading(cell):
    r"""{'docstring': 'Remove any background shading from a table cell.\n\nThis function operates on the low-level OOXML structure of the table cell.\nIt accesses the `w:tcPr` (Table Cell Properties) element and removes any\n`w:shd` (Shading) child elements present. This modification is performed\nin-place on the `cell` object.\n\nArgs:\n    cell (docx.table._Cell): The cell object whose shading will be removed.'}."""
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.split("}")[-1] == "shd":
            tcPr.remove(child)


def set_cell_text(
    cell,
    text,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    font_size=10.5,
    vertical=WD_ALIGN_VERTICAL.CENTER,
    space_after=2,
):
    r"""{'docstring': "Sets the text and formatting for a DOCX table cell.\n\nResets the cell's content to a single paragraph containing the specified text.\nThis function first clears any pre-existing content, which ensures the cell\ncontains exactly one paragraph. It then applies formatting for horizontal\nand vertical alignment, font size, bolding, and paragraph spacing. Note that\nparagraph `space_before` is set to 0 points and `line_spacing` is set to 1.0.\n\nArgs:\n    cell (docx.table._Cell): The `python-docx` cell object to modify.\n    text (str): The string content to be placed in the cell.\n    bold (bool): If True, the text is formatted as bold. Defaults to False.\n    align (int): Horizontal paragraph alignment. Must be a member of\n        `docx.enum.text.WD_ALIGN_PARAGRAPH`. Defaults to\n        `WD_ALIGN_PARAGRAPH.JUSTIFY`.\n    font_size (Union[int, float]): The font size in points. Defaults to 10.5.\n    vertical (int): Vertical alignment of content within the cell. Must be a\n        member of `docx.enum.table.WD_ALIGN_VERTICAL`. Defaults to\n        `WD_ALIGN_VERTICAL.CENTER`.\n    space_after (Union[int, float]): The space after the paragraph in points.\n        Defaults to 2.\n\nReturns:\n    None.\n\nRaises:\n    AttributeError: If `cell` is not a valid `python-docx` cell object.\n    IndexError: If the cell contains no paragraphs after its content is cleared."}."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(clean_text(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = vertical


def set_repeat_table_header(row):
    """Marks a table row to repeat as a header on subsequent pages.

    This function modifies the underlying OOXML of the row by adding a
    `w:tblHeader` element to its properties. This instructs the rendering
    application to repeat this row at the top of each new page that the
    table spans.

    Args:
        row (docx.table._Row): The table row object to be configured as a
            repeating header.
    """
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def prevent_row_break(row):
    r"""{'docstring': 'Prevents a table row from breaking across a page.\n\nThis function modifies the underlying OOXML of a `docx.table._Row` object\nto insert a `w:cantSplit` element within its `w:trPr` (table row\nproperties). This property instructs rendering applications, such as Microsoft\nWord, to treat the row as an atomic unit that should not be split by a\npage break. The input `row` object is modified in-place.\n\nArgs:\n    row (docx.table._Row): The table row object from the `python-docx` library\n        to modify.\n\nReturns:\n    None\n\nRaises:\n    AttributeError: If the `row` object does not have the expected internal\n        `_tr` attribute, which contains the core XML element.'}."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)


def finalize_table(table):
    """Apply 'Table Grid' style, left alignment, and autofit to a docx Table object."""
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True


def autofit_table_to_contents(table):
    """Forces a `python-docx` table to autofit its contents via direct OOXML manipulation.

    This function provides a more assertive autofit capability than the default
    `table.autofit = True` property by directly modifying the underlying
    Office Open XML (OOXML). It systematically removes any pre-existing
    table grid (`w:tblGrid`), table-level width (`w:tblW`), table layout
    (`w:tblLayout`), and cell-level width (`w:tcW`) properties. Subsequently,
    it reapplies layout properties, setting the table layout type to 'autofit'
    and both table and cell width types to 'auto'.

    This process is effective for overriding fixed widths that may be inherited
    from a document template or set by prior operations, ensuring the table
    dimensions conform strictly to its content.

    Args:
        table (docx.table.Table): The `python-docx` Table object to modify in-place.

    Returns:
        None

    Raises:
        AttributeError: If the provided object lacks the expected internal
            attributes of a `python-docx` Table (e.g., `_tbl`, `_tc`).
    """
    table.autofit = True
    tbl = table._tbl

    for child in list(tbl):
        if child.tag.split("}")[-1] == "tblGrid":
            tbl.remove(child)

    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    for child in list(tblPr):
        tag = child.tag.split("}")[-1]
        if tag in {"tblW", "tblLayout"}:
            tblPr.remove(child)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "auto")
    tblW.set(qn("w:w"), "0")
    tblPr.append(tblW)

    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for child in list(tcPr):
                if child.tag.split("}")[-1] == "tcW":
                    tcPr.remove(child)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:type"), "auto")
            tcW.set(qn("w:w"), "0")
            tcPr.append(tcW)


def render_note_box(doc, placeholder, text):
    """Replaces a placeholder paragraph in a document with a styled note box.

    This function locates the first paragraph containing the exact `placeholder`
    text and substitutes it with a single-cell table formatted as a note box.
    The note box is styled with a light blue background shade (#F2F6FA),
    left-aligned text with a 9.5 pt font, and vertical centering within the cell.

    If the `placeholder` is not found, the document is not modified. If the
    provided `text` is empty or becomes empty after sanitization, the placeholder
    paragraph is removed, and no table is inserted. The table is automatically
    sized to fit its contents.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to modify.
        placeholder (str): The text content of the paragraph to find and replace.
        text (str): The text to be placed inside the styled note box.

    Returns:
        None. The document object is modified in-place.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return
    text = clean_text(text)
    if not text:
        remove_paragraph(paragraph)
        return

    table = doc.add_table(rows=1, cols=1)
    finalize_table(table)
    move_table_after_paragraph(paragraph, table)
    cell = table.cell(0, 0)
    set_cell_text(
        cell,
        text,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=9.5,
        vertical=WD_ALIGN_VERTICAL.CENTER,
        space_after=0,
    )
    shade_cell(cell, "F2F6FA")
    autofit_table_to_contents(table)
    remove_paragraph(paragraph)


def render_pillar_score_table(doc, placeholder, rows):
    """Renders and inserts a formatted pillar score table into a DOCX document.

    Locates a paragraph matching the `placeholder` text, then inserts a
    formatted table populated with pillar score data immediately after it. The
    original placeholder paragraph is subsequently removed.

    The generated table includes a shaded header row with fixed Spanish titles
    ("Pilar", "Score", "Nivel", "Lectura ejecutiva") that is configured to
    repeat across page breaks. Specific formatting for alignment, font size,
    and cell spacing is applied throughout the table.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to be
            modified in-place.
        placeholder (str): The exact text content of the paragraph that serves
            as the insertion point for the table.
        rows (List[Dict[str, Any]]): A list of dictionaries, where each
            dictionary represents a data row. Expected keys include
            'pillar_label', 'score_display', 'maturity_band', and
            'executive_reading'.

    Returns:
        None. The function modifies the input `doc` object directly.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return

    table = doc.add_table(rows=1, cols=4)
    finalize_table(table)
    move_table_after_paragraph(paragraph, table)

    headers = ["Pilar", "Score", "Nivel", "Lectura ejecutiva"]
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    set_repeat_table_header(table.rows[0])

    for item in rows:
        row = table.add_row()
        prevent_row_break(row)
        set_cell_text(
            row.cells[0],
            item.get("pillar_label", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10.2,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[1],
            item.get("score_display", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10.2,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[2],
            item.get("maturity_band", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10.2,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[3],
            item.get("executive_reading", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10.2,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )

    autofit_table_to_contents(table)
    remove_paragraph(paragraph)


def render_risks_table(doc, placeholder, rows):
    """Inserts a formatted, 4-column table of risks into a DOCX document.

    Locates a paragraph containing the specified placeholder text and inserts a
    table immediately after it. The table is populated with risk data and
    features styled, repeating headers ('Riesgo', 'Impacto', 'Probabilidad',
    'Mitigación resumida'). After populating the table, the function
    autofits the column widths to the content and removes the original
    placeholder paragraph.

    If the placeholder is not found, the document remains unmodified.

    Args:
        doc (docx.document.Document): The document object to be modified.
        placeholder (str): The text within a paragraph that serves as the
            anchor for table insertion.
        rows (List[Dict[str, str]]): A list of dictionaries, where each
            dictionary represents a row. Expected keys are 'risk', 'impact',
            'probability', and 'mitigation_summary'.

    Returns:
        None. The `doc` object is modified in-place.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return

    table = doc.add_table(rows=1, cols=4)
    finalize_table(table)
    move_table_after_paragraph(paragraph, table)

    headers = ["Riesgo", "Impacto", "Probabilidad", "Mitigación resumida"]
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    set_repeat_table_header(table.rows[0])

    for item in rows:
        row = table.add_row()
        prevent_row_break(row)
        set_cell_text(
            row.cells[0],
            item.get("risk", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[1],
            item.get("impact", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[2],
            item.get("probability", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[3],
            item.get("mitigation_summary", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )

    autofit_table_to_contents(table)
    remove_paragraph(paragraph)


def render_gap_table(doc, placeholder, rows):
    """Renders and inserts a gap analysis table into a Word document.

    Locates a paragraph containing the specified placeholder text, removes the
    paragraph, and inserts a formatted four-column table in its place. The table
    is populated with the provided row data. A styled header row with the column
    titles "Pilar", "Situación actual", "Estado objetivo", and "Brecha clave" is
    added and configured to repeat across page breaks. Cell formatting, including
    alignment, font size, and shading, is applied. The table columns are autofit
    to their contents.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to be
            modified.
        placeholder (str): The text within a single paragraph that marks the
            insertion point for the table.
        rows (list[dict[str, str]]): A list of dictionaries, where each
            dictionary represents a data row. Each dictionary should contain the
            keys 'pillar', 'as_is_summary', 'target_state', and 'key_gap'.
            Missing keys will result in empty cells.

    Returns:
        None: The `doc` object is modified in-place.
    """
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return

    table = doc.add_table(rows=1, cols=4)
    finalize_table(table)
    move_table_after_paragraph(paragraph, table)

    headers = ["Pilar", "Situación actual", "Estado objetivo", "Brecha clave"]
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    set_repeat_table_header(table.rows[0])

    for item in rows:
        row = table.add_row()
        prevent_row_break(row)
        set_cell_text(
            row.cells[0],
            item.get("pillar", ""),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[1],
            item.get("as_is_summary", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[2],
            item.get("target_state", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        set_cell_text(
            row.cells[3],
            item.get("key_gap", ""),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )

    autofit_table_to_contents(table)
    remove_paragraph(paragraph)


def render_initiative_cards(doc, placeholder, cards):
    r"""{'docstring': "Replaces a placeholder paragraph in a Word document with formatted tables.\n\nFinds the first paragraph containing the `placeholder` text within the\nprovided `doc` object. It then iterates through the `cards` list,\ngenerating and inserting a distinct, formatted table for each card at the\nplaceholder's original position.\n\nIf the `cards` list is empty, the placeholder paragraph is removed. If the\nplaceholder is not found, the document remains unmodified. Upon successful\ninsertion of all tables, the original placeholder paragraph is deleted. The\n`doc` object is modified in-place.\n\nArgs:\n    doc (docx.document.Document): The python-docx Document object to modify.\n    placeholder (str): The text within a paragraph that marks the insertion\n        point for the tables.\n    cards (list[dict]): A list of dictionaries, each representing an\n        initiative card. Expected keys include 'sequence', 'initiative',\n        'objective', 'priority', 'expected_outcome', and\n        'dependencies_display'."}."""
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return
    if not cards:
        remove_paragraph(paragraph)
        return

    anchor = paragraph
    for idx, item in enumerate(cards, start=1):
        table = doc.add_table(rows=5, cols=2)
        finalize_table(table)
        move_table_after_paragraph(anchor, table)

        title_row = table.rows[0]
        merged = title_row.cells[0].merge(title_row.cells[1])
        set_cell_text(
            merged,
            f"{item.get('sequence', idx)}. {item.get('initiative', '')}",
            bold=True,
            font_size=11,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(merged, "D9EAF7")
        prevent_row_break(title_row)

        labels = ["Objetivo", "Prioridad", "Resultado esperado", "Dependencias"]
        values = [
            item.get("objective", ""),
            item.get("priority", ""),
            item.get("expected_outcome", ""),
            item.get("dependencies_display", ""),
        ]
        for row_idx in range(4):
            row = table.rows[row_idx + 1]
            set_cell_text(
                row.cells[0],
                labels[row_idx],
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=10,
                vertical=WD_ALIGN_VERTICAL.CENTER,
                space_after=0,
            )
            shade_cell(row.cells[0], "D9EAF7")
            set_cell_text(
                row.cells[1],
                values[row_idx],
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=10,
                vertical=WD_ALIGN_VERTICAL.CENTER,
            )
            prevent_row_break(row)

        autofit_table_to_contents(table)
        spacer = insert_paragraph_after_block(table)
        spacer.paragraph_format.space_after = Pt(6)
        add_run(spacer, "", font_size=1)
        anchor = spacer

    remove_paragraph(paragraph)


def render_radar_chart(doc, placeholder, image_path):
    r"""{'docstring': "Replaces a placeholder paragraph in a DOCX document with a radar chart image.\n\nSearches for the first paragraph containing the `placeholder` text. If found,\nthe paragraph's content is cleared, and the specified image is inserted. The\nimage is center-aligned, set to a fixed width of 4.35 inches, and the\nparagraph is given 6 points of spacing after.\n\nIf the `placeholder` text is not found, the document remains unmodified. If\nthe `image_path` cannot be resolved to a valid local file, the entire\nplaceholder paragraph is removed from the document.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` Document object to be\n        modified.\n    placeholder (str): The text content that identifies the paragraph to be\n        replaced.\n    image_path (str): The file system path or resolvable resource identifier\n        for the image file.\n\nReturns:\n    None\n\nRaises:\n    FileNotFoundError: If the resolved path from `image_path` points to a\n        file that does not exist."}."""
    paragraph = find_first_paragraph(doc, placeholder)
    if not paragraph:
        return

    path = resolve_radar_chart_image(image_path)
    if path is None:
        remove_paragraph(paragraph)
        return

    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(4.35))
    paragraph.paragraph_format.space_after = Pt(6)


def add_heading_paragraph(doc, text, level=1):
    """Adds a custom-formatted heading paragraph to a `docx.document.Document`.

    The appearance of the heading is controlled by the `level` parameter, which
    applies specific styling rules:
      - Level 1: 20pt Georgia font in blue (RGB 0, 114, 188) with 18pt spacing
        before and after.
      - Level 2: 14pt Georgia font in blue (RGB 0, 114, 188) with 12pt spacing
        before and after.
      - Other levels: 11.5pt default font, bold, with 3pt spacing before and
        4pt after.

    Args:
        doc (docx.document.Document): The document object to which the heading
            will be added.
        text (str): The text content of the heading. This text is processed by
            the `clean_text` function before insertion.
        level (int): The heading level used for styling. Defaults to 1.

    Returns:
        docx.text.paragraph.Paragraph: The newly created and formatted heading
            paragraph object.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(clean_text(text))

    if level == 1:
        run.bold = False
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(18)
        paragraph.paragraph_format.page_break_before = False
        run.font.name = "Georgia"
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 114, 188)
    elif level == 2:
        run.bold = False
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        run.font.name = "Georgia"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 114, 188)
    else:
        run.bold = True
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(11.5)

    return paragraph


def add_body_paragraph(doc, text, justify=True, space_after=6, color_rgb=None):
    """Adds a formatted paragraph to the document body.

    The function first cleans the input text. If the resulting string is empty,
    no paragraph is added and the function returns None. Otherwise, it adds a new
    paragraph, populates it with the text, applies standard formatting such as
    justification and spacing, and optionally sets the font color.

    Args:
        doc (docx.document.Document): The document object to which the paragraph
            will be added.
        text (str): The text content for the paragraph.
        justify (bool): If True, the paragraph alignment is set to justified.
            Defaults to True.
        space_after (int): The spacing, in points, to apply after the paragraph.
            Defaults to 6.
        color_rgb (docx.shared.RGBColor, optional): An `RGBColor` object to set
            the font color. If None, the default font color is used.
            Defaults to None.

    Returns:
        docx.text.paragraph.Paragraph | None: The newly created `Paragraph`
            object, or None if the input text is empty after cleaning.
    """
    text = clean_text(text)
    if not text:
        return None
    paragraph = doc.add_paragraph()
    add_run(paragraph, text, font_size=10.5)
    apply_body_format(paragraph, justify=justify, space_after=space_after)
    if color_rgb:
        for run in paragraph.runs:
            run.font.color.rgb = color_rgb
    return paragraph


def add_bullet_list(doc, items):
    """Adds a formatted bulleted list to a `python-docx` document.

    This function processes a list of strings, filters out any empty or
    whitespace-only items after cleaning, and adds the remaining items as a
    bulleted list to the provided document. Each list item is formatted with a
    standard bullet point style and a font size of 10 points.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to which
            the bulleted list will be appended.
        items (List[str]): A list of strings, where each string represents a list
            item. Items that are empty or contain only whitespace after cleaning
            are omitted from the final list.
    """
    items = [clean_text(item) for item in items if clean_text(item)]
    for item in items:
        paragraph = doc.add_paragraph()
        strip_numbering_and_indents(paragraph)
        apply_bullet_list_format(paragraph)
        text = paragraph.add_run(item)
        text.font.size = Pt(10)


def add_label_value_paragraph(doc, label, value):
    """Adds a formatted paragraph containing a label and its associated value.

    The paragraph is structured as 'Label: Value', where the label is rendered
    in bold. The value is sanitized via an internal utility before rendering.
    If the sanitized value is an empty string, no paragraph is added to the
    document. Both the label and value text runs are set to a font size of
    10.5 points. Specific paragraph formatting, including left alignment and
    4-point spacing after the paragraph, is also applied.

    Args:
        doc (docx.document.Document): The document object to which the new
            paragraph will be appended.
        label (str): The text for the label. A colon and a space are
            automatically suffixed.
        value (str): The raw text for the value, which will be cleaned before
            being added.

    Returns:
        Optional[docx.text.paragraph.Paragraph]: The newly created paragraph
        object if the cleaned value is non-empty, otherwise None.
    """
    value = clean_text(value)
    if not value:
        return None
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10.5)
    value_run = paragraph.add_run(value)
    value_run.font.size = Pt(10.5)
    apply_body_format(paragraph, justify=False, space_after=4)
    return paragraph


def add_long_detail_table(doc, title, rows):
    r"""{'docstring': 'Adds a formatted two-column table with a merged header to a document.\n\nThe table consists of a single merged header row displaying the title,\nfollowed by data rows. Each data row features a shaded label in the first\ncolumn and its corresponding value in the second. Input rows are filtered\nto exclude any entries where the value is empty or consists only of\nwhitespace after cleaning. The table is automatically sized to fit its\ncontents.\n\nArgs:\n    doc (docx.document.Document): The `python-docx` Document object to which\n        the table will be added.\n    title (str): The text to display in the merged header row.\n    rows (List[Tuple[str, str]]): A list of tuples, where each tuple\n        represents a row containing a (label, value) pair.\n\nReturns:\n    Optional[docx.table.Table]: The created `docx.table.Table` object\n        appended to the document, or `None` if all provided rows are\n        filtered out due to empty values.\n\nRaises:\n    ValueError: If an element in the `rows` list does not contain exactly\n        two elements and cannot be unpacked into a (label, value) pair.'}."""
    rows = [
        (clean_text(label), clean_text(value))
        for label, value in rows
        if clean_text(value)
    ]
    if not rows:
        return None

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    finalize_table(table)

    header = table.rows[0]
    merged = header.cells[0].merge(header.cells[1])
    set_cell_text(
        merged,
        title,
        bold=True,
        font_size=11,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        vertical=WD_ALIGN_VERTICAL.CENTER,
        space_after=0,
    )
    shade_cell(merged, "D9EAF7")
    prevent_row_break(header)

    for idx, (label, value) in enumerate(rows, start=1):
        row = table.rows[idx]
        set_cell_text(
            row.cells[0],
            label,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
            space_after=0,
        )
        shade_cell(row.cells[0], "EEF5FB")
        set_cell_text(
            row.cells[1],
            value,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            vertical=WD_ALIGN_VERTICAL.CENTER,
        )
        prevent_row_break(row)

    autofit_table_to_contents(table)
    return table


def remove_page_break_only_paragraphs(doc):
    """Remove paragraphs consisting solely of a manual page break and whitespace.

    Iterates through a document's paragraphs to identify and remove those that
    contain a manual page break but have no visible text content, effectively
    treating whitespace-only paragraphs as empty. Detection of a manual page
    break is performed by inspecting the underlying Open XML for a `<w:br>`
    element with its `type` attribute set to `page`.

    This operation is performed in-place on the provided Document object and is
    useful for cleaning up formatting artifacts that can cause unwanted blank
    pages.

    Args:
        doc (docx.document.Document): The `python-docx` Document object to be
            modified in-place.

    Returns:
        None: The function modifies the document object directly.
    """
    to_remove = []
    for paragraph in doc.paragraphs:
        has_text = bool(clean_text(paragraph.text))
        has_page_break = False
        for run in paragraph.runs:
            for br in run._r.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
            ):
                if (
                    br.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                    )
                    == "page"
                ):
                    has_page_break = True
        if has_page_break and not has_text:
            to_remove.append(paragraph)
    for paragraph in to_remove:
        remove_paragraph(paragraph)
