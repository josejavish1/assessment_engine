from pathlib import Path

from docx import Document


def verify_text_content() -> None:
    """Compare the textual content of two DOCX files and print a report.

    This function performs a two-phase analysis to verify textual fidelity
    between a primary source document and a shadow document, whose paths are
    hardcoded. The first phase provides a quantitative comparison by calculating
    the percentage of verbatim paragraphs shared between the two files. The
    second phase provides a qualitative comparison by extracting and printing
    text samples from key sections (e.g., Executive Summary, SWOT table) to
    facilitate manual visual inspection.

    If either file path is not found, an error is printed to the console and the
    function terminates.

    Returns:
        None. The analysis is printed directly to standard output.

    Raises:
        docx.opc.exceptions.PackageNotFoundError: If either of the hardcoded paths
            does not reference a valid DOCX file.
    """
    original_path = Path("working/redeia_v3/T2/AS-IS_Anexo_Tecnico_T2.docx")
    shadow_path = Path("working/redeia_v3/T2/AS-IS_Anexo_Tecnico_T2.shadow.docx")

    if not original_path.exists() or not shadow_path.exists():
        print("❌ Uno o ambos archivos no existen.")
        return

    doc_orig = Document(str(original_path))
    doc_shadow = Document(str(shadow_path))

    orig_paras = [p.text.strip() for p in doc_orig.paragraphs if p.text.strip()]
    shadow_paras = [p.text.strip() for p in doc_shadow.paragraphs if p.text.strip()]

    print("======================================================================")
    print("      VERIFICACIÓN SOTA: ANÁLISIS DE FIDELIDAD TEXTUAL COMPLETA")
    print("======================================================================")
    print(f"Párrafos en Original (V3): {len(orig_paras)}")
    print(f"Párrafos en Sombra (V4):   {len(shadow_paras)}")
    print("-" * 70)

    # Phase 1: Establishes a baseline of verbatim paragraph correspondence between the source and target documents.
    intersection = set(orig_paras) & set(shadow_paras)
    print(f"✅ Párrafos idénticos en ambos documentos: {len(intersection)}")
    print(f"   Porcentaje de coincidencia textual de párrafos: {len(intersection) / len(orig_paras) * 100:.1f}%")
    print("-" * 70)

    # Phase 2: Produces contextualized textual excerpts for manual, side-by-side visual inspection of confirmed paragraph matches.
    print("👀 MUESTRAS DE TEXTO EXTRAÍDAS (COMPARACIÓN DIRECTA):")
    print("\n--- [Original V3] Primeros 3 párrafos del Resumen Ejecutivo:")
    #
    p_orig_count = 0
    for p in doc_orig.paragraphs:
        if "Redeia" in p.text and len(p.text) > 50:
            print(f"   * {p.text[:120]}...")
            p_orig_count += 1
            if p_orig_count >= 3:
                break

    print("\n--- [Shadow V4] Primeros 3 párrafos del Resumen Ejecutivo:")
    p_shadow_count = 0
    for p in doc_shadow.paragraphs:
        if "Redeia" in p.text and len(p.text) > 50:
            print(f"   * {p.text[:120]}...")
            p_shadow_count += 1
            if p_shadow_count >= 3:
                break

    print("\n--- [Original V3] Primeras 2 Fortalezas o Brechas en SWOT:")
    orig_swot_text = []
    for table in doc_orig.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0 and "Fortalezas" in table.rows[0].cells[0].text:
            for row in table.rows[1:3]:
                if len(row.cells) >= 2:
                    orig_swot_text.append((row.cells[0].text[:80], row.cells[1].text[:80]))
    for idx, (s, g) in enumerate(orig_swot_text):
        print(f"   Punto {idx+1}:")
        print(f"     └─ Fortaleza: {s}...")
        print(f"     └─ Brecha:    {g}...")

    print("\n--- [Shadow V4] Primeras 2 Fortalezas o Brechas en SWOT:")
    shadow_swot_text = []
    for table in doc_shadow.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0 and "Fortalezas" in table.rows[0].cells[0].text:
            for row in table.rows[1:3]:
                if len(row.cells) >= 2:
                    shadow_swot_text.append((row.cells[0].text[:80], row.cells[1].text[:80]))
    for idx, (s, g) in enumerate(shadow_swot_text):
        print(f"   Punto {idx+1}:")
        print(f"     └─ Fortaleza: {s}...")
        print(f"     └─ Brecha:    {g}...")

    print("\n======================================================================")


def doc_paragraphs_list(doc):
    """Extract the text from all non-empty paragraphs in a document object."""
    return [p.text for p in doc.paragraphs if p.text.strip()]


def paragraphs_with_keyword(doc_orig, keyword):
    """Return a list of paragraphs from an iterable that contain a given keyword."""
    return [p for p in doc_orig if keyword in p]


if __name__ == "__main__":
    verify_text_content()
