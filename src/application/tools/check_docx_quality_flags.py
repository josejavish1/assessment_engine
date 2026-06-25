from typing import Any

"""
Módulo check_docx_quality_flags.py.
Contiene la lógica y utilidades principales para el pipeline de Assessment Engine.
"""

import logging
import re
import sys
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

BAD_PATTERNS = [
    r"\{\{[^{}]+\}\}",
    r"\{'statement':",
    r'"statement":',
]


def collect_text(doc) -> Any:
    """Extracts all textual content from a `python-docx` document object.

    This function traverses all paragraphs in the main document body and within
    all table cells. Text from each non-empty paragraph is collected. The
    resulting text segments are then joined into a single string, with each
    segment separated by a newline character.

    Args:
        doc (docx.document.Document): The `python-docx` document object from
            which to extract text.

    Returns:
        str: A single string containing all text from the document's
            paragraphs and tables, separated by newlines.

    Raises:
        AttributeError: If the `doc` object does not conform to the expected
            `python-docx` Document interface and lacks `paragraphs` or `tables`
            attributes.
    """
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


def main(argv: list[str] | None = None) -> None:
    """Scan a DOCX file for predefined quality degradation patterns.

    Serves as the main entry point for a command-line script that validates a
    Microsoft Word document (.docx) against a set of quality criteria. The
    script expects a single command-line argument specifying the path to the
    document. It extracts the full text content and scans for matches against
    predefined regular expression patterns and for an excessive number of
    ellipsis occurrences (more than 3).

    The script's outcome is communicated via logging and exit codes. If quality
    issues are detected, it logs 'DOCX_QUALITY_FLAGS=YES' with details of the
    findings and exits with a status code of 1. If the document passes all
    checks, it logs 'DOCX_QUALITY_FLAGS=NO' and exits successfully.

    Args:
        argv: A list of command-line arguments. If None, `sys.argv` is used.
            Primarily intended for testing. The list is expected to contain
            the script name followed by a single path to a .docx file.

    Raises:
        SystemExit: 
            - If the number of command-line arguments is not two (script name
              and path).
            - If the specified DOCX file path does not exist.
            - If quality flags are found, exiting with a status code of 1.
    """
    if len(argv if argv is not None else sys.argv) != 2:
        raise SystemExit(
            "Uso: python -m scripts.tools.check_docx_quality_flags <docx_path>"
        )

    path = Path((argv if argv is not None else sys.argv)[1]).resolve()
    if not path.exists():
        raise SystemExit(f"No existe el archivo: {path}")

    doc = Document(str(path))
    text = collect_text(doc)

    flags = []
    for pattern in BAD_PATTERNS:
        matches = re.findall(pattern, text)
        if matches:
            flags.append((pattern, len(matches)))

    excessive_ellipsis = text.count("...")

    if flags or excessive_ellipsis > 3:
        logger.info("DOCX_QUALITY_FLAGS=YES")
        for pattern, count in flags:
            logger.info(f"pattern={pattern} count={count}")
        if excessive_ellipsis > 3:
            logger.info(f"ellipsis_count={excessive_ellipsis}")
        raise SystemExit(1)

    logger.info("DOCX_QUALITY_FLAGS=NO")


if __name__ == "__main__":
    main()
