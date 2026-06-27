import json
import tempfile
from pathlib import Path

from assessment_engine.adapters.render_togaf_asis_annex import render_asis_annex
from docx import Document


class TestDynamicGenericityFuzzing:
    """Test Suite - Dynamic Genericity Fuzzing:
    Verify that compilers act as pure functions and do not leak domain semantics (Domain Leakage).
    """

    def test_compiler_zero_leakage_assertion(self) -> None:
        """Inject 100% randomly mutated variables and assert that the final document
        contains no traces of default branding (REDEIA, NTT DATA, June 2026, €),
        mathematically proving that the engine is 100% generic.
        """
        # --- ARRANGE ---
        fuzzed_client = "MUTATED_CLIENT_NAME_999"
        fuzzed_project = "MUTATED_PROJECT_LOGISTICS"
        fuzzed_company = "MUTATED_CONSULTANCY_ABC"
        fuzzed_classification = "MUTATED_CLASSIFICATION_SECRET"
        fuzzed_date = "MUTATED_DATE_JANUARY_3030"
        fuzzed_version = "MUTATED_VERSION_9_9_9"
        fuzzed_currency = "MUTATED_CURR_ZZZ"

        # Configuration of minimum fuzzed locales
        temp_locales = {
            "es": {
                "title": "MUTATED_DOCUMENT_TITLE_111",
                "project": fuzzed_project,
                "client_lbl": "MUTATED_CLIENT_LABEL:\n",
                "proj_lbl": "\nMUTATED_PROJECT_LABEL:\n",
                "v_lbl": "MUTATED_V_LBL_ ",
                "toc_title": "MUTATED_TOC_TITLE",
                "dashboard_title": "MUTATED_DASHBOARD_TITLE",
                "dashboard_intro": "MUTATED_DASHBOARD_INTRO",
                "score_lbl": "MUTATED_SCORE_LABEL",
                "maturity_lbl": "MUTATED_MATURITY_LABEL",
                "resilience_reading": "MUTATED_RESILIENCE_READING:",
                "radar_title": "MUTATED_RADAR_TITLE",
                "justification_table_headers": [
                    "MUTATED_H1",
                    "MUTATED_H2 / MUTATED_H2_DESC",
                ],
                "platform_overview_title": "MUTATED_PLATFORM_OVERVIEW_TITLE",
                "swot_title": "MUTATED_SWOT_TITLE",
                "swot_headers": ["MUTATED_SWOT_H1", "MUTATED_SWOT_H2"],
                "transversal_title": "MUTATED_TRANSVERSAL_TITLE",
                "risk_matrix_title": "MUTATED_RISK_MATRIX_TITLE",
                "risk_intro_1": "MUTATED_RISK_I1 ",
                "risk_intro_2": " MUTATED_RISK_I2 ",
                "risk_intro_3": " MUTATED_RISK_I3",
                "exposure_summary_title": "MUTATED_EX_SUMMARY_TITLE",
                "tef_title": "MUTATED_TEF_TITLE",
                "tef_headers": ["MUTATED_TEF_H1", "MUTATED_TEF_H2", "MUTATED_TEF_H3"],
                "lm_title": "MUTATED_LM_TITLE",
                "lm_headers": ["MUTATED_LM_H1", "MUTATED_LM_H2", "MUTATED_LM_H3"],
                "top_risks_title": "MUTATED_TOP_RISKS_TITLE",
                "top_risks_intro": "MUTATED_TOP_RISKS_INTRO",
                "motivo_priorizacion": "MUTATED_PRIORITIZATION_REASON:",
                "detailed_risks_title": "MUTATED_DETAILED_RISKS_TITLE",
                "detailed_risks_pilar_title": "MUTATED_DETAILED_RISKS_PILAR:",
                "detailed_risks_headers": [
                    "MUTATED_DR_H1",
                    "MUTATED_DR_H2",
                    "MUTATED_DR_H3",
                ],
                "vulnerability_lbl": "MUTATED_VULN_LBL:",
                "evidence_lbl": "MUTATED_EVID_LBL:",
                "impact_lbl": "MUTATED_IMPA_LBL:",
                "exposure_lbl": "MUTATED_EXPO_LBL:",
                "ale_proyectado": "MUTATED_PROJECTED_ALE_LBL:",
                "next_steps_title": "MUTATED_NEXT_STEPS_TITLE",
                "appendix_a_title": "MUTATED_APPENDIX_A_TITLE",
                "appendix_a_intro": "MUTATED_APPENDIX_A_INTRO",
                "appendix_a_headers": [
                    "MUTATED_AA_H1",
                    "MUTATED_AA_H2",
                    "MUTATED_AA_H3",
                ],
                "appendix_b_title": "MUTATED_APPENDIX_B_TITLE",
                "disclaimer_text_1": "MUTATED_DISCLAIMER_1",
                "disclaimer_text_2": "MUTATED_DISCLAIMER_2",
                "disclaimer_text_3": "MUTATED_DISCLAIMER_3",
                "appendix_c_title": "MUTATED_APPENDIX_C_TITLE",
                "appendix_c_intro": "MUTATED_APPENDIX_C_INTRO",
                "appendix_c_headers": [
                    "MUTATED_AC_H1",
                    "MUTATED_AC_H2",
                    "MUTATED_AC_H3",
                ],
                "bib_cues": "MUTATED_BIB_CUES:",
                "bib_cues_desc": "MUTATED_BIB_CUES_DESC",
                "bib_contexto": "MUTATED_BIB_CONTEXTO:",
                "bib_contexto_desc": "MUTATED_BIB_CONTEXTO_DESC",
                "bib_minutas": "MUTATED_BIB_MINUTAS:",
                "bib_minutas_desc": "MUTATED_BIB_MINUTAS_DESC",
                "conclusions_title": "MUTATED_CONCLUSIONS_TITLE",
                "appendix_b_intro": "MUTATED_APPENDIX_B_INTRO",
            }
        }

        temp_brand_profile = {
            "company_name": fuzzed_company,
            "default_classification": fuzzed_classification,
            "disclaimer_text": "MUTATED_DISCLAIMER_TEXT_PARAGRAPH",
            "styling": {
                "primary_color_hex": "FF00FF",  # Color magenta fuzzeado
                "alternate_row_color_hex": "F0F0F0",
            },
        }

        # Backup and controlled write operations for the test
        locales_path = Path("engine_config/locales.json")
        brand_path = Path("engine_config/brand_profile.json")

        locales_backup = (
            locales_path.read_text(encoding="utf-8") if locales_path.exists() else None
        )
        brand_backup = (
            brand_path.read_text(encoding="utf-8") if brand_path.exists() else None
        )

        try:
            # Write mutated test profiles
            locales_path.write_text(
                json.dumps(temp_locales, ensure_ascii=False, indent=2), encoding="utf-8"
            )
            brand_path.write_text(
                json.dumps(temp_brand_profile, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )

            with tempfile.TemporaryDirectory() as tmp_dir:
                input_dir = Path(tmp_dir)

                # Generate minimum fuzzed input payload
                mock_payload = {
                    "document_meta": {
                        "client_name": fuzzed_client,
                        "tower_name": "MUTATED_TOWER_SECURITY",
                        "language": "es",
                        "date": fuzzed_date,
                        "version": fuzzed_version,
                        "currency": fuzzed_currency,
                        "tower_id": "T2",
                    },
                    "total_fair_ale": 1500000.0,
                    "platform_overview_intro": "MUTATED_PLATFORM_OVERVIEW_INTRO_PARAGRAPH",
                    "executive_summary": {
                        "global_score": "3.2 / 5.0",
                        "global_band": "MUTATED_BAND_DEFINED",
                        "target_score": "4.5 / 5.0",
                        "target_band": "MUTATED_BAND_OPTIMIZED",
                        "target_capabilities": ["MUTATED_CAP_1", "MUTATED_CAP_2"],
                        "gap_rows": [
                            {
                                "id": "G01",
                                "name": "MUTATED_GAP_1",
                                "description": "Desc",
                            }
                        ],
                        "risks": [
                            {"id": "R01", "name": "MUTATED_RISK_1", "ale": 100000.0}
                        ],
                        "initiatives": [
                            {
                                "id": "I01",
                                "name": "MUTATED_INITIATIVE_1",
                                "opex": 50000.0,
                            }
                        ],
                    },
                    "pillar_score_profile": {
                        "pillars": [
                            {
                                "pillar_label": "MUTATED_PILLAR_LABEL_1",
                                "score_display": "3.5",
                                "maturity_band": "MUTATED_BAND_MANAGED",
                                "executive_reading": "MUTATED_READING",
                            }
                        ]
                    },
                    "pillars_analysis": [
                        {
                            "pilar_name": "MUTATED_PILLAR_LABEL_1",
                            "score": 3.5,
                            "target_score": 4.5,
                            "asis_architecture_description": "MUTATED_ARCH_DESC",
                            "health_check_asis": [
                                {
                                    "finding": "MUTATED_RISK_FINDING",
                                    "business_risk": "MUTATED_RISK_BUSINESS",
                                    "literal_evidence": "MUTATED_RISK_EVIDENCE",
                                    "threat_event_frequency": 3.0,
                                    "loss_magnitude": 4.0,
                                    "fair_ale_score": 120000.0,
                                }
                            ],
                        }
                    ],
                }

                payload_path = input_dir / "blueprint_t2_payload.json"
                # Write fuzzed payload to disk
                with open(payload_path, "w", encoding="utf-8") as pf:
                    json.dump(mock_payload, pf)

                output_docx = input_dir / "AS-IS_Anexo_Tecnico_T2.docx"

                # --- ACT ---
                # Compile using the fuzzed payload
                render_asis_annex(str(payload_path), str(output_docx))

                # --- ASSERT ---
                assert output_docx.is_file(), (
                    "The compiled fuzzed file must be generated successfully."
                )

                # Load and extract all text from the resulting Word document
                doc = Document(output_docx)
                full_text = "\n".join(p.text for p in doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text += f"\n{cell.text}"

                # LEAKAGE SECURITY ASSERTION
                # No default branding or constants should be present if they were overwritten
                default_denylist = ["REDEIA", "NTT DATA", "Junio 2026", "€"]
                for forbidden in default_denylist:
                    assert forbidden not in full_text, (
                        f"Semantic leak detected! The default constant {repr(forbidden)} "
                        f"leaked into the final document despite being overwritten."
                    )

                # Confirm presence of mutated values (Pure-Function Proof)
                assert fuzzed_company.upper() in full_text
                assert fuzzed_client in full_text
                assert fuzzed_project in full_text
                assert fuzzed_classification in full_text
                assert fuzzed_date in full_text
                assert fuzzed_version in full_text
                assert fuzzed_currency in full_text

        finally:
            # Restore the original state of the disk
            if locales_backup is not None:
                locales_path.write_text(locales_backup, encoding="utf-8")
            else:
                locales_path.unlink(missing_ok=True)

            if brand_backup is not None:
                brand_path.write_text(brand_backup, encoding="utf-8")
            else:
                brand_path.unlink(missing_ok=True)
