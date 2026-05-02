import pytest
from docx import Document

from assessment_engine.scripts.render_global_report_from_template import (
    clean_t_codes,
    render_cover,
)


def test_clean_t_codes():
    """Verifica que el formateador de texto elimina códigos técnicos y limpia espacios."""
    text = "Este es un T2.P3.K1 de prueba y otro T6."
    cleaned = clean_t_codes(text)
    assert "T2.P3.K1" not in cleaned
    assert "T6" not in cleaned
    assert "  " not in cleaned  # No quedan espacios dobles


def test_render_cover_logic():
    """
    Verifica que la lógica de renderizado de la portada funciona
    y no lanza excepciones al recibir un documento vacío y un payload mockeado.
    """
    doc = Document()

    # Creamos un payload mockeado usando Pydantic
    # Necesitamos mockear el resto de la estructura de GlobalReportPayload o usar construct/mocking avanzado
    # Usaremos mock object que respete la interfaz
    class MockMeta:
        client = "TestCorp"
        date = "01 de Enero de 2026"
        version = "v1.0 Test"

    class MockPayload:
        meta = MockMeta()

    payload = MockPayload()

    # Ejecutamos la función de renderizado
    try:
        render_cover(doc, payload)
        success = True
    except Exception as e:
        success = False
        pytest.fail(f"render_cover falló con excepción: {e}")

    assert success

    # Comprobamos mínimamente que el documento ya no está vacío
    assert len(doc.paragraphs) > 0

    # Verificamos que los datos del cliente se insertaron en algún lugar (puede estar en mayúsculas)
    text_content = "\n".join([p.text for p in doc.paragraphs])
    assert "TESTCORP" in text_content.upper()
    assert "01 de Enero de 2026" in text_content
