from __future__ import annotations

import pytest
from docx import Document

from assessment_engine.schemas.blueprint import BlueprintPayload
from assessment_engine.scripts.render_tower_blueprint import (
    render_maturity_profile,
    render_snapshot_page,
)
from assessment_engine.scripts.render_web_presentation import (
    _build_strategy,
    _build_tower_nexus,
    _render_html,
)
from assessment_engine.scripts.run_scoring import build_scoring, resolve_band


def _tower_definition() -> dict:
    return {
        "working_rules": {
            "score_indicator": "mean",
            "score_pillar": "mean",
            "score_tower": "weighted_mean",
        },
        "pillars": [
            {
                "pillar_id": "T1.P1",
                "pillar_name": "Pillar One",
                "weight_pct": 100,
                "kpis": [{"kpi_id": "T1.P1.K1"}],
            }
        ],
        "score_bands": [
            {"min": 1.0, "max": 1.8, "label": "Level 1"},
            {"min": 1.8, "max": 2.6, "label": "Level 2"},
            {"min": 2.6, "max": 3.4, "label": "Level 3"},
            {"min": 3.4, "max": 4.2, "label": "Level 4"},
            {"min": 4.2, "max": 5.0, "label": "Level 5"},
        ],
        "maturity_scale": [],
    }


def _blueprint_payload() -> BlueprintPayload:
    return BlueprintPayload.model_validate(
        {
            "_generation_metadata": {
                "artifact_type": "blueprint_payload",
                "artifact_version": "1.0.0",
            },
            "document_meta": {
                "client_name": "characterization-client",
                "tower_name": "Test Tower 1",
                "tower_code": "T1",
                "financial_tier": "mid",
                "transformation_horizon": "12m",
            },
            "executive_snapshot": {
                "bottom_line": "BP Summary",
                "decisions": ["Decision 1"],
                "cost_of_inaction": "Cost",
                "structural_risks": ["Risk"],
                "business_impact": "Impact",
                "operational_benefits": ["Benefit"],
                "transformation_complexity": "Media",
            },
            "cross_capabilities_analysis": {
                "common_deficiency_patterns": ["Pattern"],
                "transformation_paradigm": "Paradigm",
                "critical_technical_debt": "Debt",
            },
            "roadmap": [{"wave": "Wave 1", "projects": ["Project 1"]}],
            "external_dependencies": [
                {
                    "project": "Project 1",
                    "depends_on": "Dependency 1",
                    "reason": "Because",
                }
            ],
            "pillars_analysis": [
                {
                    "pilar_id": "P1",
                    "pilar_name": "Pillar 1",
                    "score": 3.5,
                    "target_score": 4.0,
                    "health_check_asis": [
                        {
                            "capability": "Capability 1",
                            "finding": "Finding 1",
                            "business_risk": "Risk 1",
                        }
                    ],
                    "target_architecture_tobe": {
                        "vision": "Vision 1",
                        "design_principles": ["Principle 1"],
                    },
                    "projects_todo": [
                        {
                            "name": "Project 1",
                            "business_case": "Business case 1",
                            "tech_objective": "Objective 1",
                            "deliverables": ["Deliverable 1"],
                            "sizing": "M",
                            "duration": "3 months",
                        }
                    ],
                }
            ],
        }
    )


@pytest.mark.parametrize(
    ("score", "expected_label"),
    [
        (1.0, "Level 1"),
        (1.8, "Level 1"),
        (1.8001, "Level 2"),
        (2.6, "Level 2"),
        (2.60004, "Level 2"),
        (2.60005, "Level 2"),
        (2.60015, "Level 3"),
        (3.4, "Level 3"),
        (4.2, "Level 4"),
        (5.1, "Level 5"),
        (0.5, "Level 5"),
    ],
)
def test_scoring_band_mapping_characterizes_boundaries_rounding_and_fallback(
    score: float,
    expected_label: str,
) -> None:
    assert resolve_band(score, _tower_definition())["label"] == expected_label


def test_scoring_build_uses_exact_internal_score_instead_of_display_rounding() -> None:
    scoring = build_scoring(
        case_input={
            "case_id": "case-1",
            "tower_id": "T1",
            "tower_name": "Test Tower",
            "answers": [{"kpi_id": "T1.P1.K1", "value": 2.64994}],
        },
        tower_definition=_tower_definition(),
    )

    assert scoring["tower_score_exact"] == pytest.approx(2.6499)
    assert scoring["tower_score_display_1d"] == 2.6
    assert scoring["maturity_band_from_exact"]["label"] == "Level 3"


def test_blueprint_snapshot_uses_annex_global_band_without_recomputing() -> None:
    doc = Document()

    render_snapshot_page(
        doc,
        _blueprint_payload(),
        client_intelligence={},
        annex_data={
            "executive_summary": {
                "global_score": "2.6 / 5.0",
                "global_band": "Band From Annex",
                "target_maturity": "4.0",
            }
        },
    )

    assert doc.tables[0].rows[1].cells[1].text == "Band From Annex"


def test_blueprint_snapshot_derives_annex_global_band_from_shared_policy() -> None:
    doc = Document()

    render_snapshot_page(
        doc,
        _blueprint_payload(),
        client_intelligence={},
        annex_data={
            "executive_summary": {
                "global_score": "3.5 / 5.0",
                "target_maturity": "4.0",
            }
        },
    )

    assert doc.tables[0].rows[1].cells[1].text == "Gestionado"


def test_blueprint_maturity_profile_uses_annex_pillar_band_without_recomputing() -> (
    None
):
    doc = Document()

    render_maturity_profile(
        doc,
        {
            "pillar_score_profile": {
                "profile_intro": "",
                "scoring_method_note": "",
                "pillars": [
                    {
                        "pillar_label": "Pillar 1",
                        "score_display": "2.6",
                        "maturity_band": "Band From Annex Pillar",
                        "executive_reading": "Reading",
                    }
                ],
            },
            "sections": {
                "asis": {
                    "narrative": "",
                    "strengths": [],
                    "gaps": [],
                    "operational_impacts": [],
                }
            },
        },
    )

    assert doc.tables[0].rows[1].cells[2].text == "Band From Annex Pillar"


def test_blueprint_maturity_profile_derives_pillar_band_from_shared_policy() -> None:
    doc = Document()

    render_maturity_profile(
        doc,
        {
            "pillar_score_profile": {
                "profile_intro": "",
                "scoring_method_note": "",
                "pillars": [
                    {
                        "pillar_label": "Pillar 1",
                        "score_display": "2.6 / 5.0",
                        "executive_reading": "Reading",
                    }
                ],
            },
            "sections": {
                "asis": {
                    "narrative": "",
                    "strengths": [],
                    "gaps": [],
                    "operational_impacts": [],
                }
            },
        },
    )

    assert doc.tables[0].rows[1].cells[2].text == "Definido"


def test_web_dashboard_derives_band_from_score_when_missing() -> None:
    tower_meta = {
        "id": "T1",
        "name": "Test Tower 1",
        "score": "3.5",
        "status_color": "93C47D",
        "executive_message": "Bottom line",
        "target_maturity": "4.0",
    }
    strategy = _build_strategy(
        {
            "executive_summary": {
                "headline": "",
                "narrative": "",
                "key_business_impacts": [],
            },
            "target_vision": {},
            "execution_roadmap": {},
            "heatmap": [tower_meta],
            "meta": {},
        },
        "client",
    )
    tower_nexus = _build_tower_nexus(tower_meta, blueprint_payload=None)

    assert strategy["heatmap"][0]["band"] == "Optimizada"
    assert tower_nexus["meta"]["band"] == "Optimizada"


def test_web_dashboard_keeps_band_as_passthrough_until_template_uppercases_it() -> None:
    tower_meta = {
        "id": "T1",
        "name": "Test Tower 1",
        "score": "3.0",
        "band": "Band MiXa",
        "status_color": "FFD966",
        "executive_message": "Bottom line",
        "target_maturity": "4.0",
    }
    tower_nexus = _build_tower_nexus(tower_meta, blueprint_payload=None)

    html = _render_html(
        "client",
        {
            "meta": {"client": "CLIENT", "version": "v5.1 Strategic Ops"},
            "strategy": {
                "headline": "",
                "narrative": "",
                "global_score": "3.0",
                "burning_platform": [],
                "principles": [],
                "architecture_principles": [],
                "operating_model": [],
                "gtm_strategy": {},
                "stakeholders": [],
                "business_impacts": [],
                "estimated_tam": "TBD",
            },
            "roadmap": {
                "programs": [],
                "horizons": {},
                "proactive_proposals": [],
            },
            "heatmap": [tower_meta],
            "towers": {"T1": tower_nexus},
        },
    )

    assert tower_nexus["meta"]["band"] == "Band MiXa"
    assert '"band": "Band MiXa"' in html
    assert "(t.meta.band || '---').toUpperCase()" in html
