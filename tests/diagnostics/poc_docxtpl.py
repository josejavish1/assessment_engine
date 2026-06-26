import json
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate

def create_dummy_template():
    """We create a dummy template programmatically since we don't have MS Word open."""
    doc = Document()
    doc.add_heading('Anexo AS-IS: {{ tower_name }}', 0)
    
    doc.add_heading('Resumen de Madurez', level=1)
    doc.add_paragraph('Score Global: {{ global_score }}')
    
    doc.add_heading('Registro de Riesgos Técnicos (FAIR)', level=1)
    doc.add_paragraph('{% for pilar in pillars %}')
    doc.add_heading('Dominio: {{ pilar.pilar_name }}', level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Hallazgo (Evidencia)'
    hdr_cells[2].text = 'Coordenadas FAIR'
    hdr_cells[3].text = 'ALE Estimado (€)'
    
    doc.add_paragraph('{% for risk in pilar.risks %}')
    row_cells = table.add_row().cells
    row_cells[0].text = '{{ risk.id }}'
    row_cells[1].text = '{{ risk.text }}'
    row_cells[2].text = '{{ risk.fair }}'
    row_cells[3].text = '{{ risk.ale }}'
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endfor %}')
    
    doc.save('poc_template.docx')

def render_docxtpl():
    # 1. Load data
    payload_path = Path("working/redeia_v3/T2/blueprint_t2_payload.json")
    with open(payload_path, 'r', encoding='utf-8-sig') as f:
        data = json.load(f)
        
    context = {
        "tower_name": data.get("document_meta", {}).get("tower_name", "T2"),
        "global_score": "4.5 / 5.0", # Mock for POC
        "pillars": []
    }
    
    risk_counter = 1
    for p in data.get("pillars_analysis", []):
        pilar_data = {
            "pilar_name": p.get("pilar_name"),
            "risks": []
        }
        for finding in p.get("health_check_asis", []):
            tef = finding.get("threat_event_frequency", 0)
            lm = finding.get("loss_magnitude", 0)
            ale = finding.get("fair_ale_score", 0.0)
            if tef > 0 and lm > 0:
                pilar_data["risks"].append({
                    "id": f"R{risk_counter:02d}",
                    "text": finding.get("risk_observed", ""),
                    "fair": f"TEF:{tef} x LM:{lm}",
                    "ale": f"{ale:,.2f} €"
                })
                risk_counter += 1
        context["pillars"].append(pilar_data)

    # 2. Render Template
    tpl = DocxTemplate('poc_template.docx')
    tpl.render(context)
    tpl.save('poc_docxtpl_output.docx')
    print("✅ docxtpl generó el archivo: poc_docxtpl_output.docx")

if __name__ == "__main__":
    create_dummy_template()
    render_docxtpl()
