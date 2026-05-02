import json
import zipfile

import pytest
from docx import Document
from pydantic import ValidationError

from assessment_engine.scripts.render_tower_annex_from_template import (
    _ensure_radar_chart_path,
    main,
)
from tests.artifact_helpers import ROOT, require_artifact_es

PAYLOAD_PATH = (
    ROOT / "working" / "ivirma" / "T5" / "approved_annex_t5.template_payload.json"
)
TEMPLATE_PATH = (
    ROOT / "templates" / "Template_Documento_Anexos_Alpha_v06_Tower_Annex_v2_6.docx"
)


def test_render_tower_annex_from_real_payload(tmp_path):
    output_path = tmp_path / "annex_t5_test.docx"

    main(
        [
            "render_tower_annex_from_template",
            str(require_artifact_es(PAYLOAD_PATH)),
            str(TEMPLATE_PATH),
            str(output_path),
        ]
    )

    assert output_path.exists()

    doc = Document(output_path)
    text_content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    assert "ANEXO T5" in text_content
    assert "Resumen ejecutivo de la torre" in text_content

    title = next(
        p
        for p in doc.paragraphs
        if p.text.strip() == "ANEXO T5 – Resilience & Continuity"
    )
    assert title.style.style_id == "Ttulo"


def test_render_tower_annex_semantic_mode_uses_word_styles(tmp_path):
    output_path = tmp_path / "annex_t5_semantic.docx"

    main(
        [
            "render_tower_annex_from_template",
            str(require_artifact_es(PAYLOAD_PATH)),
            str(TEMPLATE_PATH),
            str(output_path),
            "--semantic-styles",
        ]
    )

    doc = Document(output_path)
    paragraphs = {p.text.strip(): p for p in doc.paragraphs if p.text.strip()}
    subtitle = next(
        p
        for p in doc.paragraphs
        if p.text.strip() and "Fast Infrastructure Assessment" in p.text
    )

    assert paragraphs["ANEXO T5 – Resilience & Continuity"].style.style_id == "Ttulo"
    assert subtitle.style.style_id == "Subttulo"
    assert paragraphs["Resumen ejecutivo de la torre"].style.style_id == "Ttulo1"
    assert paragraphs["Perfil de madurez por pilar"].style.style_id == "Ttulo1"
    assert not any(
        p.text.strip().startswith("• ") for p in doc.paragraphs if p.text.strip()
    )

    with zipfile.ZipFile(output_path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8", "ignore")
        settings_xml = zf.read("word/settings.xml").decode("utf-8", "ignore")

    assert 'TOC \\o "1-2" \\h \\z \\u' in document_xml
    assert "<w:numPr>" in document_xml
    assert "updateFields" in settings_xml


def test_render_tower_annex_can_opt_into_legacy_mode(tmp_path):
    output_path = tmp_path / "annex_t5_legacy.docx"

    main(
        [
            "render_tower_annex_from_template",
            str(require_artifact_es(PAYLOAD_PATH)),
            str(TEMPLATE_PATH),
            str(output_path),
            "--legacy-styles",
        ]
    )

    doc = Document(output_path)
    title = next(
        p
        for p in doc.paragraphs
        if p.text.strip() == "ANEXO T5 – Resilience & Continuity"
    )

    assert title.style.style_id != "Ttulo"


def test_ensure_radar_chart_path_generates_missing_chart(tmp_path):
    payload_path = tmp_path / "approved_annex_t2.template_payload.json"
    profile = {
        "radar_chart": "BASE64_ENCODED_IMAGE_PLACEHOLDER",
        "pillars": [
            {"pillar_label": "P1", "score_display": "3.4"},
            {"pillar_label": "P2", "score_display": "3.7"},
        ],
    }

    chart_path = _ensure_radar_chart_path(payload_path, profile)

    assert chart_path.endswith("pillar_radar_chart.generated.png")
    assert (tmp_path / "pillar_radar_chart.generated.png").exists()


def test_render_tower_annex_fails_on_invalid_payload(tmp_path):
    payload_path = tmp_path / "invalid_annex.json"
    payload_path.write_text(
        json.dumps(
            {
                "document_meta": {"tower_code": "T5", "tower_name": "Resilience"},
                "executive_summary": {
                    "global_score": "3.5",
                    "global_band": "Managed",
                },
            }
        ),
        encoding="utf-8",
    )

    with pytest.raises(ValidationError):
        main(
            [
                "render_tower_annex_from_template",
                str(payload_path),
                str(TEMPLATE_PATH),
                str(tmp_path / "unused.docx"),
            ]
        )
