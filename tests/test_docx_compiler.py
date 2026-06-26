import tempfile
from pathlib import Path

from docx import Document

from assessment_engine.adapters.compilers.docx_compiler import DocxCompiler
from assessment_engine.domain.schemas.ast import (
    CellNode,
    DocumentAST,
    HeadingNode,
    PageBreakNode,
    ParagraphNode,
    SpacerNode,
    TableNode,
    TableRowNode,
)


def test_docx_compiler_e2e() -> None:
    """
    Verify e2e that the DocxCompiler can compile a complex DocumentAST
    into a physical .docx file while accurately preserving styles and cell properties.
    """
    # --- ARRANGE ---
    ast = DocumentAST(
        title="Diagnostic Technical Report",
        metadata={"client": "Acme Corp", "version": "1.0"},
        nodes=[
            HeadingNode(
                text="1. Executive Summary", level=1, primary_color_rgb=[0, 114, 188]
            ),
            ParagraphNode(
                text="This is a fuzzed diagnostic paragraph to evaluate AST stability.",
                bold=True,
                style="Normal",
                space_after=12,
            ),
            SpacerNode(points=18),
            TableNode(
                rows=[
                    TableRowNode(
                        cells=[
                            CellNode(
                                text="Pilar",
                                bold=True,
                                bg_color="D9EAF7",
                                align="CENTER",
                            ),
                            CellNode(
                                text="Score",
                                bold=False,
                                bg_color="FFFFFF",
                                align="RIGHT",
                            ),
                        ],
                        is_header=True,
                    ),
                    TableRowNode(
                        cells=[
                            CellNode(
                                text="Compute",
                                bold=False,
                                bg_color="F2F2F2",
                                align="LEFT",
                            ),
                            CellNode(
                                text="4.5", bold=True, bg_color="FFF3CD", align="RIGHT"
                            ),
                        ]
                    ),
                ]
            ),
            PageBreakNode(),
        ],
    )

    compiler = DocxCompiler()

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_file = Path(tmp_dir) / "output.docx"

        # --- ACT ---
        compiler.compile(ast, str(output_file))

        # --- ASSERT ---
        assert output_file.is_file(), (
            "The compiled file must be physically saved on disk."
        )

        # Load generated document to validate internal structure
        doc = Document(str(output_file))

        # Verify headings and paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "1. Executive Summary" in paragraphs
        assert (
            "This is a fuzzed diagnostic paragraph to evaluate AST stability."
            in paragraphs
        )

        # Verify bold paragraph formatting
        bold_para = next(p for p in doc.paragraphs if "stability" in p.text)
        assert len(bold_para.runs) > 0
        assert bold_para.runs[0].bold is True

        # Verify tables and cells
        assert len(doc.tables) == 1, "Exactly one compiled table must exist."
        table = doc.tables[0]
        assert len(table.rows) == 2, "The table must have exactly 2 rows."
        assert len(table.columns) == 2, "The table must have exactly 2 columns."
        assert len(table.rows[0].cells) == 2, "The first row must have 2 cells."
        assert len(table.rows[1].cells) == 2, "The second row must have 2 cells."

        # Row 1 (Header)
        assert table.rows[0].cells[0].text == "Pilar"
        assert table.rows[0].cells[1].text == "Score"

        # Row 2 (Data)
        assert table.rows[1].cells[0].text == "Compute"
        assert table.rows[1].cells[1].text == "4.5"
