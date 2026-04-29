import hashlib
import json
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from assessment_engine.schemas.annex_synthesis import AnnexPayload
from assessment_engine.schemas.blueprint import BlueprintPayload


ROOT = Path(__file__).resolve().parents[1]
T5_DIR = ROOT / "working" / "smoke_ivirma" / "T5"
BLUEPRINT_PAYLOAD = T5_DIR / "blueprint_t5_payload.json"
ANNEX_PAYLOAD = T5_DIR / "approved_annex_t5.template_payload.json"
BLUEPRINT_DOCX = T5_DIR / "Blueprint_Transformacion_T5_smoke_ivirma.docx"
ANNEX_DOCX = T5_DIR / "annex_t5_smoke_ivirma_final.docx"
RADAR_PNG = T5_DIR / "pillar_radar_chart.generated.png"


def _require(path: Path) -> Path:
    if not path.exists():
        pytest.skip(f"Missing artifact: {path}")
    return path


def _load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def _docx_xml_and_rels(path: Path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", "ignore")
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8", "ignore")
    return xml, rels


def _docx_words(path: Path) -> int:
    doc = Document(path)
    total = 0
    for paragraph in doc.paragraphs:
        total += len(paragraph.text.split())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += len(cell.text.split())
    return total


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _sha256_file(path: Path) -> str:
    return _sha256_bytes(path.read_bytes())


def test_t5_blueprint_payload_schema_and_shape():
    data = _load_json(_require(BLUEPRINT_PAYLOAD))
    payload = BlueprintPayload.model_validate(data)
    assert payload.document_meta.tower_code == "T5"
    assert payload.document_meta.tower_name == "Resilience & Continuity"
    assert len(payload.pillars_analysis) == 5
    assert len(payload.executive_snapshot.decisions) >= 3
    assert len(payload.roadmap) >= 3
    # Check versioning metadata
    assert payload.generation_metadata is not None
    assert payload.generation_metadata.artifact_type == "blueprint_payload"


def test_t5_annex_payload_schema_and_executive_limits():
    data = _load_json(_require(ANNEX_PAYLOAD))
    payload = AnnexPayload.model_validate(data)
    assert payload.document_meta["tower_code"] == "T5"
    assert payload.document_meta["tower_name"] == "Resilience & Continuity"
    # Check versioning metadata
    assert payload.generation_metadata is not None
    assert payload.generation_metadata.artifact_type in ["annex_payload", "annex_template_payload"]
    assert payload.pillar_score_profile.radar_chart.endswith(
        "pillar_radar_chart.generated.png"
    )
    assert len(data["sections"]["risks"]["risks"]) <= 6
    assert len(data["sections"]["gap"]["gap_rows"]) <= 6
    assert len(data["sections"]["gap"]["target_capabilities"]) <= 5
    assert len(data["sections"]["todo"]["priority_initiatives"]) <= 6
    summary = data["executive_summary"]["summary_body"]
    if isinstance(summary, list):
        summary = " ".join(summary)
    summary_lower = summary.lower()
    assert any(x in summary_lower for x in ["m&a", "adquis", "crecimiento", "expansión"])
    assert any(x in summary_lower for x in ["ia", "datos", "digital", "tecnológ"])


def test_t5_annex_docx_has_no_functional_placeholders():
    xml, _ = _docx_xml_and_rels(_require(ANNEX_DOCX))
    forbidden = [
        "{{RISKS_TABLE}}",
        "{{GAP_TABLE}}",
        "{{PRIORITY_INITIATIVES_CARDS}}",
        "{{RADAR_CHART_BLOCK}}",
        "[Fecha]",
        "T[X]",
        "[Nombre de la Torre]",
    ]
    for token in forbidden:
        assert token not in xml


def test_t5_annex_docx_embeds_real_radar_chart():
    annex_path = _require(ANNEX_DOCX)
    radar_path = _require(RADAR_PNG)
    with zipfile.ZipFile(annex_path) as zf:
        document = ET.fromstring(zf.read("word/document.xml"))
        rels = ET.fromstring(zf.read("word/_rels/document.xml.rels"))
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        relmap = {rel.attrib["Id"]: rel.attrib["Target"] for rel in rels}
        paragraphs = document.findall(".//w:p", ns)

        radar_target = None
        for idx, paragraph in enumerate(paragraphs):
            text = "".join(t.text or "" for t in paragraph.findall(".//w:t", ns))
            if "Gráfico radial" in text:
                for next_paragraph in paragraphs[idx : idx + 4]:
                    for blip in next_paragraph.findall(".//a:blip", ns):
                        embed = blip.attrib.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                        if embed:
                            radar_target = relmap.get(embed)
                            break
                    if radar_target:
                        break
                break

        assert radar_target is not None
        embedded = zf.read(f"word/{radar_target}")
    assert _sha256_bytes(embedded) == _sha256_file(radar_path)


def test_t5_blueprint_and_annex_roles_are_distinct():
    blueprint_words = _docx_words(_require(BLUEPRINT_DOCX))
    annex_words = _docx_words(_require(ANNEX_DOCX))
    assert blueprint_words > annex_words
    assert annex_words < 5000
    assert blueprint_words > 7000


def test_t5_blueprint_opening_contains_business_layer():
    doc = Document(_require(BLUEPRINT_DOCX))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    assert "Por qué importa al negocio" in text
    assert "Riesgos de negocio más materiales" in text
    assert "Decisiones prioritarias" in text

