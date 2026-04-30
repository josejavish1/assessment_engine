from pathlib import Path
from xml.etree import ElementTree as ET
import zipfile

from docx import Document
import pytest


ROOT = Path(__file__).resolve().parents[1]


def _resolve_t5_docx_pair() -> tuple[Path, Path]:
    preferred_pairs = [
        (
            ROOT / "working" / "smoke_ivirma" / "T5"
            / "Blueprint_Transformacion_T5_smoke_ivirma.docx",
            ROOT / "working" / "smoke_ivirma" / "T5"
            / "annex_t5_smoke_ivirma_final.docx",
        ),
        (
            ROOT / "working" / "smoke_moeve" / "T5"
            / "Blueprint_Transformacion_T5_smoke_moeve.docx",
            ROOT / "working" / "smoke_moeve" / "T5"
            / "annex_t5_smoke_moeve_final.docx",
        ),
        (
            ROOT / "working" / "ivirma" / "T5"
            / "Blueprint_Transformacion_T5_ivirma.docx",
            ROOT / "working" / "ivirma" / "T5"
            / "annex_t5_ivirma_final.docx",
        ),
    ]
    for blueprint, annex in preferred_pairs:
        if blueprint.exists() and annex.exists():
            return blueprint, annex
    pytest.skip(
        "No hay artefactos DOCX T5 disponibles en working/ para validar integridad.",
        allow_module_level=True,
    )


BLUEPRINT_DOCX, ANNEX_DOCX = _resolve_t5_docx_pair()


def _parse_docx(path: Path):
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        document = ET.fromstring(zf.read("word/document.xml"))
        rels = ET.fromstring(zf.read("word/_rels/document.xml.rels"))
    return names, document, rels


def test_t5_word_documents_open_cleanly():
    for path in [BLUEPRINT_DOCX, ANNEX_DOCX]:
        doc = Document(path)
        assert len(doc.paragraphs) + len(doc.tables) > 0


def test_t5_docx_relationship_targets_exist():
    for path in [BLUEPRINT_DOCX, ANNEX_DOCX]:
        names, _, rels = _parse_docx(path)
        for rel in rels:
            target = rel.attrib.get("Target", "")
            if not target or "://" in target or target.startswith("http"):
                continue
            if target.startswith("/"):
                continue
            internal = f"word/{target}" if not target.startswith("word/") else target
            if target.startswith("../"):
                internal = target.replace("../", "", 1)
            assert internal in names, f"Missing relationship target {target} in {path.name}"


def test_t5_documents_contain_expected_executive_headings():
    blueprint_doc = Document(BLUEPRINT_DOCX)
    annex_doc = Document(ANNEX_DOCX)

    blueprint_text = "\n".join(p.text for p in blueprint_doc.paragraphs if p.text.strip())
    annex_text = "\n".join(p.text for p in annex_doc.paragraphs if p.text.strip())

    assert "Executive Snapshot" in blueprint_text
    assert (
        "Impacto de Inacción (Cost of Inaction)" in blueprint_text
        or "Por qué importa al negocio" in blueprint_text
    )
    assert (
        "Decisiones Ejecutivas Bloqueantes" in blueprint_text
        or "Decisiones prioritarias" in blueprint_text
    )

    assert "Resumen ejecutivo de la torre" in annex_text
    assert "Riesgos principales" in annex_text
    assert "Iniciativas prioritarias" in annex_text
