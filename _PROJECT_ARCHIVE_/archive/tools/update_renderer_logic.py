import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def process_footnotes(doc, dossier):
    evidences = dossier.get("evidences", [])
    if not evidences:
        return
    
    ref_pattern = re.compile(r'\[\[REF:(\d+)\]\]')
    found_refs = set()
    
    def process_paragraph_refs(p):
        if "[[" in p.text:
            matches = ref_pattern.findall(p.text)
            for m in matches:
                idx = int(m)
                found_refs.add(idx)
                old_text = f"[[REF:{idx}]]"
                for run in p.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, f" [{idx}]")
                        run.font.superscript = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 114, 188)

    for p in doc.paragraphs:
        process_paragraph_refs(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph_refs(p)

    if found_refs:
        doc.add_page_break()
        title = doc.add_paragraph()
        run = title.add_run("Fuentes y Evidencias de Inteligencia")
        run.bold = True; run.font.size = Pt(14); run.font.name = "Georgia"; run.font.color.rgb = RGBColor(0, 114, 188)
        
        for idx in sorted(list(found_refs)):
            if idx <= len(evidences):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(4)
                r_idx = p.add_run(f"[{idx}] ")
                r_idx.bold = True; r_idx.font.size = Pt(9); r_idx.font.color.rgb = RGBColor(46, 64, 77)
                r_text = p.add_run(evidences[idx-1])
                r_text.font.size = Pt(9); r_text.font.color.rgb = RGBColor(127, 127, 127)
