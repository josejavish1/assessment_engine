from docx import Document
from pathlib import Path


def replace_text_in_doc(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            p.text = p.text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        p.text = p.text.replace(old_text, new_text)


print("Función lista.")
